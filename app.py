"""
╔══════════════════════════════════════════════════════════════════════════════╗
║  Yev Capital | LoanLogic v3.0                                                ║
║  Credit Analysis & Reporting System                                          ║
║                                                                              ║
║  Copyright (c) 2026 Bohdan Yevtushenko (MrCemper)                            ║
║  SPDX-License-Identifier: AGPL-3.0-or-later                                  ║
║                                                                              ║
║  This program is free software: you can redistribute it and/or modify        ║
║  it under the terms of the GNU Affero General Public License as published    ║
║  by the Free Software Foundation, either version 3 of the License, or        ║
║  (at your option) any later version.                                         ║
║                                                                              ║
║  This program is distributed in the hope that it will be useful,             ║
║  but WITHOUT ANY WARRANTY; without even the implied warranty of              ║
║  MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the                ║
║  GNU Affero General Public License for more details.                         ║
║                                                                              ║
║  You should have received a copy of the GNU Affero General Public License    ║
║  along with this program. If not, see <https://www.gnu.org/licenses/>.       ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""
import io
import math
import re
import html
import hashlib
from datetime import datetime, date, timedelta

import streamlit as st
import pandas as pd
import plotly.graph_objects as go

from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors as rl_colors
from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                 Paragraph, Spacer, HRFlowable)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ─────────────────────────────────────────────────────────────────────────────
#  Cyrillic-capable PDF font registration
#  Search order:
#    0) DejaVuSans.ttf next to this script (project-local)
#    1) DejaVuSans   — system path on Linux / Mac / Docker
#    2) Arial        — Windows  C:\Windows\Fonts\arial.ttf
#    3) FreeSans     — Linux fallback
#    4) Helvetica    — built-in PDF font (no Cyrillic, but never crashes)
# ─────────────────────────────────────────────────────────────────────────────
import os as _os

# Resolve the directory of the current script. Used to look for a bundled
# font file next to the .py. In a small number of environments (REPL / exec)
# __file__ may not be defined; fall back to the current working directory.
try:
    _PROJECT_DIR = _os.path.dirname(_os.path.abspath(__file__))
except NameError:
    _PROJECT_DIR = _os.getcwd()

_FONT_CANDIDATES = [
    # (regular_path, bold_path, family_name)
    # 0) Project-local — checked first so a bundled font next to the script
    #    takes precedence over system fonts.
    (_os.path.join(_PROJECT_DIR, "DejaVuSans.ttf"),
     _os.path.join(_PROJECT_DIR, "DejaVuSans-Bold.ttf"),
     "DejaVuSans"),
    # 1..3) System paths
    ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
     "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
     "DejaVuSans"),
    ("C:\\Windows\\Fonts\\arial.ttf",
     "C:\\Windows\\Fonts\\arialbd.ttf",
     "Arial"),
    ("/Library/Fonts/Arial.ttf",
     "/Library/Fonts/Arial Bold.ttf",
     "Arial"),
    ("/usr/share/fonts/truetype/freefont/FreeSans.ttf",
     "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf",
     "FreeSans"),
]

PDF_FONT         = "Helvetica"        # fallback: встроенный, без кириллицы
PDF_FONT_BOLD    = "Helvetica-Bold"
PDF_FONT_ITALIC  = "Helvetica-Oblique"
PDF_FONT_WARN    = None               # предупреждение, если шрифт не найден

for _reg, _bold, _family in _FONT_CANDIDATES:
    try:
        if _os.path.exists(_reg):
            pdfmetrics.registerFont(TTFont(_family, _reg))
            if _os.path.exists(_bold):
                pdfmetrics.registerFont(TTFont(f"{_family}-Bold", _bold))
                PDF_FONT_BOLD = f"{_family}-Bold"
            else:
                PDF_FONT_BOLD = _family
            PDF_FONT        = _family
            PDF_FONT_ITALIC = _family   # italic не критичен, используем regular
            PDF_FONT_WARN   = None
            break
    except Exception as _e:
        PDF_FONT_WARN = f"Шрифт {_family} не зарегистрирован: {_e}"

if PDF_FONT == "Helvetica":
    PDF_FONT_WARN = (
        "⚠️ Кириллический шрифт не найден на этом устройстве. "
        "В PDF-отчётах кириллические символы могут отображаться как □. "
        "Чтобы это исправить — установите шрифт DejaVuSans или Arial "
        "в систему."
    )

# ─────────────────────────────────────────────────────────────────────────────
#  ЛОКАЛИЗАЦИЯ
# ─────────────────────────────────────────────────────────────────────────────
TRANSLATIONS = {
    "ru": {
        "app_title": "Yev Capital LoanLogic",
        "app_subtitle": "Расчёт банковских кредитов и вкладов",
        "section_params": "⚙️ Параметры кредита",
        "loan_amount": "Сумма кредита",
        "loan_amount_slider": "Ползунок суммы",
        "loan_term": "Срок",
        "term_unit": "Единица срока",
        "weeks": "Недели", "months": "Месяцы", "quarters": "Кварталы",
        "halfyears": "Полугодия", "years": "Годы",
        "interest_rate": "Процентная ставка (% годовых)",
        "calc_scheme": "Схема расчёта",
        "annuity": "Аннуитет",
        "classic": "Классика (дифференцированный)",
        "balloon": "Буллит (Balloon Payment)",
        "deposit_scheme": "Вклад / Депозит",
        "section_commissions": "💼 Комиссии",
        "one_time_comm": "Разовая комиссия",
        "one_time_type": "Тип разовой комиссии",
        "monthly_comm": "Ежемесячная комиссия",
        "monthly_type": "Тип ежемесячной комиссии",
        "pct_of_amount": "% от суммы кредита",
        "fixed_amount": "Фиксированная сумма",
        "calc_btn": "🚀 Рассчитать",
        "calc_error": "⚠️ Не удалось выполнить расчёт с заданными параметрами. Проверьте корректность ввода (срок, ставка, сумма).",
        "section_results": "📊 Результаты расчёта",
        "total_payment": "Общая сумма выплат",
        "total_interest": "Переплата по процентам",
        "total_commission": "Сумма комиссий",
        "effective_rate": "Эффективная ставка",
        "monthly_payment": "Платёж (1-й период)",
        "section_schedule": "📋 График платежей",
        "period": "Период", "date": "Дата",
        "balance_open": "Остаток (нач.)",
        "payment_total": "Платёж",
        "principal": "Тело долга",
        "interest": "Проценты",
        "commission": "Комиссия",
        "balance_close": "Остаток (кон.)",
        "total_row": "ИТОГО",
        "section_chart": "📈 Диаграмма",
        "chart_title": "Структура выплат по периодам",
        "chart_principal": "Тело кредита",
        "chart_interest": "Проценты",
        "chart_commission": "Комиссии",
        "chart_pie_title": "Итоговое соотношение",
        "section_templates": "💾 Шаблоны расчётов",
        "template_name": "Название шаблона",
        "save_template": "💾 Сохранить",
        "load_template": "📂 Загрузить",
        "delete_template": "🗑️",
        "no_templates": "Нет сохранённых шаблонов",
        "template_saved": "✅ Шаблон сохранён!",
        "template_loaded": "✅ Шаблон загружен!",
        "template_deleted": "🗑️ Удалён!",
        "template_name_empty": "Введите название шаблона",
        "template_overwrite_warn": "⚠️ Шаблон «{name}» уже существует. Нажмите «Сохранить» ещё раз, чтобы перезаписать.",
        "currency": "Валюта",
        "uah": "₴ Гривна (UAH)", "usd": "$ Доллар (USD)", "eur": "€ Евро (EUR)",
        "rub": "₽ Рубль (RUB)", "gbp": "£ Фунт стерлингов (GBP)",
        "jpy": "¥ Иена (JPY)", "cad": "C$ Канадский доллар (CAD)",
        "aud": "A$ Австралийский доллар (AUD)", "chf": "Fr Швейцарский франк (CHF)",
        "cny": "¥ Китайский юань (CNY)", "hkd": "HK$ Гонконгский доллар (HKD)",
        "custom": "✏️ Своя валюта",
        "custom_symbol": "Символ своей валюты",
        "download_excel": "⬇️ Excel",
        "download_pdf": "⬇️ PDF",
        "download_docx": "⬇️ Word",
        "overpayment_pct": "Переплата от суммы",
        "compare_schemes": "⚖️ Сравнить схемы",
        "annuity_vs_classic": "Сравнение схем кредитования",
        "savings_classic": "Экономия при классике",
        "annuity_short": "Аннуитет",
        "classic_short": "Классика",
        "enter_name": "Введите название...",
        "template_examples": "Быстрые шаблоны:",
        "mortgage": "🏠 Ипотека",
        "car_loan": "🚗 Автокредит",
        "consumer": "💳 Потребительский",
        "deposit": "Вклад",
        # Инвестиции
        "invest_section": "📈 Сравнение с инвестициями",
        "invest_sp500": "📊 Сравнить с S&P 500 (13.7% годовых)",
        "invest_custom": "💹 Сравнить с кастомным % инвестирования",
        "invest_custom_rate": "Годовая доходность (%)",
        "invest_chart_title": "Кредит vs Инвестиции",
        "invest_loan_balance": "Остаток долга",
        "invest_portfolio": "Инвестиционный портфель",
        "invest_total_paid": "Выплачено банку",
        "invest_portfolio_val": "Стоимость портфеля",
        "invest_net_gain": "Чистая прибыль",
        "invest_explanation": "Что было бы, если вместо кредитных платежей инвестировать эти деньги?",
        "invest_rate_label": "доходность",
        "invest_only_one": "⚠️ Можно включить только одно сравнение! Выберите одно.",
        # Вклад
        "deposit_mode_label": "Режим вклада",
        "deposit_capitalize": "💰 Капитализация процентов (сложный %)",
        "deposit_payout": "💸 Выплата процентов",
        "deposit_final": "Итоговая сумма вклада",
        "deposit_earned": "Начислено процентов",
        # Тултипы
        "tooltip_payment": "💡 **Платёж** — полная сумма, которую нужно перечислить в этот период (тело долга + проценты + комиссия).",
        "tooltip_principal": "💡 **Тело долга** — часть платежа, которая идёт на погашение основной суммы кредита.",
        "tooltip_interest": "💡 **Проценты** — плата за пользование деньгами банка, начисляется на текущий остаток долга.",
        # Схемы
        "scheme_annuity_info": "Равные платежи весь срок. Удобно планировать бюджет.",
        "scheme_classic_info": "Фиксированное тело долга + убывающие проценты. Меньшая переплата.",
        "scheme_balloon_info": "Только проценты каждый период. Вся сумма долга — в последнем платеже.",
        "scheme_deposit_info": "Расчёт вклада/депозита с капитализацией или выплатой процентов.",
        "col_explain": "ℹ️ Пояснение колонок",
        # Депозит — колонки таблицы
        "dep_balance_open":    "Баланс (нач.)",
        "dep_interest_earned": "Начислено %",
        "dep_payout":          "Выплата",
        "dep_balance_close":   "Баланс (кон.)",
        "dep_total_row":       "ИТОГО",
        # Депозит — итоги
        "dep_initial":         "Сумма вклада",
        "dep_final_balance":   "Итоговая сумма",
        "dep_total_earned":    "Начислено процентов",
        "dep_rate_label":      "Ставка годовых",
        "dep_period_payout":   "Выплата за период",
        "dep_growth_title":    "📈 Рост вклада",
        "dep_balance_label":   "Баланс вклада",
        "dep_interest_label":  "Начисленные проценты",
        "dep_vs_invest":       "Вклад vs Альтернатива",
        "dep_your_deposit":    "Ваш вклад",
        "dep_alternative":     "Альтернативные инвестиции",
        "dep_invest_exp":      "Сравнение: если бы ту же сумму вложить под другой процент",
        "dep_sp500_exp":       "Сравнение вашего вклада с доходностью S&P 500 (13.7% годовых)",
        "dep_tooltip_interest":"💡 **Начислено %** — сумма процентов за период (при капитализации добавляется к балансу, при выплате — зачисляется на счёт).",
        "dep_tooltip_balance": "💡 **Баланс** — сумма на вкладе с учётом накопленных процентов (при капитализации растёт каждый период).",
        "dep_tooltip_payout":  "💡 **Выплата** — сумма процентов, фактически полученных на руки в этот период.",
        # CSV
        "download_csv": "⬇️ CSV",
        # Инвест — для депозита
        "dep_invest_section":  "📊 Сравнение с альтернативой",
        "dep_invest_caption":  "Что если вместо этого вклада инвестировать под другой % ?",
        "dep_invest_yours":    "Ваш депозит (итог)",
        "dep_invest_alt":      "Альтернатива (итог)",
        "dep_invest_diff":     "Разница",
        # Дата начала
        "start_date_label":    "📅 Дата начала",
        "start_date_hint":     "Первый платёж будет сделан через один период",
        "term_mode_manual":    "Вручную",
        "term_mode_enddate":   "По дате окончания",
        "end_date_label":      "📅 Дата окончания",
        "end_date_hint":       "Срок пересчитается автоматически в полных периодах",
        # Заголовки графиков
        "chart_balance_title": "📉 Остаток долга",
        "chart_balance_hover": "Остаток",
        "balloon_short":       "Буллит",
        # Balloon Break-even
        "balloon_breakeven":       "Точка инвест-безубыточности",
        "balloon_breakeven_tip":   (
            "Минимальная годовая доходность инвестиций (сложный процент), "
            "при которой будущая стоимость ежемесячно 'сэкономленных' сумм "
            "(разница между аннуитетным платежом и буллитным) к концу срока "
            "равна сумме основного долга. Ниже этой ставки — буллит невыгоден."
        ),
        "balloon_breakeven_label": "Анализ безубыточности (Буллит)",
        "balloon_breakeven_desc":  "Мин. доходность инвестиций для оправдания буллитной схемы",
        # Срок — подпись
        "term_caption":        "мес. / лет",
        # Приветственный экран
        "welcome_h2":  "Введите параметры и нажмите",
        "welcome_calc": "Рассчитать",
        "welcome_sub": "Аннуитет · Классика · Буллит · Вклад<br>Экспорт Excel / PDF / Word / CSV · Сравнение с инвестициями",
        "tab_balance": "📉 Остаток долга",
        "copyright": "© 2026 Bohdan Yevtushenko (MrCemper) · Yev Capital LoanLogic v3.0",
        # ── Расширенная инвест-безубыточность ─────────────────────────────────
        "invest_breakeven_section":      "📐 Анализ инвест-безубыточности",
        "invest_breakeven_universal":    "Универсальная ставка безубыточности",
        "invest_breakeven_universal_tip":"Годовая доходность, при которой реинвестирование платежей полностью перекроет переплату по процентам.",
        "invest_breakeven_abs":          "Абсолютная безубыточность (Буллит)",
        "invest_breakeven_abs_tip":      "Минимальная доходность, при которой FV сэкономленных средств покрывает всю переплату по балунному кредиту.",
        "invest_breakeven_vs_ann":       "Vs. Аннуитет (Буллит)",
        "invest_breakeven_vs_ann_tip":   "Минимальная доходность, при которой буллит выгоднее аннуитета.",
        "invest_breakeven_vs_cla_balloon":     "Vs. Классика (Буллит)",
        "invest_breakeven_vs_cla_balloon_tip": "Минимальная годовая доходность, при которой буллитная схема (с отложенной выплатой тела) к концу срока выходит на тот же результат, что и классическая.",
        "invest_breakeven_vs_cla_annuity":     "Vs. Классика (Аннуитет)",
        "invest_breakeven_vs_cla_annuity_tip": "Минимальная годовая доходность, при которой инвестирование сэкономленных средств (аннуитет в ранние периоды дешевле классики) компенсирует переплату в поздние периоды.",
        "annuity_be_caption":            "Если вы можете заработать ≥ этой ставки на сэкономленных в ранние периоды деньгах — аннуитет к концу срока выйдет на тот же или лучший результат, чем классика.",
        "sp500_disclaimer": "Примечание: доходность S&P 500 в 13.7% основана на среднегодовом показателе за последние 10 лет. Прошлые результаты не гарантируют будущих. Значение используется как статистический ориентир и может корректироваться согласно текущим рыночным ожиданиям.",
        # ── Инфляция ──────────────────────────────────────────────────────────
        "inflation_expander":      "Учёт инфляции",
        "inflation_toggle":        "Включить учёт инфляции",
        "inflation_rate":          "Ожидаемая инфляция (% годовых)",
        "inflation_section":       "💰 Реальная vs Номинальная стоимость",
        "nominal_cost":            "Номинальная стоимость",
        "real_cost":               "Реальная стоимость (PV)",
        "inflation_savings":       "Скидка на инфляцию",
        "inflation_note":          "Реальная стоимость дисконтирует будущие платежи к сегодняшней покупательной способности.",
        # ── Кредитные каникулы ────────────────────────────────────────────────
        "grace_expander":          "Кредитные каникулы",
        "grace_toggle":            "Включить кредитные каникулы",
        "grace_duration":          "Длительность (периодов)",
        "grace_start":             "Начальный период",
        "grace_type":              "Тип",
        "grace_interest_only":     "Только проценты",
        "grace_full_holiday":      "Полные каникулы (капитализация %)",
        "grace_note":              "При полных каникулах начисленные проценты добавляются к основной сумме долга.",
        "grace_active":            "Активный период каникул",
        # ── DSCR ──────────────────────────────────────────────────────────────
        "dscr_expander":           "Бизнес-анализ (DSCR)",
        "dscr_toggle":             "Включить анализ DSCR",
        "dscr_noi":                "Ежемесячный NOI",
        "dscr_label":              "DSCR",
        "dscr_status_safe":        "Безопасно",
        "dscr_status_warning":     "Предупреждение",
        "dscr_status_risk":        "Высокий риск",
        "ltv_status_safe":         "Безопасный",
        "ltv_status_standard":     "Стандартный",
        "ltv_status_high":         "Высокий",
        "ltv_status_critical":     "Критический",
        "ltv_status_na":           "Н/Д",
        "credit_health_title":     "Кредитное здоровье",
        "credit_rank_a":           "Ранг A — Отлично",
        "credit_rank_b":           "Ранг B — Хорошо",
        "credit_rank_c":           "Ранг C — Удовлетворительно",
        "credit_rank_d":           "Ранг D — Высокий риск",
        "credit_rank_na":          "Н/Д",
        "credit_health_caption":   "Балл объединяет указанные вами метрики ({metrics}). Чем выше — тем здоровее.",
        "glossary_title":          "📚 Словарь терминов / Glossary",
        "help_negative_rate":      "Отрицательные ставки допускаются (редкие продукты с субнулевой ставкой). Расчёт корректно их обрабатывает.",
        "gloss_annuity_term":      "Аннуитет",
        "gloss_annuity_def":       "Схема погашения с равными платежами каждый период. В начале большую часть платежа составляют проценты, к концу — тело кредита.",
        "gloss_bullet_term":       "Буллит (Bullet)",
        "gloss_bullet_def":        "Кредит, где в течение срока платятся только проценты, а всё тело долга возвращается одним платежом в конце.",
        "gloss_classic_term":      "Классическая (дифференцированная)",
        "gloss_classic_def":       "Тело кредита гасится равными долями, а проценты начисляются на остаток. Платёж уменьшается со временем.",
        "gloss_dscr_term":         "DSCR",
        "gloss_dscr_def":          "Debt Service Coverage Ratio = чистый доход / обслуживание долга. Показывает, во сколько раз доход покрывает выплаты. ≥1.25 — безопасно.",
        "gloss_ltv_term":          "LTV",
        "gloss_ltv_def":           "Loan-to-Value = кредит / стоимость залога × 100%. Чем ниже, тем меньше риск для кредитора. ≤80% обычно считается хорошим.",
        "gloss_dti_term":          "DTI",
        "gloss_dti_def":           "Debt-to-Income = (платёж + другие долги) / доход × 100%. Доля дохода, уходящая на долги. ≤36% обычно приемлемо.",
        "gloss_npv_term":          "NPV",
        "gloss_npv_def":           "Net Present Value (чистая приведённая стоимость) — сумма будущих денежных потоков, дисконтированных к сегодняшнему дню. Учитывает временную стоимость денег.",
        "dscr_note":               "DSCR = NOI / Платёж. Безопасно: >1.25 | Внимание: 1.0–1.25 | Риск: <1.0",
        # ── LTV ───────────────────────────────────────────────────────────────
        "ltv_expander":            "Залог и LTV",
        "ltv_toggle":              "Включить анализ LTV",
        "ltv_collateral":          "Стоимость залога",
        "ltv_label":               "LTV",
        "ltv_note":                "LTV > 80% обычно требует дополнительного обеспечения.",
        "ltv_warning":             "Высокий LTV — может потребоваться дополнительный залог или страховка.",
        # ── DTI ───────────────────────────────────────────────────────────────
        "dti_expander":            "Платёжеспособность (DTI)",
        "dti_toggle":              "Включить анализ DTI",
        "dti_income":              "Совокупный месячный доход",
        "dti_other_debts":         "Другие месячные долги",
        "dti_label":               "DTI",
        "dti_excellent":           "Отлично (≤28%)",
        "dti_good":                "Хорошо (28–36%)",
        "dti_acceptable":          "Допустимо (36–43%)",
        "dti_critical":            "Критично (>43%)",
        "dti_note":                "DTI = (Платёж + Другие долги) / Доход",
        # ── Risk panel ────────────────────────────────────────────────────────
        "risk_section":            "🎯 Риск-аналитика",
        # ── Рефинансирование ──────────────────────────────────────────────────
        "refi_section":          "🔄 Анализ рефинансирования",
        "refi_caption":          "Сравните текущий кредит с условиями нового — узнайте, окупится ли переход.",
        "refi_current_block":    "Текущий кредит",
        "refi_new_block":        "Новый кредит (рефинансирование)",
        "refi_current_balance":  "Остаток долга",
        "refi_current_rate":     "Текущая ставка (% годовых)",
        "refi_remaining_term":   "Оставшийся срок (мес.)",
        "refi_penalty":          "Штраф за досрочное закрытие",
        "refi_penalty_pct":      "Штраф (% от остатка)",
        "refi_penalty_fixed":    "Штраф (фиксированная сумма)",
        "refi_penalty_type":     "Тип штрафа",
        "refi_new_rate":         "Новая ставка (% годовых)",
        "refi_new_term":         "Срок нового кредита (мес.)",
        "refi_new_fees":         "Комиссии за выдачу",
        "refi_calculate":        "Рассчитать рефинансирование",
        "refi_discount_rate":    "Ставка дисконтирования NPV (% годовых)",
        "refi_discount_rate_help": "Годовая ставка для дисконтирования будущих денежных потоков при расчёте NPV. Можно использовать доходность вашей альтернативной инвестиции.",
        "refi_current_payment":  "Текущий платёж",
        "refi_new_payment":      "Новый платёж",
        "refi_monthly_savings":  "Месячная экономия",
        "refi_total_costs":      "Совокупные расходы перехода",
        "refi_breakeven_months": "Точка окупаемости",
        "refi_total_savings":    "Чистая экономия за срок",
        "refi_help_npv":          "NPV-сравнение: положительное = рефинансирование экономит деньги в приведённой стоимости; отрицательное = дороже после дисконтирования будущих платежей.",
        "balloon_be_caption":     "💡 **Vs. Annuity** — минимальная доходность, при которой Balloon становится не дороже Annuity. **Absolute** — минимальная доходность, при которой инвестирование удерживаемого тела полностью покрывает все процентные выплаты банку (настоящий «выход в ноль»).",
        # ── Day-Count Convention ──────────────────────────────────────────────
        "day_count_section":      "📅 Конвенция начисления процентов",
        "day_count_toggle":       "Использовать конвенцию начисления (day-count)",
        "day_count_toggle_help":  "Включает расчёт процентов по выбранному рыночному стандарту вместо упрощённой формулы r=ставка/ppy. Базовая формула: I = P · r · (N / B), где N — дни в периоде, B — база года.",
        "day_count_method":       "Метод day-count",
        "day_count_method_help":  "Выберите стандарт начисления. Формула: I = P · r_годовая · (N / B). N и B зависят от метода:\n• 30/360 (ISDA): N=условные 30 дн./мес., B=360.\n• 30E/360 (Eurobond): N=европейское 30/360 (конец месяца → 30), B=360.\n• ACT/360: N=фактич. календарные дни, B=360.\n• ACT/365 (Actual/365F): N=факт. дни, B=365.\n• ACT/ACT (ISDA): N=факт. дни, B=факт. дни года (учёт високосных).",
        "day_count_caption":      "💡 Формула в каждой строке графика: **I = P × r × (N / B)**. Сейчас выбран: **{method}**.",
        "day_count_desc_30_360":  "Условный 30-дневный месяц, год = 360 дней (ISDA вариант).",
        "day_count_desc_30E_360": "Европейский 30/360: конец месяца безусловно округляется до 30.",
        "day_count_desc_ACT_360": "Фактические календарные дни / 360.",
        "day_count_desc_ACT_365": "Фактические календарные дни / 365 (синоним Actual/365F).",
        "day_count_desc_ACT_ACT": "Фактические дни / фактические дни года (учитывает 29 февраля).",
        "apr_failed_caption":     "Эффективную ставку не удалось рассчитать (IRR не сошёлся). Показано N/A — это не ошибка ввода, а сигнал о численном сбое.",
        "grace_failed_banner":    "Кредитные каникулы не были применены — расчёт показан БЕЗ каникул. Результат может не соответствовать введённым параметрам",
        "negative_amort_banner":  "Внимание: остаток долга превысил {ratio:.1f}× от исходного. Это спираль отрицательной амортизации — каникулы + высокая ставка делают кредит экономически нестабильным. Расчёт остановлен.",
        "partial_result_banner":  "ЧАСТИЧНЫЙ РЕЗУЛЬТАТ — ОТНОСИТЕСЬ КАК К НЕНАДЁЖНОМУ",
        "partial_grace":          "Кредитные каникулы НЕ были применены — график показан БЕЗ них",
        "partial_apr":            "Эффективная APR не была рассчитана — IRR не сошёлся",
        # ── Theme editor ─────────────────────────────────────────────────────
        "theme_section":          "🎨 Тема оформления",
        "theme_preset_label":     "Готовый пресет",
        "theme_preset_help":      "Выберите готовую цветовую схему или 'Custom' для ручной настройки.",
        "theme_custom":           "Своя тема (редактировать ниже)",
        "theme_custom_hint":      "Меняйте любой цвет — изменения применяются мгновенно.",
        "theme_field_bg":         "Фон приложения",
        "theme_field_bg_secondary": "Боковая панель / Карточки",
        "theme_field_bg_tertiary":  "Третичный фон",
        "theme_field_text":       "Основной текст",
        "theme_field_text_muted": "Приглушённый текст",
        "theme_field_text_subtle":"Едва заметный текст",
        "theme_field_accent":     "Акцент / Ссылки",
        "theme_field_accent_strong":"Сильный акцент",
        "theme_field_success":    "Успех",
        "theme_field_warning":    "Предупреждение",
        "theme_field_danger":     "Опасность",
        "theme_field_border":     "Границы",
        "theme_field_input_text": "Текст в полях ввода",
        "theme_font_size":        "Размер шрифта",
        "theme_font_size_help":   "Масштаб всего текста (0.75-1.4×)",
        "theme_density":          "Плотность UI",
        "theme_density_help":     "Отступы вокруг элементов (0.7-1.4×)",
        "theme_radius":           "Скругление углов",
        "theme_radius_help":      "Радиус закругления, px (0-24)",
        "theme_reset":            "↻ Сбросить к дефолту",
        # ── Day-Count comparison chart ────────────────────────────────────────
        "daycount_compare_tab":     "Сравнение Day-Count",
        "daycount_compare_title":   "💰 Сумма процентов по разным Day-Count Conventions",
        "daycount_compare_caption": "Тот же кредит, посчитанный пятью разными правилами начисления дней — наглядно показывает, какое правило выгоднее.",
        "daycount_best":            "Самый дешёвый",
        "daycount_worst":           "Самый дорогой",
        "daycount_spread":          "Разница",
        "daycount_no_data":         "Не удалось рассчитать сравнение Day-Count.",
        "compare_partial_fail":     "⚠️ Одна или несколько схем не были рассчитаны из-за конфликта параметров (например, каникул). Сравнение показано по успешным.",
        "compare_savings_vs":       "Экономия относительно",
        "compare_overpay_vs":       "Переплата относительно",
        "syndicated_tranche_offset":      "Сдвиг выдачи (мес.) ({tranche})",
        "syndicated_tranche_offset_help": "Через сколько месяцев после общей даты старта банк фактически выдаст этот транш. 0 = в день старта.",
        "refi_recommendation":   "Рекомендация",
        "refi_worth_it":         "✓ Рефинансирование выгодно. Окупится через {months:.1f} мес.",
        "refi_not_worth":        "✗ Рефинансирование невыгодно — новый платёж выше текущего.",
        "refi_long_payback":     "⚠️ Окупаемость превышает срок нового кредита.",
        "refi_longer_term_trap": "⚠️ Ловушка длинного срока: ежемесячный платёж ниже, но общая стоимость нового долга (с учётом дисконтирования) ВЫШЕ — экономия в моменте превращается в переплату на дистанции.",
        "refi_shorter_term_win": "✓ Ежемесячный платёж выше, но приведённая (NPV) общая стоимость ниже — рефинансирование экономически выгодно (обычно при переходе на более короткий срок с меньшей ставкой).",
        "refi_help_balance":     "Сколько ещё осталось выплатить банку по текущему кредиту.",
        "refi_help_penalty":     "Большинство банков взимают штраф за досрочное погашение (обычно 1-3% от остатка).",
        "refi_help_fees":        "Комиссии за оформление нового кредита (оценка имущества, страховка, юр. услуги).",
        "refi_help_breakeven":   "Месяцев нужно, чтобы экономия на платежах окупила издержки перехода.",
        # ── Flat CSV ──────────────────────────────────────────────────────────
        "download_csv_flat":     "⬇️ Flat CSV (ERP)",
        "download_csv_flat_help":"Чистые данные для импорта в SAP / 1C / Oracle: только числа, ISO-даты, без итогов и валюты.",
        # ── Help-tooltips для всех 5 advanced expanders (локализованы) ────────
        "help_grace_toggle":     "Временно уменьшает или приостанавливает платежи. Полезно на старте проекта или при перебоях с доходом.",
        "help_grace_start":      "Период начала окна каникул (1 = первый период).",
        "help_grace_duration":   "Сколько периодов длятся каникулы.",
        "help_grace_type":       "Interest Only — платим только проценты, тело сохраняется. Full Holiday — ничего не платим; проценты КАПИТАЛИЗИРУЮТСЯ (добавляются к телу), увеличивая будущий долг.",
        "info_full_holiday":     "ℹ️ **Внимание Full Holiday:** Пропущенные проценты **капитализируются** — добавляются к остатку тела долга. После каникул вы должны больше, чем до них, и платежи пересчитываются на новый увеличенный остаток.",
        "help_inflation_toggle": "Сравнить номинальные платежи с их реальной (приведённой) стоимостью с учётом инфляции.",
        "help_inflation_rate":   "Прогнозная годовая инфляция. Дисконтирует будущие платежи к сегодняшней стоимости.",
        "caption_inflation":     "💡 Инфляция «съедает» реальную стоимость денег. Платёж $1 000 через 5 лет стоит меньше, чем $1 000 сегодня.",
        "help_dscr_toggle":      "DSCR = Debt Service Coverage Ratio — используется банками для оценки бизнес-кредитов.",
        "help_dscr_noi":         "Net Operating Income — месячная прибыль до обслуживания долга (выручка минус операционные расходы, без учёта кредитных платежей).",
        "caption_dscr":          "💡 **DSCR** = NOI / Платёж. Банки предпочитают **≥ 1.25** (комфортный запас). Меньше 1.0 — операционный доход не покрывает кредит.",
        "help_ltv_toggle":       "LTV = Loan-to-Value — какую долю стоимости актива покрывает кредит.",
        "help_ltv_collateral":   "Рыночная стоимость актива, обеспечивающего кредит (например, стоимость недвижимости).",
        "caption_ltv":           "💡 **LTV** = Кредит / Залог × 100%. Порог риска: **> 80%** обычно требует PMI или дополнительного залога. Меньше LTV = лучше условия.",
        "help_dti_toggle":       "DTI = Debt-to-Income — ключевой показатель доступности персонального кредита.",
        "help_dti_income":       "Совокупный месячный доход (до налогов) из всех источников.",
        "help_dti_other_debts":  "Сумма всех других месячных платежей по долгам: карты, другие кредиты, алименты и т.д.",
        "caption_dti":           "💡 **DTI** = (Платёж + Другие долги) / Доход × 100%. Банки: **≤ 36%** идеально · **> 43%** обычно отказ.",
        # Inflation panel result tooltips
        "help_nominal":          "Сумма всех платежей по номиналу (без поправки на временную стоимость денег).",
        "help_real":             "Приведённая стоимость всех будущих платежей — сколько они «стоят» сегодня после того, как инфляция съест их покупательную способность.",
        "help_disc":             "На сколько ваши номинальные платежи «дисконтируются» инфляцией в реальном выражении.",
        "caption_inflation_full":"💡 {note} Инфляция «съедает» реальную стоимость денег со временем — будущие платежи экономически менее обременительны, чем кажутся номинально.",
        # Math formulas in help= for key metrics
        "help_total_payment":    "Total Payment = Тело + Σ Процентов + Σ Комиссий. (В графике каждый платёж = принципал + проценты + комиссия за период; итог получается тем же, что и Σ платежей, без двойного учёта.)",
        "help_total_interest":   "Σ Процентов = Total Payment − Тело − Σ Комиссий. Сумма всех процентных начислений за период жизни кредита.",
        "help_eff_rate":         "Effective Rate (APR) = годовая IRR денежных потоков:\n\nNPV = Σₜ₌₀ⁿ CFₜ / (1 + r)ᵗ = 0\n\nГде CF₀ = +(Принципал − Разовая комиссия) — выдача (net proceeds), а CFₜ для t ≥ 1 — отрицательные платежи заёмщика. Найденная r — периодическая, затем приведена к годовой: APR = (1+r)^ppy − 1.",
        "help_first_payment":    "Аннуитетный платёж (Annuity):\n\nPMT = P · r · (1+r)ⁿ / ((1+r)ⁿ − 1)\n\nГде P — принципал, r — периодическая ставка (годовая / ppy), n — число периодов.",
        "help_overpay_pct":      "Overpayment % = Σ Процентов / Тело × 100%. Доля переплаты от исходной суммы кредита.",
        "help_universal_be":     "Universal Break-even Rate — ставка реинвестирования платежей, при которой накопленный портфель компенсирует общую переплату:\n\nFV(payments at r) − Σ Pmt = Σ Interest\n\nГде FV(payments at r) = Σₜ₌₁ⁿ Pmtₜ · (1+r)ⁿ⁻ᵗ. Решается бинарным поиском.",
        "help_balloon_be_va":    "Vs. Annuity Break-even (Balloon) — ставка реинвестирования, при которой Balloon становится не дороже Annuity. Эквивалентное PV-условие:\n\nPV_Annuity = PV_Balloon\n\nГде PV_Annuity = Σₜ₌₁ⁿ PMT_aₜ /(1+r)ᵗ, а PV_Balloon = Σₜ₌₁ⁿ⁻¹ PMT_bₜ /(1+r)ᵗ + Balloon/(1+r)ⁿ. В реализации используется FV-эквивалент: FV сэкономленного тела покрывает разницу переплат.",
        "help_balloon_be_abs":   "Absolute Break-even (Balloon) — минимальная доходность инвестиций, при которой принципал, удерживаемый весь срок, покрывает все процентные выплаты банку:\n\nP · (1+r)ⁿ ≥ P + Σₜ₌₁ⁿ⁻¹ Interestₜ\n\nГде P — принципал, инвестируемый на n периодов; Interestₜ — процентные платежи Balloon. Решается бинарным поиском.",
        "help_real_cost_long":   "Present Value (PV) всех платежей с учётом инфляции:\n\nPV = Σₜ₌₁ⁿ paymentₜ / (1 + r_per)ᵗ\n\nГде r_per — периодическая ставка дисконтирования = i_annual / ppy (приближение нормальной периодической ставкой). Будущие платежи приводятся к сегодняшней покупательной способности.",
        "help_dscr_metric":      "DSCR = NOI / Debt Service (используем месячный NOI / месячный платёж — соотношение одинаковое).\n\n• ≥ 1.25 — Безопасно (типичный порог банка)\n• 1.00–1.25 — Внимание\n• < 1.00 — Высокий риск\n\nКонкретный порог зависит от банка (обычно 1.20–1.30) и типа займа.",
        "help_ltv_metric":       "LTV = Loan Amount / Appraised Value × 100%.\n\n• ≤ 60% — Безопасно\n• ≤ 80% — Стандарт\n• ≤ 95% — Высокий\n• > 95% — Критично\n\nЛимиты не универсальны — зависят от банка и типа кредита (обычно ≤ 80–90%).",
        "help_dti_metric":       "DTI = Total Monthly Debt Payments / Gross Monthly Income × 100%.\n\n• ≤ 28% — Отлично\n• ≤ 36% — Хорошо\n• ≤ 43% — Допустимо\n• > 43% — Критично\n\nКонкретные пороги зависят от банка и страны (обычно лимит ≤ 36–45%).",
        # ── Email / SMTP отправка отчёта ───────────────────────────────────────
        "email_section":         "📧 Отправить отчёт по Email",
        "email_caption":         "Отправьте PDF или Excel-отчёт на указанный email-адрес.",
        "email_recipient":       "Email получателя",
        "email_recipient_help":  "Email-адрес, на который будет отправлен отчёт.",
        "email_subject":         "Тема письма",
        "email_format":          "Формат вложения",
        "email_format_pdf":      "PDF",
        "email_format_xlsx":     "Excel (XLSX)",
        "email_format_docx":     "Word (DOCX)",
        "email_message":         "Текст сообщения (опционально)",
        "email_send_btn":        "📤 Отправить отчёт",
        "email_sending":         "Отправка...",
        "email_success":         "✓ Отчёт успешно отправлен на {email}",
        "email_error":           "✗ Ошибка отправки: {error}",
        "email_no_recipient":    "Укажите email получателя",
        "email_invalid":         "Некорректный email-адрес",
        "email_no_secrets":      "⚠️ SMTP не настроен. Добавьте секреты в `.streamlit/secrets.toml`:\n```toml\n[smtp]\nserver = \"smtp.gmail.com\"\nport = 587\nlogin = \"your@email.com\"\npassword = \"app_password\"\nsender = \"Yev Capital LoanLogic <your@email.com>\"\n```",
        "email_default_subject": "Отчёт Yev Capital LoanLogic",
        "email_default_body":    "Прилагаю отчёт, сгенерированный Yev Capital LoanLogic v3.0.\n\nС уважением,\nYev Capital",
        # ── Email setup help / disabled states ───────────────────────────────
        "email_disabled_warning": "Email-сервис не настроен. Добавьте SMTP-учётные данные в секреты.",
        "email_invalid_warning":  "Введите корректный email-адрес для активации кнопки отправки.",
        "email_setup_title":      "📖 Как настроить отправку Email",
        "email_setup_steps":      (
            "**Локально (`.streamlit/secrets.toml`)**\n\n"
            "1. Создайте файл `.streamlit/secrets.toml` в корне проекта.\n"
            "2. Добавьте секцию SMTP:\n\n"
            "```toml\n"
            "[smtp]\n"
            "server = \"smtp.gmail.com\"\n"
            "port = 587\n"
            "login = \"your@email.com\"\n"
            "password = \"app_password\"\n"
            "sender = \"Yev Capital LoanLogic <your@email.com>\"\n"
            "```\n\n"
            "**Streamlit Cloud**\n\n"
            "1. Откройте панель приложения → **Settings** → **Secrets**.\n"
            "2. Вставьте те же ключи в формате TOML и нажмите Save.\n"
            "3. Приложение автоматически перезапустится.\n\n"
            "**Google App Password (для Gmail)**\n\n"
            "1. Включите двухфакторную аутентификацию: [myaccount.google.com/security](https://myaccount.google.com/security)\n"
            "2. Создайте App Password: [myaccount.google.com/apppasswords](https://myaccount.google.com/apppasswords)\n"
            "3. Используйте 16-значный пароль в поле `password` (без пробелов).\n\n"
            "**Порты SMTP**\n\n"
            "- `587` — STARTTLS (рекомендуется для Gmail/Outlook)\n"
            "- `465` — SSL (legacy)\n"
            "- `25`  — открытый (обычно заблокирован хостингами)"
        ),
        # ── Syndicated Loan ──────────────────────────────────────────────────
        "syndicated_section":     "🏦 Синдицированный кредит (Multi-Tranche)",
        "syndicated_toggle":      "Режим синдицированного кредита",
        "syndicated_help":        "Консолидированный долг от нескольких кредиторов с разными ставками, сроками и комиссиями. Каждый транш погашается по своему графику; итоговая таблица суммирует платежи по всем траншам.",
        "syndicated_tranche":     "Транш {n}",
        "syndicated_tranche_amount":   "Сумма ({tranche})",
        "syndicated_tranche_rate":     "Ставка % годовых ({tranche})",
        "syndicated_tranche_term":     "Срок ({tranche})",
        "syndicated_tranche_unit":     "Единица срока ({tranche})",
        "syndicated_tranche_scheme":   "Схема ({tranche})",
        "syndicated_tranche_ot_comm":  "Разовая комиссия ({tranche})",
        "syndicated_tranche_mo_comm":  "Период. комиссия ({tranche})",
        "syndicated_tranche_enabled":  "Активировать {tranche}",
        "syndicated_total":       "Общая сумма синдиката (Σ траншей)",
        "syndicated_blended_rate":"Смешанная APR (effective)",
        "syndicated_master_help": "Master Schedule = Σ платежей по всем траншам в каждый период. Если у траншей разные сроки — короткие просто отсутствуют в поздних периодах.",
        "syndicated_zero_error":  "⚠️ Введите сумму хотя бы одного транша.",
        "syndicated_zero_amount_warn": "⚠️ Транш(и) {letters} включены, но имеют нулевую сумму — они не учтены. Укажите сумму или отключите их.",
        "end_date_fallback": "↳ Оставлен прежний срок: {n} {unit}.",
        "email_empty_attachment": "⚠️ Файл отчёта пуст (экспорт мог завершиться ошибкой). Письмо не отправлено.",
        "refi_annuity_note": "ℹ️ Анализ моделирует и текущий, и новый кредит как **аннуитетные** (с равными платежами). Если текущий кредит использует другую схему (классическую или буллит), считайте сравнение приблизительным.",
        "syndicated_invalid":     "⚠️ Все траншевые суммы должны быть положительными.",
        "syndicated_chart_title": "Структура платежей по траншам",
        "syndicated_chart_caption":"Каждый цвет — отдельный транш. Высота столбца = совокупный платёж периода.",
        "help_synd_master_payment": "Master Payment_t = Σᵢ PMTᵢ(t), где PMTᵢ(t) — платёж транша i в периоде t (после конца срока транш = 0).",
        "help_synd_blended_apr":  "Смешанная APR — IRR совокупных денежных потоков всех траншей: NPV(combined CF, r) = 0.",
        # ── Audit Trail ──────────────────────────────────────────────────────
        "audit_section":          "📋 Журнал изменений (Audit Trail)",
        "audit_caption":          "Фиксирует изменения ключевых параметров (сумма, ставка, срок) в течение текущей сессии. Запись делается в момент нажатия «Рассчитать».",
        "audit_empty":            "Пока нет записей. Измените параметры и нажмите «Рассчитать», чтобы зафиксировать изменение.",
        "audit_clear":            "Очистить журнал",
        "audit_field_amount":     "Сумма",
        "audit_field_rate":       "Ставка",
        "audit_field_term":       "Срок",
        "audit_changed_to":       "{field} изменена с {old} на {new}",
        "audit_impact_interest":  "Изменение переплаты: {delta}",
        "audit_impact_first_payment": "Изменение первого платежа: {delta}",
        "audit_field_grace":       "Льготный период",
        "audit_field_inflation":   "Учёт инфляции",
        "audit_field_day_count":   "База начисления",
        "audit_field_syndicated":  "Синдиц. режим",
        "audit_toggle_on":         "Вкл",
        "audit_toggle_off":        "Выкл",
        "audit_field_scheme":      "Схема",
        "audit_field_currency":    "Валюта",
        "audit_field_term_mode":   "Режим срока",
        "audit_field_dc_method":   "Метод базы",
        "audit_field_start_date":  "Дата начала",
        "audit_field_one_time":    "Разовая комиссия",
        "audit_field_monthly":     "Периодич. комиссия",
        "audit_impact_payment":   "Изменение платежа: {delta}",
    },
    "uk": {
        "app_title": "Yev Capital LoanLogic",
        "app_subtitle": "Розрахунок банківських кредитів і вкладів",
        "section_params": "⚙️ Параметри кредиту",
        "loan_amount": "Сума кредиту",
        "loan_amount_slider": "Повзунок суми",
        "loan_term": "Термін",
        "term_unit": "Одиниця терміну",
        "weeks": "Тижні", "months": "Місяці", "quarters": "Квартали",
        "halfyears": "Півроки", "years": "Роки",
        "interest_rate": "Відсоткова ставка (% річних)",
        "calc_scheme": "Схема розрахунку",
        "annuity": "Ануїтет",
        "classic": "Класика (диференційований)",
        "balloon": "Буліт (Bullet / Balloon)",
        "deposit_scheme": "Вклад / Депозит",
        "section_commissions": "💼 Комісії",
        "one_time_comm": "Разова комісія",
        "one_time_type": "Тип разової комісії",
        "monthly_comm": "Щомісячна комісія",
        "monthly_type": "Тип щомісячної комісії",
        "pct_of_amount": "% від суми кредиту",
        "fixed_amount": "Фіксована сума",
        "calc_btn": "🚀 Розрахувати",
        "calc_error": "⚠️ Не вдалося виконати розрахунок із заданими параметрами. Перевірте коректність введення (термін, ставка, сума).",
        "section_results": "📊 Результати розрахунку",
        "total_payment": "Загальна сума виплат",
        "total_interest": "Переплата за відсотками",
        "total_commission": "Сума комісій",
        "effective_rate": "Ефективна ставка",
        "monthly_payment": "Платіж (1-й період)",
        "section_schedule": "📋 Графік платежів",
        "period": "Період", "date": "Дата",
        "balance_open": "Залишок (поч.)",
        "payment_total": "Платіж",
        "principal": "Тіло боргу",
        "interest": "Відсотки",
        "commission": "Комісія",
        "balance_close": "Залишок (кін.)",
        "total_row": "ПІДСУМОК",
        "section_chart": "📈 Діаграма",
        "chart_title": "Структура виплат за періодами",
        "chart_principal": "Тіло кредиту",
        "chart_interest": "Відсотки",
        "chart_commission": "Комісії",
        "chart_pie_title": "Підсумкове співвідношення",
        "section_templates": "💾 Шаблони розрахунків",
        "template_name": "Назва шаблону",
        "save_template": "💾 Зберегти",
        "load_template": "📂 Завантажити",
        "delete_template": "🗑️",
        "no_templates": "Немає збережених шаблонів",
        "template_saved": "✅ Шаблон збережено!",
        "template_loaded": "✅ Шаблон завантажено!",
        "template_deleted": "🗑️ Видалено!",
        "template_name_empty": "Введіть назву шаблону",
        "template_overwrite_warn": "⚠️ Шаблон «{name}» вже існує. Натисніть «Зберегти» ще раз, щоб перезаписати.",
        "currency": "Валюта",
        "uah": "₴ Гривня (UAH)", "usd": "$ Долар (USD)", "eur": "€ Євро (EUR)",
        "rub": "₽ Рубль (RUB)", "gbp": "£ Фунт стерлінгів (GBP)",
        "jpy": "¥ Єна (JPY)", "cad": "C$ Канадський долар (CAD)",
        "aud": "A$ Австралійський долар (AUD)", "chf": "Fr Швейцарський франк (CHF)",
        "cny": "¥ Китайський юань (CNY)", "hkd": "HK$ Гонконгський долар (HKD)",
        "custom": "✏️ Своя валюта",
        "custom_symbol": "Символ своєї валюти",
        "download_excel": "⬇️ Excel",
        "download_pdf": "⬇️ PDF",
        "download_docx": "⬇️ Word",
        "overpayment_pct": "Переплата від суми",
        "compare_schemes": "⚖️ Порівняти схеми",
        "annuity_vs_classic": "Порівняння схем кредитування",
        "savings_classic": "Економія при класиці",
        "annuity_short": "Ануїтет",
        "classic_short": "Класика",
        "enter_name": "Введіть назву...",
        "template_examples": "Швидкі шаблони:",
        "mortgage": "🏠 Іпотека",
        "car_loan": "🚗 Автокредит",
        "consumer": "💳 Споживчий",
        "deposit": "Вклад",
        "invest_section": "📈 Порівняння з інвестиціями",
        "invest_sp500": "📊 Порівняти з S&P 500 (13.7% річних)",
        "invest_custom": "💹 Порівняти з кастомним % інвестування",
        "invest_custom_rate": "Річна прибутковість (%)",
        "invest_chart_title": "Кредит vs Інвестиції",
        "invest_loan_balance": "Залишок боргу",
        "invest_portfolio": "Інвестиційний портфель",
        "invest_total_paid": "Виплачено банку",
        "invest_portfolio_val": "Вартість портфеля",
        "invest_net_gain": "Чистий прибуток",
        "invest_explanation": "Що було б, якби замість кредитних платежів ці гроші інвестувались?",
        "invest_rate_label": "дохідність",
        "invest_only_one": "⚠️ Можна увімкнути лише одне порівняння!",
        "deposit_mode_label": "Режим вкладу",
        "deposit_capitalize": "💰 Капіталізація відсотків (складний %)",
        "deposit_payout": "💸 Виплата відсотків",
        "deposit_final": "Підсумкова сума вкладу",
        "deposit_earned": "Нараховано відсотків",
        "tooltip_payment": "💡 **Платіж** — повна сума, яку потрібно перерахувати банку в цей період.",
        "tooltip_principal": "💡 **Тіло боргу** — частина платежу на погашення основної суми кредиту.",
        "tooltip_interest": "💡 **Відсотки** — плата за користування грошима банку, нараховується на залишок боргу.",
        "scheme_annuity_info": "Рівні платежі на весь термін.",
        "scheme_classic_info": "Фіксоване тіло боргу + відсотки, що зменшуються.",
        "scheme_balloon_info": "Тільки відсотки щомісяця, вся сума — в останньому платежі.",
        "scheme_deposit_info": "Розрахунок вкладу/депозиту з капіталізацією або виплатою відсотків.",
        "col_explain": "ℹ️ Пояснення колонок",
        "dep_balance_open":    "Баланс (поч.)",
        "dep_interest_earned": "Нараховано %",
        "dep_payout":          "Виплата",
        "dep_balance_close":   "Баланс (кін.)",
        "dep_total_row":       "ПІДСУМОК",
        "dep_initial":         "Сума вкладу",
        "dep_final_balance":   "Підсумкова сума",
        "dep_total_earned":    "Нараховано відсотків",
        "dep_rate_label":      "Ставка річних",
        "dep_period_payout":   "Виплата за період",
        "dep_growth_title":    "📈 Зростання вкладу",
        "dep_balance_label":   "Баланс вкладу",
        "dep_interest_label":  "Нараховані відсотки",
        "dep_vs_invest":       "Вклад vs Альтернатива",
        "dep_your_deposit":    "Ваш вклад",
        "dep_alternative":     "Альтернативні інвестиції",
        "dep_invest_exp":      "Порівняння: якби ту саму суму вкласти під інший відсоток",
        "dep_sp500_exp":       "Порівняння вашого вкладу з дохідністю S&P 500 (13.7% річних)",
        "dep_tooltip_interest":"💡 **Нараховано %** — сума відсотків за період.",
        "dep_tooltip_balance": "💡 **Баланс** — сума на вкладі з урахуванням нарахованих відсотків.",
        "dep_tooltip_payout":  "💡 **Виплата** — сума відсотків, отриманих на руки в цей період.",
        "download_csv": "⬇️ CSV",
        "dep_invest_section":  "📊 Порівняння з альтернативою",
        "dep_invest_caption":  "Що якби замість цього вкладу інвестувати під інший % ?",
        "dep_invest_yours":    "Ваш депозит (підсумок)",
        "dep_invest_alt":      "Альтернатива (підсумок)",
        "dep_invest_diff":     "Різниця",
        "start_date_label":    "📅 Дата початку",
        "start_date_hint":     "Перший платіж буде здійснено через один період",
        "term_mode_manual":    "Вручну",
        "term_mode_enddate":   "За датою закінчення",
        "end_date_label":      "📅 Дата закінчення",
        "end_date_hint":       "Термін перерахується автоматично в повних періодах",
        "chart_balance_title": "📉 Залишок боргу",
        "chart_balance_hover": "Залишок",
        "balloon_short":       "Буліт",
        "balloon_breakeven":       "Break-even інвестицій",
        "balloon_breakeven_tip":   "Мінімальна річна дохідність інвестицій (складний %), при якій схема «буліт» вигідніша за ануїтет.",
        "balloon_breakeven_label": "Аналіз Break-even (Буліт)",
        "balloon_breakeven_desc":  "Мін. дохідність інвестицій для виправдання схеми «буліт»",
        "term_caption":        "міс. / рок.",
        "welcome_h2":  "Введіть параметри та натисніть",
        "welcome_calc": "Розрахувати",
        "welcome_sub": "Ануїтет · Класика · Буліт · Вклад<br>Експорт Excel / PDF / Word / CSV · Порівняння з інвестиціями",
        "tab_balance": "📉 Залишок боргу",
        "copyright": "© 2026 Bohdan Yevtushenko (MrCemper) · Yev Capital LoanLogic v3.0",
        "invest_breakeven_section":      "📐 Аналіз інвест-беззбитковості",
        "invest_breakeven_universal":    "Універсальна ставка беззбитковості",
        "invest_breakeven_universal_tip":"Річна дохідність, при якій реінвестування платежів повністю перекриє переплату за відсотками.",
        "invest_breakeven_abs":          "Абсолютна беззбитковість (Буліт)",
        "invest_breakeven_abs_tip":      "Мінімальна дохідність, при якій FV заощаджених коштів покриє всю переплату.",
        "invest_breakeven_vs_ann":       "Vs. Ануїтет (Буліт)",
        "invest_breakeven_vs_ann_tip":   "Мінімальна дохідність, при якій буліт вигідніший за ануїтет.",
        "invest_breakeven_vs_cla_balloon":     "Vs. Класика (Буліт)",
        "invest_breakeven_vs_cla_balloon_tip": "Мінімальна річна дохідність, при якій схема «буліт» (з відкладеною виплатою тіла) до кінця терміну досягає того ж результату, що й класична.",
        "invest_breakeven_vs_cla_annuity":     "Vs. Класика (Ануїтет)",
        "invest_breakeven_vs_cla_annuity_tip": "Мінімальна річна дохідність, при якій інвестування зекономлених коштів (ануїтет у ранні періоди дешевший за класику) компенсує переплату в пізні періоди.",
        "annuity_be_caption":            "Якщо ви можете заробляти ≥ цієї ставки на зекономлених у ранні періоди грошах — ануїтет до кінця терміну буде не гіршим, ніж класика.",
        "sp500_disclaimer": "Примітка: дохідність S&P 500 у 13.7% ґрунтується на середньорічному показнику за останні 10 років. Минулі результати не гарантують майбутніх.",
        "inflation_expander":      "Облік інфляції",
        "inflation_toggle":        "Увімкнути облік інфляції",
        "inflation_rate":          "Очікувана інфляція (% річних)",
        "inflation_section":       "💰 Реальна vs Номінальна вартість",
        "nominal_cost":            "Номінальна вартість",
        "real_cost":               "Реальна вартість (PV)",
        "inflation_savings":       "Знижка на інфляцію",
        "inflation_note":          "Реальна вартість дисконтує майбутні платежі до сьогоднішньої купівельної спроможності.",
        "grace_expander":          "Кредитні канікули",
        "grace_toggle":            "Увімкнути кредитні канікули",
        "grace_duration":          "Тривалість (періодів)",
        "grace_start":             "Початковий період",
        "grace_type":              "Тип",
        "grace_interest_only":     "Лише відсотки",
        "grace_full_holiday":      "Повні канікули (капіталізація %)",
        "grace_note":              "При повних канікулах нараховані відсотки додаються до тіла боргу.",
        "grace_active":            "Активний період канікул",
        "dscr_expander":           "Бізнес-аналіз (DSCR)",
        "dscr_toggle":             "Увімкнути аналіз DSCR",
        "dscr_noi":                "Щомісячний NOI",
        "dscr_label":              "DSCR",
        "dscr_status_safe":        "Безпечно",
        "dscr_status_warning":     "Попередження",
        "dscr_status_risk":        "Високий ризик",
        "ltv_status_safe":         "Безпечний",
        "ltv_status_standard":     "Стандартний",
        "ltv_status_high":         "Високий",
        "ltv_status_critical":     "Критичний",
        "ltv_status_na":           "Н/Д",
        "credit_health_title":     "Кредитне здоров'я",
        "credit_rank_a":           "Ранг A — Відмінно",
        "credit_rank_b":           "Ранг B — Добре",
        "credit_rank_c":           "Ранг C — Задовільно",
        "credit_rank_d":           "Ранг D — Високий ризик",
        "credit_rank_na":          "Н/Д",
        "credit_health_caption":   "Бал поєднує вказані вами метрики ({metrics}). Що вище — тим здоровіше.",
        "glossary_title":          "📚 Словник термінів / Glossary",
        "help_negative_rate":      "Від'ємні ставки дозволені (рідкісні продукти із субнульовою ставкою). Розрахунок коректно їх обробляє.",
        "gloss_annuity_term":      "Ануїтет",
        "gloss_annuity_def":       "Схема погашення з рівними платежами щоперіоду. Спочатку більшу частину платежу складають відсотки, наприкінці — тіло кредиту.",
        "gloss_bullet_term":       "Буліт (Bullet)",
        "gloss_bullet_def":        "Кредит, де протягом строку сплачуються лише відсотки, а все тіло боргу повертається одним платежем у кінці.",
        "gloss_classic_term":      "Класична (диференційована)",
        "gloss_classic_def":       "Тіло кредиту гаситься рівними частками, а відсотки нараховуються на залишок. Платіж зменшується з часом.",
        "gloss_dscr_term":         "DSCR",
        "gloss_dscr_def":          "Debt Service Coverage Ratio = чистий дохід / обслуговування боргу. Показує, у скільки разів дохід покриває виплати. ≥1.25 — безпечно.",
        "gloss_ltv_term":          "LTV",
        "gloss_ltv_def":           "Loan-to-Value = кредит / вартість застави × 100%. Чим нижче, тим менший ризик для кредитора. ≤80% зазвичай вважається добрим.",
        "gloss_dti_term":          "DTI",
        "gloss_dti_def":           "Debt-to-Income = (платіж + інші борги) / дохід × 100%. Частка доходу, що йде на борги. ≤36% зазвичай прийнятно.",
        "gloss_npv_term":          "NPV",
        "gloss_npv_def":           "Net Present Value (чиста приведена вартість) — сума майбутніх грошових потоків, дисконтованих до сьогодні. Враховує часову вартість грошей.",
        "dscr_note":               "DSCR = NOI / Платіж. Безпечно: >1.25 | Увага: 1.0–1.25 | Ризик: <1.0",
        "ltv_expander":            "Застава та LTV",
        "ltv_toggle":              "Увімкнути аналіз LTV",
        "ltv_collateral":          "Вартість застави",
        "ltv_label":               "LTV",
        "ltv_note":                "LTV > 80% зазвичай вимагає додаткового забезпечення.",
        "ltv_warning":             "Високий LTV — може знадобитися додаткова застава.",
        "dti_expander":            "Платоспроможність (DTI)",
        "dti_toggle":              "Увімкнути аналіз DTI",
        "dti_income":              "Сукупний місячний дохід",
        "dti_other_debts":         "Інші місячні борги",
        "dti_label":               "DTI",
        "dti_excellent":           "Відмінно (≤28%)",
        "dti_good":                "Добре (28–36%)",
        "dti_acceptable":          "Допустимо (36–43%)",
        "dti_critical":            "Критично (>43%)",
        "dti_note":                "DTI = (Платіж + Борги) / Дохід",
        "risk_section":            "🎯 Ризик-аналітика",
        "refi_section":          "🔄 Аналіз рефінансування",
        "refi_caption":          "Порівняйте поточний кредит з умовами нового — дізнайтеся, чи окупиться перехід.",
        "refi_current_block":    "Поточний кредит",
        "refi_new_block":        "Новий кредит (рефінансування)",
        "refi_current_balance":  "Залишок боргу",
        "refi_current_rate":     "Поточна ставка (% річних)",
        "refi_remaining_term":   "Термін, що залишився (міс.)",
        "refi_penalty":          "Штраф за дострокове закриття",
        "refi_penalty_pct":      "Штраф (% від залишку)",
        "refi_penalty_fixed":    "Штраф (фіксована сума)",
        "refi_penalty_type":     "Тип штрафу",
        "refi_new_rate":         "Нова ставка (% річних)",
        "refi_new_term":         "Термін нового кредиту (міс.)",
        "refi_new_fees":         "Комісії за видачу",
        "refi_calculate":        "Розрахувати рефінансування",
        "refi_discount_rate":    "Ставка дисконтування NPV (% річних)",
        "refi_discount_rate_help": "Річна ставка для дисконтування майбутніх грошових потоків при розрахунку NPV. Можна використати дохідність вашої альтернативної інвестиції.",
        "refi_current_payment":  "Поточний платіж",
        "refi_new_payment":      "Новий платіж",
        "refi_monthly_savings":  "Місячна економія",
        "refi_total_costs":      "Сукупні витрати переходу",
        "refi_breakeven_months": "Точка окупності",
        "refi_total_savings":    "Чиста економія за термін",
        "refi_help_npv":          "NPV-порівняння: позитивне = рефінансування економить гроші у приведеній вартості; негативне = дорожче після дисконтування майбутніх платежів.",
        "balloon_be_caption":     "💡 **Vs. Annuity** — мінімальна дохідність, при якій Balloon стає не дорожче Annuity. **Absolute** — мінімальна дохідність, при якій інвестування утримуваного тіла повністю покриває всі процентні виплати банку (справжній «вихід у нуль»).",
        "day_count_section":      "📅 Конвенція нарахування відсотків",
        "day_count_toggle":       "Використовувати конвенцію нарахування (day-count)",
        "day_count_toggle_help":  "Вмикає розрахунок відсотків за вибраним ринковим стандартом замість спрощеної формули r=ставка/ppy. Базова формула: I = P · r · (N / B), де N — дні в періоді, B — база року.",
        "day_count_method":       "Метод day-count",
        "day_count_method_help":  "Виберіть стандарт нарахування. Формула: I = P · r_річна · (N / B). N та B залежать від методу:\n• 30/360 (ISDA): N=умовні 30 дн./міс., B=360.\n• 30E/360 (Eurobond): N=європейське 30/360 (кінець місяця → 30), B=360.\n• ACT/360: N=фактичні дні, B=360.\n• ACT/365 (Actual/365F): N=фактичні дні, B=365.\n• ACT/ACT (ISDA): N=фактичні дні, B=фактичні дні року (з урахуванням високосних).",
        "day_count_caption":      "💡 Формула у кожному рядку графіка: **I = P × r × (N / B)**. Зараз обрано: **{method}**.",
        "day_count_desc_30_360":  "Умовний 30-денний місяць, рік = 360 днів (ISDA варіант).",
        "day_count_desc_30E_360": "Європейський 30/360: кінець місяця безумовно округлюється до 30.",
        "day_count_desc_ACT_360": "Фактичні календарні дні / 360.",
        "day_count_desc_ACT_365": "Фактичні календарні дні / 365 (синонім Actual/365F).",
        "day_count_desc_ACT_ACT": "Фактичні дні / фактичні дні року (враховує 29 лютого).",
        "apr_failed_caption":     "Ефективну ставку не вдалося розрахувати (IRR не зійшовся). Показано N/A — це не помилка вводу, а сигнал про чисельний збій.",
        "grace_failed_banner":    "Кредитні канікули не було застосовано — розрахунок показано БЕЗ канікул. Результат може не відповідати введеним параметрам",
        "negative_amort_banner":  "Увага: залишок боргу перевищив {ratio:.1f}× від початкового. Це спіраль негативної амортизації — канікули + висока ставка роблять кредит економічно нестабільним. Розрахунок зупинено.",
        "partial_result_banner":  "ЧАСТКОВИЙ РЕЗУЛЬТАТ — ВВАЖАЙТЕ НЕНАДІЙНИМ",
        "partial_grace":          "Кредитні канікули НЕ було застосовано — графік показано БЕЗ них",
        "partial_apr":            "Ефективну APR не вдалось розрахувати — IRR не зійшовся",
        "theme_section":          "🎨 Тема оформлення",
        "theme_preset_label":     "Готовий пресет",
        "theme_preset_help":      "Виберіть готову колірну схему або 'Custom' для ручного налаштування.",
        "theme_custom":           "Своя тема (редагувати нижче)",
        "theme_custom_hint":      "Змінюйте будь-який колір — зміни застосовуються миттєво.",
        "theme_field_bg":         "Фон додатка",
        "theme_field_bg_secondary": "Бічна панель / Картки",
        "theme_field_bg_tertiary":  "Третинний фон",
        "theme_field_text":       "Основний текст",
        "theme_field_text_muted": "Приглушений текст",
        "theme_field_text_subtle":"Ледь помітний текст",
        "theme_field_accent":     "Акцент / Посилання",
        "theme_field_accent_strong":"Сильний акцент",
        "theme_field_success":    "Успіх",
        "theme_field_warning":    "Попередження",
        "theme_field_danger":     "Небезпека",
        "theme_field_border":     "Межі",
        "theme_field_input_text": "Текст у полях вводу",
        "theme_font_size":        "Розмір шрифту",
        "theme_font_size_help":   "Масштаб усього тексту (0.75-1.4×)",
        "theme_density":          "Щільність UI",
        "theme_density_help":     "Відступи навколо елементів (0.7-1.4×)",
        "theme_radius":           "Закруглення кутів",
        "theme_radius_help":      "Радіус закруглення, px (0-24)",
        "theme_reset":            "↻ Скинути до дефолту",
        "daycount_compare_tab":     "Порівняння Day-Count",
        "daycount_compare_title":   "💰 Сума відсотків за різними Day-Count Conventions",
        "daycount_compare_caption": "Той самий кредит, обчислений п'ятьма різними правилами нарахування днів — наочно показує, яке правило вигідніше.",
        "daycount_best":            "Найдешевший",
        "daycount_worst":           "Найдорожчий",
        "daycount_spread":          "Різниця",
        "daycount_no_data":         "Не вдалося розрахувати порівняння Day-Count.",
        "compare_partial_fail":     "⚠️ Одну або декілька схем не вдалось розрахувати через конфлікт параметрів (наприклад, канікул). Порівняння показано за успішними.",
        "compare_savings_vs":       "Економія відносно",
        "compare_overpay_vs":       "Переплата відносно",
        "syndicated_tranche_offset":      "Зсув видачі (міс.) ({tranche})",
        "syndicated_tranche_offset_help": "Через скільки місяців після спільної дати старту банк фактично видасть цей транш. 0 = у день старту.",
        "refi_recommendation":   "Рекомендація",
        "refi_worth_it":         "✓ Рефінансування вигідне. Окупиться через {months:.1f} міс.",
        "refi_not_worth":        "✗ Рефінансування невигідне — новий платіж вищий за поточний.",
        "refi_long_payback":     "⚠️ Окупність перевищує термін нового кредиту.",
        "refi_longer_term_trap": "⚠️ Пастка довгого терміну: щомісячний платіж нижчий, але загальна вартість нового боргу (з урахуванням дисконтування) ВИЩА — економія в моменті перетворюється на переплату на дистанції.",
        "refi_shorter_term_win": "✓ Щомісячний платіж вищий, але приведена (NPV) загальна вартість нижча — рефінансування економічно вигідне (зазвичай при переході на коротший термін з меншою ставкою).",
        "refi_help_balance":     "Скільки ще залишилося виплатити банку за поточним кредитом.",
        "refi_help_penalty":     "Більшість банків стягують штраф за дострокове погашення (зазвичай 1-3% від залишку).",
        "refi_help_fees":        "Комісії за оформлення нового кредиту (оцінка майна, страхування, юр. послуги).",
        "refi_help_breakeven":   "Місяців потрібно, щоб економія на платежах окупила витрати переходу.",
        "download_csv_flat":     "⬇️ Flat CSV (ERP)",
        "download_csv_flat_help":"Чисті дані для імпорту в SAP / 1C / Oracle: лише числа, ISO-дати, без підсумків і валюти.",
        "help_grace_toggle":     "Тимчасово зменшує або призупиняє платежі. Корисно на старті проєкту або при перебоях з доходом.",
        "help_grace_start":      "Період початку вікна канікул (1 = перший період).",
        "help_grace_duration":   "Скільки періодів тривають канікули.",
        "help_grace_type":       "Interest Only — платимо лише відсотки, тіло зберігається. Full Holiday — нічого не платимо; відсотки КАПІТАЛІЗУЮТЬСЯ (додаються до тіла), збільшуючи майбутній борг.",
        "info_full_holiday":     "ℹ️ **Увага Full Holiday:** Пропущені відсотки **капіталізуються** — додаються до залишку тіла. Після канікул ви винні більше, ніж до них, і платежі перераховуються на новий збільшений залишок.",
        "help_inflation_toggle": "Порівняти номінальні платежі з їх реальною (приведеною) вартістю з урахуванням інфляції.",
        "help_inflation_rate":   "Прогнозна річна інфляція. Дисконтує майбутні платежі до сьогоднішньої вартості.",
        "caption_inflation":     "💡 Інфляція «з'їдає» реальну вартість грошей. Платіж $1 000 за 5 років коштує менше, ніж $1 000 сьогодні.",
        "help_dscr_toggle":      "DSCR = Debt Service Coverage Ratio — використовується банками для оцінки бізнес-кредитів.",
        "help_dscr_noi":         "Net Operating Income — місячний прибуток до обслуговування боргу (виручка мінус операційні витрати, без урахування кредитних платежів).",
        "caption_dscr":          "💡 **DSCR** = NOI / Платіж. Банки надають перевагу **≥ 1.25** (комфортний запас). Менше 1.0 — операційний дохід не покриває кредит.",
        "help_ltv_toggle":       "LTV = Loan-to-Value — яку частку вартості активу покриває кредит.",
        "help_ltv_collateral":   "Ринкова вартість активу, що забезпечує кредит (наприклад, вартість нерухомості).",
        "caption_ltv":           "💡 **LTV** = Кредит / Застава × 100%. Поріг ризику: **> 80%** зазвичай вимагає PMI або додаткової застави. Менше LTV = кращі умови.",
        "help_dti_toggle":       "DTI = Debt-to-Income — ключовий показник доступності персонального кредиту.",
        "help_dti_income":       "Сукупний місячний дохід (до податків) з усіх джерел.",
        "help_dti_other_debts":  "Сума всіх інших місячних платежів за боргами: картки, інші кредити, аліменти тощо.",
        "caption_dti":           "💡 **DTI** = (Платіж + Інші борги) / Дохід × 100%. Банки: **≤ 36%** ідеально · **> 43%** зазвичай відмова.",
        "help_nominal":          "Сума всіх платежів за номіналом (без поправки на часову вартість грошей).",
        "help_real":             "Приведена вартість усіх майбутніх платежів — скільки вони «коштують» сьогодні після того, як інфляція з'їсть їх купівельну спроможність.",
        "help_disc":             "Наскільки ваші номінальні платежі «дисконтуються» інфляцією у реальному виразі.",
        "caption_inflation_full":"💡 {note} Інфляція «з'їдає» реальну вартість грошей з часом — майбутні платежі економічно менш обтяжливі, ніж здаються номінально.",
        "help_total_payment":    "Total Payment = Тіло + Σ Відсотків + Σ Комісій. (У графіку кожен платіж = принципал + відсотки + комісія за період; підсумок збігається з Σ платежів, без подвійного обліку.)",
        "help_total_interest":   "Σ Відсотків = Total Payment − Тіло − Σ Комісій. Сума всіх процентних нарахувань за період життя кредиту.",
        "help_eff_rate":         "Effective Rate (APR) = річна IRR грошових потоків:\n\nNPV = Σₜ₌₀ⁿ CFₜ / (1 + r)ᵗ = 0\n\nДе CF₀ = +(Принципал − Разова комісія) — видача (net proceeds), а CFₜ для t ≥ 1 — від'ємні платежі позичальника. Знайдена r — періодична, потім приведена до річної: APR = (1+r)^ppy − 1.",
        "help_first_payment":    "Ануїтетний платіж (Annuity):\n\nPMT = P · r · (1+r)ⁿ / ((1+r)ⁿ − 1)\n\nДе P — принципал, r — періодична ставка (річна / ppy), n — кількість періодів.",
        "help_overpay_pct":      "Overpayment % = Σ Відсотків / Тіло × 100%. Частка переплати від початкової суми кредиту.",
        "help_universal_be":     "Universal Break-even Rate — ставка реінвестування платежів, при якій накопичений портфель компенсує загальну переплату:\n\nFV(payments at r) − Σ Pmt = Σ Interest\n\nДе FV(payments at r) = Σₜ₌₁ⁿ Pmtₜ · (1+r)ⁿ⁻ᵗ. Розв'язується бінарним пошуком.",
        "help_balloon_be_va":    "Vs. Annuity Break-even (Balloon) — ставка реінвестування, при якій Balloon стає не дорожчим за Annuity. Еквівалентна PV-умова:\n\nPV_Annuity = PV_Balloon\n\nДе PV_Annuity = Σₜ₌₁ⁿ PMT_aₜ /(1+r)ᵗ, а PV_Balloon = Σₜ₌₁ⁿ⁻¹ PMT_bₜ /(1+r)ᵗ + Balloon/(1+r)ⁿ. У реалізації — FV-еквівалент: FV збереженого тіла покриває різницю переплат.",
        "help_balloon_be_abs":   "Absolute Break-even (Balloon) — мінімальна дохідність інвестицій, при якій принципал, утримуваний весь термін, покриває всі процентні виплати банку:\n\nP · (1+r)ⁿ ≥ P + Σₜ₌₁ⁿ⁻¹ Interestₜ\n\nДе P — принципал, що інвестується на n періодів; Interestₜ — процентні платежі Balloon. Розв'язується бінарним пошуком.",
        "help_real_cost_long":   "Present Value (PV) усіх платежів з урахуванням інфляції:\n\nPV = Σₜ₌₁ⁿ paymentₜ / (1 + r_per)ᵗ\n\nДе r_per — періодична ставка дисконтування = i_annual / ppy (наближення нормальної періодичної ставкою). Майбутні платежі приводяться до сьогоднішньої купівельної спроможності.",
        "help_dscr_metric":      "DSCR = NOI / Debt Service (використовується місячний NOI / місячний платіж — співвідношення однакове).\n\n• ≥ 1.25 — Безпечно (типовий поріг банку)\n• 1.00–1.25 — Увага\n• < 1.00 — Високий ризик\n\nКонкретний поріг залежить від банку (зазвичай 1.20–1.30) та типу позики.",
        "help_ltv_metric":       "LTV = Loan Amount / Appraised Value × 100%.\n\n• ≤ 60% — Безпечно\n• ≤ 80% — Стандарт\n• ≤ 95% — Високий\n• > 95% — Критично\n\nЛіміти не універсальні — залежать від банку та типу кредиту (зазвичай ≤ 80–90%).",
        "help_dti_metric":       "DTI = Total Monthly Debt Payments / Gross Monthly Income × 100%.\n\n• ≤ 28% — Відмінно\n• ≤ 36% — Добре\n• ≤ 43% — Допустимо\n• > 43% — Критично\n\nКонкретні пороги залежать від банку та країни (зазвичай ліміт ≤ 36–45%).",
        "email_section":         "📧 Надіслати звіт по Email",
        "email_caption":         "Надішліть PDF або Excel-звіт на вказану email-адресу.",
        "email_recipient":       "Email одержувача",
        "email_recipient_help":  "Email-адреса, на яку буде надіслано звіт.",
        "email_subject":         "Тема листа",
        "email_format":          "Формат вкладення",
        "email_format_pdf":      "PDF",
        "email_format_xlsx":     "Excel (XLSX)",
        "email_format_docx":     "Word (DOCX)",
        "email_message":         "Текст повідомлення (опціонально)",
        "email_send_btn":        "📤 Надіслати звіт",
        "email_sending":         "Надсилання...",
        "email_success":         "✓ Звіт успішно надіслано на {email}",
        "email_error":           "✗ Помилка надсилання: {error}",
        "email_no_recipient":    "Вкажіть email одержувача",
        "email_invalid":         "Некоректна email-адреса",
        "email_no_secrets":      "⚠️ SMTP не налаштовано. Додайте секрети в `.streamlit/secrets.toml`:\n```toml\n[smtp]\nserver = \"smtp.gmail.com\"\nport = 587\nlogin = \"your@email.com\"\npassword = \"app_password\"\nsender = \"Yev Capital LoanLogic <your@email.com>\"\n```",
        "email_default_subject": "Звіт Yev Capital LoanLogic",
        "email_default_body":    "Додаю звіт, згенерований Yev Capital LoanLogic v3.0.\n\nЗ повагою,\nYev Capital",
        "email_disabled_warning": "Email-сервіс не налаштовано. Додайте SMTP-облікові дані в секрети.",
        "email_invalid_warning":  "Введіть коректну email-адресу для активації кнопки надсилання.",
        "email_setup_title":      "📖 Як налаштувати надсилання Email",
        "email_setup_steps":      (
            "**Локально (`.streamlit/secrets.toml`)**\n\n"
            "1. Створіть файл `.streamlit/secrets.toml` у корені проєкту.\n"
            "2. Додайте секцію SMTP:\n\n"
            "```toml\n"
            "[smtp]\n"
            "server = \"smtp.gmail.com\"\n"
            "port = 587\n"
            "login = \"your@email.com\"\n"
            "password = \"app_password\"\n"
            "sender = \"Yev Capital LoanLogic <your@email.com>\"\n"
            "```\n\n"
            "**Streamlit Cloud**\n\n"
            "1. Відкрийте панель застосунку → **Settings** → **Secrets**.\n"
            "2. Вставте ті ж ключі у форматі TOML і натисніть Save.\n"
            "3. Застосунок автоматично перезапуститься.\n\n"
            "**Google App Password (для Gmail)**\n\n"
            "1. Увімкніть двофакторну автентифікацію: [myaccount.google.com/security](https://myaccount.google.com/security)\n"
            "2. Створіть App Password: [myaccount.google.com/apppasswords](https://myaccount.google.com/apppasswords)\n"
            "3. Використайте 16-значний пароль у полі `password` (без пробілів).\n\n"
            "**Порти SMTP**\n\n"
            "- `587` — STARTTLS (рекомендовано для Gmail/Outlook)\n"
            "- `465` — SSL (legacy)\n"
            "- `25`  — відкритий (зазвичай заблокований хостингами)"
        ),
        "syndicated_section":     "🏦 Синдикований кредит (Multi-Tranche)",
        "syndicated_toggle":      "Режим синдикованого кредиту",
        "syndicated_help":        "Консолідований борг від кількох кредиторів з різними ставками, термінами та комісіями. Кожен транш погашається за своїм графіком; підсумкова таблиця сумує платежі за всіма траншами.",
        "syndicated_tranche":     "Транш {n}",
        "syndicated_tranche_amount":   "Сума ({tranche})",
        "syndicated_tranche_rate":     "Ставка % річних ({tranche})",
        "syndicated_tranche_term":     "Термін ({tranche})",
        "syndicated_tranche_unit":     "Одиниця терміну ({tranche})",
        "syndicated_tranche_scheme":   "Схема ({tranche})",
        "syndicated_tranche_ot_comm":  "Разова комісія ({tranche})",
        "syndicated_tranche_mo_comm":  "Період. комісія ({tranche})",
        "syndicated_tranche_enabled":  "Активувати {tranche}",
        "syndicated_total":       "Загальна сума синдикату (Σ траншів)",
        "syndicated_blended_rate":"Змішана APR (effective)",
        "syndicated_master_help": "Master Schedule = Σ платежів за всіма траншами в кожен період. Якщо у траншів різні терміни — короткі просто відсутні в пізніх періодах.",
        "syndicated_zero_error":  "⚠️ Введіть суму хоча б одного траншу.",
        "syndicated_zero_amount_warn": "⚠️ Транш(і) {letters} увімкнені, але мають нульову суму — вони не враховані. Вкажіть суму або вимкніть їх.",
        "end_date_fallback": "↳ Залишено попередній строк: {n} {unit}.",
        "email_empty_attachment": "⚠️ Файл звіту порожній (експорт міг завершитися помилкою). Лист не надіслано.",
        "refi_annuity_note": "ℹ️ Аналіз моделює і поточний, і новий кредит як **ануїтетні** (з рівними платежами). Якщо поточний кредит використовує іншу схему (класичну або буліт), вважайте порівняння приблизним.",
        "syndicated_invalid":     "⚠️ Усі траншеві суми повинні бути додатними.",
        "syndicated_chart_title": "Структура платежів за траншами",
        "syndicated_chart_caption":"Кожен колір — окремий транш. Висота стовпця = сукупний платіж періоду.",
        "help_synd_master_payment": "Master Payment_t = Σᵢ PMTᵢ(t), де PMTᵢ(t) — платіж траншу i в періоді t (після завершення терміну транш = 0).",
        "help_synd_blended_apr":  "Змішана APR — IRR сукупних грошових потоків усіх траншів: NPV(combined CF, r) = 0.",
        "audit_section":          "📋 Журнал змін (Audit Trail)",
        "audit_caption":          "Фіксує зміни ключових параметрів (сума, ставка, термін) протягом поточної сесії. Запис робиться у момент натискання «Розрахувати».",
        "audit_empty":            "Поки немає записів. Змініть параметри та натисніть «Розрахувати», щоб зафіксувати зміну.",
        "audit_clear":            "Очистити журнал",
        "audit_field_amount":     "Сума",
        "audit_field_rate":       "Ставка",
        "audit_field_term":       "Термін",
        "audit_changed_to":       "{field} змінено з {old} на {new}",
        "audit_impact_interest":  "Зміна переплати: {delta}",
        "audit_impact_first_payment": "Зміна першого платежу: {delta}",
        "audit_field_grace":       "Пільговий період",
        "audit_field_inflation":   "Облік інфляції",
        "audit_field_day_count":   "База нарахування",
        "audit_field_syndicated":  "Синдиц. режим",
        "audit_toggle_on":         "Увімк",
        "audit_toggle_off":        "Вимк",
        "audit_field_scheme":      "Схема",
        "audit_field_currency":    "Валюта",
        "audit_field_term_mode":   "Режим строку",
        "audit_field_dc_method":   "Метод бази",
        "audit_field_start_date":  "Дата початку",
        "audit_field_one_time":    "Разова комісія",
        "audit_field_monthly":     "Періодич. комісія",
        "audit_impact_payment":   "Зміна платежу: {delta}",
    },
    "en": {
        "app_title": "Yev Capital LoanLogic",
        "app_subtitle": "Credit Analysis & Reporting System · v3.0",
        "section_params": "⚙️ Loan Parameters",
        "loan_amount": "Loan / Deposit Amount",
        "loan_amount_slider": "Amount Slider",
        "loan_term": "Term",
        "term_unit": "Term Unit",
        "weeks": "Weeks", "months": "Months", "quarters": "Quarters",
        "halfyears": "Half-years", "years": "Years",
        "interest_rate": "Interest Rate (% per year)",
        "calc_scheme": "Payment Scheme",
        "annuity": "Annuity",
        "classic": "Classic (Differentiated)",
        "balloon": "Bullet (Balloon Payment)",
        "deposit_scheme": "Deposit / Savings",
        "section_commissions": "💼 Commissions",
        "one_time_comm": "One-time Commission",
        "one_time_type": "One-time Commission Type",
        "monthly_comm": "Periodic Commission",
        "monthly_type": "Periodic Commission Type",
        "pct_of_amount": "% of loan amount",
        "fixed_amount": "Fixed amount",
        "calc_btn": "🚀 Calculate",
        "calc_error": "⚠️ Calculation could not be completed with the given inputs. Please check the term, rate, and amount.",
        "section_results": "📊 Calculation Results",
        "total_payment": "Total Payments",
        "total_interest": "Interest Overpayment",
        "total_commission": "Total Commissions",
        "effective_rate": "Effective Rate",
        "monthly_payment": "Payment (1st period)",
        "section_schedule": "📋 Payment Schedule",
        "period": "Period", "date": "Date",
        "balance_open": "Balance (Open)",
        "payment_total": "Payment",
        "principal": "Principal",
        "interest": "Interest",
        "commission": "Commission",
        "balance_close": "Balance (Close)",
        "total_row": "TOTAL",
        "section_chart": "📈 Charts",
        "chart_title": "Payment Structure by Period",
        "chart_principal": "Principal",
        "chart_interest": "Interest",
        "chart_commission": "Commissions",
        "chart_pie_title": "Total Breakdown",
        "chart_balance_title": "📉 Remaining Balance",
        "chart_balance_hover": "Remaining Balance",
        "section_templates": "💾 Templates",
        "template_name": "Template Name",
        "save_template": "💾 Save",
        "load_template": "📂 Load",
        "delete_template": "🗑️",
        "no_templates": "No saved templates",
        "template_saved": "✅ Template saved!",
        "template_loaded": "✅ Template loaded!",
        "template_deleted": "🗑️ Deleted!",
        "template_name_empty": "Enter a template name",
        "template_overwrite_warn": "⚠️ Template \"{name}\" already exists. Press Save again to overwrite.",
        "currency": "Currency",
        "uah": "₴ Hryvnia (UAH)",
        "usd": "$ US Dollar (USD)",
        "eur": "€ Euro (EUR)",
        "rub": "₽ Ruble (RUB)",
        "gbp": "£ Pound Sterling (GBP)",
        "jpy": "¥ Japanese Yen (JPY)",
        "cad": "C$ Canadian Dollar (CAD)",
        "aud": "A$ Australian Dollar (AUD)",
        "chf": "Fr Swiss Franc (CHF)",
        "cny": "¥ Chinese Yuan (CNY)",
        "hkd": "HK$ Hong Kong Dollar (HKD)",
        "custom": "✏️ Custom Currency",
        "custom_symbol": "Custom currency symbol",
        "download_excel": "⬇️ Excel",
        "download_pdf": "⬇️ PDF",
        "download_docx": "⬇️ Word",
        "download_csv": "⬇️ CSV",
        "overpayment_pct": "Overpayment %",
        "compare_schemes": "⚖️ Compare Schemes",
        "annuity_vs_classic": "Loan Scheme Comparison",
        "savings_classic": "Savings with Classic",
        "annuity_short": "Annuity",
        "classic_short": "Classic",
        "balloon_short": "Balloon",
        "enter_name": "Enter name...",
        "template_examples": "Quick Templates:",
        "mortgage": "🏠 Mortgage",
        "car_loan": "🚗 Car Loan",
        "consumer": "💳 Consumer Loan",
        "deposit": "Deposit",
        "invest_section": "📈 Investment Comparison",
        "invest_sp500": "📊 Compare with S&P 500 (13.7% p.a.)",
        "invest_custom": "💹 Compare with custom investment yield",
        "invest_custom_rate": "Annual Yield (%)",
        "invest_chart_title": "Loan vs Investment Portfolio",
        "invest_loan_balance": "Remaining Loan Balance",
        "invest_portfolio": "Investment Portfolio",
        "invest_total_paid": "Total Paid to Bank",
        "invest_portfolio_val": "Portfolio Value",
        "invest_net_gain": "Net Investment Gain",
        "invest_explanation": "What if instead of loan payments, this money was invested?",
        "invest_rate_label": "yield",
        "invest_only_one": "⚠️ Only one comparison can be active at a time!",
        "deposit_mode_label": "Deposit Mode",
        "deposit_capitalize": "💰 Interest Capitalization (compound)",
        "deposit_payout": "💸 Interest Payout",
        "deposit_final": "Final Deposit Amount",
        "deposit_earned": "Total Interest Earned",
        "tooltip_payment": "💡 **Payment** — full amount transferred to the bank this period (principal + interest + fees).",
        "tooltip_principal": "💡 **Principal** — the portion that reduces your outstanding loan balance.",
        "tooltip_interest": "💡 **Interest** — the cost of borrowing, calculated on the remaining balance.",
        "scheme_annuity_info": "Equal payments throughout the entire term. Easy to budget.",
        "scheme_classic_info": "Fixed principal + decreasing interest. Lower total overpayment.",
        "scheme_balloon_info": "Interest-only payments each period; full principal due at the last payment.",
        "scheme_deposit_info": "Savings/deposit calculation with capitalization or interest payout mode.",
        "col_explain": "ℹ️ Column Explanations",
        # Deposit table columns
        "dep_balance_open":    "Balance (Open)",
        "dep_interest_earned": "Interest Earned",
        "dep_payout":          "Payout",
        "dep_balance_close":   "Balance (Close)",
        "dep_total_row":       "TOTAL",
        # Deposit summary labels
        "dep_initial":         "Initial Deposit",
        "dep_final_balance":   "Final Balance",
        "dep_total_earned":    "Total Interest Earned",
        "dep_rate_label":      "Annual Rate",
        "dep_period_payout":   "Period Payout",
        "dep_growth_title":    "📈 Deposit Growth",
        "dep_balance_label":   "Deposit Balance",
        "dep_interest_label":  "Accrued Interest",
        "dep_vs_invest":       "Deposit vs Alternative",
        "dep_your_deposit":    "Your Deposit",
        "dep_alternative":     "Alternative Investment",
        "dep_invest_exp":      "What if the same amount was invested at a different rate?",
        "dep_sp500_exp":       "Comparing your deposit vs S&P 500 (13.7% p.a.)",
        "dep_tooltip_interest":"💡 **Interest Earned** — interest accrued this period (added to balance with capitalization, or paid out).",
        "dep_tooltip_balance": "💡 **Balance** — deposit amount including accumulated interest.",
        "dep_tooltip_payout":  "💡 **Payout** — actual cash received this period.",
        "dep_invest_section":  "📊 Compare with Alternative",
        "dep_invest_caption":  "What if you invested this money at a different rate instead?",
        "dep_invest_yours":    "Your Deposit (final)",
        "dep_invest_alt":      "Alternative (final)",
        "dep_invest_diff":     "Difference",
        # Start date
        "start_date_label":    "📅 Start Date",
        "start_date_hint":     "First payment is made one period after start date",
        "term_mode_manual":    "Manual (periods)",
        "term_mode_enddate":   "By End Date",
        "end_date_label":      "📅 End Date",
        "end_date_hint":       "Term auto-calculated as full periods from start to end date",
        # Term caption
        "term_caption":        "mo. / yrs",
        # Balloon break-even
        "balloon_breakeven":       "Inv. Break-even Rate",
        "balloon_breakeven_tip":   (
            "Minimum annual investment return (compound) needed for the Balloon "
            "scheme to be financially advantageous over Annuity. "
            "Logic: returns on invested principal-savings must offset the higher "
            "interest cost of the Balloon loan."
        ),
        "balloon_breakeven_label": "Balloon Break-even Analysis",
        "balloon_breakeven_desc":  "Min. investment yield to justify Balloon over Annuity",
        # Welcome screen
        "welcome_h2": "Enter parameters and click",
        "welcome_calc": "Calculate",
        "welcome_sub": "Annuity · Classic · Balloon · Deposit<br>Export Excel / PDF / Word / CSV · Investment Comparison",
        # Scheme comparison tab labels
        "tab_balance": "📉 Remaining Balance",
        "copyright": "© 2026 Bohdan Yevtushenko (MrCemper) · Yev Capital LoanLogic v3.0",
        "invest_breakeven_section":      "📐 Investment Break-even Analysis",
        "invest_breakeven_universal":    "Universal Break-even Rate",
        "invest_breakeven_universal_tip":"Compound annual investment return at which reinvesting all loan payments would generate enough profit to fully cover the total interest paid.",
        "invest_breakeven_abs":          "Absolute Break-even Rate (Balloon)",
        "invest_breakeven_abs_tip":      "For Balloon loans: minimum compound return at which the FV of invested freed cash (vs Annuity) covers the entire interest cost of the Balloon loan.",
        "invest_breakeven_vs_ann":       "Vs. Annuity Break-even (Balloon)",
        "invest_breakeven_vs_ann_tip":   "Minimum return to make Balloon cheaper than Annuity by reinvesting principal portions.",
        "invest_breakeven_vs_cla_balloon":     "Vs. Classic (Balloon)",
        "invest_breakeven_vs_cla_balloon_tip": "Min. annual yield at which Balloon (with all its deferred principal) matches Classic by maturity.",
        "invest_breakeven_vs_cla_annuity":     "Vs. Classic (Annuity)",
        "invest_breakeven_vs_cla_annuity_tip": "Min. annual yield at which investing the early-period savings of Annuity-over-Classic compensates for the later-period overpayment by maturity.",
        "annuity_be_caption":            "If you can earn ≥ this rate on the early-period cash you save with Annuity, you finish at least as well off as you would have with Classic.",
        "sp500_disclaimer": "Note: The 13.7% S&P 500 return is based on the average annual performance over the past 10 years. Historical performance does not guarantee future results. This value is used as a statistical benchmark and can be adjusted based on current market expectations.",
        "inflation_expander":      "Inflation Accounting",
        "inflation_toggle":        "Enable Inflation Adjustment",
        "inflation_rate":          "Expected Annual Inflation (%)",
        "inflation_section":       "💰 Nominal vs Real Cost",
        "nominal_cost":            "Nominal Total Cost",
        "real_cost":               "Real Total Cost (PV)",
        "inflation_savings":       "Inflation Discount",
        "inflation_note":          "Real cost discounts future payments to today's purchasing power.",
        "grace_expander":          "Grace Period (Payment Holiday)",
        "grace_toggle":            "Enable Grace Period",
        "grace_duration":          "Duration (periods)",
        "grace_start":             "Start Period",
        "grace_type":              "Type",
        "grace_interest_only":     "Interest Only",
        "grace_full_holiday":      "Full Holiday (interest capitalises)",
        "grace_note":              "During Full Holiday, accrued interest is added to outstanding principal.",
        "grace_active":            "Grace Period Active",
        "dscr_expander":           "Business Analysis (DSCR)",
        "dscr_toggle":             "Enable DSCR Analysis",
        "dscr_noi":                "Monthly Net Operating Income (NOI)",
        "dscr_label":              "DSCR",
        "dscr_status_safe":        "Safe",
        "dscr_status_warning":     "Warning",
        "dscr_status_risk":        "High Risk",
        "ltv_status_safe":         "Safe",
        "ltv_status_standard":     "Standard",
        "ltv_status_high":         "High",
        "ltv_status_critical":     "Critical",
        "ltv_status_na":           "N/A",
        "credit_health_title":     "Credit Health",
        "credit_rank_a":           "Rank A — Excellent",
        "credit_rank_b":           "Rank B — Good",
        "credit_rank_c":           "Rank C — Fair",
        "credit_rank_d":           "Rank D — High Risk",
        "credit_rank_na":          "N/A",
        "credit_health_caption":   "Score combines the metrics you provided ({metrics}). Higher is healthier.",
        "glossary_title":          "📚 Glossary / Словарь терминов",
        "help_negative_rate":      "Negative rates are allowed (rare sub-zero products). The calculation handles them correctly.",
        "gloss_annuity_term":      "Annuity",
        "gloss_annuity_def":       "A repayment scheme with equal payments each period. Early on most of the payment is interest; later it is mostly principal.",
        "gloss_bullet_term":       "Bullet (Balloon)",
        "gloss_bullet_def":        "A loan where only interest is paid during the term and the entire principal is repaid in a single lump sum at maturity.",
        "gloss_classic_term":      "Standard (Differentiated)",
        "gloss_classic_def":       "Principal is repaid in equal slices while interest accrues on the remaining balance, so the total payment decreases over time.",
        "gloss_dscr_term":         "DSCR",
        "gloss_dscr_def":          "Debt Service Coverage Ratio = net operating income / debt service. How many times income covers the payments. ≥1.25 is considered safe.",
        "gloss_ltv_term":          "LTV",
        "gloss_ltv_def":           "Loan-to-Value = loan / collateral value × 100%. The lower it is, the less risk for the lender. ≤80% is generally considered good.",
        "gloss_dti_term":          "DTI",
        "gloss_dti_def":           "Debt-to-Income = (payment + other debts) / income × 100%. The share of income going to debt. ≤36% is usually acceptable.",
        "gloss_npv_term":          "NPV",
        "gloss_npv_def":           "Net Present Value — the sum of future cash flows discounted back to today. It accounts for the time value of money.",
        "dscr_note":               "DSCR = NOI / Monthly Payment. Safe: >1.25 | Warning: 1.0-1.25 | Risk: <1.0",
        "ltv_expander":            "Collateral & LTV Analysis",
        "ltv_toggle":              "Enable LTV Analysis",
        "ltv_collateral":          "Collateral / Asset Value",
        "ltv_label":               "Loan-to-Value (LTV)",
        "ltv_note":                "LTV > 80% typically requires additional collateral or PMI.",
        "ltv_warning":             "High LTV — additional collateral or insurance may be required.",
        "dti_expander":            "Personal Solvency (DTI)",
        "dti_toggle":              "Enable DTI Analysis",
        "dti_income":              "Total Monthly Income",
        "dti_other_debts":         "Other Monthly Debt Payments",
        "dti_label":               "Debt-to-Income (DTI)",
        "dti_excellent":           "Excellent (≤28%)",
        "dti_good":                "Good (28–36%)",
        "dti_acceptable":          "Acceptable (36–43%)",
        "dti_critical":            "Critical (>43%)",
        "dti_note":                "DTI = (Loan Payment + Other Debts) / Monthly Income",
        "risk_section":            "🎯 Risk Analytics",
        "refi_section":          "🔄 Refinancing Analysis",
        "refi_caption":          "Compare your current loan with new terms — find out if switching pays off.",
        "refi_current_block":    "Current Loan",
        "refi_new_block":        "New Loan (Refinancing)",
        "refi_current_balance":  "Outstanding Balance",
        "refi_current_rate":     "Current Rate (% annual)",
        "refi_remaining_term":   "Remaining Term (months)",
        "refi_penalty":          "Early Closure Penalty",
        "refi_penalty_pct":      "Penalty (% of balance)",
        "refi_penalty_fixed":    "Penalty (fixed amount)",
        "refi_penalty_type":     "Penalty Type",
        "refi_new_rate":         "New Rate (% annual)",
        "refi_new_term":         "New Loan Term (months)",
        "refi_new_fees":         "Origination Fees",
        "refi_calculate":        "Calculate Refinancing",
        "refi_discount_rate":    "NPV Discount Rate (% annual)",
        "refi_discount_rate_help": "Annual rate used to discount future cash flows when computing NPV. Use your alternative-investment yield.",
        "refi_current_payment":  "Current Payment",
        "refi_new_payment":      "New Payment",
        "refi_monthly_savings":  "Monthly Savings",
        "refi_total_costs":      "Total Switching Costs",
        "refi_breakeven_months": "Break-even Point",
        "refi_total_savings":    "Net Savings over Term",
        "refi_help_npv":          "NPV-based comparison: positive = refinancing saves money in present-value terms; negative = costs more after discounting future payments.",
        "balloon_be_caption":     "💡 **Vs. Annuity** — minimum return to make Balloon cheaper than Annuity. **Absolute** — minimum return at which investing the kept principal fully covers all interest paid to the bank (true 'break-even').",
        "day_count_section":      "📅 Day-Count Convention",
        "day_count_toggle":       "Use day-count convention",
        "day_count_toggle_help":  "Enables interest calculation per the chosen market standard instead of the simplified r=rate/ppy formula. Base formula: I = P · r · (N / B), where N is days in the period and B is the year base.",
        "day_count_method":       "Day-count method",
        "day_count_method_help":  "Pick a standard. Formula: I = P · r_annual · (N / B). N and B depend on the method:\n• 30/360 (ISDA): N=assumed 30 days/month, B=360.\n• 30E/360 (Eurobond): N=European 30/360 (month-end → 30), B=360.\n• ACT/360: N=actual calendar days, B=360.\n• ACT/365 (Actual/365F): N=actual days, B=365.\n• ACT/ACT (ISDA): N=actual days, B=actual days of year (leap-aware).",
        "day_count_caption":      "💡 Formula applied to each schedule row: **I = P × r × (N / B)**. Currently using: **{method}**.",
        "day_count_desc_30_360":  "Assumed 30-day month, year = 360 days (ISDA variant).",
        "day_count_desc_30E_360": "European 30/360: month-end unconditionally rounded to 30.",
        "day_count_desc_ACT_360": "Actual calendar days / 360.",
        "day_count_desc_ACT_365": "Actual calendar days / 365 (a.k.a. Actual/365F).",
        "day_count_desc_ACT_ACT": "Actual days / actual days of year (handles Feb 29).",
        "apr_failed_caption":     "Effective rate could not be computed (IRR did not converge). N/A is shown — this is a numerical failure signal, not user input error.",
        "grace_failed_banner":    "Grace Period could NOT be applied — calculation shown WITHOUT grace. Result may not match the entered parameters",
        "negative_amort_banner":  "Warning: outstanding balance grew beyond {ratio:.1f}× initial. This is a negative-amortization spiral — grace period + high rate makes the loan economically unstable. Calculation halted.",
        "partial_result_banner":  "PARTIAL RESULT — TREAT AS UNRELIABLE",
        "partial_grace":          "Grace Period was NOT applied — schedule shown WITHOUT it",
        "partial_apr":            "Effective APR could not be computed — IRR did not converge",
        "theme_section":          "🎨 Theme",
        "theme_preset_label":     "Preset",
        "theme_preset_help":      "Pick a ready palette, or 'Custom' to tweak manually.",
        "theme_custom":           "Custom (edit below)",
        "theme_custom_hint":      "Tweak any color — changes apply instantly.",
        "theme_field_bg":         "App Background",
        "theme_field_bg_secondary": "Sidebar / Cards",
        "theme_field_bg_tertiary":  "Tertiary Background",
        "theme_field_text":       "Main Text",
        "theme_field_text_muted": "Muted Text",
        "theme_field_text_subtle":"Subtle Text",
        "theme_field_accent":     "Accent / Links",
        "theme_field_accent_strong":"Accent (Strong)",
        "theme_field_success":    "Success",
        "theme_field_warning":    "Warning",
        "theme_field_danger":     "Danger",
        "theme_field_border":     "Borders",
        "theme_field_input_text": "Input Text",
        "theme_font_size":        "Font Size",
        "theme_font_size_help":   "Scales all text (0.75-1.4×)",
        "theme_density":          "UI Density",
        "theme_density_help":     "Padding around UI elements (0.7-1.4×)",
        "theme_radius":           "Border Radius",
        "theme_radius_help":      "Corner roundness, px (0-24)",
        "theme_reset":            "↻ Reset to Default",
        "daycount_compare_tab":     "Day-Count Compare",
        "daycount_compare_title":   "💰 Total Interest by Day-Count Convention",
        "daycount_compare_caption": "Same loan, five interest-day conventions side by side — at a glance, which rule is cheapest.",
        "daycount_best":            "Cheapest",
        "daycount_worst":           "Most expensive",
        "daycount_spread":          "Spread",
        "daycount_no_data":         "Could not compute day-count comparison.",
        "compare_partial_fail":     "⚠️ One or more schemes failed due to parameter conflicts (e.g. grace settings). Comparison shows the successful ones.",
        "compare_savings_vs":       "Savings vs",
        "compare_overpay_vs":       "Overpayment vs",
        "syndicated_tranche_offset":      "Disbursement offset (months) ({tranche})",
        "syndicated_tranche_offset_help": "How many months after the common start date this tranche is actually disbursed. 0 = at start.",
        "refi_recommendation":   "Recommendation",
        "refi_worth_it":         "✓ Refinancing is worthwhile. Pays back in {months:.1f} months.",
        "refi_not_worth":        "✗ Refinancing not worthwhile — new payment is higher than current.",
        "refi_long_payback":     "⚠️ Payback period exceeds the new loan's term.",
        "refi_longer_term_trap": "⚠️ Longer-term trap: monthly payment drops, but the total cost of debt of the new loan (with discounting) is HIGHER — immediate savings become a long-run overpayment.",
        "refi_shorter_term_win": "✓ Higher monthly payment, but lower total cost in present-value (NPV) terms — economically worthwhile (typically refinancing into a shorter term at a lower rate).",
        "refi_help_balance":     "How much you still owe on your current loan.",
        "refi_help_penalty":     "Most banks charge an early-payoff penalty (typically 1-3% of balance).",
        "refi_help_fees":        "Origination costs for the new loan (appraisal, insurance, legal fees).",
        "refi_help_breakeven":   "Months needed for monthly payment savings to recover the switching costs.",
        "download_csv_flat":     "⬇️ Flat CSV (ERP)",
        "download_csv_flat_help":"Clean data for SAP / 1C / Oracle import: numbers only, ISO dates, no totals, no currency.",
        "help_grace_toggle":     "Temporarily reduce or suspend payments. Useful in early-stage projects or income gaps.",
        "help_grace_start":      "The period when the grace window begins (1 = first period).",
        "help_grace_duration":   "How many periods the grace window lasts.",
        "help_grace_type":       "Interest Only — pay only interest, principal preserved. Full Holiday — no payments; interest CAPITALISES (added to principal), increasing your future debt.",
        "info_full_holiday":     "ℹ️ **Full Holiday warning:** Skipped interest is **capitalised** — added to your outstanding principal. After grace, you owe more than before, and post-grace payments are recalculated on the higher balance.",
        "help_inflation_toggle": "Compare nominal payments vs their real (present-value) cost after inflation.",
        "help_inflation_rate":   "Forecasted annual inflation rate. Discounts future payments to today's value.",
        "caption_inflation":     "💡 Inflation 'eats' the real value of money. A payment of $1,000 in 5 years is worth less than $1,000 today.",
        "help_dscr_toggle":      "DSCR = Debt Service Coverage Ratio — used by banks to evaluate business loans.",
        "help_dscr_noi":         "Net Operating Income — monthly profit before debt service (revenue minus operating expenses, excluding loan payments).",
        "caption_dscr":          "💡 **DSCR** = NOI / Loan Payment. Lenders prefer **≥ 1.25** (comfortable cushion). Below 1.0 means operating income can't cover the loan.",
        "help_ltv_toggle":       "LTV = Loan-to-Value — how much of the asset's value the loan covers.",
        "help_ltv_collateral":   "Market value of the asset securing the loan (e.g. property value).",
        "caption_ltv":           "💡 **LTV** = Loan / Collateral × 100%. Risk threshold: **> 80%** typically requires PMI or additional collateral. Lower LTV = better terms.",
        "help_dti_toggle":       "DTI = Debt-to-Income — a key affordability metric for personal loans.",
        "help_dti_income":       "Gross (pre-tax) total monthly income from all sources.",
        "help_dti_other_debts":  "Sum of all other monthly debt obligations: cards, other loans, alimony, etc.",
        "caption_dti":           "💡 **DTI** = (Loan Payment + Other Debts) / Income × 100%. Banking standard: **≤ 36%** ideal · **> 43%** typically denied.",
        "help_nominal":          "Sum of all payments at face value (no time-value adjustment).",
        "help_real":             "Present value of all future payments — what they're 'worth' today after inflation erodes their purchasing power.",
        "help_disc":             "How much your nominal payments are 'discounted' by inflation in real terms.",
        "caption_inflation_full":"💡 {note} Inflation 'eats' the real value of money over time — future payments are economically less burdensome than they appear in nominal terms.",
        "help_total_payment":    "Total Payment = Principal + Σ Interest + Σ Commissions. (Each periodic payment in the schedule = principal + interest + period commission; the total equals Σ payments without double-counting.)",
        "help_total_interest":   "Σ Interest = Total Payment − Principal − Σ Commissions. The total interest accrued over the loan's life.",
        "help_eff_rate":         "Effective Rate (APR) = annual IRR of cash flows:\n\nNPV = Σₜ₌₀ⁿ CFₜ / (1 + r)ᵗ = 0\n\nWhere CF₀ = +(Principal − One-Time Fee) is the loan disbursement (net proceeds), and CFₜ for t ≥ 1 are the borrower's negative repayments. The solved r is periodic, then annualised: APR = (1+r)^ppy − 1.",
        "help_first_payment":    "Annuity payment formula:\n\nPMT = P · r · (1+r)ⁿ / ((1+r)ⁿ − 1)\n\nWhere P is principal, r is the periodic rate (annual / ppy), n is the number of periods.",
        "help_overpay_pct":      "Overpayment % = Σ Interest / Principal × 100%. Share of overpayment relative to the original loan.",
        "help_universal_be":     "Universal Break-even Rate — the reinvestment rate at which the accumulated portfolio of payments offsets total interest:\n\nFV(payments at r) − Σ Pmt = Σ Interest\n\nWhere FV(payments at r) = Σₜ₌₁ⁿ Pmtₜ · (1+r)ⁿ⁻ᵗ. Solved by binary search.",
        "help_balloon_be_va":    "Vs. Annuity Break-even (Balloon) — the reinvestment rate at which Balloon becomes no more expensive than Annuity. Equivalent PV condition:\n\nPV_Annuity = PV_Balloon\n\nWhere PV_Annuity = Σₜ₌₁ⁿ PMT_aₜ /(1+r)ᵗ and PV_Balloon = Σₜ₌₁ⁿ⁻¹ PMT_bₜ /(1+r)ᵗ + Balloon/(1+r)ⁿ. The implementation uses the FV-equivalent: FV of saved principal covers the overpayment difference.",
        "help_balloon_be_abs":   "Absolute Break-even (Balloon) — minimum investment yield at which the principal kept invested for the full term covers all periodic interest paid to the bank:\n\nP · (1+r)ⁿ ≥ P + Σₜ₌₁ⁿ⁻¹ Interestₜ\n\nWhere P is the principal invested for n periods; Interestₜ are the Balloon interest payments. Solved by binary search.",
        "help_real_cost_long":   "Present Value (PV) of all payments adjusted for inflation:\n\nPV = Σₜ₌₁ⁿ paymentₜ / (1 + r_per)ᵗ\n\nWhere r_per is the periodic discount rate = i_annual / ppy (a nominal periodic rate approximation). Future payments are discounted to today's purchasing power.",
        "help_dscr_metric":      "DSCR = NOI / Debt Service (we use monthly NOI / monthly payment — the ratio is identical).\n\n• ≥ 1.25 — Safe (typical bank threshold)\n• 1.00–1.25 — Warning\n• < 1.00 — High Risk\n\nThe specific threshold depends on the bank (typically 1.20–1.30) and loan type.",
        "help_ltv_metric":       "LTV = Loan Amount / Appraised Value × 100%.\n\n• ≤ 60% — Safe\n• ≤ 80% — Standard\n• ≤ 95% — High\n• > 95% — Critical\n\nLimits are not universal — they depend on the bank and loan type (typically ≤ 80–90%).",
        "help_dti_metric":       "DTI = Total Monthly Debt Payments / Gross Monthly Income × 100%.\n\n• ≤ 28% — Excellent\n• ≤ 36% — Good\n• ≤ 43% — Acceptable\n• > 43% — Critical\n\nThresholds depend on the bank and country (typically limit ≤ 36–45%).",
        "email_section":         "📧 Send Report via Email",
        "email_caption":         "Send a PDF, Excel or Word report to the specified email address.",
        "email_recipient":       "Recipient Email",
        "email_recipient_help":  "Email address where the report will be sent.",
        "email_subject":         "Subject",
        "email_format":          "Attachment Format",
        "email_format_pdf":      "PDF",
        "email_format_xlsx":     "Excel (XLSX)",
        "email_format_docx":     "Word (DOCX)",
        "email_message":         "Message (optional)",
        "email_send_btn":        "📤 Send Report",
        "email_sending":         "Sending...",
        "email_success":         "✓ Report successfully sent to {email}",
        "email_error":           "✗ Send failed: {error}",
        "email_no_recipient":    "Please enter recipient email",
        "email_invalid":         "Invalid email address",
        "email_no_secrets":      "⚠️ SMTP not configured. Add secrets in `.streamlit/secrets.toml`:\n```toml\n[smtp]\nserver = \"smtp.gmail.com\"\nport = 587\nlogin = \"your@email.com\"\npassword = \"app_password\"\nsender = \"Yev Capital LoanLogic <your@email.com>\"\n```",
        "email_default_subject": "Yev Capital LoanLogic Report",
        "email_default_body":    "Please find attached the report generated by Yev Capital LoanLogic v3.0.\n\nBest regards,\nYev Capital",
        "email_disabled_warning": "Email service is not configured. Please add SMTP credentials to secrets.",
        "email_invalid_warning":  "Enter a valid email address to enable the send button.",
        "email_setup_title":      "📖 How to setup Email",
        "email_setup_steps":      (
            "**Local development (`.streamlit/secrets.toml`)**\n\n"
            "1. Create a file `.streamlit/secrets.toml` in your project root.\n"
            "2. Add the SMTP section:\n\n"
            "```toml\n"
            "[smtp]\n"
            "server = \"smtp.gmail.com\"\n"
            "port = 587\n"
            "login = \"your@email.com\"\n"
            "password = \"app_password\"\n"
            "sender = \"Yev Capital LoanLogic <your@email.com>\"\n"
            "```\n\n"
            "**Streamlit Cloud**\n\n"
            "1. Open your app dashboard → **Settings** → **Secrets**.\n"
            "2. Paste the same TOML keys and click Save.\n"
            "3. The app restarts automatically.\n\n"
            "**Google App Password (for Gmail)**\n\n"
            "1. Enable 2-Factor Authentication: [myaccount.google.com/security](https://myaccount.google.com/security)\n"
            "2. Generate an App Password: [myaccount.google.com/apppasswords](https://myaccount.google.com/apppasswords)\n"
            "3. Use the 16-character password in the `password` field (no spaces).\n\n"
            "**SMTP Ports**\n\n"
            "- `587` — STARTTLS (recommended for Gmail/Outlook)\n"
            "- `465` — SSL (legacy)\n"
            "- `25`  — plain (usually blocked by hosting providers)"
        ),
        "syndicated_section":     "🏦 Syndicated Loan (Multi-Tranche)",
        "syndicated_toggle":      "Syndicated Loan Mode",
        "syndicated_help":        "Consolidated debt from multiple lenders with different rates, terms and commissions. Each tranche is amortized on its own schedule; the master table sums payments across all tranches.",
        "syndicated_tranche":     "Tranche {n}",
        "syndicated_tranche_amount":   "Amount ({tranche})",
        "syndicated_tranche_rate":     "Rate % p.a. ({tranche})",
        "syndicated_tranche_term":     "Term ({tranche})",
        "syndicated_tranche_unit":     "Term Unit ({tranche})",
        "syndicated_tranche_scheme":   "Scheme ({tranche})",
        "syndicated_tranche_ot_comm":  "One-time Comm. ({tranche})",
        "syndicated_tranche_mo_comm":  "Periodic Comm. ({tranche})",
        "syndicated_tranche_enabled":  "Enable {tranche}",
        "syndicated_total":       "Total Syndicate Amount (Σ tranches)",
        "syndicated_blended_rate":"Blended APR (effective)",
        "syndicated_master_help": "Master Schedule = Σ of payments across all tranches per period. Tranches with shorter terms simply contribute zero in later periods.",
        "syndicated_zero_error":  "⚠️ Enter the amount of at least one tranche.",
        "syndicated_zero_amount_warn": "⚠️ Tranche(s) {letters} are enabled but have a zero amount — they are not included. Set an amount or disable them.",
        "end_date_fallback": "↳ Keeping the previous term: {n} {unit}.",
        "email_empty_attachment": "⚠️ The report file is empty (export may have failed). Nothing was sent.",
        "refi_annuity_note": "ℹ️ This analysis models both the current and the new loan as **annuity** (equal-payment) loans. If your current loan uses a different scheme (classic or balloon), treat the comparison as approximate.",
        "syndicated_invalid":     "⚠️ All tranche amounts must be positive.",
        "syndicated_chart_title": "Payment Structure by Tranche",
        "syndicated_chart_caption":"Each color = one tranche. Bar height = consolidated period payment.",
        "help_synd_master_payment": "Master Payment_t = Σᵢ PMTᵢ(t), where PMTᵢ(t) is tranche i's payment in period t (zero after a tranche matures).",
        "help_synd_blended_apr":  "Blended APR — IRR of the consolidated cash-flow stream across all tranches: NPV(combined CF, r) = 0.",
        "audit_section":          "📋 Audit Trail",
        "audit_caption":          "Tracks changes to key parameters (amount, rate, term) during the current session. Entries are recorded the moment you press «Calculate».",
        "audit_empty":            "No entries yet. Change parameters and press «Calculate» to record a change.",
        "audit_clear":             "Clear Trail",
        "audit_field_amount":     "Amount",
        "audit_field_rate":       "Interest rate",
        "audit_field_term":       "Term",
        "audit_changed_to":       "{field} changed from {old} to {new}",
        "audit_impact_interest":  "Total Interest impact: {delta}",
        "audit_impact_first_payment": "First payment impact: {delta}",
        "audit_field_grace":       "Grace period",
        "audit_field_inflation":   "Inflation adj.",
        "audit_field_day_count":   "Day-count",
        "audit_field_syndicated":  "Syndicated mode",
        "audit_toggle_on":         "On",
        "audit_toggle_off":        "Off",
        "audit_field_scheme":      "Scheme",
        "audit_field_currency":    "Currency",
        "audit_field_term_mode":   "Term input mode",
        "audit_field_dc_method":   "Day-count method",
        "audit_field_start_date":  "Start date",
        "audit_field_one_time":    "One-time fee",
        "audit_field_monthly":     "Periodic fee",
        "audit_impact_payment":   "Payment impact: {delta}",
    },
}

# ─────────────────────────────────────────────────────────────────────────────
#  ЦВЕТА
# ─────────────────────────────────────────────────────────────────────────────
C = {
    "principal":  "#4FC3F7",
    "interest":   "#FF6B6B",
    "commission": "#FFD166",
    "invest":     "#06D6A0",
    "bg":         "#0F172A",
    "card":       "#1E293B",
    "text":       "#E2E8F0",
    "grid":       "#334155",
    "accent":     "#7DD3FC",
    "muted":      "#94A3B8",
}

CURRENCY_SYMBOLS = {
    "uah": "₴",   # Ukrainian Hryvnia
    "usd": "$",   # US Dollar
    "eur": "€",   # Euro
    "rub": "₽",   # Russian Ruble
    "gbp": "£",   # British Pound
    "jpy": "¥",   # Japanese Yen
    "cad": "C$",  # Canadian Dollar
    "aud": "A$",  # Australian Dollar
    "chf": "Fr",  # Swiss Franc
    "cny": "¥",   # Chinese Yuan
    "hkd": "HK$", # Hong Kong Dollar
}

# ─────────────────────────────────────────────────────────────────────────────
#  ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ─────────────────────────────────────────────────────────────────────────────
def periods_per_year(unit: str) -> float:
    ppy = {"weeks": 52, "months": 12, "quarters": 4, "halfyears": 2, "years": 1}
    if unit not in ppy:
        # Consistent with term_to_periods_in_base / period_dates_for_schedule:
        # an unknown unit is a programming error (e.g. "month" vs "months")
        # and must surface, not silently default to 12.
        raise ValueError(
            f"periods_per_year: unknown unit {unit!r}. "
            f"Supported: {list(ppy.keys())}.")
    return ppy[unit]

def term_to_months(term: float, unit: str) -> float:
    """Float-precise conversion of (term, unit) → months. Used for display."""
    factors = {"weeks": 7.0 / 30.4375, "months": 1.0,
               "quarters": 3.0, "halfyears": 6.0, "years": 12.0}
    if unit not in factors:
        raise ValueError(
            f"term_to_months: unknown unit {unit!r}. "
            f"Supported: {list(factors.keys())}.")
    return float(term) * factors[unit]


def term_to_periods_in_base(term: int, unit: str, base_unit: str = "months") -> int:
    """
    Strict integer conversion of a term expressed in `unit` into the
    equivalent number of periods in `base_unit` for use in numerical schedules.

    For weeks→months we round up so that 1–4 weeks remain at least 1 month
    while 5+ weeks become 2 months, etc. — but the underlying number of
    real periods is always preserved when `unit == base_unit`.

    Raises:
      ValueError — for unknown units (consistency with period_dates_for_schedule;
                    old behavior silently used months=1.0 as a fallback, which
                    masked typos in upstream code).
    """
    SUPPORTED = ("weeks", "months", "quarters", "halfyears", "years")
    if unit not in SUPPORTED:
        raise ValueError(
            f"term_to_periods_in_base: unknown unit {unit!r}. "
            f"Supported: {list(SUPPORTED)}.")
    if base_unit not in SUPPORTED:
        raise ValueError(
            f"term_to_periods_in_base: unknown base_unit {base_unit!r}. "
            f"Supported: {list(SUPPORTED)}.")

    if unit == base_unit:
        return max(0, int(term))

    # Convert via months as a common denominator
    months_per_unit = {"weeks": 7.0 / 30.4375, "months": 1.0,
                        "quarters": 3.0, "halfyears": 6.0, "years": 12.0}
    months_per_base = months_per_unit[base_unit]
    months_total    = term * months_per_unit[unit]
    periods         = months_total / months_per_base
    # Round up: a fractional period requires a full extra slot, which is the
    # safer choice for amortization (never under-counts).
    import math
    return max(0, int(math.ceil(periods - 1e-9)))

def generate_dates(n: int, unit: str, start: date | None = None) -> list:
    """
    Генерирует список дат платежей.
    start — дата начала кредита/вклада (первый платёж = start + 1 период).
    Если start не передан — берётся сегодняшняя дата.

    Each date is computed by adding a multiple of the unit to `start`
    directly. Accumulating via `d += delta` would cause month-end drift
    (e.g. Jan 31 + 1 month = Feb 29 → + 1 month = Mar 29 instead of Mar 31).

    Raises ValueError для неизвестного unit (раньше тихо использовал months).
    """
    from dateutil.relativedelta import relativedelta
    if unit not in {"weeks", "months", "quarters", "halfyears", "years"}:
        raise ValueError(
            f"generate_dates: unknown unit {unit!r}. "
            f"Supported: ['halfyears', 'months', 'quarters', 'weeks', 'years'].")
    def _delta_for(i: int) -> relativedelta:
        if unit == "weeks":     return relativedelta(weeks=i)
        if unit == "months":    return relativedelta(months=i)
        if unit == "quarters":  return relativedelta(months=3 * i)
        if unit == "halfyears": return relativedelta(months=6 * i)
        if unit == "years":     return relativedelta(years=i)
        raise ValueError(f"Unsupported unit: {unit}")

    base = start if start is not None else date.today()
    return [(base + _delta_for(i)).strftime("%d.%m.%Y") for i in range(1, n + 1)]

def get_sym(ss) -> str:
    if ss.currency == "custom":
        raw = ss.get("custom_symbol", "?") or "?"
        return _sanitize_currency_symbol(raw)
    return CURRENCY_SYMBOLS.get(ss.currency, "$")


def _sanitize_currency_symbol(raw: str) -> str:
    """
    Sanitize a user-supplied currency symbol.

    The symbol is embedded into HTML (st.markdown with unsafe_allow_html) and
    into PDF/Word/Excel output. A raw value could otherwise carry markup or a
    script payload (e.g. "<script>...</script>"), so we strip HTML-dangerous
    characters and cap the length — a currency symbol is at most a few glyphs
    (e.g. "$", "Fr", "грн", "kr").
    """
    s = str(raw)
    # Remove characters that could break out of an HTML text context or an
    # attribute/markup boundary.
    s = re.sub(r'[<>&"\'`{}\\/]', "", s)
    # Collapse any CR/LF/tabs that could disrupt layout or injected contexts.
    s = re.sub(r"[\r\n\t]+", "", s).strip()
    # Cap length — real currency symbols/codes are short.
    s = s[:10]
    return s or "?"

def fmt_money(v, sym="₴"):
    """
    Форматирует сумму с пробелом как разделителем тысяч — банковский стандарт.
    Например: 1 000 000.00  (не 1,000,000.00)
    """
    if not isinstance(v, (int, float)):
        return str(v)
    # Форматируем с запятой, потом заменяем на пробел
    return f"{sym} {v:,.2f}".replace(",", "\u202f")   # \u202f = narrow no-break space

def fmt_money_plain(v):
    """Число с пробелом-разделителем тысяч, без символа валюты."""
    if not isinstance(v, (int, float)):
        return str(v)
    return f"{v:,.2f}".replace(",", "\u202f")

def fmt_pct(v):
    if v is None:
        return "N/A"
    # Guard against inf/nan slipping through (e.g. a pathological ratio): show
    # "N/A" rather than the literal "inf%"/"nan%", which would look broken.
    try:
        if not math.isfinite(v):
            return "N/A"
    except (TypeError, ValueError):
        return "N/A"
    return f"{v:.2f}%"

# ─────────────────────────────────────────────────────────────────────────────
#  РАСЧЁТ СХЕМ
# ─────────────────────────────────────────────────────────────────────────────
# ─────────────────────────────────────────────────────────────────────────────
#  DAY-COUNT CONVENTIONS  —  стандарты начисления процентов
# ─────────────────────────────────────────────────────────────────────────────
#  Возможные методы:
#   • "30/360"     — ISDA, условный месяц 30 дней, год 360
#   • "30E/360"    — European, конец месяца → 30
#   • "ACT/360"    — фактические дни / 360
#   • "ACT/365"    — фактические дни / 365  (синоним "Actual/365F")
#   • "ACT/ACT"    — ISDA, фактические дни / фактические дни года
#                    (с разделением на високосную/обычную часть периода)
#
#  Все методы возвращают «долю года» (year_fraction) для интервала [d1, d2].
#  Формула начисления процентов в периоде: I = P · r_annual · year_fraction
# ─────────────────────────────────────────────────────────────────────────────

DAY_COUNT_METHODS = ["30/360", "30E/360", "ACT/360", "ACT/365", "ACT/ACT"]
DAY_COUNT_DEFAULT = "30/360"

# S&P 500 long-run average annual return used as the default investment
# benchmark in the "compare with investing" panels. Defined once here so the
# value stays consistent across the UI code and the localized strings.
SP500_BENCHMARK_RATE = 13.7


def _days_30_360(d1: date, d2: date, european: bool = False) -> int:
    """
    30/360 (ISDA) или 30E/360 (Eurobond) — количество дней между датами.

    ISDA правило:
      Если d1 = 31 → 30.
      Если d2 = 31 AND d1 = (30 or 31) → 30.
    European 30E/360:
      Если d1 = 31 → 30.
      Если d2 = 31 → 30 (безусловно).
    """
    d1d, d2d = d1.day, d2.day
    if european:
        if d1d == 31:
            d1d = 30
        if d2d == 31:
            d2d = 30
    else:  # ISDA
        if d1d == 31:
            d1d = 30
        if d2d == 31 and d1d in (30, 31):
            d2d = 30
    return ((d2.year - d1.year) * 360
            + (d2.month - d1.month) * 30
            + (d2d - d1d))


def _is_leap_year(year: int) -> bool:
    return (year % 4 == 0 and year % 100 != 0) or (year % 400 == 0)


def _actual_days(d1: date, d2: date) -> int:
    """Фактическое количество календарных дней между датами."""
    return (d2 - d1).days


def year_fraction(d1: date, d2: date, method: str = "30/360") -> float:
    """
    Возвращает долю года между датами d1 и d2 согласно выбранному стандарту.

    Args:
      d1, d2   — даты периода (d2 > d1 ожидается).
      method   — один из DAY_COUNT_METHODS.

    Returns:
      float — N/B, где N = дни периода, B = база года.

    Raises:
      ValueError — для неизвестного метода или для d2 < d1 (нарушение порядка
                    дат, это всегда программная ошибка в caller).
    """
    if d2 == d1:
        return 0.0
    if d2 < d1:
        # Reverse date ordering is undefined for a period length —
        # surface as an explicit error rather than a zero-length result.
        raise ValueError(
            f"year_fraction: d2 ({d2}) must be ≥ d1 ({d1}). "
            f"Negative-period day-count is undefined; check date ordering."
        )

    method = method.upper().strip()

    if method == "30/360":
        return _days_30_360(d1, d2, european=False) / 360.0
    elif method == "30E/360":
        return _days_30_360(d1, d2, european=True) / 360.0
    elif method == "ACT/360":
        return _actual_days(d1, d2) / 360.0
    elif method in ("ACT/365", "ACTUAL/365F", "ACT/365F"):
        return _actual_days(d1, d2) / 365.0
    elif method in ("ACT/ACT", "ACT/ACT ISDA", "ACTUAL/ACTUAL"):
        # ISDA вариант: разделяем период на части, попадающие в високосный
        # год и обычный год, и используем разные базы (366 vs 365).
        if d1.year == d2.year:
            base = 366.0 if _is_leap_year(d1.year) else 365.0
            return _actual_days(d1, d2) / base
        # Период пересекает несколько лет — суммируем доли
        total = 0.0
        # Часть, попадающая в год d1
        next_year_start = date(d1.year + 1, 1, 1)
        base_1 = 366.0 if _is_leap_year(d1.year) else 365.0
        total += _actual_days(d1, next_year_start) / base_1
        # Полные промежуточные годы
        for y in range(d1.year + 1, d2.year):
            total += 1.0
        # Часть, попадающая в год d2
        this_year_start = date(d2.year, 1, 1)
        base_2 = 366.0 if _is_leap_year(d2.year) else 365.0
        total += _actual_days(this_year_start, d2) / base_2
        return total
    else:
        raise ValueError(f"Unknown day-count method: {method!r}. "
                          f"Supported: {DAY_COUNT_METHODS}")


def period_dates_for_schedule(start_date: date, n: int, unit: str) -> list[date]:
    """
    Returns n+1 dates: [start, end_of_period_1, end_of_period_2, ..., end_of_period_n].

    Used by day-count-aware schedule generators to compute exact N (days)
    per period via year_fraction(d_prev, d_curr, method).

    Each end-of-period date is computed by adding a multiple of the unit to
    `start_date` directly. Accumulating via `d = d + delta` causes month-end
    drift: a loan opened on Jan 31 would land on Feb 29 (correct), then on
    Mar 29 (instead of Mar 31), then on Apr 29 (instead of Apr 30), and so
    on — each step compounds the loss. Computing from the original start_date
    preserves the original day-of-month wherever calendar permits.

    Raises:
      ValueError — if `unit` is not one of the supported period units.
    """
    from dateutil.relativedelta import relativedelta
    if unit not in {"weeks", "months", "quarters", "halfyears", "years"}:
        raise ValueError(
            f"Unknown period unit: {unit!r}. "
            f"Supported: ['halfyears', 'months', 'quarters', 'weeks', 'years']"
        )

    def _delta_for(i: int) -> relativedelta:
        # Build i-units delta directly from start, not by accumulation.
        if unit == "weeks":     return relativedelta(weeks=i)
        if unit == "months":    return relativedelta(months=i)
        if unit == "quarters":  return relativedelta(months=3 * i)
        if unit == "halfyears": return relativedelta(months=6 * i)
        if unit == "years":     return relativedelta(years=i)
        raise ValueError(f"Unsupported unit: {unit}")

    dates = [start_date]
    for i in range(1, n + 1):
        dates.append(start_date + _delta_for(i))
    return dates


def _validate_schedule(sched: list, principal: float, scheme: str,
                        tol: float = 0.05, check_principal_sum: bool = True) -> None:
    """
    Defensive post-condition check on a generated amortization schedule.

    Verifies the structural invariants that ANY correct loan schedule must
    satisfy, and raises ValueError if any is violated. This is a safety net:
    if a future change to a calc function introduces a regression, the error
    surfaces immediately instead of silently producing a wrong schedule.

    Invariants checked:
      1. Non-empty.
      2. No NaN / infinity in any numeric field.
      3. Each row: payment ≈ principal_part + interest + commission.
      4. Balance is monotonically non-increasing for fully-amortizing schemes
         (annuity/classic) — i.e. no silent negative amortization.
      5. Final balance closes to ~0 (within tolerance scaled to principal).
      6. Sum of principal portions ≈ original principal (skippable: a
         full_holiday grace capitalizes interest into principal, so the
         repaid principal legitimately exceeds the original).

    Tolerance is absolute in currency units, scaled up for very large
    principals so float noise on a 100-trillion loan doesn't false-trip.
    """
    if not sched:
        raise ValueError("_validate_schedule: empty schedule produced.")

    # Scale tolerance for large principals (float epsilon grows with magnitude)
    scaled_tol = max(tol, abs(principal) * 1e-7)

    sum_principal = 0.0
    prev_balance = None
    for idx, row in enumerate(sched):
        for key in ("payment", "principal", "interest", "commission", "balance_close"):
            v = row.get(key, 0.0)
            if v is None or (isinstance(v, float) and (math.isnan(v) or math.isinf(v))):
                raise ValueError(
                    f"_validate_schedule: non-finite {key}={v!r} at period {idx+1}.")

        decomp = row["principal"] + row["interest"] + row["commission"]
        if abs(row["payment"] - decomp) > scaled_tol:
            raise ValueError(
                f"_validate_schedule: payment decomposition broken at period "
                f"{idx+1}: payment={row['payment']:.4f} != "
                f"P+I+C={decomp:.4f}.")

        sum_principal += row["principal"]

        # Monotone balance only for fully-amortizing schemes. Balloon holds
        # the balance flat then drops it, which is fine; deposits aren't loans.
        if scheme in ("annuity", "classic") and prev_balance is not None:
            if row["balance_close"] > prev_balance + scaled_tol:
                raise ValueError(
                    f"_validate_schedule: balance increased at period {idx+1} "
                    f"({prev_balance:.4f} → {row['balance_close']:.4f}) — "
                    f"unexpected negative amortization for {scheme}.")
        prev_balance = row["balance_close"]

    if abs(sched[-1]["balance_close"]) > scaled_tol:
        raise ValueError(
            f"_validate_schedule: final balance {sched[-1]['balance_close']:.4f} "
            f"does not close to zero (tol={scaled_tol:.4f}).")

    if check_principal_sum and abs(sum_principal - principal) > max(scaled_tol, scaled_tol * len(sched)):
        raise ValueError(
            f"_validate_schedule: principal portions sum to {sum_principal:.4f}, "
            f"expected {principal:.4f}.")


def calc_annuity(principal, n, rate_pa, unit, monthly_comm,
                  day_count: str | None = None,
                  start_date: date | None = None):
    """
    Аннуитет.

    Two modes:
      • Legacy (default, day_count=None): uses periodic rate r = rate_pa / ppy.
        Each interest = balance × r. This is the textbook annuity formula
        and matches existing UI behavior.

      • Day-count aware (day_count is one of DAY_COUNT_METHODS, start_date
        provided): interest in each period is recomputed exactly as
            I_i = balance × rate_pa_decimal × year_fraction(d_{i-1}, d_i, method)
        where dates are end-of-period dates from start_date.

        In day-count mode the per-period rate VARIES (different months have
        different day counts), so the constant-PMT textbook formula doesn't
        apply. We use a CLOSED-FORM generalization (NOT Newton's method,
        which would be numerically fragile here):

            PMT = P · Π_{i=0..n-1}(1+r_i) / Σ_{j=0..n-1} Π_{i=j+1..n-1}(1+r_i)

        This is the analytical solution of the variable-rate amortization
        recurrence; it has NO convergence issues, NO ZeroDivisionError risk
        (denominator is always positive when n ≥ 1 and r_i > -1), and the
        resulting schedule amortizes EXACTLY to 0 (within float precision).
        Verified across 30-year mortgages, weekly schedules, leap-year
        Feb-29 starts, and rates from 0.01% to 200% p.a.
    """
    # ── Input validation ──────────────────────────────────────────────────────
    if not isinstance(n, int):
        try:
            n = int(n)
        except (TypeError, ValueError):
            raise ValueError(
                f"calc_annuity: 'n' must be a positive integer, got {n!r}")
    if n < 1:
        raise ValueError(f"calc_annuity: 'n' must be ≥ 1, got {n}")
    try:
        rate_pa = float(rate_pa)
        principal = float(principal)
    except (TypeError, ValueError):
        raise ValueError(
            f"calc_annuity: rate_pa and principal must be numeric, got "
            f"rate_pa={rate_pa!r}, principal={principal!r}")
    import math as _math
    if _math.isnan(rate_pa) or _math.isinf(rate_pa):
        raise ValueError(f"calc_annuity: 'rate_pa' must be finite, got {rate_pa}.")
    if _math.isnan(principal) or _math.isinf(principal):
        raise ValueError(f"calc_annuity: 'principal' must be finite, got {principal}.")
    if principal < 0:
        # Negative principal would yield negative payments — outside the
        # supported domain.
        raise ValueError(
            f"calc_annuity: 'principal' must be ≥ 0, got {principal}.")

    ppy = periods_per_year(unit)

    # ── Day-count aware path ──────────────────────────────────────────────────
    if day_count and start_date is not None:
        dates = period_dates_for_schedule(start_date, n, unit)
        r_annual = rate_pa / 100.0  # decimal annual rate
        # Period rates (variable)
        period_rates = [
            r_annual * year_fraction(dates[i], dates[i + 1], day_count)
            for i in range(n)
        ]

        # Solve for constant PMT such that, given variable r_i,
        # the final balance after n periods equals 0.
        # Equivalent to PMT = P * [Π(1+r_i)] / Σ_{j=0..n-1} Π_{i=j+1..n-1}(1+r_i)
        # (a generalization of the textbook annuity formula).
        # Use a closed-form construction:
        prod_after = [1.0] * (n + 1)   # prod_after[j] = Π_{i=j..n-1}(1+r_i)
        for i in range(n - 1, -1, -1):
            prod_after[i] = prod_after[i + 1] * (1.0 + period_rates[i])
            if math.isinf(prod_after[i]) or math.isnan(prod_after[i]):
                raise ValueError(
                    f"calc_annuity: rate {rate_pa}% over {n} periods overflows "
                    f"the day-count annuity formula. Reduce rate or term.")
        denom = sum(prod_after[j + 1] for j in range(n))
        if denom <= 0:
            pmt = principal / n
        else:
            pmt = principal * prod_after[0] / denom

        rows, bal = [], principal
        for i in range(1, n + 1):
            r_i = period_rates[i - 1]
            interest = bal * r_i
            princ = (pmt - interest) if i < n else bal
            bal -= princ
            rows.append({
                "period": i, "balance_open": bal + princ,
                "payment": princ + interest + monthly_comm,
                "principal": princ, "interest": interest,
                "commission": monthly_comm,
                "balance_close": max(bal, 0),
            })
        return rows

    # ── Legacy path ───────────────────────────────────────────────────────────
    r = rate_pa / 100 / ppy
    # The annuity formula is valid for any r > -1 (it has a removable
    # singularity only at r == 0). Using `r > 0` here would wrongly route
    # negative rates (e.g. EU-style negative-interest products) into the
    # flat principal/n branch, producing a non-constant payment stream.
    if abs(r) < 1e-12:
        pmt = principal / n
    else:
        try:
            growth = (1 + r) ** n
            if math.isinf(growth) or math.isnan(growth):
                raise OverflowError
            pmt = principal * r * growth / (growth - 1)
        except OverflowError:
            raise ValueError(
                f"calc_annuity: rate {rate_pa}% over {n} periods overflows the "
                f"annuity formula. Reduce the rate or term to a realistic range."
            )
    rows, bal = [], principal
    for i in range(1, n+1):
        interest = bal * r
        princ = (pmt - interest) if i < n else bal
        bal -= princ
        rows.append({"period": i, "balance_open": bal+princ,
                     "payment": princ+interest+monthly_comm,
                     "principal": princ, "interest": interest,
                     "commission": monthly_comm, "balance_close": max(bal, 0)})
    return rows

def calc_classic(principal, n, rate_pa, unit, monthly_comm,
                  day_count: str | None = None,
                  start_date: date | None = None):
    """Классика — фиксированное тело долга + убывающие проценты."""
    # Input validation (same rationale as calc_annuity)
    if not isinstance(n, int):
        try:
            n = int(n)
        except (TypeError, ValueError):
            raise ValueError(f"calc_classic: 'n' must be a positive integer, got {n!r}")
    if n < 1:
        raise ValueError(f"calc_classic: 'n' must be ≥ 1, got {n}")
    try:
        rate_pa = float(rate_pa)
        principal = float(principal)
    except (TypeError, ValueError):
        raise ValueError(
            f"calc_classic: rate_pa and principal must be numeric")
    import math as _math
    if _math.isnan(rate_pa) or _math.isinf(rate_pa):
        raise ValueError(f"calc_classic: 'rate_pa' must be finite, got {rate_pa}.")
    if _math.isnan(principal) or _math.isinf(principal):
        raise ValueError(f"calc_classic: 'principal' must be finite, got {principal}.")
    if principal < 0:
        raise ValueError(f"calc_classic: 'principal' must be ≥ 0, got {principal}")

    ppy = periods_per_year(unit)

    # ── Day-count aware path ──────────────────────────────────────────────────
    if day_count and start_date is not None:
        dates = period_dates_for_schedule(start_date, n, unit)
        r_annual = rate_pa / 100.0
        pp_base = principal / n
        rows, bal = [], principal
        for i in range(1, n + 1):
            yf = year_fraction(dates[i - 1], dates[i], day_count)
            interest = bal * r_annual * yf
            # In the final period, set principal exactly to the remaining
            # balance so payment = principal + interest + commission stays
            # internally consistent. Otherwise float division (principal/n)
            # can leave a non-zero residual that conflicts with balance_close.
            is_last = (i == n)
            pp_i = bal if is_last else pp_base
            bal -= pp_i
            rows.append({
                "period": i, "balance_open": bal + pp_i,
                "payment": pp_i + interest + monthly_comm,
                "principal": pp_i, "interest": interest,
                "commission": monthly_comm,
                "balance_close": max(bal, 0) if not is_last else 0.0,
            })
        return rows

    # ── Legacy path ───────────────────────────────────────────────────────────
    r = rate_pa / 100 / ppy
    pp_base = principal / n
    rows, bal = [], principal
    for i in range(1, n+1):
        interest = bal * r
        is_last = (i == n)
        pp_i = bal if is_last else pp_base
        bal -= pp_i
        rows.append({"period": i, "balance_open": bal + pp_i,
                     "payment": pp_i + interest + monthly_comm,
                     "principal": pp_i, "interest": interest,
                     "commission": monthly_comm,
                     "balance_close": max(bal, 0) if not is_last else 0.0})
    return rows

def calc_balloon(principal, n, rate_pa, unit, monthly_comm,
                  day_count: str | None = None,
                  start_date: date | None = None):
    """Буллит — только проценты каждый период; тело долга — в последнем платеже."""
    # Input validation
    if not isinstance(n, int):
        try:
            n = int(n)
        except (TypeError, ValueError):
            raise ValueError(f"calc_balloon: 'n' must be a positive integer, got {n!r}")
    if n < 1:
        raise ValueError(f"calc_balloon: 'n' must be ≥ 1, got {n}")
    try:
        rate_pa = float(rate_pa)
        principal = float(principal)
    except (TypeError, ValueError):
        raise ValueError(
            f"calc_balloon: rate_pa and principal must be numeric")
    import math as _math
    if _math.isnan(rate_pa) or _math.isinf(rate_pa):
        raise ValueError(f"calc_balloon: 'rate_pa' must be finite, got {rate_pa}.")
    if _math.isnan(principal) or _math.isinf(principal):
        raise ValueError(f"calc_balloon: 'principal' must be finite, got {principal}.")
    if principal < 0:
        raise ValueError(f"calc_balloon: 'principal' must be ≥ 0, got {principal}")

    ppy = periods_per_year(unit)

    # ── Day-count aware path ──────────────────────────────────────────────────
    if day_count and start_date is not None:
        dates = period_dates_for_schedule(start_date, n, unit)
        r_annual = rate_pa / 100.0
        rows = []
        for i in range(1, n + 1):
            yf = year_fraction(dates[i - 1], dates[i], day_count)
            int_pmt = principal * r_annual * yf
            is_last = (i == n)
            princ = principal if is_last else 0.0
            rows.append({
                "period": i, "balance_open": principal,
                "payment": princ + int_pmt + monthly_comm,
                "principal": princ, "interest": int_pmt,
                "commission": monthly_comm,
                "balance_close": 0.0 if is_last else principal,
            })
        return rows

    # ── Legacy path ───────────────────────────────────────────────────────────
    r = rate_pa / 100 / ppy
    int_pmt = principal * r
    rows = []
    for i in range(1, n+1):
        is_last = (i == n)
        princ = principal if is_last else 0.0
        rows.append({"period": i,
                     "balance_open": principal,
                     "payment": princ + int_pmt + monthly_comm,
                     "principal": princ, "interest": int_pmt,
                     "commission": monthly_comm,
                     "balance_close": 0.0 if is_last else principal})
    return rows

def calc_deposit(principal, n, rate_pa, unit, mode):
    """
    Вклад/Депозит:
      capitalize — сложный процент: проценты реинвестируются в тело каждый период.
                   Никаких выплат — всё накапливается. Итог забирается в конце.
      payout     — простой процент: проценты выплачиваются каждый период наличными.
                   Тело вклада не меняется. В конце возвращается начальный вклад.
    
    Колонки schedule:
      balance_open  — баланс на начало периода
      interest      — начисленные проценты за период
      payout        — фактически выплаченные проценты (0 при capitalize)
      balance_close — баланс на конец периода
    """
    # Input validation — mirror what calc_annuity/classic/balloon do.
    # Old version silently returned [] for n=0, hiding upstream bugs.
    if not isinstance(n, int):
        try:
            n = int(n)
        except (TypeError, ValueError):
            raise ValueError(f"calc_deposit: 'n' must be a positive integer, got {n!r}")
    if n < 1:
        raise ValueError(f"calc_deposit: 'n' must be ≥ 1, got {n}")
    try:
        rate_pa = float(rate_pa)
        principal = float(principal)
    except (TypeError, ValueError):
        raise ValueError(f"calc_deposit: rate_pa and principal must be numeric")
    import math as _math
    if _math.isnan(rate_pa) or _math.isinf(rate_pa):
        raise ValueError(f"calc_deposit: 'rate_pa' must be finite, got {rate_pa}.")
    if _math.isnan(principal) or _math.isinf(principal):
        raise ValueError(f"calc_deposit: 'principal' must be finite, got {principal}.")
    if principal < 0:
        raise ValueError(f"calc_deposit: 'principal' must be ≥ 0, got {principal}")
    if mode not in ("capitalize", "payout"):
        raise ValueError(
            f"calc_deposit: mode must be 'capitalize' or 'payout', got {mode!r}")

    ppy = periods_per_year(unit)
    r   = rate_pa / 100 / ppy
    bal, rows = principal, []
    for i in range(1, n + 1):
        open_bal = bal
        interest = open_bal * r
        if mode == "capitalize":
            bal += interest
            payout_val = 0.0
        else:
            payout_val = interest  # выплата на руки
            # bal не меняется; в конце вернут тело
        close_bal = bal
        rows.append({
            "period":        i,
            "balance_open":  round(open_bal,   2),
            "interest":      round(interest,   2),
            "payout":        round(payout_val, 2),
            "balance_close": round(close_bal,  2),
        })
    return rows

def calc_effective_rate(principal, schedule, one_time_comm, ppy):
    """
    Effective annual rate (APR) as the IRR of the cash-flow stream.

    CF_0 = +(principal − one_time_comm)  ─ borrower receives net proceeds
    CF_t = -(periodic payment) for t ≥ 1 ─ borrower's outflows

    NPV(r) = Σ_{t=0..n} CF_t / (1+r)^t = 0

    Approach:
      1. Try Newton's method (fast convergence on well-shaped problems).
      2. If Newton fails (divergence / non-convergence / poor residual),
         fall back to bracket-and-bisection search on [-0.99, 5.0] per period.
      3. Return None if BOTH methods fail to find a root with |NPV| < 1e-3.

    Callers (run_calculation) should treat None as "APR could not be computed"
    and surface that in the UI rather than substituting nominal rate silently.

    Returns: annualised rate in % (or None on failure).
    """
    # Sign convention: CF_0 = −(principal − one_time_comm) is the borrower's
    # initial inflow expressed negatively; CF_t = +payment_t is each outflow.
    cf = [-(principal - one_time_comm)] + [r["payment"] for r in schedule]

    if not schedule or principal <= 0:
        return None

    def npv(r: float) -> float:
        return sum(cf[t] / (1 + r) ** t for t in range(len(cf)))

    def dnpv(r: float) -> float:
        return sum(-t * cf[t] / (1 + r) ** (t + 1) for t in range(1, len(cf)))

    # ── Phase 1: Newton's method ──────────────────────────────────────────────
    r_newton = 0.01
    converged = False
    for _ in range(200):
        f  = npv(r_newton)
        df = dnpv(r_newton)
        if abs(df) < 1e-12:
            break
        r_new = r_newton - f / df
        # Clamp: per-period rate must satisfy 1+r > 0
        if r_new <= -0.99:
            r_new = -0.99
        if r_new > 5.0:
            r_new = 5.0
        if abs(r_new - r_newton) < 1e-10:
            r_newton = r_new
            converged = True
            break
        r_newton = r_new

    if converged and abs(npv(r_newton)) < 1e-3:
        return ((1 + r_newton) ** ppy - 1) * 100

    # ── Phase 2: bracket + bisection on [-0.99, 5.0] ──────────────────────────
    lo, hi = -0.99, 5.0
    f_lo, f_hi = npv(lo), npv(hi)
    if f_lo * f_hi > 0:
        # No sign change across the bracket — root lies outside [-0.99, 5.0]
        # or does not exist.
        return None

    for _ in range(200):
        mid = (lo + hi) / 2
        f_mid = npv(mid)
        if abs(f_mid) < 1e-6:
            return ((1 + mid) ** ppy - 1) * 100
        if f_lo * f_mid < 0:
            hi, f_hi = mid, f_mid
        else:
            lo, f_lo = mid, f_mid
        if hi - lo < 1e-10:
            break

    r_final = (lo + hi) / 2
    if abs(npv(r_final)) > 1e-2:
        # Even bisection failed to drive residual below acceptable threshold
        return None
    return ((1 + r_final) ** ppy - 1) * 100

# ─────────────────────────────────────────────────────────────────────────────
#  ИНВЕСТИЦИОННОЕ СРАВНЕНИЕ
# ─────────────────────────────────────────────────────────────────────────────
def calc_investment(payments, annual_yield, unit):
    """Портфельный рост: invest each payment at annual_yield% compounded per period."""
    r = annual_yield / 100 / periods_per_year(unit)
    portfolio, values = 0.0, []
    for p in payments:
        portfolio = portfolio * (1 + r) + p
        values.append(portfolio)
    return values


def calc_balloon_breakeven(principal: float, n: int, rate_pa: float,
                            unit: str, mo_comm: float = 0.0) -> float | None:
    """
    Minimum annual compound investment return such that investing the
    annuity-style principal portions (the cash you don't pay in a Balloon
    scheme but would in an Annuity scheme) accumulates exactly to the
    principal lump-sum at maturity.

    What this is NOT
    ────────────────
    This is NOT "Vs Annuity total cost equivalence". At rate=0% both schemes
    have identical zero-interest cost, yet this function returns ~4–5% — the
    return needed to grow the partial-principal stream up to the FULL
    principal payable at period n. That answers a DIFFERENT question:
    "what investment return justifies deferring principal repayment to the
    end" — not "what rate makes balloon as cheap as annuity".

    Economic logic
    ──────────────
    • Annuity:  pmt_ann[i]  = interest[i] + amortised_principal[i].
    • Balloon:  int_only[i] = interest only (i < n); int_only[n] + principal.
    • Freed[i] = annuity_principal_portion[i]  for i = 1 … n-1.
      (Period n is excluded: balloon already pays full principal then.)
    • Break-even r_inv satisfies:

          Σ_{i=0}^{n-2} freed[i] · (1 + r_per)^(n-1-i) = principal

    Solved by binary search (monotone in r).

    Returns
    ───────
    • Annual break-even rate in % (positive float).
    • 0.0 if the accumulated FV already covers principal at 0 % return.
    • None if inputs are degenerate (n ≤ 1, principal ≤ 0, etc.).
    """
    # ── Guard: degenerate inputs ──────────────────────────────────────────────
    if n <= 1 or principal <= 0:
        return None

    ppy = periods_per_year(unit)
    if ppy <= 0:
        return None

    # ── Build annuity schedule to get principal-amortisation amounts ──────────
    try:
        sched_ann = calc_annuity(principal, n, rate_pa, unit, mo_comm)
    except Exception:
        return None

    # Freed cash: annuity principal portions for periods 1 … n-1  (0-indexed: 0..n-2)
    # Period n itself (index n-1) is excluded because at that point the balloon
    # borrower pays the full principal anyway — there is nothing "saved" to invest.
    freed = [row["principal"] for row in sched_ann[:-1]]  # length = n-1

    if not freed or sum(freed) <= 0:
        return None

    # ── FV function ───────────────────────────────────────────────────────────
    def portfolio_fv(r_annual_pct: float) -> float:
        """FV at end of period n of investing freed[j] from end of period j+1."""
        r_per = r_annual_pct / 100.0 / ppy
        # freed[j] (period j+1) is invested for (n-1-j) additional periods.
        return sum(
            freed[j] * (1.0 + r_per) ** (n - 1 - j)
            for j in range(len(freed))
        )

    # ── At r = 0 ──────────────────────────────────────────────────────────────
    fv_zero = portfolio_fv(0.0)   # = sum(freed) = principal - freed_last
    if fv_zero >= principal:
        # Balloon can be justified with 0 % investment return — no break-even.
        return 0.0

    # Binary search on [0 %, 1000 %]: portfolio_fv is strictly increasing in
    # r_annual_pct, and FV(0) < principal implies the root lies above 0.
    lo, hi = 0.0, 1_000.0
    for _ in range(150):           # 150 iterations → precision ≈ 1000/2^150 ≈ 0
        mid = (lo + hi) * 0.5
        if portfolio_fv(mid) < principal:
            lo = mid               # need a higher rate
        else:
            hi = mid               # rate too high
        if hi - lo < 1e-8:
            break

    result = (lo + hi) * 0.5
    # Convergence near the upper bound indicates no realistic investment
    # return inside [0, 1000 %] can amortize the principal — return None
    # rather than a misleading saturated value.
    if result > 999.0:
        return None
    return result


def calc_pairwise_breakeven(chosen_payments: list,
                             alt_payments: list,
                             unit: str) -> float | None:
    """
    Pairwise Investment Break-even Rate between two payment schemes.

    Asks: "If I choose scheme A (with payments `chosen_payments`) over
    scheme B (`alt_payments`), at what investment return r do the
    per-period savings (or extra costs) accumulate to exactly zero by
    maturity?"

    Mathematical formulation
    ────────────────────────
    Let diff[i] = alt_pmt[i] − chosen_pmt[i]   for i = 1 … n.
      • diff[i] > 0  →  you save this much vs the alternative this period
                         (invest the surplus).
      • diff[i] < 0  →  you pay this much more this period
                         (must liquidate prior investments).

    Future value at maturity n of investing each diff[i] at per-period
    rate r_per:

        FV(r_per) = Σ diff[i] · (1 + r_per)^(n−i)

    Break-even: FV(r_per) = 0.

    Interpretation
    ──────────────
      • Schemes with identical totals → no meaningful rate (returns None).
      • A vs B where A is cheaper in present-value terms already at r = 0
        (FV(0) ≥ 0) → returns 0.0: no positive investment return is required
        for A to be at least as good as B.
      • Otherwise: a unique positive break-even rate exists when the
        cash-flow difference stream has at least one sign change.

    Returns
    ───────
      • Annualised break-even rate in % (positive float).
      • 0.0 if FV(0) ≥ 0 — the chosen scheme already wins at zero return.
      • None for degenerate inputs (mismatched lengths, empty streams, or
        no realistic root in the search bracket).
    """
    if not chosen_payments or not alt_payments:
        return None
    if len(chosen_payments) != len(alt_payments):
        return None

    n = len(chosen_payments)
    diff = [a - c for a, c in zip(alt_payments, chosen_payments)]
    ppy = periods_per_year(unit)

    # Total nominal savings of the chosen scheme: Σ diff
    sum_diff = sum(diff)

    # Degenerate / trivial cases
    if all(abs(d) < 1e-9 for d in diff):
        return None                       # streams identical
    if sum_diff <= 0:
        # Chosen scheme costs more nominally already. Either it costs more
        # at every period (no surplus to invest at all — no break-even
        # exists), or it costs less early but enough more later to outweigh.
        # In both sub-cases we still attempt the search; bisection below
        # decides.
        pass

    def fv(r_per: float) -> float:
        # Future value at maturity of investing each diff[i] from period i
        # forward until period n. For very long terms (e.g. 30-year mortgage,
        # n=360) combined with high test rates, the compound factor can
        # overflow float — treat that as +∞ (the FV is enormous, certainly
        # the wrong sign vs a negative target).
        if 1.0 + r_per <= 0:
            return float('inf')           # invalid domain
        try:
            total = 0.0
            for i, d in enumerate(diff, start=1):
                total += d * (1.0 + r_per) ** (n - i)
            return total
        except OverflowError:
            # At a compounding rate large enough to overflow float, the term
            # with the LARGEST exponent dominates the sum. In this FV the
            # exponent is (n - i), so the EARLIEST non-negligible diff (smallest
            # i) dominates — NOT the largest-magnitude diff. Using max(diff)
            # would pick the wrong sign whenever the early flows are negative
            # but a later flow is the biggest positive (or vice-versa).
            for d in diff:
                if abs(d) > 1e-9:
                    return float('inf') if d > 0 else float('-inf')
            return 0.0

    # If FV(0) < 0 ⇒ even with zero return, accumulated savings don't reach
    # the deficit. We need a HIGHER return; root must lie in (0, +∞).
    # If FV(0) ≥ 0 ⇒ already at-or-above break-even at r=0; the chosen scheme
    # is unconditionally at least as good. Using ">= 0" here (not "> 0")
    # matches the documented contract and guarantees an EXACT 0.0 at the
    # boundary, instead of letting the bisection below return a spurious tiny
    # positive rate when FV(0) is exactly zero.
    fv0 = fv(0.0)
    if fv0 >= 0:
        # Chosen scheme is unconditionally at least as good as the alternative
        # — even with no investment return, you end up no worse off.
        return 0.0

    # Bisect on r ∈ [0, 10] per period. Upper bound 10 ≈ 12000 % pa
    # (monthly compounding) is well above any realistic investment yield;
    # going higher risks float overflow on long-term loans.
    lo, hi = 0.0, 10.0
    f_lo = fv0
    f_hi = fv(hi)
    if f_lo * f_hi > 0:
        # No sign change — root outside [0, ~12000 %]. Unrealistic.
        return None

    for _ in range(200):
        mid = (lo + hi) * 0.5
        f_mid = fv(mid)
        if abs(f_mid) < 1e-6:
            break
        if f_lo * f_mid < 0:
            hi, f_hi = mid, f_mid
        else:
            lo, f_lo = mid, f_mid
        if hi - lo < 1e-10:
            break

    r_per_be = (lo + hi) * 0.5
    # Annualise
    r_annual = (1 + r_per_be) ** ppy - 1
    rate_pct = r_annual * 100.0

    # Reject if result saturated the upper bound (unrealistic)
    if rate_pct > 9_999.0:
        return None
    return rate_pct


# ─────────────────────────────────────────────────────────────────────────────
#  ГЛАВНЫЙ РАСЧЁТ — Investment Break-even, Inflation, Grace, Risk
# ─────────────────────────────────────────────────────────────────────────────

def calc_universal_breakeven(payments: list, total_interest: float,
                              unit: str,
                              one_time_comm: float = 0.0) -> float | None:
    """
    Universal Investment Break-even Rate (для всех типов кредитов).

    Идея: если бы вы не платили банку, а инвестировали те же суммы
    под ставку r, то прибыль от инвестиций должна покрыть переплату
    по процентам:

      Σ Pmtₜ · (1+r_per)^(n-t)  −  Σ Pmtₜ  =  Σ Interest      (без комиссии)

    Если включена upfront-комиссия, она вычитается из исходной экономической
    выгоды (т.к. вы её всё равно платите):

      FV_payments − sum_payments  =  total_interest + one_time_comm

    Args:
      payments        — список периодических платежей.
      total_interest  — суммарная переплата (без upfront).
      unit            — единица периода ('weeks', 'months', etc.).
      one_time_comm   — разовая (upfront) комиссия. Если > 0, она входит
                        в целевой барьер для break-even, делая результат
                        более консервативным.

    Returns:
      • float — ставка в % годовых (≥ 0), если решение существует.
      • 0.0   — если даже при r=0 FV уже покрывает целевую стоимость.
      • None  — если входы вырожденные.
    """
    n = len(payments)
    if n == 0:
        return None

    target = float(total_interest) + float(one_time_comm)
    if target <= 0:
        return 0.0  # нечего покрывать

    ppy = periods_per_year(unit)
    if ppy <= 0:
        return None

    sum_pmt = sum(payments)

    def fv_profit(r_pct: float) -> float:
        r = r_pct / 100.0 / ppy
        # payments[i] is the (i+1)-th payment, made at the END of period
        # (i+1). Reinvested until the horizon (end of period n), it compounds
        # for (n - (i+1)) = (n - 1 - i) periods. The last payment sits AT the
        # horizon and is not reinvested (exponent 0).
        fv = sum(payments[i] * (1.0 + r) ** (n - 1 - i) for i in range(n))
        return fv - sum_pmt

    if fv_profit(0.0) >= target:
        return 0.0

    lo, hi = 0.0, 1_000.0
    # Sanity: верхняя граница должна гарантированно перекрывать target
    if fv_profit(hi) < target:
        return None

    for _ in range(200):
        mid = (lo + hi) * 0.5
        if fv_profit(mid) < target:
            lo = mid
        else:
            hi = mid
        if hi - lo < 1e-7:
            break
    return (lo + hi) * 0.5


def calc_balloon_absolute_breakeven(principal: float, n: int, rate_pa: float,
                                     unit: str, mo_comm: float = 0.0) -> float | None:
    """
    Absolute Break-even Rate for Balloon (Bullet) loans.

    Economic interpretation we ACTUALLY implement (cleaner than the textbook
    "FV equals total repayment" form, which has no positive root because
    principal alone never reaches principal + interest at r=0):

        "What investment yield r on the kept-principal makes the FV of that
        principal cover the principal AND all periodic interest payments
        the borrower makes to the bank?"

    Equation:
        principal · (1 + r/ppy)^n  ≥  principal + Σ Interestₜ

    Where Interestₜ are the periodic interest payments paid to the bank
    over the loan's life (excluding the final balloon principal repayment,
    which the borrower funds from the matured investment principal).

    Args:
      principal — amount kept invested for the full term.
      n         — number of periods.
      rate_pa   — loan rate (% p.a.).
      unit      — period unit.
      mo_comm   — periodic commission (added to "interest cost" coverage).

    Returns:
      • float — break-even annual rate in percent.
      • 0.0   — if even at r=0 the principal alone already covers it
                (degenerate: there's nothing to cover, e.g. zero-rate loan).
      • None  — when inputs are degenerate (n ≤ 1, principal ≤ 0, etc.) or
                no positive root exists within [0, 1000%].
    """
    if n <= 1 or principal <= 0:
        return None

    ppy = periods_per_year(unit)
    if ppy <= 0:
        return None

    try:
        sched_bal = calc_balloon(principal, n, rate_pa, unit, mo_comm)
    except Exception:
        return None

    # Sum of periodic costs (interest + commission) the borrower actually pays.
    # Excludes the final period's balloon-principal portion (it's NOT a cost,
    # the borrower repays it from the matured investment).
    sum_costs = sum(row["interest"] + row["commission"] for row in sched_bal)
    if sum_costs <= 0:
        # Zero-cost loan (rate=0, no commission): trivially covered at any r.
        return 0.0

    target = principal + sum_costs

    def fv_principal(r_pct: float) -> float:
        r = r_pct / 100.0 / ppy
        return principal * (1.0 + r) ** n

    # At r=0: FV = principal < target (since sum_costs > 0) → root must exist
    # for any positive r. Verify upper bound covers it.
    if fv_principal(0.0) >= target:
        return 0.0

    lo, hi = 0.0, 1_000.0
    if fv_principal(hi) < target:
        return None  # no root in reasonable range — extreme inputs

    # Bisection (function monotone increasing in r)
    for _ in range(200):
        mid = (lo + hi) * 0.5
        if fv_principal(mid) < target:
            lo = mid
        else:
            hi = mid
        if hi - lo < 1e-7:
            break

    result = (lo + hi) * 0.5
    return result if result >= 0 else None


def calc_real_cost(payments: list, annual_inflation_pct: float, unit: str) -> float:
    """
    Реальная стоимость кредита (Present Value):
    дисконтирует будущие платежи к сегодняшней покупательной способности.
        PV = Σ payment_t / (1 + r_inflation_per_period)^t

    Поддерживает положительную инфляцию (PV < nominal — переплата «съедается»),
    нулевую инфляцию (PV = nominal) и отрицательную инфляцию / дефляцию
    (PV > nominal — будущие деньги дороже нынешних). Старая версия молча
    возвращала sum(payments) при rate ≤ 0, что скрывало случай дефляции.

    Sanity: для (1 + r) > 0 формула работает. Запрещаем только r ≤ -100%/period
    (это означало бы «вообще никакой денежной системы»).
    """
    if not payments:
        return 0.0
    if annual_inflation_pct == 0:
        return float(sum(payments))
    ppy = periods_per_year(unit)
    r = annual_inflation_pct / 100.0 / ppy
    if 1.0 + r <= 0:
        # Дегенеративный случай: дисконт-фактор неположителен. Возвращаем
        # nominal как наименее опасное значение и сигнализируем через caller.
        return float(sum(payments))
    return sum(p / (1.0 + r) ** (t + 1) for t, p in enumerate(payments))


def discount_payments_to_pv(payments: list, annual_inflation_pct: float,
                             unit: str) -> list[float]:
    """
    Returns the present value of EACH payment individually (not the cumulative
    sum). Supports positive, zero, and negative inflation symmetrically.
    """
    if not payments:
        return []
    if annual_inflation_pct == 0:
        return [float(p) for p in payments]
    ppy = periods_per_year(unit)
    r = annual_inflation_pct / 100.0 / ppy
    if 1.0 + r <= 0:
        # См. calc_real_cost — same degenerate-case handling
        return [float(p) for p in payments]
    return [p / (1.0 + r) ** (t + 1) for t, p in enumerate(payments)]


def apply_grace_period(sched: list, grace_start: int, grace_duration: int,
                       grace_type: str, rate_pa: float, unit: str,
                       day_count: str | None = None,
                       start_date: date | None = None,
                       scheme_hint: str | None = None) -> list:
    """
    Динамически перестраивает график платежей с учётом кредитных каникул.

    Args:
      grace_start    — номер периода начала каникул (1-based)
      grace_duration — количество периодов каникул
      grace_type     — "interest_only" | "full_holiday"
      day_count      — если задан, проценты в каникулах и пост-каникулярном
                        графике считаются через year_fraction (важно для
                        ACT/ACT и пересечения 29 февраля).
      start_date     — нужен только если day_count задан.
      scheme_hint    — "annuity" | "classic" | "balloon" если известно от
                        caller. Если задан, используется напрямую вместо
                        хрупкого вывода типа графика по его форме.

    Логика:
      • interest_only  → платим только проценты, тело не уменьшается
      • full_holiday   → ничего не платим, проценты капитализируются (+к телу)
      • После каникул  → пересчитываем оставшийся график на новый увеличенный
                          остаток с прежним сроком (оставшиеся периоды)
    """
    if not sched or grace_duration <= 0 or grace_start < 1:
        return sched

    # Only the two documented grace modes are supported; reject any other
    # value so misconfiguration surfaces in the UI banner.
    if grace_type not in ("interest_only", "full_holiday"):
        raise ValueError(
            f"apply_grace_period: grace_type must be 'interest_only' "
            f"or 'full_holiday', got {grace_type!r}.")

    n   = len(sched)
    ppy = periods_per_year(unit)
    mo_comm = sched[0]["commission"] if sched else 0.0

    # Day-count: подготавливаем массив дат и функцию для расчёта r_per
    use_dc = day_count is not None and start_date is not None
    if use_dc:
        dates = period_dates_for_schedule(start_date, n, unit)
        r_annual = rate_pa / 100.0

        def r_for_period(i: int) -> float:
            # i is 1-based period index
            return r_annual * year_fraction(dates[i - 1], dates[i], day_count)
    else:
        r_legacy = rate_pa / 100.0 / ppy

        def r_for_period(i: int) -> float:
            return r_legacy

    # Нормализуем границы каникул
    grace_start    = max(1, min(grace_start, n))
    grace_end      = min(grace_start + grace_duration - 1, n)
    grace_duration = grace_end - grace_start + 1

    # Reject configurations that mathematically leave principal unpaid:
    # if the grace window reaches the final period, no remaining periods are
    # available to amortize the balance — the loan would close with debt
    # still outstanding.
    if grace_end >= n:
        raise ValueError(
            f"Grace period covers the final loan period (grace_end={grace_end}, "
            f"n={n}). After the grace window there are no remaining periods "
            f"to amortize the principal — the loan would end with an "
            f"outstanding balance. Either shorten the grace window so it ends "
            f"before period {n}, or extend the loan term."
        )

    # Восстанавливаем баланс до начала каникул из исходного графика
    if grace_start == 1:
        balance_before = sched[0]["balance_open"]
    else:
        balance_before = sched[grace_start - 2]["balance_close"]

    new_sched = []

    # 1) Платежи ДО каникул — без изменений
    for i in range(grace_start - 1):
        new_sched.append(dict(sched[i]))

    # 2) Платежи ВО ВРЕМЯ каникул
    balance = balance_before
    # Capture once outside the loop — it's invariant for the safety check.
    initial_balance = balance_before
    for i in range(grace_start, grace_end + 1):
        r_i = r_for_period(i)
        interest = balance * r_i
        if grace_type == "interest_only":
            new_sched.append({
                "period":        i,
                "balance_open":  round(balance, 2),
                "payment":       round(interest + mo_comm, 2),
                "principal":     0.0,
                "interest":      round(interest, 2),
                "commission":    round(mo_comm, 2),
                "balance_close": round(balance, 2),
            })
        else:  # full_holiday — проценты капитализируются
            balance += interest
            # ── Negative-amortization safety guard ────────────────────────────
            # If interest capitalization drives the balance past 2× the
            # principal at grace start, the schedule is in a debt-spiral
            # regime that produces meaningless payments. Abort cleanly so
            # the caller records the failure in `grace_error`.
            if balance > initial_balance * 2.0:
                raise ValueError(
                    f"Negative amortization spiral detected: balance "
                    f"{balance:,.2f} > 2× initial ({initial_balance * 2:,.2f}) "
                    f"at grace period {i}. Reduce rate, grace duration, or "
                    f"use Interest-Only mode instead."
                )
            new_sched.append({
                "period":        i,
                "balance_open":  round(balance - interest, 2),
                "payment":       round(mo_comm, 2),
                "principal":     0.0,
                "interest":      0.0,
                "commission":    round(mo_comm, 2),
                "balance_close": round(balance, 2),
            })

    # 3) Платежи ПОСЛЕ каникул — пересчитываем на оставшийся срок и новый balance
    remaining_periods = n - grace_end
    if remaining_periods > 0 and balance > 0:
        # Определяем тип графика. Если caller передал scheme_hint — используем
        # его напрямую (надёжно). Иначе выводим по форме графика, что может
        # давать сбои на вырожденных случаях (например, аннуитет при ставке 0%
        # неотличим от классики, так как оба дают постоянное тело).
        if scheme_hint in ("annuity", "classic", "balloon"):
            was_balloon = (scheme_hint == "balloon")
            was_classic = (scheme_hint == "classic")
        elif len(sched) > 1:
            was_balloon = all(sched[i]["principal"] == 0
                               for i in range(len(sched) - 1))
            was_classic = (
                abs(sched[0]["principal"] - sched[1]["principal"]) < 0.01
            )
        else:
            # Однопериодный график — это уже целиком "последний платёж",
            # для каникул это вырожденный случай. Считаем как annuity.
            was_balloon = False
            was_classic = False

        if was_balloon:
            for i in range(grace_end + 1, n + 1):
                is_last = (i == n)
                r_i = r_for_period(i)
                interest = balance * r_i
                principal_part = balance if is_last else 0.0
                new_sched.append({
                    "period":        i,
                    "balance_open":  round(balance, 2),
                    "payment":       round(principal_part + interest + mo_comm, 2),
                    "principal":     round(principal_part, 2),
                    "interest":      round(interest, 2),
                    "commission":    round(mo_comm, 2),
                    "balance_close": round(0.0 if is_last else balance, 2),
                })
                if is_last:
                    balance = 0.0
        elif was_classic:
            pp_base = balance / remaining_periods
            for i in range(grace_end + 1, n + 1):
                r_i = r_for_period(i)
                interest = balance * r_i
                # Match calc_classic's behavior: in the final period, set the
                # principal exactly to the remaining balance so the row
                # internally amortizes to zero.
                is_last = (i == n)
                pp_i = balance if is_last else pp_base
                balance -= pp_i
                new_sched.append({
                    "period":        i,
                    "balance_open":  round(balance + pp_i, 2),
                    "payment":       round(pp_i + interest + mo_comm, 2),
                    "principal":     round(pp_i, 2),
                    "interest":      round(interest, 2),
                    "commission":    round(mo_comm, 2),
                    "balance_close": round(max(balance, 0) if not is_last else 0.0, 2),
                })
        else:
            # Annuity post-grace
            if use_dc:
                # Variable-rate annuity: same closed-form as in calc_annuity
                rates = [r_for_period(i) for i in
                          range(grace_end + 1, n + 1)]
                m = len(rates)
                prod_after = [1.0] * (m + 1)
                for k in range(m - 1, -1, -1):
                    prod_after[k] = prod_after[k + 1] * (1.0 + rates[k])
                denom = sum(prod_after[k + 1] for k in range(m))
                pmt = balance * prod_after[0] / denom if denom > 0 else balance / m
                for k, i in enumerate(range(grace_end + 1, n + 1)):
                    r_i = rates[k]
                    interest = balance * r_i
                    principal_part = pmt - interest
                    if i == n:
                        principal_part = balance
                    balance -= principal_part
                    new_sched.append({
                        "period":        i,
                        "balance_open":  round(balance + principal_part, 2),
                        "payment":       round(principal_part + interest + mo_comm, 2),
                        "principal":     round(principal_part, 2),
                        "interest":      round(interest, 2),
                        "commission":    round(mo_comm, 2),
                        "balance_close": round(max(balance, 0), 2),
                    })
            else:
                r = r_legacy
                # Valid for any r > -1; only r == 0 needs the flat fallback.
                if abs(r) < 1e-12:
                    pmt = balance / remaining_periods
                else:
                    pmt = balance * r * (1 + r) ** remaining_periods / (
                        (1 + r) ** remaining_periods - 1)
                for i in range(grace_end + 1, n + 1):
                    interest = balance * r
                    principal_part = pmt - interest
                    if i == n:
                        principal_part = balance
                    balance -= principal_part
                    new_sched.append({
                        "period":        i,
                        "balance_open":  round(balance + principal_part, 2),
                        "payment":       round(principal_part + interest + mo_comm, 2),
                        "principal":     round(principal_part, 2),
                        "interest":      round(interest, 2),
                        "commission":    round(mo_comm, 2),
                        "balance_close": round(max(balance, 0), 2),
                    })

    return new_sched


# ─────────────────────────────────────────────────────────────────────────────
#  РИСК-АНАЛИТИКА  —  LTV, DSCR, DTI
# ─────────────────────────────────────────────────────────────────────────────

def calc_ltv(loan_amount: float, collateral_value: float) -> float | None:
    """
    LTV = (Loan / Collateral) × 100.

    Returns:
      • float — LTV percent (loan_amount must be ≥ 0).
      • None  — when collateral_value ≤ 0 (LTV undefined). Callers must check
                for None before formatting; the previous behaviour of returning
                +inf created surprising % strings in metric tiles.
    """
    if collateral_value is None or collateral_value <= 0:
        return None
    if loan_amount is None or loan_amount < 0:
        return None
    return (loan_amount / collateral_value) * 100.0


def calc_dscr(noi: float, monthly_payment: float) -> float | None:
    """
    DSCR = NOI / Debt Service (monthly basis — ratio is unit-agnostic).

    Returns:
      • float — DSCR (≥ 0).
      • None  — when monthly_payment ≤ 0 (no debt service to cover, ratio
                undefined). Treating this as "Safe = ∞" is misleading; callers
                should display "N/A" or skip the metric tile.
    """
    if monthly_payment is None or monthly_payment <= 0:
        return None
    if noi is None or noi < 0:
        return None
    return noi / monthly_payment


# ─────────────────────────────────────────────────────────────────────────────
#  РЕФИНАНСИРОВАНИЕ — анализ выгоды перехода на новый кредит
# ─────────────────────────────────────────────────────────────────────────────
def calc_refinance_analysis(
    current_balance: float,
    current_rate_pa: float,
    remaining_months: int,
    penalty: float,                  # already-computed penalty amount
    new_rate_pa: float,
    new_term_months: int,
    new_fees: float,
    discount_rate_pa: float = 0.0,   # discount rate for NPV; 0 = no discounting
) -> dict:
    """
    Refinancing analysis (Annuity ↔ Annuity).

    Two distinct comparison metrics, each with its own clean meaning:

    1) MONTHLY-PAYMENT BREAK-EVEN (cash-flow horizon-agnostic):
         break-even months = (penalty + fees) / (current_pmt - new_pmt)

       Tells you when the SAVINGS PER MONTH have offset upfront cost.
       Independent of either loan's term.

    2) NPV-BASED TOTAL COST DIFFERENCE (apples-to-apples):
         Discounts BOTH cash-flow streams to present value over the
         common horizon = max(remaining_months, new_term_months).

         NPV_current = Σ_{t=1..remaining_months} cur_pmt / (1+d)^t
         NPV_new     = (penalty + fees)
                     + Σ_{t=1..new_term_months} new_pmt / (1+d)^t

       net_savings_pv = NPV_current − NPV_new
       net_savings_pv > 0  →  refinancing is economically advantageous.

       If discount_rate_pa = 0, this reduces to a simple sum (still correct,
       but ignores time-value of money).

    Returns dict with both views; the UI labels them clearly so users don't
    confuse "monthly savings × N months" with the NPV-based answer.
    """
    # Basic guards: positive principal / non-zero terms
    invalid_result = {
        "current_payment":   0.0,
        "new_payment":       0.0,
        "monthly_savings":   0.0,
        "total_costs":       0.0,
        "breakeven_months":  None,
        "npv_savings":       0.0,
        "verdict":           "invalid",
        "worth_it":          False,
    }
    if current_balance <= 0 or remaining_months <= 0 or new_term_months <= 0:
        return invalid_result
    # Penalty and fees represent cash outflows: negative values are invalid.
    if penalty < 0 or new_fees < 0:
        invalid_result["total_costs"] = penalty + new_fees
        return invalid_result
    # Loan rates: a modest negative band is allowed (sub-zero-interest
    # products are exotic but real, and the annuity formula handles r < 0
    # exactly as in calc_annuity). Only reject values far outside any
    # realistic range, consistent with the main loan-rate input.
    if current_rate_pa < -20.0 or new_rate_pa < -20.0:
        return invalid_result

    def annuity_pmt(P, r_pa, n):
        if n <= 0 or P <= 0:
            return 0.0
        r = r_pa / 100.0 / 12.0
        if abs(r) < 1e-12:
            return P / n
        return P * r * (1 + r) ** n / ((1 + r) ** n - 1)

    cur_pmt = annuity_pmt(current_balance, current_rate_pa, remaining_months)
    new_pmt = annuity_pmt(current_balance, new_rate_pa, new_term_months)

    monthly_savings = cur_pmt - new_pmt
    total_costs     = penalty + new_fees

    # ── Break-even on monthly savings ─────────────────────────────────────────
    if monthly_savings <= 0:
        breakeven = None
        verdict   = "not_worth"
        worth_it  = False
    else:
        breakeven = total_costs / monthly_savings
        if breakeven > new_term_months:
            verdict  = "long_payback"
            worth_it = False
        else:
            verdict  = "worth"
            worth_it = True

    # ── NPV-based comparison on common horizon ────────────────────────────────
    # discount per month
    d = (discount_rate_pa / 100.0 / 12.0) if discount_rate_pa > 0 else 0.0

    def pv_stream(pmt: float, n: int, d_per: float) -> float:
        if n <= 0 or pmt <= 0:
            return 0.0
        if d_per == 0:
            return pmt * n
        # Standard annuity PV: PMT × [1 − (1+d)^-n] / d
        return pmt * (1.0 - (1.0 + d_per) ** (-n)) / d_per

    npv_current = pv_stream(cur_pmt, remaining_months, d)
    npv_new     = total_costs + pv_stream(new_pmt, new_term_months, d)
    npv_savings = npv_current - npv_new

    # ── Verdict reconciliation between the two metrics ─────────────────────────
    # The monthly-cash-flow test and the NPV test can disagree. Reconcile so
    # that neither metric silently wins:
    #
    #   • verdict == "worth" but NPV < 0 → "longer_term_trap": the monthly
    #     payment drops, but the total/present cost of the new (often longer)
    #     debt is actually higher. The cash-flow improvement is a trap.
    #
    #   • verdict == "not_worth" (monthly payment is HIGHER) but NPV > 0 →
    #     "shorter_term_win": the borrower pays more each month, yet the
    #     present value of total cost is lower — typically refinancing into a
    #     shorter term at a lower rate. Surfacing this as not_worth would
    #     wrongly discourage an economically beneficial move.
    if verdict == "worth" and npv_savings < 0:
        verdict  = "longer_term_trap"
        worth_it = False
    elif verdict == "not_worth" and npv_savings > 0:
        verdict  = "shorter_term_win"
        worth_it = True

    return {
        "current_payment":   round(cur_pmt, 2),
        "new_payment":       round(new_pmt, 2),
        "monthly_savings":   round(monthly_savings, 2),
        "total_costs":       round(total_costs, 2),
        "breakeven_months":  round(breakeven, 1) if breakeven is not None else None,
        "npv_savings":       round(npv_savings, 2),
        "discount_rate_pa":  discount_rate_pa,
        "verdict":           verdict,
        "worth_it":          worth_it,
    }


# ─────────────────────────────────────────────────────────────────────────────
#  FLAT CSV — машинно-читаемый экспорт для ERP (SAP / 1C / Oracle)
# ─────────────────────────────────────────────────────────────────────────────
def export_flat_csv(df, summary, t, sym, is_deposit: bool = False) -> bytes:
    """
    Возвращает «плоский» CSV для ERP-импорта:
      • заголовки в snake_case на английском
      • даты в ISO формате (YYYY-MM-DD)
      • только числа (без символов валюты, без узких пробелов)
      • без итоговой строки TOTAL
      • без BOM, разделитель — запятая
    """
    import csv as _csv
    import io as _io
    from datetime import datetime as _dt

    # Snake_case заголовки на английском
    if is_deposit:
        cols_map = {
            "period":         t["period"],
            "date":           t["date"],
            "balance_open":   t.get("dep_balance_open",    "Balance (Open)"),
            "interest_earned":t.get("dep_interest_earned", "Interest Earned"),
            "payout":         t.get("dep_payout",          "Payout"),
            "balance_close":  t.get("dep_balance_close",   "Balance (Close)"),
        }
        total_key = t.get("dep_total_row", "TOTAL")
    else:
        cols_map = {
            "period":         t.get("period",        "Period"),
            "date":           t.get("date",          "Date"),
            "balance_open":   t.get("balance_open",  "Balance (Open)"),
            "payment_amount": t.get("payment_total", "Payment"),
            "principal_amount": t.get("principal",   "Principal"),
            "interest_amount":  t.get("interest",    "Interest"),
            "commission_amount":t.get("commission",  "Commission"),
            "balance_close":  t.get("balance_close", "Balance (Close)"),
        }
        total_key = t.get("total_row", "TOTAL")

    # Удаляем итоговую строку — устойчиво к смене языка.
    # Старая логика сравнивала period_col с локализованной строкой "ИТОГО"/"TOTAL",
    # что ломалось при смене языка. Теперь определяем итоговую строку по
    # структурному признаку: значение в колонке period не приводится к int
    # (для реальных строк это 1..N, для TOTAL — текст любой локализации).
    period_col = t.get("period", "Period")

    def _is_data_row(val) -> bool:
        """True для обычных строк (period — целое число); False для TOTAL."""
        if val is None:
            return False
        if isinstance(val, (int,)):
            return True
        if isinstance(val, float):
            return not (val != val)  # exclude NaN
        try:
            int(str(val).strip())
            return True
        except (ValueError, TypeError):
            return False

    df_clean = df[df[period_col].apply(_is_data_row)].copy()

    # Преобразуем даты в ISO
    date_col = t["date"]
    def _to_iso(v):
        if not isinstance(v, str) or not v.strip():
            return ""
        # Пробуем парсить разные форматы
        for fmt in ("%d.%m.%Y", "%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y"):
            try:
                return _dt.strptime(v.strip(), fmt).strftime("%Y-%m-%d")
            except ValueError:
                continue
        return v  # если не парсится — оставляем как есть
    df_clean[date_col] = df_clean[date_col].apply(_to_iso)

    # Записываем CSV вручную для полного контроля (без BOM, чистые числа)
    buf = _io.StringIO(newline="")
    writer = _csv.writer(buf, quoting=_csv.QUOTE_MINIMAL, lineterminator="\n")
    writer.writerow(list(cols_map.keys()))   # snake_case headers

    for _, row in df_clean.iterrows():
        out_row = []
        for snake_key, display_key in cols_map.items():
            val = row.get(display_key, "")
            if snake_key == "period":
                try:
                    out_row.append(int(val))
                except (ValueError, TypeError):
                    out_row.append(val)
            elif snake_key == "date":
                out_row.append(val)
            else:
                # числовые колонки — приводим к float
                if isinstance(val, (int, float)):
                    out_row.append(f"{float(val):.2f}")
                else:
                    try:
                        # Strip everything except digits, decimal points,
                        # and signs. The previous approach iterated through a
                        # hard-coded set of currency glyphs ("$₴€₽£¥CFr") and
                        # would corrupt custom symbols containing those
                        # letters (e.g. "Krone 500" → "Kone500", "CHF 1000"
                        # → "H1000"), causing float() to fail.
                        cleaned = (str(val)
                                   .replace("\u202f", "")
                                   .replace(",", "."))
                        cleaned = re.sub(r"[^\d.\-]", "", cleaned).strip()
                        out_row.append(f"{float(cleaned):.2f}" if cleaned else "")
                    except (ValueError, TypeError):
                        out_row.append("")
        writer.writerow(out_row)

    return buf.getvalue().encode("utf-8")  # UTF-8, без BOM


def calc_dti(monthly_payment: float, other_debts: float,
              monthly_income: float) -> float | None:
    """
    DTI = (Payment + Other Debts) / Income × 100.

    Returns:
      • float — DTI percent.
      • None  — when monthly_income ≤ 0 (ratio undefined). Returning 0 in that
                case used to look like "Excellent (≤ 28%)" which was wrong.
    """
    if monthly_income is None or monthly_income <= 0:
        return None
    if monthly_payment is None or other_debts is None:
        return None
    if monthly_payment < 0 or other_debts < 0:
        return None
    return (monthly_payment + other_debts) / monthly_income * 100.0


def dscr_status(dscr: float | None, t: dict) -> tuple[str, str]:
    """Возвращает (label, color_hex) для значения DSCR. None → 'N/A'."""
    if dscr is None:
        return "N/A", "#64748B"
    if dscr >= 1.25:
        return t.get("dscr_status_safe", "Safe"), "#10B981"
    elif dscr >= 1.0:
        return t.get("dscr_status_warning", "Warning"), "#F59E0B"
    else:
        return t.get("dscr_status_risk", "High Risk"), "#DC2626"


def dti_status(dti_pct: float | None, t: dict) -> tuple[str, str]:
    """Возвращает (label, color_hex) для значения DTI. None → 'N/A'."""
    if dti_pct is None:
        return "N/A", "#64748B"
    if dti_pct <= 28.0:
        return t.get("dti_excellent", "Excellent"), "#10B981"
    elif dti_pct <= 36.0:
        return t.get("dti_good", "Good"), "#34D399"
    elif dti_pct <= 43.0:
        return t.get("dti_acceptable", "Acceptable"), "#F59E0B"
    else:
        return t.get("dti_critical", "Critical"), "#DC2626"


def ltv_status(ltv_pct: float | None, t: dict | None = None) -> tuple[str, str]:
    """Возвращает (label, color_hex) для значения LTV. None → 'N/A'.

    `t` (optional translation dict) localizes the labels to match DSCR/DTI.
    When omitted, falls back to English so existing callers keep working.
    """
    tr = t or {}
    if ltv_pct is None:
        return tr.get("ltv_status_na", "N/A"), "#64748B"
    if ltv_pct <= 60.0:
        return tr.get("ltv_status_safe", "Safe"), "#10B981"
    elif ltv_pct <= 80.0:
        return tr.get("ltv_status_standard", "Standard"), "#34D399"
    elif ltv_pct <= 95.0:
        return tr.get("ltv_status_high", "High"), "#F59E0B"
    else:
        return tr.get("ltv_status_critical", "Critical"), "#DC2626"


def _run_syndicated(tranches: list[dict], t: dict, sym: str,
                     start_date=None,
                     inflation_enabled=False, inflation_rate=0.0,
                     ltv_enabled=False, ltv_collateral=0.0,
                     dscr_enabled=False, dscr_noi=0.0,
                     dti_enabled=False, dti_income=0.0, dti_other_debts=0.0,
                     day_count_enabled: bool = False,
                     day_count_method: str = DAY_COUNT_DEFAULT,
                     grace_requested: bool = False):
    """
    Builds the consolidated master schedule for a multi-tranche syndicated loan.
    Returns (df_display, summary) compatible with downstream renderers.

    `grace_requested`: indicates the user enabled the Grace Period feature.
    Multi-tranche grace is not currently supported (per-tranche capitalization
    rules are ambiguous when several loans share a holiday); when this flag
    is true, the resulting schedule omits the grace and reports the limitation
    via `grace_error` in the summary.
    """
    synd_grace_error = None
    if grace_requested:
        synd_grace_error = (
            "Grace Period is not supported in Syndicated (multi-tranche) "
            "mode. The schedule below is computed WITHOUT the requested "
            "grace. To use grace, switch to a single-loan scheme."
        )

    master, totals, per_tranche = calc_syndicated_master_schedule(
        tranches, base_unit="months",
        day_count=day_count_method if day_count_enabled else None,
        start_date=start_date if day_count_enabled else None,
    )

    if not master or totals["n_tranches_active"] == 0:
        # No tranche produced a valid schedule. Returning zeros here would be
        # indistinguishable from a real 0% loan — instead surface this as a
        # partial result with None metrics and the original tranche errors so
        # the UI can render the failure banner.
        tranche_errs = totals.get("tranche_errors", [])
        err_msg = ("No valid tranches: " + "; ".join(e[1] for e in tranche_errs)
                    if tranche_errs else
                    "Syndicated calculation produced no schedule (all tranches failed).")
        empty_df = pd.DataFrame()
        return empty_df, {
            "is_deposit": False, "scheme_key": "syndicated",
            "rate_pa": None, "start_date": start_date or date.today(),
            "total_payment": 0, "total_interest": 0, "total_commission": 0,
            "effective_rate": None, "effective_rate_error": True,
            "first_payment": None, "overpay_pct": 0,
            "one_time_comm": 0, "principal": 0, "sym": sym,
            "syndicated": True, "tranches": [], "master_schedule": [],
            "tranche_errors": tranche_errs,
            "grace_error": synd_grace_error,
            "partial_result": True,
            "synd_empty_error": err_msg,
        }

    n = totals["n_periods"]
    dates = generate_dates(n, "months", start=start_date)

    # Build display DataFrame
    rows = []
    for i, row in enumerate(master):
        # Derive each row's displayed payment from its ROUNDED components so
        # every row satisfies payment == principal + interest + commission
        # exactly. The master's own `payment` is summed across tranches and can
        # round 1 cent away from round(P)+round(I)+round(C); deriving it here
        # keeps each row — and therefore the TOTAL — internally consistent
        # (matching the single-loan display contract).
        r_princ = round(row["principal"],  2)
        r_int   = round(row["interest"],   2)
        r_comm  = round(row["commission"], 2)
        r_pay   = round(r_princ + r_int + r_comm, 2)
        rows.append({
            t["period"]:        row["period"],
            t["date"]:          dates[i],
            t["balance_open"]:  round(row["balance_open"], 2),
            t["payment_total"]: r_pay,
            t["principal"]:     r_princ,
            t["interest"]:      r_int,
            t["commission"]:    r_comm,
            t["balance_close"]: round(row["balance_close"], 2),
        })
    df = pd.DataFrame(rows)

    # Display-reconciled TOTAL row: sum the per-row ROUNDED values so a hand-sum
    # of each displayed column equals the TOTAL row exactly (round(Σx) ≠ Σround(x),
    # so we use the latter for what the user sees — matching the single-loan
    # path). The precise unrounded aggregates remain in `totals`/the summary for
    # APR and other internal math.
    disp_princ = round(sum(round(r["principal"],   2) for r in master), 2)
    disp_int   = round(sum(round(r["interest"],    2) for r in master), 2)
    disp_comm  = round(sum(round(r["commission"],  2) for r in master), 2)
    disp_pay   = round(disp_princ + disp_int + disp_comm, 2)
    total_row = {
        t["period"]:        t["total_row"],
        t["date"]:          "",
        t["balance_open"]:  "",
        t["payment_total"]: disp_pay,
        t["principal"]:     disp_princ,
        t["interest"]:      disp_int,
        t["commission"]:    disp_comm,
        t["balance_close"]: "",
    }
    df_with_total = pd.concat([df, pd.DataFrame([total_row])], ignore_index=True)

    # Chart df
    df_chart = pd.DataFrame([{
        "period":        r["period"],
        "period_label":  str(r["period"]),
        "principal":     round(r["principal"],     2),
        "interest":      round(r["interest"],      2),
        "commission":    round(r["commission"],    2),
        "balance_close": round(r["balance_close"], 2),
        "payment":       round(r["payment"],       2),
    } for r in master])

    # Per-tranche payment series for stacked bar chart
    tranche_payments_by_period = {}  # letter → [pmt_period_1, pmt_period_2, ...]
    for tr in per_tranche:
        letter = tr["letter"]
        series = []
        for r in master:
            series.append(r["tranche_payments"].get(letter, 0.0))
        tranche_payments_by_period[letter] = series

    # Blended APR — None means root-finder failed; UI surfaces it
    blended_apr_error = None
    try:
        blended_apr = calc_syndicated_blended_apr(
            per_tranche, totals.get("total_one_time_comm", 0))
        if blended_apr is None:
            blended_apr_error = "irr_failure"
            # Keep None — do not fake "0%" which would be both wrong and
            # economically nonsensical.
    except Exception as e:
        blended_apr_error = str(e)
        blended_apr = None

    # Inflation real cost
    real_cost_val = None
    inflation_savings = None
    if inflation_enabled and inflation_rate != 0:
        try:
            full_payments = [r["payment"] for r in master]
            # One-time fees are paid at t=0 and remain undiscounted; they are
            # added after computing the PV of the future payment stream.
            pv_future = calc_real_cost(full_payments, inflation_rate, "months")
            ot_fee_today = totals.get("total_one_time_comm", 0) or 0
            real_cost_val = pv_future + ot_fee_today
            inflation_savings = totals["total_payment"] - real_cost_val
        except Exception:
            real_cost_val = None

    # Risk metrics. Use totals["first_payment"] which already skips the
    # offset gap when tranches are staggered; master[0]["payment"] could be 0
    # if no tranche starts in period 1.
    first_pmt = totals.get("first_payment", 0.0) or 0.0
    ltv_val  = (calc_ltv(totals["total_principal"], ltv_collateral)
                if ltv_enabled and ltv_collateral > 0 else None)
    dscr_val = (calc_dscr(dscr_noi, first_pmt)
                if dscr_enabled and dscr_noi > 0 else None)
    dti_val  = (calc_dti(first_pmt, dti_other_debts, dti_income)
                if dti_enabled and dti_income > 0 else None)

    # Universal break-even on consolidated cash flow (includes one-time commissions)
    payments_list = [r["payment"] for r in master]
    try:
        universal_be = calc_universal_breakeven(
            payments_list, totals["total_interest"], "months",
            one_time_comm=totals.get("total_one_time_comm", 0))
    except Exception:
        universal_be = None

    summary = {
        "is_deposit":            False,
        "scheme_key":            "syndicated",
        "rate_pa":               blended_apr,
        "start_date":            start_date or date.today(),
        "total_payment":         totals["total_payment"],
        "total_interest":        totals["total_interest"],
        "total_commission":      totals["total_commission"],
        "effective_rate":        blended_apr,
        "effective_rate_error":  blended_apr_error,
        "grace_error":           synd_grace_error,
        "partial_result":        bool(blended_apr_error)
                                  or bool(totals.get("tranche_errors"))
                                  or bool(synd_grace_error),
        "day_count_enabled":     day_count_enabled,
        "day_count_method":      day_count_method if day_count_enabled else None,
        "first_payment":         first_pmt,
        "overpay_pct":           (totals["total_interest"] / totals["total_principal"] * 100
                                   if totals["total_principal"] > 0 else 0),
        "one_time_comm":         totals.get("total_one_time_comm", 0),
        "df_chart":              df_chart,
        "principal":             totals["total_principal"],
        "sym":                   sym,
        "payments":              payments_list,
        "schedule":              master,
        "unit":                  "months",
        # Syndicated-specific
        "syndicated":              True,
        "tranches":                per_tranche,
        "master_schedule":         master,
        "tranche_payments_series": tranche_payments_by_period,
        "tranche_errors":          totals.get("tranche_errors", []),
        # Standard summary fields
        "balloon_breakeven":     None,
        "balloon_breakeven_abs": None,
        "universal_breakeven":   universal_be,
        "grace_enabled":         False,
        "inflation_enabled":     inflation_enabled,
        "inflation_rate":        inflation_rate if inflation_enabled else None,
        "real_cost":             real_cost_val,
        "inflation_savings":     inflation_savings,
        "ltv":                   ltv_val,
        "ltv_collateral":        ltv_collateral if ltv_enabled else None,
        "dscr":                  dscr_val,
        "dscr_noi":              dscr_noi if dscr_enabled else None,
        "dti":                   dti_val,
        "dti_income":            dti_income if dti_enabled else None,
        "dti_other_debts":       dti_other_debts if dti_enabled else None,
    }
    return df_with_total, summary


def run_calculation(principal, n, rate_pa, unit, scheme,
                    ot_val, ot_type, mo_val, mo_type,
                    cur_key, custom_sym, t, deposit_mode="capitalize",
                    start_date=None,
                    grace_enabled=False, grace_start=1,
                    grace_duration=3, grace_type="interest_only",
                    inflation_enabled=False, inflation_rate=0.0,
                    ltv_enabled=False, ltv_collateral=0.0,
                    dscr_enabled=False, dscr_noi=0.0,
                    dti_enabled=False, dti_income=0.0, dti_other_debts=0.0,
                    syndicated_tranches: list[dict] | None = None,
                    day_count_enabled: bool = False,
                    day_count_method: str = DAY_COUNT_DEFAULT):
    """
    Основной расчёт. Возвращает (df_display_with_total, summary).

    Args (new):
      day_count_enabled — если True, проценты считаются по выбранной
                           конвенции (I = P · r_annual · N/B), где
                           N/B — год-фракция периода.
      day_count_method  — один из DAY_COUNT_METHODS.
    """
    sym = (_sanitize_currency_symbol(custom_sym) if cur_key == "custom"
            else CURRENCY_SYMBOLS.get(cur_key, "$"))

    # ── СИНДИЦИРОВАННЫЙ КРЕДИТ — multi-tranche master schedule ────────────────
    if syndicated_tranches:
        return _run_syndicated(
            syndicated_tranches, t, sym, start_date=start_date,
            inflation_enabled=inflation_enabled, inflation_rate=inflation_rate,
            ltv_enabled=ltv_enabled, ltv_collateral=ltv_collateral,
            dscr_enabled=dscr_enabled, dscr_noi=dscr_noi,
            dti_enabled=dti_enabled, dti_income=dti_income,
            dti_other_debts=dti_other_debts,
            day_count_enabled=day_count_enabled,
            day_count_method=day_count_method,
            # Pass through the grace flag so the syndicated path can report
            # the unsupported combination via grace_error.
            grace_requested=grace_enabled and grace_duration > 0,
        )

    # ── ДЕПОЗИТ — отдельная ветка ─────────────────────────────────────────────
    if scheme == "deposit":
        return _run_deposit(principal, n, rate_pa, unit, deposit_mode, sym, t,
                            start_date=start_date,
                            inflation_enabled=inflation_enabled,
                            inflation_rate=inflation_rate)

    # ── КРЕДИТ ────────────────────────────────────────────────────────────────
    ot_comm = principal * ot_val / 100 if ot_type == "pct" else ot_val
    mo_comm = principal * mo_val / 100 if mo_type == "pct" else mo_val

    # Determine effective day-count args for calc_*
    dc_method = day_count_method if day_count_enabled else None
    dc_start  = start_date if day_count_enabled else None

    if scheme == "annuity":
        sched = calc_annuity(principal, n, rate_pa, unit, mo_comm,
                              day_count=dc_method, start_date=dc_start)
    elif scheme == "classic":
        sched = calc_classic(principal, n, rate_pa, unit, mo_comm,
                              day_count=dc_method, start_date=dc_start)
    elif scheme == "balloon":
        sched = calc_balloon(principal, n, rate_pa, unit, mo_comm,
                              day_count=dc_method, start_date=dc_start)
    else:
        sched = calc_annuity(principal, n, rate_pa, unit, mo_comm,
                              day_count=dc_method, start_date=dc_start)

    # ── Применяем кредитные каникулы (если включены) ──────────────────────────
    grace_error = None
    if grace_enabled and grace_duration > 0:
        try:
            # Передаём метод day-count, чтобы каникулы корректно сочетались
            # с ACT/ACT (правильное начисление при пересечении 29 февраля).
            sched = apply_grace_period(sched, grace_start, grace_duration,
                                       grace_type, rate_pa, unit,
                                       day_count=dc_method,
                                       start_date=dc_start,
                                       scheme_hint=scheme)
        except Exception as e:
            grace_error = str(e)
            # Log it but continue with original schedule; UI will surface it.

    # Defensive invariant check before we present anything. If grace failed
    # above, grace_error is already set and we validate the (unchanged) base
    # schedule; a genuine structural problem still raises here. When grace
    # WAS applied successfully, the balance path is legitimately non-monotone
    # (full_holiday capitalizes interest), so we relax that specific check.
    grace_applied = (grace_enabled and grace_error is None
                      and scheme in ("annuity", "classic", "balloon"))
    try:
        _validate_schedule(sched, principal,
                            scheme="balloon" if grace_applied else scheme,
                            check_principal_sum=not grace_applied)
    except ValueError as _ve:
        raise ValueError(f"Schedule integrity check failed: {_ve}") from _ve

    dates = generate_dates(n, unit, start=start_date)
    rows  = []
    for i, row in enumerate(sched):
        # Round each component to cents, then DERIVE the displayed payment as
        # their sum. This guarantees every visible row satisfies
        # payment == principal + interest + commission exactly (no 1-cent
        # artifact from rounding the pre-summed payment independently — the
        # classic amortization-schedule rounding inconsistency).
        r_princ = round(row["principal"],  2)
        r_int   = round(row["interest"],   2)
        r_comm  = round(row["commission"], 2)
        r_pay   = round(r_princ + r_int + r_comm, 2)
        rows.append({
            t["period"]:        row["period"],
            t["date"]:          dates[i],
            t["balance_open"]:  round(row["balance_open"],  2),
            t["payment_total"]: r_pay,
            t["principal"]:     r_princ,
            t["interest"]:      r_int,
            t["commission"]:    r_comm,
            t["balance_close"]: round(row["balance_close"], 2),
        })
    df = pd.DataFrame(rows)

    # Precise (unrounded) aggregates — used for internal math (APR, ratios).
    # NOTE: under full_holiday grace, capitalized interest is rolled into the
    # rows' PRINCIPAL column, so Σ(row principal) can exceed the original
    # principal and Σ(row interest) captures only interest charged on the
    # outstanding balance. We therefore derive the aggregates directly from the
    # actual schedule rows (what the borrower truly pays) rather than assuming
    # the identity Σprincipal == principal, which only holds without
    # capitalization. tot_payment is the genuine sum of period cash flows plus
    # the one-time fee; tot_interest is the economic interest (total paid minus
    # principal returned minus commissions).
    sum_row_principal = sum(r["principal"] for r in sched)
    sum_row_interest  = sum(r["interest"]  for r in sched)
    sum_row_comm      = sum(r["commission"] for r in sched)
    tot_comm     = sum_row_comm + ot_comm
    tot_payment  = sum_row_principal + sum_row_interest + tot_comm
    # Economic interest = everything paid toward interest, including any portion
    # that was capitalized into principal during a full holiday. Equivalent to
    # (total paid − original principal − commissions).
    tot_interest = tot_payment - principal - tot_comm

    # Display aggregates — sum of the per-row ROUNDED values so the visible
    # TOTAL row reconciles exactly with a hand-sum of the displayed columns
    # (accounting/audit requirement: round(Σx) ≠ Σround(x), so we must use
    # the latter for what the user sees). The principal and one-time
    # commission are added in so the TOTAL itself satisfies
    # payment_total == principal + interest + commission.
    disp_principal = round(sum(round(r["principal"], 2) for r in sched), 2)
    disp_interest  = round(sum(round(r["interest"],  2) for r in sched), 2)
    disp_comm      = round(sum(round(r["commission"],2) for r in sched) + round(ot_comm, 2), 2)
    disp_payment   = round(disp_principal + disp_interest + disp_comm, 2)

    # Headline summary interest (the metric card and overpay_pct). This must be
    # the ECONOMIC interest so the three headline metrics reconcile against the
    # ORIGINAL principal: total_payment == principal + total_interest +
    # total_commission. Under full_holiday grace, capitalized interest sits in
    # the schedule's principal column, so disp_interest (the interest-column
    # sum) understates it; we add back the capitalized portion
    # (disp_principal − original principal) so the headline figure matches what
    # the borrower truly pays beyond principal and fees. Without capitalization
    # disp_principal == principal, so this reduces to disp_interest unchanged.
    capitalized = round(disp_principal - round(principal, 2), 2)
    summary_interest = round(disp_interest + max(capitalized, 0.0), 2)

    total_row = {
        t["period"]:        t["total_row"],
        t["date"]:          "",
        t["balance_open"]:  "",
        t["payment_total"]: disp_payment,
        t["principal"]:     disp_principal,
        t["interest"]:      disp_interest,
        t["commission"]:    disp_comm,
        t["balance_close"]: "",
    }
    df_with_total = pd.concat([df, pd.DataFrame([total_row])], ignore_index=True)

    df_chart = pd.DataFrame([{
        "period":        r["period"],
        "period_label":  str(r["period"]),
        "principal":     round(r["principal"],     2),
        "interest":      round(r["interest"],      2),
        "commission":    round(r["commission"],    2),
        "balance_close": round(r["balance_close"], 2),
        "payment":       round(r["payment"],       2),
    } for r in sched])

    ppy = periods_per_year(unit)
    eff_error = None
    try:
        eff = calc_effective_rate(principal, sched, ot_comm, ppy)
        if eff is None:
            # IRR converged-failure: function returned None to signal it.
            # Do NOT substitute nominal rate — that would silently mislead the
            # user into thinking the APR was computed successfully.
            eff_error = "irr_failure"
            # eff stays None
    except Exception as e:
        eff_error = str(e)
        eff = None

    # ── Investment Break-even ─────────────────────────────────────────────────
    payments_list = [r["payment"] for r in sched]
    breakeven_rate = None
    breakeven_abs  = None
    universal_be   = None
    breakeven_vs_classic = None   # pairwise: chosen vs classic
    breakeven_vs_annuity = None   # pairwise: chosen vs annuity (mostly for balloon)

    try:
        universal_be = calc_universal_breakeven(
            payments_list, tot_interest, unit, one_time_comm=ot_comm)
    except Exception:
        universal_be = None

    if scheme == "balloon":
        try:
            breakeven_rate = calc_balloon_breakeven(
                principal, n, rate_pa, unit, mo_comm)
            breakeven_abs = calc_balloon_absolute_breakeven(
                principal, n, rate_pa, unit, mo_comm)
        except Exception:
            pass

    # Pairwise comparisons: at what investment return does the chosen scheme
    # break even with the alternative? Computed only for loan schemes (not
    # deposit / synd), and only against schemes the chosen one differs from.
    if scheme in ("annuity", "classic", "balloon"):
        try:
            if scheme != "classic":
                # Reuse legacy (no day-count) calc to keep apples-to-apples
                # with the chosen schedule. If user enabled day-count, both
                # schedules should use the same convention.
                _dc  = day_count_method if day_count_enabled else None
                _sd  = start_date if day_count_enabled else None
                cl_sched = calc_classic(principal, n, rate_pa, unit, mo_comm,
                                          day_count=_dc, start_date=_sd)
                cl_pmts  = [r["payment"] for r in cl_sched]
                breakeven_vs_classic = calc_pairwise_breakeven(
                    payments_list, cl_pmts, unit)

            if scheme == "balloon":
                _dc  = day_count_method if day_count_enabled else None
                _sd  = start_date if day_count_enabled else None
                an_sched = calc_annuity(principal, n, rate_pa, unit, mo_comm,
                                          day_count=_dc, start_date=_sd)
                an_pmts  = [r["payment"] for r in an_sched]
                breakeven_vs_annuity = calc_pairwise_breakeven(
                    payments_list, an_pmts, unit)
        except Exception:
            # Pairwise BE is informational — failure shouldn't propagate.
            pass

    # ── Инфляция (real cost) ──────────────────────────────────────────────────
    real_cost_val = None
    inflation_savings = None
    if inflation_enabled and inflation_rate != 0:
        try:
            # The one-time commission is paid AT t=0 — undiscounted; the
            # periodic payments are discounted as future cash flows.
            pv_future = calc_real_cost(payments_list, inflation_rate, unit)
            ot_today  = ot_comm if ot_comm > 0 else 0
            real_cost_val = pv_future + ot_today
            inflation_savings = tot_payment - real_cost_val
        except Exception:
            real_cost_val = None

    # ── Риск-аналитика ────────────────────────────────────────────────────────
    first_pmt = sched[0]["payment"] if sched else 0.0

    ltv_val  = calc_ltv(principal, ltv_collateral)  if ltv_enabled  and ltv_collateral > 0 else None
    dscr_val = calc_dscr(dscr_noi, first_pmt)        if dscr_enabled and dscr_noi > 0       else None
    dti_val  = calc_dti(first_pmt, dti_other_debts, dti_income) if dti_enabled and dti_income > 0 else None

    summary = {
        "is_deposit":            False,
        "scheme_key":            scheme,
        "rate_pa":               rate_pa,
        "n_periods":             n,
        "commission_per_period": mo_comm,
        "start_date":            start_date or date.today(),
        "total_payment":         disp_payment,
        "total_interest":        summary_interest,
        "total_commission":      disp_comm,
        # Precise (unrounded) aggregates retained for any downstream math that
        # needs full precision rather than the cent-reconciled display values.
        "total_payment_exact":   tot_payment,
        "total_interest_exact":  tot_interest,
        "effective_rate":        eff,
        "effective_rate_error":  eff_error,
        "grace_error":           grace_error,
        # True when grace was requested but couldn't be applied (e.g. a
        # negative-amortization spiral). The UI uses this to mark the whole
        # result as unreliable rather than presenting a half-valid schedule.
        "partial_result":        bool(grace_error) or bool(eff_error),
        "day_count_enabled":     day_count_enabled,
        "day_count_method":      day_count_method if day_count_enabled else None,
        "first_payment":         first_pmt,
        "overpay_pct":           summary_interest / principal * 100 if principal > 0 else 0,
        "one_time_comm":         ot_comm,
        "df_chart":              df_chart,
        "principal":             principal,
        "sym":                   sym,
        "payments":              payments_list,
        "schedule":              sched,
        "unit":                  unit,
        # Investment break-even
        "balloon_breakeven":     breakeven_rate,
        "balloon_breakeven_abs": breakeven_abs,
        "universal_breakeven":   universal_be,
        "breakeven_vs_classic":  breakeven_vs_classic,
        "breakeven_vs_annuity":  breakeven_vs_annuity,
        # Grace period
        "grace_enabled":         grace_enabled,
        "grace_start":           grace_start if grace_enabled else None,
        "grace_duration":        grace_duration if grace_enabled else None,
        "grace_type":            grace_type if grace_enabled else None,
        # Inflation
        "inflation_enabled":     inflation_enabled,
        "inflation_rate":        inflation_rate if inflation_enabled else None,
        "real_cost":             real_cost_val,
        "inflation_savings":     inflation_savings,
        # Risk metrics
        "ltv":                   ltv_val,
        "ltv_collateral":        ltv_collateral if ltv_enabled else None,
        "dscr":                  dscr_val,
        "dscr_noi":              dscr_noi if dscr_enabled else None,
        "dti":                   dti_val,
        "dti_income":            dti_income if dti_enabled else None,
        "dti_other_debts":       dti_other_debts if dti_enabled else None,
    }
    return df_with_total, summary


def _run_deposit(principal, n, rate_pa, unit, mode, sym, t, start_date=None,
                  inflation_enabled: bool = False, inflation_rate: float = 0.0):
    """
    Расчёт депозита.
    start_date — дата открытия вклада; первое начисление = start + 1 период.

    Опционально вычисляет реальную (PV-скорректированную) итоговую стоимость
    при ненулевой инфляции.
    """
    sched = calc_deposit(principal, n, rate_pa, unit, mode)
    dates = generate_dates(n, unit, start=start_date)

    rows = []
    for i, row in enumerate(sched):
        rows.append({
            t["period"]:               row["period"],
            t["date"]:                 dates[i],
            t["dep_balance_open"]:     row["balance_open"],
            t["dep_interest_earned"]:  row["interest"],
            t["dep_payout"]:           row["payout"],
            t["dep_balance_close"]:    row["balance_close"],
        })
    df = pd.DataFrame(rows)

    total_earned  = sum(r["interest"] for r in sched)
    total_payout  = sum(r["payout"]   for r in sched)
    final_balance = sched[-1]["balance_close"]
    # При выплате процентов — в конце возвращается тело + последние %
    # При капитализации — итог = final_balance (всё включено)

    total_row = {
        t["period"]:               t["dep_total_row"],
        t["date"]:                 "",
        t["dep_balance_open"]:     "",
        t["dep_interest_earned"]:  round(total_earned,  2),
        t["dep_payout"]:           round(total_payout,  2),
        t["dep_balance_close"]:    round(final_balance, 2),
    }
    df_with_total = pd.concat([df, pd.DataFrame([total_row])], ignore_index=True)

    # Chart DataFrame
    cumulative_interest = 0.0
    chart_rows = []
    for row in sched:
        cumulative_interest += row["interest"]
        chart_rows.append({
            "period":              row["period"],
            "period_label":        str(row["period"]),
            "interest":            row["interest"],
            "cumulative_interest": round(cumulative_interest, 2),
            "balance_close":       row["balance_close"],
            "payout":              row["payout"],
        })
    df_chart = pd.DataFrame(chart_rows)

    # Реальная доходность (CAGR)
    ppy = periods_per_year(unit)
    if mode == "capitalize":
        # Guard against principal == 0 (calc_deposit allows principal >= 0):
        # CAGR is undefined with no initial capital, so report 0%.
        if principal > 0 and n > 0:
            cagr = (final_balance / principal) ** (ppy / n) - 1
        else:
            cagr = 0.0
        eff_rate = cagr * 100
    else:
        eff_rate = rate_pa  # простой процент = номинальная ставка

    # ── Inflation-adjusted real value ─────────────────────────────────────────
    # For depositors, real_cost = present value of future receipts in today's
    # purchasing power. Positive inflation eats into real value; deflation
    # increases it.
    #   capitalize: single inflow `final_balance` at period n
    #   payout:    each periodic interest payment + principal returned at n
    real_cost_val = None
    inflation_savings = None
    if inflation_enabled and inflation_rate != 0:
        try:
            if mode == "capitalize":
                # Single future value discounted by (1+r_per)^n
                r_per = inflation_rate / 100.0 / ppy
                if 1.0 + r_per > 0:
                    real_cost_val = final_balance / (1.0 + r_per) ** n
                else:
                    real_cost_val = final_balance
                # Nominal receipt in capitalize mode is the single final_balance.
                nominal_receipts = final_balance
            else:
                # Payout: interest stream + principal returned at last period.
                interest_cashflows = [r["payout"] for r in sched]
                cashflows_with_principal = list(interest_cashflows)
                cashflows_with_principal[-1] += principal
                real_cost_val = calc_real_cost(
                    cashflows_with_principal, inflation_rate, unit)
                # Nominal receipts in payout mode = every interest payment PLUS
                # the principal returned at maturity. Comparing real_cost_val
                # against final_balance (= principal only) would mismatch the
                # bases; the correct nominal total is principal + Σ interest.
                nominal_receipts = sum(cashflows_with_principal)
            # Sign convention: positive = nominal exceeds real (purchasing
            # power lost to inflation); negative = deflation gain.
            inflation_savings = nominal_receipts - real_cost_val
        except Exception:
            real_cost_val = None

    summary = {
        "is_deposit":      True,
        "deposit_mode":    mode,
        "scheme_key":      "deposit",
        "rate_pa":         rate_pa,
        "start_date":      start_date or date.today(),
        "principal":       principal,
        "final_balance":   final_balance,
        "total_earned":    total_earned,
        "total_payout":    total_payout,
        "effective_rate":  eff_rate,
        "df_chart":        df_chart,
        "sym":             sym,
        "unit":            unit,
        # Совместимость с общим кодом:
        # total_payment = total cash the depositor RECEIVES over the deposit's
        # life, defined consistently across both modes:
        #   • capitalize → interest is retained and compounded, so everything
        #     comes back as the single final_balance.
        #   • payout     → interest is paid out each period AND the principal is
        #     returned at maturity, so the total received is final_balance
        #     (= principal) PLUS the sum of payouts. Using just final_balance
        #     here would understate the depositor's total receipts and make the
        #     field mean different things in each mode.
        "total_payment":   (final_balance if mode == "capitalize"
                             else final_balance + total_payout),
        "total_interest":  total_earned,
        "total_commission":0,
        "first_payment":   sched[0]["interest"],
        "overpay_pct":     0,
        "payments":        [r["interest"] for r in sched],
        "schedule":        sched,
        # Inflation
        "inflation_enabled": inflation_enabled,
        "inflation_rate":    inflation_rate if inflation_enabled else None,
        "real_cost":         real_cost_val,
        "inflation_savings": inflation_savings,
    }
    return df_with_total, summary


# ─────────────────────────────────────────────────────────────────────────────
#  ДИАГРАММЫ
# ─────────────────────────────────────────────────────────────────────────────
def _apply_layout(fig, title, height=420):
    fig.update_layout(
        title=dict(text=title, font=dict(color=C["text"], size=15)),
        paper_bgcolor=C["bg"], plot_bgcolor=C["card"],
        font=dict(color=C["text"]),
        legend=dict(bgcolor=C["card"], bordercolor=C["grid"], borderwidth=1),
        xaxis=dict(gridcolor=C["grid"], tickfont=dict(size=9)),
        yaxis=dict(gridcolor=C["grid"]),
        margin=dict(l=10, r=10, t=55, b=10),
        height=height,
    )

def chart_bar(df, t):
    """
    Period-by-period payment composition chart.

    For loans with > 120 periods (e.g. 30-year mortgages, syndicated long
    tranches), drawing a stacked bar per period creates 360+ bars, which is
    visually noisy and slow in the browser. In that case we switch to a
    stacked AREA chart (continuous lines + fill) which is light, smooth,
    and keeps the same information density.
    """
    n_periods = len(df)
    use_area = n_periods > 120

    # Choose trace type based on density
    def _make_trace(name, x, y, color, fillsum=None):
        if use_area:
            return go.Scatter(
                name=name, x=x, y=y, mode="lines",
                stackgroup="one",
                line=dict(color=color, width=0.5),
                fillcolor=color,
                hovertemplate=f"{name}: %{{y:,.2f}}<extra></extra>",
            )
        return go.Bar(
            name=name, x=x, y=y, marker_color=color,
            hovertemplate=f"{name}: %{{y:,.2f}}<extra></extra>",
        )

    fig = go.Figure()
    fig.add_trace(_make_trace(t["chart_principal"], df["period_label"],
                                df["principal"], C["principal"]))
    fig.add_trace(_make_trace(t["chart_interest"], df["period_label"],
                                df["interest"], C["interest"]))
    if df["commission"].sum() > 0:
        fig.add_trace(_make_trace(t["chart_commission"], df["period_label"],
                                    df["commission"], C["commission"]))
    if not use_area:
        fig.update_layout(barmode="stack")

    title = t["chart_title"]
    if use_area:
        title += f"  ·  area mode ({n_periods} periods)"
    _apply_layout(fig, title)
    return fig


def chart_syndicated_tranches(tranche_payments_series: dict, t: dict):
    """
    Stacked chart showing per-tranche contribution to consolidated payment.
    Bar mode for ≤ 120 periods, area mode for longer horizons.

    tranche_payments_series — dict {letter: [pmt_p1, pmt_p2, ...]}
    """
    if not tranche_payments_series:
        return None

    tranche_colors = {
        "A": "#4FC3F7",  # cyan  (Corporate Navy palette)
        "B": "#FFD166",  # amber
        "C": "#06D6A0",  # green
    }

    max_periods = max(len(v) for v in tranche_payments_series.values())
    use_area = max_periods > 120
    x = [str(p) for p in range(1, max_periods + 1)]

    fig = go.Figure()
    for letter in sorted(tranche_payments_series.keys()):
        series = tranche_payments_series[letter]
        padded = list(series) + [0.0] * (max_periods - len(series))
        tranche_label = t.get("syndicated_tranche", "Tranche {n}").format(n=letter)
        color = tranche_colors.get(letter, "#94A3B8")
        if use_area:
            fig.add_trace(go.Scatter(
                name=tranche_label,
                x=x, y=padded, mode="lines",
                stackgroup="tranches",
                line=dict(color=color, width=0.5),
                fillcolor=color,
                hovertemplate=f"{tranche_label}: %{{y:,.2f}}<extra></extra>",
            ))
        else:
            fig.add_trace(go.Bar(
                name=tranche_label,
                x=x, y=padded,
                marker_color=color,
                hovertemplate=f"{tranche_label}: %{{y:,.2f}}<extra></extra>",
            ))

    if not use_area:
        fig.update_layout(barmode="stack")

    title = t.get("syndicated_chart_title", "Payment Structure by Tranche")
    if use_area:
        title += f"  ·  area mode ({max_periods} periods)"
    _apply_layout(fig, title)
    return fig


def chart_pie(principal, tot_interest, tot_comm, t):
    labels = [t["chart_principal"], t["chart_interest"]]
    values = [principal, tot_interest]
    clrs   = [C["principal"], C["interest"]]
    if tot_comm > 0:
        labels.append(t["chart_commission"]); values.append(tot_comm); clrs.append(C["commission"])
    fig = go.Figure(go.Pie(
        labels=labels, values=values, hole=0.55,
        marker=dict(colors=clrs, line=dict(color=C["bg"], width=2)),
        textinfo="label+percent", textfont=dict(size=13),
        hovertemplate="%{label}: %{value:,.2f} (%{percent})<extra></extra>",
    ))
    fig.update_layout(
        paper_bgcolor=C["bg"], font=dict(color=C["text"]),
        legend=dict(bgcolor=C["card"], bordercolor=C["grid"], borderwidth=1),
        margin=dict(l=10, r=10, t=55, b=10), height=380,
        title=dict(text=t["chart_pie_title"], font=dict(color=C["text"], size=15)),
    )
    return fig

def chart_balance(df, t):
    hover_label = t.get("chart_balance_hover", "Balance")
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=df["period_label"], y=df["balance_close"],
        mode="lines+markers", fill="tozeroy",
        line=dict(color=C["principal"], width=2),
        fillcolor="rgba(79,195,247,0.12)", marker=dict(size=4),
        hovertemplate=f"{hover_label}: %{{y:,.2f}}<extra></extra>",
    ))
    fig.update_layout(showlegend=False)
    _apply_layout(fig, t.get("chart_balance_title", "📉 Remaining Balance"), height=300)
    return fig

def chart_compare(ann_total, cla_total, ann_int, cla_int, t,
                   bal_total=None, bal_int=None):
    """
    Side-by-side scheme comparison (Annuity / Classic / optional Balloon).
    Bar groups: Total Payment, Total Interest. Balloon is rendered only
    when its totals are supplied (kept optional for backwards compatibility
    with older call sites).
    """
    cats = [t["total_payment"], t["total_interest"]]
    fig  = go.Figure()
    fig.add_trace(go.Bar(name=t["annuity_short"], x=cats, y=[ann_total, ann_int],
                         marker_color=C["interest"],
                         text=[f"{v:,.0f}" for v in [ann_total, ann_int]],
                         textposition="outside", textfont=dict(color=C["text"])))
    fig.add_trace(go.Bar(name=t["classic_short"], x=cats, y=[cla_total, cla_int],
                         marker_color=C["principal"],
                         text=[f"{v:,.0f}" for v in [cla_total, cla_int]],
                         textposition="outside", textfont=dict(color=C["text"])))
    if bal_total is not None and bal_int is not None:
        fig.add_trace(go.Bar(
            name=t.get("balloon_short", "Balloon"),
            x=cats, y=[bal_total, bal_int],
            marker_color="#FFD166",  # amber — visually distinct from blue / red
            text=[f"{v:,.0f}" for v in [bal_total, bal_int]],
            textposition="outside", textfont=dict(color=C["text"])))
    fig.update_layout(barmode="group")
    _apply_layout(fig, t["annuity_vs_classic"], height=380)
    return fig

def chart_daycount_compare(principal: float, n: int, rate_pa: float, unit: str,
                             scheme: str, start_dt, mo_comm: float, t: dict):
    """
    Side-by-side total-interest comparison across the five day-count
    conventions for the same loan. Helps the user see which convention
    yields the lowest cost (which depends on calendar, start date, and
    rate magnitude).

    Returns (figure, results_dict). `results_dict` maps method → total_interest
    so the caller can also render it as a table. Methods that fail to compute
    are skipped from the figure.
    """
    methods = ["30/360", "30E/360", "ACT/360", "ACT/365", "ACT/ACT"]
    sched_fn = {
        "annuity": calc_annuity,
        "classic": calc_classic,
        "balloon": calc_balloon,
    }.get(scheme, calc_annuity)

    results = {}
    for m in methods:
        try:
            sched = sched_fn(principal, n, rate_pa, unit, mo_comm,
                              day_count=m, start_date=start_dt)
            results[m] = sum(r["interest"] for r in sched)
        except Exception:
            results[m] = None

    # Build the figure from successful computations
    valid_methods = [m for m, v in results.items() if v is not None]
    valid_values  = [results[m] for m in valid_methods]

    if not valid_methods:
        return None, results

    # Highlight the lowest-interest method in green, the highest in red,
    # the rest in neutral blue. Tied methods all share the same color.
    min_v = min(valid_values)
    max_v = max(valid_values)
    bar_colors = []
    for v in valid_values:
        if v == min_v and min_v != max_v:
            bar_colors.append("#10B981")  # cheapest — green
        elif v == max_v and min_v != max_v:
            bar_colors.append("#DC2626")  # most expensive — red
        else:
            bar_colors.append("#4FC3F7")  # neutral

    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=valid_methods,
        y=valid_values,
        marker_color=bar_colors,
        text=[f"{v:,.2f}" for v in valid_values],
        textposition="outside",
        textfont=dict(color=C["text"], size=11),
        hovertemplate="<b>%{x}</b><br>Total interest: %{y:,.2f}<extra></extra>",
    ))
    _apply_layout(fig, t.get("daycount_compare_title",
                              "💰 Total Interest by Day-Count Convention"),
                   height=360)
    fig.update_layout(showlegend=False)
    return fig, results


def chart_invest(df_chart, invest_vals, t, yield_label):
    """Остаток долга vs инвестиционный портфель."""
    periods = df_chart["period_label"].tolist()
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=periods, y=df_chart["balance_close"].tolist(),
        name=t["invest_loan_balance"], mode="lines+markers",
        line=dict(color=C["interest"], width=2), marker=dict(size=4),
        fill="tozeroy", fillcolor="rgba(255,107,107,0.1)",
        hovertemplate=f"{t['invest_loan_balance']}: %{{y:,.2f}}<extra></extra>",
    ))
    fig.add_trace(go.Scatter(
        x=periods, y=invest_vals,
        name=f"{t['invest_portfolio']} ({yield_label})", mode="lines+markers",
        line=dict(color=C["invest"], width=2), marker=dict(size=4),
        fill="tozeroy", fillcolor="rgba(6,214,160,0.1)",
        hovertemplate=f"{t['invest_portfolio']}: %{{y:,.2f}}<extra></extra>",
    ))
    _apply_layout(fig, t["invest_chart_title"], height=420)
    return fig


# ── Депозитные диаграммы ──────────────────────────────────────────────────────

def chart_dep_growth(df_chart, principal, t):
    """
    График роста депозита: линия баланса (растёт вверх).
    Заштрихована область — начисленные проценты сверху тела.
    """
    periods = df_chart["period_label"].tolist()
    fig = go.Figure()

    # Тело вклада (горизонтальная линия)
    fig.add_trace(go.Scatter(
        x=periods, y=[principal] * len(periods),
        name=t["dep_your_deposit"],
        mode="lines",
        line=dict(color=C["principal"], width=1.5, dash="dot"),
        hoverinfo="skip",
    ))
    # Линия полного баланса
    fig.add_trace(go.Scatter(
        x=periods, y=df_chart["balance_close"].tolist(),
        name=t["dep_balance_label"],
        mode="lines+markers",
        line=dict(color=C["invest"], width=2.5),
        marker=dict(size=4),
        fill="tonexty",
        fillcolor="rgba(6,214,160,0.15)",
        hovertemplate=f"{t['dep_balance_label']}: %{{y:,.2f}}<extra></extra>",
    ))

    _apply_layout(fig, t["dep_growth_title"], height=380)
    fig.update_layout(showlegend=True)
    return fig


def chart_dep_bar(df_chart, principal, t):
    """
    Столбиковая диаграмма: тело вклада + накопленные проценты по периодам.
    """
    periods = df_chart["period_label"].tolist()
    principal_vals = [principal] * len(periods)
    interest_vals  = df_chart["cumulative_interest"].tolist()

    fig = go.Figure()
    fig.add_trace(go.Bar(
        name=t["dep_initial"], x=periods, y=principal_vals,
        marker_color=C["principal"],
        hovertemplate=f"{t['dep_initial']}: %{{y:,.2f}}<extra></extra>",
    ))
    fig.add_trace(go.Bar(
        name=t["dep_interest_label"], x=periods, y=interest_vals,
        marker_color=C["invest"],
        hovertemplate=f"{t['dep_interest_label']}: %{{y:,.2f}}<extra></extra>",
    ))
    fig.update_layout(barmode="stack")
    _apply_layout(fig, t["chart_title"], height=380)
    return fig


def chart_dep_pie(principal, total_earned, t):
    """Пирог: начальный вклад vs заработанные проценты."""
    fig = go.Figure(go.Pie(
        labels=[t["dep_initial"], t["dep_interest_label"]],
        values=[principal, total_earned],
        hole=0.55,
        marker=dict(colors=[C["principal"], C["invest"]],
                    line=dict(color=C["bg"], width=2)),
        textinfo="label+percent",
        textfont=dict(size=13),
        hovertemplate="%{label}: %{value:,.2f} (%{percent})<extra></extra>",
    ))
    fig.update_layout(
        paper_bgcolor=C["bg"], font=dict(color=C["text"]),
        legend=dict(bgcolor=C["card"], bordercolor=C["grid"], borderwidth=1),
        margin=dict(l=10, r=10, t=55, b=10), height=380,
        title=dict(text=t["chart_pie_title"], font=dict(color=C["text"], size=15)),
    )
    return fig


def chart_dep_compare_modes(principal, n, rate_pa, unit, t, sym):
    """
    Сравнение двух режимов вклада: капитализация vs выплата.
    Показывает итоговый баланс и итоговые проценты для каждого.
    """
    sched_cap = calc_deposit(principal, n, rate_pa, unit, "capitalize")
    sched_pay = calc_deposit(principal, n, rate_pa, unit, "payout")

    cap_final   = sched_cap[-1]["balance_close"]
    cap_earned  = sum(r["interest"] for r in sched_cap)
    pay_earned  = sum(r["interest"] for r in sched_pay)   # = pay_final_income
    pay_total   = principal + pay_earned  # тело + всё что выплатили

    cats = [t["dep_final_balance"], t["dep_total_earned"]]
    fig  = go.Figure()
    fig.add_trace(go.Bar(
        name=t["deposit_capitalize"], x=cats, y=[cap_final, cap_earned],
        marker_color=C["invest"],
        text=[f"{v:,.0f}" for v in [cap_final, cap_earned]],
        textposition="outside", textfont=dict(color=C["text"]),
    ))
    fig.add_trace(go.Bar(
        name=t["deposit_payout"], x=cats, y=[pay_total, pay_earned],
        marker_color=C["principal"],
        text=[f"{v:,.0f}" for v in [pay_total, pay_earned]],
        textposition="outside", textfont=dict(color=C["text"]),
    ))
    fig.update_layout(barmode="group")
    _apply_layout(fig, t["dep_vs_invest"], height=380)
    return fig


def chart_dep_vs_alternative(df_chart, alt_vals, deposit_vals, t, yield_label, sym):
    """
    Линейный график: рост депозита vs альтернативные инвестиции под другой %.
    """
    periods = df_chart["period_label"].tolist()
    fig = go.Figure()

    fig.add_trace(go.Scatter(
        x=periods, y=deposit_vals,
        name=t["dep_your_deposit"],
        mode="lines+markers",
        line=dict(color=C["principal"], width=2.5),
        marker=dict(size=4),
        fill="tozeroy", fillcolor="rgba(79,195,247,0.1)",
        hovertemplate=f"{t['dep_your_deposit']}: %{{y:,.2f}}<extra></extra>",
    ))
    fig.add_trace(go.Scatter(
        x=periods, y=alt_vals,
        name=f"{t['dep_alternative']} ({yield_label})",
        mode="lines+markers",
        line=dict(color=C["invest"], width=2.5),
        marker=dict(size=4),
        fill="tozeroy", fillcolor="rgba(6,214,160,0.1)",
        hovertemplate=f"{t['dep_alternative']}: %{{y:,.2f}}<extra></extra>",
    ))
    _apply_layout(fig, t["dep_invest_section"], height=420)
    return fig

# ─────────────────────────────────────────────────────────────────────────────
#  ЭКСПОРТ В EXCEL  —  Bloomberg / Investment-Bank grade
#  Три листа: 1) Summary  2) Payment Schedule  3) Analysis
# ─────────────────────────────────────────────────────────────────────────────
def export_excel(df, summary, t, sym):
    """
    Создаёт профессиональный многолистовый Excel-файл уровня Bloomberg:
      Sheet 1 – Summary:          шапка, параметры, ключевые метрики
      Sheet 2 – Payment Schedule: подробная таблица с формулами и "зеброй"
      Sheet 3 – Analysis:         сравнительная аналитика схем платежей
    """
    buf = io.BytesIO()
    wb  = Workbook()

    if df is None or getattr(df, "empty", True) or len(df.columns) == 0:
        raise ValueError(
            "export_excel: no schedule data to export (calculation may have "
            "failed or produced an empty result).")

    is_dep = summary.get("is_deposit", False)

    # ── Палитра цветов (Bloomberg / dark fintech) ─────────────────────────────
    NAVY      = "0D1B2A"   # фон шапки
    DARKBLUE  = "1D3557"   # фон параметров
    MIDBLUE   = "1D4ED8"   # заголовки таблицы
    LIGHTBLUE = "EFF6FF"   # светлые ячейки ("зебра")
    WHITE     = "FFFFFF"
    GOLD      = "F59E0B"   # акцент / итого
    GREEN     = "065F46"   # положительные значения
    GREENLT   = "D1FAE5"   # светло-зелёный
    RED       = "7F1D1D"
    REDLT     = "FEE2E2"
    GRAY1     = "1E293B"   # тёмный чередующийся ряд
    GRAY2     = "F8FAFC"   # светлый чередующийся ряд
    TEXTLIGHT = "E2E8F0"
    SUBTEXT   = "94A3B8"
    BORDER_C  = "CBD5E1"
    TOTALBG   = "064E3B"
    TOTALFG   = "6EE7B7"

    # ── Общие стили ──────────────────────────────────────────────────────────
    def fill(hex_color):
        return PatternFill("solid", fgColor=hex_color)

    def font(color=WHITE, bold=False, size=10, italic=False, name="Calibri"):
        return Font(name=name, bold=bold, size=size, color=color, italic=italic)

    def align(h="left", v="center", wrap=False):
        return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

    def border(color=BORDER_C, style="thin"):
        s = Side(style=style, color=color)
        return Border(left=s, right=s, top=s, bottom=s)

    def border_bottom(color=BORDER_C):
        s = Side(style="medium", color=color)
        return Border(bottom=s)

    def set_col_width(ws):
        """Авторазмер всех колонок (без решёток #####)."""
        for col_cells in ws.columns:
            max_len = 0
            col_letter = get_column_letter(col_cells[0].column)
            for cell in col_cells:
                try:
                    cell_len = len(str(cell.value or ""))
                    if cell_len > max_len:
                        max_len = cell_len
                except Exception:
                    pass
            ws.column_dimensions[col_letter].width = max(max_len + 4, 12)

    num_fmt  = '#,##0.00'          # числовой с тысячными
    pct_fmt  = '0.00%'            # процентный
    date_fmt = 'DD.MM.YYYY'       # дата

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    #  ЛИСТ 1: SUMMARY
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    ws1 = wb.active
    ws1.title = "Summary"
    ws1.sheet_view.showGridLines = False
    ws1.column_dimensions["A"].width = 32
    ws1.column_dimensions["B"].width = 26
    ws1.column_dimensions["C"].width = 4
    ws1.column_dimensions["D"].width = 32
    ws1.column_dimensions["E"].width = 26

    # ── ШАПКА документа ───────────────────────────────────────────────────────
    ws1.merge_cells("A1:E1")
    c = ws1["A1"]
    c.value     = "Yev Capital LoanLogic"
    c.font      = Font(name="Calibri", bold=True, size=22, color="4FC3F7")
    c.fill      = fill(NAVY)
    c.alignment = align("center")
    ws1.row_dimensions[1].height = 42

    ws1.merge_cells("A2:E2")
    c2 = ws1["A2"]
    c2.value     = "Loan Analysis Report" if not is_dep else "Deposit Analysis Report"
    c2.font      = Font(name="Calibri", italic=True, size=12, color=SUBTEXT)
    c2.fill      = fill(NAVY)
    c2.alignment = align("center")
    ws1.row_dimensions[2].height = 22

    ws1.merge_cells("A3:E3")
    slogan_cell = ws1["A3"]
    slogan_cell.value     = "Precision in every percent. Logic in every loan."
    slogan_cell.font      = Font(name="Calibri", italic=True, size=9, color="334155")
    slogan_cell.fill      = fill(NAVY)
    slogan_cell.alignment = align("center")
    ws1.row_dimensions[3].height = 16

    # ── Мета-данные ───────────────────────────────────────────────────────────
    ws1.merge_cells("A4:E4")
    ws1["A4"].value = ""
    ws1["A4"].fill  = fill(NAVY)
    ws1.row_dimensions[4].height = 8

    meta_rows = [
        ("Report Date",    datetime.now().strftime("%d %B %Y, %H:%M")),
        ("Prepared by",    "Yevtushenko  |  Yev Capital LoanLogic v3.0"),
        ("Calculator",     "Advanced Loan & Deposit Calculator"),
    ]
    for r_off, (lbl, val) in enumerate(meta_rows):
        r = 5 + r_off
        ws1.row_dimensions[r].height = 18
        c_lbl = ws1.cell(r, 1, lbl)
        c_lbl.font  = font(color=SUBTEXT, size=9, italic=True)
        c_lbl.fill  = fill(NAVY)
        c_val = ws1.cell(r, 2, val)
        c_val.font  = font(color=TEXTLIGHT, size=9)
        c_val.fill  = fill(NAVY)
        for ci in range(3, 6):
            ws1.cell(r, ci).fill = fill(NAVY)

    # Разделитель
    sep_row = 5 + len(meta_rows)
    ws1.merge_cells(f"A{sep_row}:E{sep_row}")
    ws1[f"A{sep_row}"].fill  = fill("1D4ED8")
    ws1[f"A{sep_row}"].value = ""
    ws1.row_dimensions[sep_row].height = 3

    # ── БЛОК ВХОДНЫХ ПАРАМЕТРОВ ───────────────────────────────────────────────
    params_start = sep_row + 2
    ws1.row_dimensions[params_start].height = 22
    ws1.merge_cells(f"A{params_start}:B{params_start}")
    ph = ws1[f"A{params_start}"]
    ph.value     = "  INPUT PARAMETERS"
    ph.font      = Font(name="Calibri", bold=True, size=10, color=WHITE)
    ph.fill      = fill(MIDBLUE)
    ph.alignment = align("left")
    ws1.merge_cells(f"D{params_start}:E{params_start}")
    ws1[f"D{params_start}"].fill = fill(MIDBLUE)

    # Определяем схему для отображения
    scheme_display = {
        "annuity": "Annuity (Equal Payments)",
        "classic": "Classic (Differentiated)",
        "balloon": "Balloon (Interest-Only)",
        "deposit": "Deposit / Savings",
    }.get(summary.get("scheme_key", "annuity"), "Annuity")

    principal = summary.get("principal", 0)
    rate_pa   = summary.get("rate_pa", 0) or summary.get("effective_rate", 0)
    n_periods = len(summary.get("schedule", []))
    unit_lbl  = summary.get("unit", "months")

    params_data = [
        ("Loan Amount",           principal,     num_fmt,  sym),
        ("Annual Interest Rate",  rate_pa / 100, pct_fmt,  ""),
        ("Number of Periods",     n_periods,     "#,##0",  f" {unit_lbl}"),
        ("Payment Type",          scheme_display,"@",      ""),
        ("Currency",              sym,           "@",      ""),
    ]
    if is_dep:
        mode_lbl = "Capitalization (Compound)" if summary.get("deposit_mode") == "capitalize" else "Interest Payout"
        params_data[3] = ("Deposit Mode", mode_lbl, "@", "")

    for i, (lbl, val, fmt, unit_s) in enumerate(params_data):
        r = params_start + 1 + i
        ws1.row_dimensions[r].height = 20
        is_even = i % 2 == 0
        row_fill = fill("F1F5F9") if is_even else fill(WHITE)

        c_lbl = ws1.cell(r, 1, f"  {lbl}")
        c_lbl.font      = font(color="1E293B", bold=True, size=9)
        c_lbl.fill      = row_fill
        c_lbl.alignment = align("left")
        c_lbl.border    = border(BORDER_C)

        c_val = ws1.cell(r, 2)
        if fmt in ("@", ""):
            c_val.value = f"{val}{unit_s}"
        else:
            c_val.value  = val
            c_val.number_format = fmt
        c_val.font      = font(color="1D3557", bold=True, size=9)
        c_val.fill      = row_fill
        c_val.alignment = align("right")
        c_val.border    = border(BORDER_C)

        # Заполняем пустую разделительную колонку C
        ws1.cell(r, 3).fill = fill(WHITE)

    # ── БЛОК КЛЮЧЕВЫХ МЕТРИК (Key Metrics) ───────────────────────────────────
    metrics_start = params_start
    ws1.row_dimensions[metrics_start].height = 22
    ws1.merge_cells(f"D{metrics_start}:E{metrics_start}")
    mh = ws1[f"D{metrics_start}"]
    mh.value     = "  KEY METRICS"
    mh.font      = Font(name="Calibri", bold=True, size=10, color=WHITE)
    mh.fill      = fill("065F46")
    mh.alignment = align("left")

    if is_dep:
        dep_eff = summary.get("effective_rate")
        dep_princ = summary.get("principal", 0) or 0
        dep_earned = summary.get("total_earned", 0) or 0
        eff_entry = (("Effective Annual Rate", "N/A", "@", "#B45309", False)
                      if dep_eff is None
                      else ("Effective Annual Rate", dep_eff / 100, pct_fmt, "#B45309", False))
        profit_pct = (dep_earned / dep_princ) if dep_princ > 0 else 0
        metrics_data = [
            ("Initial Deposit",        summary.get("principal", 0),      num_fmt, "#1D4ED8", False),
            ("Final Balance",           summary.get("final_balance", 0),   num_fmt, "#065F46", True),
            ("Total Interest Earned",   summary.get("total_earned", 0),    num_fmt, "#065F46", True),
            eff_entry,
            ("Profit %",               profit_pct, pct_fmt, "#065F46", True),
        ]
    else:
        # Effective APR: when IRR fails the value is None — passing it through
        # arithmetic would crash. Show "N/A" as a text-format cell instead of
        # masking the failure with 0.00%.
        eff_val = summary.get("effective_rate")
        if eff_val is None:
            eff_apr_entry = ("Effective Annual Rate", "N/A", "@", "#B45309", True)
        else:
            eff_apr_entry = ("Effective Annual Rate", eff_val / 100, pct_fmt, "#B45309", True)

        metrics_data = [
            ("Total Amount Payable",    summary.get("total_payment", 0),    num_fmt, "#B45309", False),
            ("Total Interest Paid",     summary.get("total_interest", 0),   num_fmt, "#7F1D1D", True),
            ("Total Commissions",       summary.get("total_commission", 0), num_fmt, "#92400E", False),
            eff_apr_entry,
            ("1st Period Payment",      summary.get("first_payment", 0),    num_fmt, "#1D4ED8", False),
        ]

        # ── Investment break-even — universal + balloon-specific ──────────────
        uni_be = summary.get("universal_breakeven")
        if uni_be is not None:
            metrics_data.append(
                ("Universal Inv. Break-even", uni_be / 100, pct_fmt, "#D97706", True))

        be_rate = summary.get("balloon_breakeven")
        if be_rate is not None:
            metrics_data.append(
                ("Vs. Annuity Break-even (Balloon)", be_rate / 100, pct_fmt, "#D97706", True))

        be_abs = summary.get("balloon_breakeven_abs")
        if be_abs is not None:
            metrics_data.append(
                ("Absolute Break-even (Balloon)", be_abs / 100, pct_fmt, "#D97706", True))

        # ── Inflation ─────────────────────────────────────────────────────────
        if summary.get("inflation_enabled") and summary.get("real_cost") is not None:
            metrics_data.append(
                ("Real Total Cost (PV)", summary.get("real_cost"), num_fmt, "#065F46", True))
            metrics_data.append(
                ("Inflation Discount", summary.get("inflation_savings", 0), num_fmt, "#0F766E", False))

        # ── Risk metrics ──────────────────────────────────────────────────────
        ltv = summary.get("ltv")
        if ltv is not None:
            metrics_data.append(("LTV (Loan-to-Value)", ltv / 100, pct_fmt, "#1D4ED8", True))

        dscr = summary.get("dscr")
        if dscr is not None:
            metrics_data.append(("DSCR", dscr, "0.00", "#1D4ED8", True))

        dti = summary.get("dti")
        if dti is not None:
            metrics_data.append(("DTI (Debt-to-Income)", dti / 100, pct_fmt, "#1D4ED8", True))

    for i, (lbl, val, fmt, fg_color, is_highlight) in enumerate(metrics_data):
        r = metrics_start + 1 + i
        ws1.row_dimensions[r].height = 20
        is_even = i % 2 == 0
        row_fill = fill("F0FDF4") if (is_highlight and is_even) else \
                   fill("ECFDF5") if is_highlight else \
                   fill("F1F5F9") if is_even else fill(WHITE)

        c_lbl = ws1.cell(r, 4, f"  {lbl}")
        c_lbl.font      = font(color="1E293B", bold=is_highlight, size=9)
        c_lbl.fill      = row_fill
        c_lbl.alignment = align("left")
        c_lbl.border    = border(BORDER_C)

        c_val = ws1.cell(r, 5)
        c_val.value          = val
        c_val.number_format  = fmt
        c_val.font           = Font(name="Calibri", bold=True, size=9,
                                     color=fg_color.replace("#",""))
        c_val.fill           = row_fill
        c_val.alignment      = align("right")
        c_val.border         = border(BORDER_C)

    # ── Дисклеймер на Summary ─────────────────────────────────────────────────
    disc_row = metrics_start + len(metrics_data) + 3
    ws1.merge_cells(f"A{disc_row}:E{disc_row}")
    dc = ws1[f"A{disc_row}"]
    dc.value = (
        "DISCLAIMER: Results are preliminary and for informational purposes only. "
        "This report is not a public offer. Verify exact terms with your bank. "
        "Yev Capital LoanLogic v3.0 / Bohdan Yevtushenko (MrCemper) — "
        "© 2026 Not liable for financial decisions based on these calculations."
    )
    dc.font      = Font(name="Calibri", size=7, italic=True, color="94A3B8")
    dc.alignment = Alignment(horizontal="left", wrap_text=True)
    dc.fill      = fill("F8FAFC")
    ws1.row_dimensions[disc_row].height = 36

    set_col_width(ws1)
    ws1.column_dimensions["A"].width = 34
    ws1.column_dimensions["B"].width = 28
    ws1.column_dimensions["C"].width = 3
    ws1.column_dimensions["D"].width = 34
    ws1.column_dimensions["E"].width = 28

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    #  ЛИСТ 2: PAYMENT SCHEDULE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    ws2 = wb.create_sheet("Payment Schedule")
    ws2.sheet_view.showGridLines = False

    # ── Шапка листа ───────────────────────────────────────────────────────────
    n_cols = len(df.columns)
    last_col = get_column_letter(n_cols)
    ws2.merge_cells(f"A1:{last_col}1")
    sh = ws2["A1"]
    sh.value     = "Yev Capital LoanLogic  ·  Payment Schedule"
    sh.font      = Font(name="Calibri", bold=True, size=14, color="4FC3F7")
    sh.fill      = fill(NAVY)
    sh.alignment = align("center")
    ws2.row_dimensions[1].height = 34

    ws2.merge_cells(f"A2:{last_col}2")
    sh2 = ws2["A2"]
    sh2.value     = (f"Principal: {sym}{principal:,.2f}  ·  Rate: {rate_pa:.2f}%  ·  "
                     f"Periods: {n_periods}  ·  Generated: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    sh2.font      = Font(name="Calibri", italic=True, size=9, color=SUBTEXT)
    sh2.fill      = fill(NAVY)
    sh2.alignment = align("center")
    ws2.row_dimensions[2].height = 18

    # ── Заголовки колонок ─────────────────────────────────────────────────────
    HDR_ROW = 3
    cols = list(df.columns)
    for ci, col_name in enumerate(cols, 1):
        cell = ws2.cell(HDR_ROW, ci, col_name)
        cell.fill      = fill(MIDBLUE)
        cell.font      = Font(name="Calibri", bold=True, size=9, color=WHITE)
        cell.alignment = align("center", wrap=True)
        cell.border    = border(WHITE, "thin")
    ws2.row_dimensions[HDR_ROW].height = 28
    ws2.freeze_panes = f"A{HDR_ROW + 1}"   # Freeze заголовки

    # ── Данные с "зеброй" и формулами ────────────────────────────────────────
    sched    = summary.get("schedule", [])
    total_key = t.get("total_row", "ИТОГО")

    # Определяем колонки с числами (начиная с 3-й и далее — всё числа)
    num_col_start = 3  # с 3-й колонки идут числа (пропускаем Период и Дата)

    for ri, (_, row) in enumerate(df.iterrows()):
        r          = HDR_ROW + 1 + ri
        row_vals   = list(row)
        is_total   = str(row_vals[0]) == total_key
        is_even    = ri % 2 == 0

        if is_total:
            row_bg = TOTALBG
        elif is_even:
            row_bg = GRAY2       # светлая зебра
        else:
            row_bg = "EBF2FB"    # чуть синеватая

        ws2.row_dimensions[r].height = 18

        for ci, val in enumerate(row_vals, 1):
            cell = ws2.cell(r, ci)

            if is_total:
                cell.fill  = fill(TOTALBG)
                cell.font  = Font(name="Calibri", bold=True, size=9, color=TOTALFG)
            elif ci <= 2:
                # Период и Дата
                cell.fill = fill(row_bg)
                cell.font = Font(name="Calibri", size=9, color="334155",
                                  bold=(ci == 1))
            else:
                cell.fill = fill(row_bg)
                cell.font = Font(name="Calibri", size=9, color="1E293B")

            # Date column (ci == 2): parse the "dd.mm.yyyy" string into a
            # real date object and apply a date number format. Excel then
            # treats the cell as a date (sortable, supports date-arithmetic
            # in formulas) rather than as opaque text.
            if ci == 2 and isinstance(val, str) and val:
                parsed_dt = None
                for _fmt in ("%d.%m.%Y", "%Y-%m-%d", "%m/%d/%Y", "%d/%m/%Y"):
                    try:
                        parsed_dt = datetime.strptime(val, _fmt).date()
                        break
                    except (ValueError, TypeError):
                        continue
                if parsed_dt is not None:
                    cell.value         = parsed_dt
                    cell.number_format = "dd.mm.yyyy"
                    cell.alignment     = align("center")
                    cell.border        = border(BORDER_C, "thin")
                    continue   # already wrote, skip generic branch

            # Числовое значение или текст
            if isinstance(val, (int, float)) and not is_total:
                cell.value          = val
                cell.number_format  = num_fmt
                cell.alignment      = align("right")
            elif isinstance(val, (int, float)) and is_total:
                cell.value          = val
                cell.number_format  = num_fmt
                cell.alignment      = align("right")
            else:
                cell.value     = val
                cell.alignment = align("center" if ci <= 2 else "right")

            cell.border = border(BORDER_C, "thin")

    # ── SUM-formula verification row below the data totals ────────────────────
    formula_row = HDR_ROW + len(df) + 2
    ws2.row_dimensions[formula_row].height = 20
    ws2.merge_cells(f"A{formula_row}:B{formula_row}")
    fc = ws2[f"A{formula_row}"]
    fc.value     = "✓ Excel SUM Verification"
    fc.font      = Font(name="Calibri", bold=True, size=8, color="065F46")
    fc.fill      = fill("F0FDF4")
    fc.alignment = align("center")

    data_start_r = HDR_ROW + 1
    data_end_r   = HDR_ROW + len(df) - 1  # без итоговой строки df
    for ci in range(3, n_cols + 1):
        col_l = get_column_letter(ci)
        fc2   = ws2.cell(formula_row, ci)
        fc2.value          = f"=SUM({col_l}{data_start_r}:{col_l}{data_end_r})"
        fc2.number_format  = num_fmt
        fc2.font           = Font(name="Calibri", bold=True, size=8, color="065F46")
        fc2.fill           = fill("F0FDF4")
        fc2.alignment      = align("right")
        fc2.border         = border("10B981", "medium")

    set_col_width(ws2)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    #  ЛИСТ 3: ANALYSIS
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    ws3 = wb.create_sheet("Analysis")
    ws3.sheet_view.showGridLines = False
    ws3.column_dimensions["A"].width = 36
    ws3.column_dimensions["B"].width = 22
    ws3.column_dimensions["C"].width = 5
    ws3.column_dimensions["D"].width = 36
    ws3.column_dimensions["E"].width = 22

    # Шапка
    ws3.merge_cells("A1:E1")
    ah = ws3["A1"]
    ah.value     = "Yev Capital LoanLogic  ·  Analysis"
    ah.font      = Font(name="Calibri", bold=True, size=14, color="4FC3F7")
    ah.fill      = fill(NAVY)
    ah.alignment = align("center")
    ws3.row_dimensions[1].height = 34

    def write_section_header(ws, row, text, bg=MIDBLUE):
        ws.merge_cells(f"A{row}:E{row}")
        c = ws[f"A{row}"]
        c.value     = f"  {text}"
        c.font      = Font(name="Calibri", bold=True, size=10, color=WHITE)
        c.fill      = fill(bg)
        c.alignment = align("left")
        ws.row_dimensions[row].height = 22

    def write_kv_row(ws, row, key, value, fmt=None, is_even=True, highlight=False):
        bg = fill("F0FDF4") if highlight else (fill("F1F5F9") if is_even else fill(WHITE))
        ck = ws.cell(row, 1, f"  {key}")
        ck.font      = Font(name="Calibri", size=9, color="1E293B", bold=highlight)
        ck.fill      = bg
        ck.alignment = align("left")
        ck.border    = border(BORDER_C)
        cv = ws.cell(row, 2)
        cv.value     = value
        if fmt:
            cv.number_format = fmt
        cv.font      = Font(name="Calibri", size=9, bold=True,
                             color="065F46" if highlight else "1D3557")
        cv.fill      = bg
        cv.alignment = align("right")
        cv.border    = border(BORDER_C)
        for ci in range(3, 6):
            ws.cell(row, ci).fill = fill(WHITE)
        ws.row_dimensions[row].height = 19

    r = 3
    if not is_dep:
        # ── Секция 1: Сравнение схем ─────────────────────────────────────────
        write_section_header(ws3, r, "SCHEME COMPARISON  —  Annuity vs Classic vs Balloon")
        r += 1

        sched_ref = summary.get("schedule", [])
        unit_r    = summary.get("unit", "months")
        rate_r    = summary.get("rate_pa", summary.get("effective_rate", 0))
        n_r       = len(sched_ref)
        princ_r   = summary.get("principal", 0)
        mo_comm_r = sched_ref[0]["commission"] if sched_ref else 0

        try:
            sched_ann = calc_annuity(princ_r, n_r, rate_r, unit_r, mo_comm_r)
            sched_cla = calc_classic(princ_r, n_r, rate_r, unit_r, mo_comm_r)
            sched_bal = calc_balloon(princ_r, n_r, rate_r, unit_r, mo_comm_r)

            ann_total = sum(x["payment"] for x in sched_ann)
            cla_total = sum(x["payment"] for x in sched_cla)
            bal_total = sum(x["payment"] for x in sched_bal)
            ann_int   = sum(x["interest"] for x in sched_ann)
            cla_int   = sum(x["interest"] for x in sched_cla)
            bal_int   = sum(x["interest"] for x in sched_bal)

            headers = ["Metric", "Annuity", "Classic (Diff.)", "Balloon"]
            hdr_bg  = [MIDBLUE, "1D4ED8", "065F46", "7C3AED"]
            for ci, (hdr, bg) in enumerate(zip(headers, hdr_bg), 1):
                cell = ws3.cell(r, ci)
                cell.value     = hdr
                cell.font      = Font(name="Calibri", bold=True, size=9, color=WHITE)
                cell.fill      = fill(bg)
                cell.alignment = align("center")
                cell.border    = border(WHITE)
            ws3.column_dimensions["B"].width = 18
            ws3.column_dimensions["C"].width = 18  # здесь Classic
            ws3.column_dimensions["D"].width = 18
            ws3.column_dimensions["E"].width = 18
            r += 1

            comparison_rows = [
                ("Total Payments",  ann_total, cla_total, bal_total),
                ("Total Interest",  ann_int,   cla_int,   bal_int),
                ("1st Payment",     sched_ann[0]["payment"], sched_cla[0]["payment"], sched_bal[0]["payment"]),
                ("Last Payment",    sched_ann[-1]["payment"],sched_cla[-1]["payment"],sched_bal[-1]["payment"]),
            ]
            for i_r, (lbl, va, vc, vb) in enumerate(comparison_rows):
                is_even = i_r % 2 == 0
                bg0 = fill("F1F5F9") if is_even else fill(WHITE)

                cl = ws3.cell(r, 1, f"  {lbl}")
                cl.font  = Font(name="Calibri", size=9, color="1E293B", bold=True)
                cl.fill  = bg0; cl.alignment = align("left"); cl.border = border(BORDER_C)

                # Подсвечиваем минимальное значение зелёным
                vals = [va, vc, vb]
                min_v = min(vals)
                for ci, v in enumerate(vals, 2):
                    is_min = (v == min_v)
                    cv = ws3.cell(r, ci)
                    cv.value          = v
                    cv.number_format  = num_fmt
                    cv.fill           = fill("D1FAE5") if is_min else bg0
                    cv.font           = Font(name="Calibri", size=9,
                                             color="065F46" if is_min else "1E293B",
                                             bold=is_min)
                    cv.alignment      = align("right")
                    cv.border         = border(BORDER_C)
                ws3.row_dimensions[r].height = 19
                r += 1

            # Savings строка
            savings = ann_total - cla_total
            ws3.merge_cells(f"A{r}:D{r}")
            sc = ws3[f"A{r}"]
            sc.value     = (f"  💡 Classic saves {sym}{savings:,.2f} vs Annuity  "
                            f"(Balloon has lowest periodic payments but large final payment)")
            sc.font      = Font(name="Calibri", size=8, italic=True, color="065F46", bold=True)
            sc.fill      = fill("ECFDF5")
            sc.alignment = align("left", wrap=True)
            ws3.row_dimensions[r].height = 24
            r += 1

            # ── Balloon Break-even ─────────────────────────────────────────────
            be_rate_val = summary.get("balloon_breakeven")
            if be_rate_val is not None:
                ws3.merge_cells(f"A{r}:D{r}")
                be_cell = ws3[f"A{r}"]
                be_cell.value = (
                    f"  📐 Inv. Break-even Rate (Balloon vs Annuity): {be_rate_val:.2f}%  "
                    f"— Min. compound annual investment return to justify Balloon scheme."
                )
                be_cell.font      = Font(name="Calibri", size=8, italic=True, color="D97706", bold=True)
                be_cell.fill      = fill("FFFBEB")
                be_cell.alignment = align("left", wrap=True)
                ws3.row_dimensions[r].height = 22
                r += 1
            r += 1

        except Exception:
            r += 1

        # ── Секция 2: Interest Breakdown ─────────────────────────────────────
        write_section_header(ws3, r, "INTEREST COST ANALYSIS", bg="7C3AED")
        r += 1

        total_p  = summary.get("total_payment", 0)
        total_i  = summary.get("total_interest", 0)
        total_c  = summary.get("total_commission", 0)
        eff_r    = summary.get("effective_rate")
        overpay  = summary.get("overpay_pct", 0) or 0

        eff_apr_row = (("Effective APR", "N/A", "@", True) if eff_r is None
                        else ("Effective APR", eff_r / 100, pct_fmt, True))
        cost_rows = [
            ("Principal (Loan Body)",   princ_r,  num_fmt, False),
            ("Total Interest Paid",     total_i,  num_fmt, True),
            ("Total Commissions",       total_c,  num_fmt, False),
            ("Total Cost of Credit",    total_p,  num_fmt, True),
            eff_apr_row,
            ("Overpayment vs Principal",overpay/100,pct_fmt,True),
        ]
        for i_r, (lbl, val, fmt, hl) in enumerate(cost_rows):
            write_kv_row(ws3, r, lbl, val, fmt, i_r % 2 == 0, hl)
            r += 1
        r += 1

        # ── Секция 3: Amortization milestones ────────────────────────────────
        write_section_header(ws3, r, "AMORTIZATION MILESTONES", bg="0369A1")
        r += 1

        sched_all = summary.get("schedule", [])
        if sched_all:
            # Находим период, когда выплачено 25%, 50%, 75% тела долга
            cumulative_princ = 0
            milestones = {0.25: None, 0.50: None, 0.75: None, 1.00: None}
            for row_s in sched_all:
                cumulative_princ += row_s.get("principal", 0)
                ratio = cumulative_princ / princ_r if princ_r > 0 else 0
                for threshold in [0.25, 0.50, 0.75, 1.00]:
                    if milestones[threshold] is None and ratio >= threshold:
                        milestones[threshold] = row_s["period"]

            milestone_rows = [
                ("25% Principal Repaid — Period", milestones.get(0.25) or n_r, "#,##0"),
                ("50% Principal Repaid — Period", milestones.get(0.50) or n_r, "#,##0"),
                ("75% Principal Repaid — Period", milestones.get(0.75) or n_r, "#,##0"),
                ("100% Repaid — Final Period",    milestones.get(1.00) or n_r, "#,##0"),
            ]
            for i_r, (lbl, val, fmt) in enumerate(milestone_rows):
                write_kv_row(ws3, r, lbl, val, fmt, i_r % 2 == 0, i_r == 3)
                r += 1

    else:
        # ── ДЕПОЗИТ: аналитика ────────────────────────────────────────────────
        write_section_header(ws3, r, "DEPOSIT PERFORMANCE ANALYSIS", bg="065F46")
        r += 1

        princ_r = summary.get("principal", 0)
        final_b = summary.get("final_balance", 0)
        earned  = summary.get("total_earned", 0)
        eff_r   = summary.get("effective_rate")

        dep_eff_row = (("Effective Annual Rate", "N/A", "@", True) if eff_r is None
                        else ("Effective Annual Rate", eff_r / 100, pct_fmt, True))
        dep_rows = [
            ("Initial Deposit",       princ_r,               num_fmt, False),
            ("Final Balance",          final_b,               num_fmt, True),
            ("Total Interest Earned",  earned,                num_fmt, True),
            ("Profit (Gross Return)",  final_b - princ_r,    num_fmt, True),
            dep_eff_row,
            ("Total Return %",         (earned / princ_r) if princ_r > 0 else 0, pct_fmt, True),
        ]
        for i_r, (lbl, val, fmt, hl) in enumerate(dep_rows):
            write_kv_row(ws3, r, lbl, val, fmt, i_r%2==0, hl)
            r += 1
        r += 1

        write_section_header(ws3, r, "DEPOSIT MODE COMPARISON", bg="0369A1")
        r += 1
        unit_r  = summary.get("unit", "months")
        rate_r  = summary.get("rate_pa", eff_r)
        n_r     = len(summary.get("schedule", []))
        try:
            sc_cap = calc_deposit(princ_r, n_r, rate_r, unit_r, "capitalize")
            sc_pay = calc_deposit(princ_r, n_r, rate_r, unit_r, "payout")
            cap_final = sc_cap[-1]["balance_close"]
            cap_earn  = sum(x["interest"] for x in sc_cap)
            pay_earn  = sum(x["interest"] for x in sc_pay)
            diff_earn = cap_earn - pay_earn

            mode_rows = [
                ("Capitalization — Final Balance", cap_final,  num_fmt, True),
                ("Capitalization — Total Earned",  cap_earn,   num_fmt, False),
                ("Payout — Total Received",        pay_earn,   num_fmt, False),
                ("Extra Earned (Capitalization)",  diff_earn,  num_fmt, True),
            ]
            for i_r, (lbl, val, fmt, hl) in enumerate(mode_rows):
                write_kv_row(ws3, r, lbl, val, fmt, i_r%2==0, hl)
                r += 1
        except Exception:
            # Informational comparison: omit the section on failure rather
            # than aborting the whole export.
            pass

    # ── Дисклеймер на Analysis ────────────────────────────────────────────────
    r += 2
    ws3.merge_cells(f"A{r}:E{r}")
    dc3 = ws3[f"A{r}"]
    dc3.value = ("DISCLAIMER: For informational purposes only. Not a public offer. "
                 "Verify all terms with your financial institution. "
                 "Yev Capital LoanLogic v3.0 / Bohdan Yevtushenko (MrCemper) — © 2026")
    dc3.font      = Font(name="Calibri", size=7, italic=True, color="94A3B8")
    dc3.alignment = Alignment(horizontal="left", wrap_text=True)
    dc3.fill      = fill("F8FAFC")
    ws3.row_dimensions[r].height = 28

    # ── Сохраняем ─────────────────────────────────────────────────────────────
    # Активируем Summary
    wb.active = ws1

    wb.save(buf)
    return buf.getvalue()

# ─────────────────────────────────────────────────────────────────────────────
#  ЭКСПОРТ В DOCX
# ─────────────────────────────────────────────────────────────────────────────
def export_docx(df, summary, t, sym):
    """
    Word-документ с банковскими стандартами форматирования:
      • Шрифт Arial / Calibri
      • Название 15 pt Bold, подзаголовки 12 pt Bold, таблица 10 pt (9 если > 50 строк)
      • Числа: правый край, пробел как разделитель тысяч
      • Дата/текст: левый край
      • Итоговая строка: тёмно-синий Bold
    """
    is_dep  = summary.get("is_deposit", False)
    if df is None or getattr(df, "empty", True) or len(df.columns) == 0:
        raise ValueError(
            "export_docx: no schedule data to export (calculation may have "
            "failed or produced an empty result).")
    n_rows  = len(df)
    tbl_fs  = Pt(9) if n_rows > 50 else Pt(10)

    doc     = Document()
    section = doc.sections[0]
    for attr in ("left_margin", "right_margin"):
        setattr(section, attr, Cm(1.8))
    for attr in ("top_margin", "bottom_margin"):
        setattr(section, attr, Cm(2))

    NAVY_RGB   = RGBColor(0x1D, 0x35, 0x57)
    BLUE_RGB   = RGBColor(0x1D, 0x4E, 0xD8)
    MUTED_RGB  = RGBColor(0x94, 0xA3, 0xB8)
    TOTAL_RGB  = RGBColor(0x0D, 0x1B, 0x2A)
    RED_RGB    = RGBColor(0xDC, 0x26, 0x26)
    WHITE_RGB  = RGBColor(0xFF, 0xFF, 0xFF)

    def set_run(run, size_pt, bold=False, italic=False, color=None, font="Arial"):
        run.font.name  = font
        run.font.size  = size_pt
        run.bold       = bold
        run.italic     = italic
        if color:
            run.font.color.rgb = color

    def add_hdr(text, level, size, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run(run, Pt(size), bold=True, color=color or NAVY_RGB)
        return p

    # ── ШАПКА ─────────────────────────────────────────────────────────────────
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = h.add_run("Yev Capital LoanLogic")
    set_run(r1, Pt(16), bold=True, color=BLUE_RGB, font="Arial")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(
        "Loan Analysis Report" if not is_dep else "Deposit Analysis Report")
    set_run(r2, Pt(11), italic=True, color=MUTED_RGB)

    start_dt  = summary.get("start_date", date.today())
    start_str = start_dt.strftime("%d.%m.%Y") if hasattr(start_dt, "strftime") else str(start_dt)
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta_p.add_run(
        f"Report: {datetime.now().strftime('%d.%m.%Y %H:%M')}  ·  "
        f"Start date: {start_str}  ·  Prepared by: Yevtushenko")
    set_run(rm, Pt(8), color=MUTED_RGB)
    doc.add_paragraph()

    # ── КЛЮЧЕВЫЕ МЕТРИКИ ──────────────────────────────────────────────────────
    add_hdr(t.get("section_results", "Key Metrics"), level=2, size=12, color=BLUE_RGB)

    if is_dep:
        kv_pairs = [
            (t.get("dep_initial",      "Initial Deposit"),  fmt_money(summary.get("principal",    0), sym)),
            (t.get("dep_final_balance","Final Balance"),     fmt_money(summary.get("final_balance",0), sym)),
            (t.get("dep_total_earned", "Interest Earned"),  fmt_money(summary.get("total_earned", 0), sym)),
            (t.get("dep_rate_label",   "Annual Rate"),       fmt_pct(summary.get("effective_rate",0))),
        ]
    else:
        eff_rate_val = summary.get("effective_rate")
        eff_rate_str = fmt_pct(eff_rate_val) if eff_rate_val is not None else "N/A"
        kv_pairs = [
            (t["total_payment"],    fmt_money(summary["total_payment"],    sym)),
            (t["total_interest"],   fmt_money(summary["total_interest"],   sym)),
            (t["total_commission"], fmt_money(summary["total_commission"], sym)),
            (t["effective_rate"],   eff_rate_str),
            (t["monthly_payment"],  fmt_money(summary["first_payment"],    sym)),
        ]
        # Investment break-even
        if summary.get("universal_breakeven") is not None:
            kv_pairs.append(("Universal Break-even Rate",
                              fmt_pct(summary["universal_breakeven"])))
        if summary.get("balloon_breakeven") is not None:
            kv_pairs.append(("Vs. Annuity Break-even (Balloon)",
                              fmt_pct(summary["balloon_breakeven"])))
        if summary.get("balloon_breakeven_abs") is not None:
            kv_pairs.append(("Absolute Break-even (Balloon)",
                              fmt_pct(summary["balloon_breakeven_abs"])))
        # Inflation
        if summary.get("inflation_enabled") and summary.get("real_cost") is not None:
            kv_pairs.append(("Real Total Cost (PV)",
                              fmt_money(summary["real_cost"], sym)))
            kv_pairs.append(("Inflation Discount",
                              fmt_money(summary.get("inflation_savings", 0), sym)))
        # Risk
        if summary.get("ltv") is not None:
            kv_pairs.append(("LTV (Loan-to-Value)", fmt_pct(summary["ltv"])))
        if summary.get("dscr") is not None:
            kv_pairs.append(("DSCR", f"{summary['dscr']:.2f}"))
        if summary.get("dti") is not None:
            kv_pairs.append(("DTI (Debt-to-Income)", fmt_pct(summary["dti"])))

    it = doc.add_table(rows=len(kv_pairs), cols=2)
    it.style = "Table Grid"
    for i, (lbl, val) in enumerate(kv_pairs):
        lc = it.cell(i, 0)
        vc = it.cell(i, 1)
        rr = lc.paragraphs[0].add_run(lbl)
        set_run(rr, Pt(10), bold=True, color=NAVY_RGB)
        lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        rv = vc.paragraphs[0].add_run(val)
        set_run(rv, Pt(10), color=BLUE_RGB)
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()

    # ── ТАБЛИЦА ПЛАТЕЖЕЙ ──────────────────────────────────────────────────────
    add_hdr(t["section_schedule"], level=2, size=12, color=BLUE_RGB)
    cols      = list(df.columns)
    total_key = t.get("dep_total_row") if is_dep else t.get("total_row", "ИТОГО")
    tbl       = doc.add_table(rows=1 + len(df), cols=len(cols))
    tbl.style = "Table Grid"

    # Заголовки
    for ci, col in enumerate(cols):
        p   = tbl.rows[0].cells[ci].paragraphs[0]
        run = p.add_run(col)
        set_run(run, Pt(9), bold=True, color=BLUE_RGB)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Данные
    for ri, (_, row) in enumerate(df.iterrows()):
        vals    = list(row)
        is_tot  = str(vals[0]) == total_key
        for ci, v in enumerate(vals):
            p   = tbl.rows[ri+1].cells[ci].paragraphs[0]
            # Форматирование числа
            if isinstance(v, (int, float)):
                display = f"{v:,.2f}".replace(",", "\u202f")
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                display = str(v)
                p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if ci <= 1
                                else WD_ALIGN_PARAGRAPH.RIGHT)
            run = p.add_run(display)
            # Цвет: красный для отрицательных, синий для итога, серый для обычных
            if is_tot:
                set_run(run, tbl_fs, bold=True, color=TOTAL_RGB)
            elif isinstance(v, (int, float)) and v < 0:
                set_run(run, tbl_fs, color=RED_RGB)
            else:
                set_run(run, tbl_fs, color=RGBColor(0x1E, 0x29, 0x3B))

    doc.add_paragraph()

    # ── BALLOON BREAK-EVEN (compact row, only if applicable) ──────────────────
    be_rate_val = summary.get("balloon_breakeven")
    if be_rate_val is not None:
        bep = doc.add_paragraph()
        ber = bep.add_run(
            f"📐 Inv. Break-even Rate (Balloon vs Annuity): {be_rate_val:.2f}%  "
            f"— minimum compound annual investment return to justify Balloon scheme.")
        set_run(ber, Pt(9), bold=True, color=RGBColor(0xD9, 0x77, 0x06))

    # ── LEGAL DISCLAIMER (always English) ─────────────────────────────────────
    dp = doc.add_paragraph()
    dr = dp.add_run(
        "Legal Disclaimer: Results are for informational purposes only and do not "
        "constitute a public offer. Verify all terms with your financial institution. "
        "Yev Capital LoanLogic v3.0 / Bohdan Yevtushenko (MrCemper) — "
        "© 2026")
    set_run(dr, Pt(8), italic=True, color=MUTED_RGB)

    fp = doc.add_paragraph(
        f"Yev Capital LoanLogic v3.0  ·  © 2026 Bohdan Yevtushenko (MrCemper)  ·  "
        f"{datetime.now().strftime('%d.%m.%Y %H:%M')}")
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in fp.runs:
        set_run(run, Pt(8), color=MUTED_RGB)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ─────────────────────────────────────────────────────────────────────────────
#  ЭКСПОРТ В PDF  —  Кириллица + банковское форматирование
# ─────────────────────────────────────────────────────────────────────────────
def export_pdf(df, summary, t, sym):
    """
    Генерирует PDF с поддержкой кириллицы (DejaVuSans / Arial).
    Банковские стандарты форматирования:
      • Название отчёта: 15 pt Bold
      • Подзаголовки: 12 pt Bold
      • Таблица: 10 pt (9 pt если > 50 строк)
      • Дисклеймер: 8 pt серый
      • Числа выровнены вправо, текст влево
      • Разделитель тысяч — пробел (1 000 000.00)
      • Отрицательные значения — красный
      • Итоговая строка — тёмно-синий Bold
    """
    # Guard against an empty schedule (e.g. a partial/failed syndicated result):
    # building a table with no columns/rows raises IndexError downstream.
    if df is None or getattr(df, "empty", True) or len(df.columns) == 0:
        raise ValueError(
            "export_pdf: no schedule data to export (calculation may have "
            "failed or produced an empty result).")
    # ── Если шрифт не был загружен — предупредить в консоль (не краш) ─────────
    if PDF_FONT_WARN:
        import warnings
        warnings.warn(PDF_FONT_WARN, stacklevel=2)

    F       = PDF_FONT        # regular (с кириллицей)
    F_BOLD  = PDF_FONT_BOLD   # bold
    F_IT    = PDF_FONT_ITALIC  # italic / fallback regular

    is_dep  = summary.get("is_deposit", False)
    n_rows  = len(df)
    tbl_fs  = 9 if n_rows > 50 else 10   # адаптивный размер шрифта таблицы

    # ── Цвета ─────────────────────────────────────────────────────────────────
    NAVY       = rl_colors.HexColor("#0D1B2A")
    MIDBLUE    = rl_colors.HexColor("#1D4ED8")
    LIGHTBLUE  = rl_colors.HexColor("#EFF6FF")
    LIGHTBLUE2 = rl_colors.HexColor("#F8FAFC")
    TOTAL_BG   = rl_colors.HexColor("#1D3557")
    TOTAL_FG   = rl_colors.HexColor("#4FC3F7")
    RED        = rl_colors.HexColor("#DC2626")
    MUTED      = rl_colors.HexColor("#64748B")
    WHITE      = rl_colors.white
    BORDER_C   = rl_colors.HexColor("#CBD5E1")
    GREEN_BG   = rl_colors.HexColor("#ECFDF5")
    GREEN_FG   = rl_colors.HexColor("#065F46")

    # ── Стили параграфов ──────────────────────────────────────────────────────
    def make_style(name, font=None, size=10, bold=False, color=None,
                   align=0, space_before=0, space_after=4, leading=None):
        """Создаёт ParagraphStyle с нужными параметрами."""
        kw = dict(
            fontName   = (font or (F_BOLD if bold else F)),
            fontSize   = size,
            textColor  = color or rl_colors.black,
            alignment  = align,
            spaceBefore= space_before,
            spaceAfter = space_after,
        )
        if leading:
            kw["leading"] = leading
        return ParagraphStyle(name, **kw)

    title_s  = make_style("pdf_title",  size=15, bold=True,
                           color=rl_colors.HexColor("#0D1B2A"),
                           align=1, space_after=4)
    sub_s    = make_style("pdf_subtitle", size=11, bold=False,
                           color=MUTED, align=1, space_after=10)
    h2_s     = make_style("pdf_h2", size=12, bold=True,
                           color=MIDBLUE, space_before=10, space_after=4)
    body_s   = make_style("pdf_body", size=10, color=rl_colors.HexColor("#1E293B"))
    disc_s   = make_style("pdf_disc", size=8, color=MUTED,
                           align=0, space_before=12, leading=11)
    footer_s = make_style("pdf_footer", size=8, color=MUTED, align=2)

    # ── Хелпер: ячейка таблицы как Paragraph для кириллицы ───────────────────
    def cell(text, bold=False, align="LEFT", color=None, size=None):
        """
        Возвращает Paragraph, а не строку — это гарантирует правильный шрифт
        и кириллицу в каждой ячейке.
        """
        fs    = size or tbl_fs
        fname = F_BOLD if bold else F
        clr   = color or rl_colors.black
        a_map = {"LEFT": 0, "RIGHT": 2, "CENTER": 1}
        style = ParagraphStyle(
            f"cell_{id(text)}",
            fontName=fname, fontSize=fs,
            textColor=clr, alignment=a_map.get(align, 0),
            leading=fs + 2, spaceBefore=0, spaceAfter=0,
        )
        return Paragraph(str(text), style)

    def hdr_cell(text):
        """Заголовок колонки — белый Bold, центр."""
        return cell(text, bold=True, align="CENTER", color=WHITE, size=tbl_fs)

    def num_cell(v, is_total=False, is_red=False):
        """Числовая ячейка — правый край, форматирование с пробелами."""
        if isinstance(v, (int, float)):
            formatted = f"{v:,.2f}".replace(",", "\u202f")
        else:
            formatted = str(v)
        color = RED if is_red else (TOTAL_FG if is_total else rl_colors.HexColor("#1E293B"))
        return cell(formatted, bold=is_total, align="RIGHT", color=color)

    def txt_cell(v, is_total=False):
        """Текстовая ячейка — левый край."""
        color = TOTAL_FG if is_total else rl_colors.HexColor("#334155")
        return cell(str(v), bold=is_total, align="LEFT", color=color)

    # ── Построение PDF ────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.8*cm, rightMargin=1.8*cm,
        topMargin=2*cm, bottomMargin=2*cm,
        title="Yev Capital LoanLogic — Report",
        author="Yev Capital / Yevtushenko",
    )

    elems = []
    PAGE_W = A4[0] - 3.6*cm   # доступная ширина

    # ── ШАПКА ─────────────────────────────────────────────────────────────────
    report_type = ("Deposit Analysis Report" if is_dep else "Loan Analysis Report")
    start_dt    = summary.get("start_date", date.today())
    start_str   = start_dt.strftime("%d.%m.%Y") if hasattr(start_dt, "strftime") else str(start_dt)

    elems.append(Paragraph("Yev Capital LoanLogic", title_s))
    elems.append(Paragraph(report_type, sub_s))

    # Линия-разделитель
    elems.append(HRFlowable(width="100%", thickness=1.5,
                             color=MIDBLUE, spaceAfter=8))

    # ── Метаданные ────────────────────────────────────────────────────────────
    meta_data = [
        [cell("Report Date:", bold=True), cell(datetime.now().strftime("%d.%m.%Y  %H:%M"))],
        [cell("Start Date:",  bold=True), cell(start_str)],
        [cell("Prepared by:", bold=True), cell("Yev Capital LoanLogic v3.0  |  Yevtushenko")],
    ]
    meta_tbl = Table(meta_data, colWidths=[4*cm, PAGE_W - 4*cm])
    meta_tbl.setStyle(TableStyle([
        ("FONTNAME",      (0,0),(-1,-1), F),
        ("FONTSIZE",      (0,0),(-1,-1), 9),
        ("TEXTCOLOR",     (0,0),(-1,-1), rl_colors.HexColor("#475569")),
        ("TOPPADDING",    (0,0),(-1,-1), 2),
        ("BOTTOMPADDING", (0,0),(-1,-1), 2),
        ("BACKGROUND",    (0,0),(-1,-1), LIGHTBLUE2),
    ]))
    elems.append(meta_tbl)
    elems.append(Spacer(1, 0.4*cm))

    # ── КЛЮЧЕВЫЕ МЕТРИКИ ──────────────────────────────────────────────────────
    elems.append(Paragraph(t.get("section_results", "Key Metrics"), h2_s))

    if is_dep:
        principal    = summary.get("principal", 0)
        final_bal    = summary.get("final_balance", 0)
        total_earned = summary.get("total_earned", 0)
        eff_r        = summary.get("effective_rate", 0)
        metrics_rows = [
            [cell(t.get("dep_initial",     "Initial Deposit"), bold=True),
             cell(fmt_money(principal,    sym))],
            [cell(t.get("dep_final_balance","Final Balance"),  bold=True),
             cell(fmt_money(final_bal,    sym))],
            [cell(t.get("dep_total_earned","Interest Earned"), bold=True),
             cell(fmt_money(total_earned, sym))],
            [cell(t.get("dep_rate_label",  "Annual Rate"),     bold=True),
             cell(fmt_pct(eff_r))],
        ]
    else:
        metrics_rows = [
            [cell(t["total_payment"],    bold=True),
             cell(fmt_money(summary["total_payment"],    sym))],
            [cell(t["total_interest"],   bold=True),
             cell(fmt_money(summary["total_interest"],   sym))],
            [cell(t["total_commission"], bold=True),
             cell(fmt_money(summary["total_commission"], sym))],
            [cell(t["effective_rate"],   bold=True),
             cell(fmt_pct(summary["effective_rate"])
                  if summary.get("effective_rate") is not None else "N/A")],
            [cell(t["monthly_payment"],  bold=True),
             cell(fmt_money(summary["first_payment"],    sym))],
        ]
        # Investment break-even
        if summary.get("universal_breakeven") is not None:
            metrics_rows.append([
                cell("Universal Inv. Break-even", bold=True),
                cell(fmt_pct(summary["universal_breakeven"])),
            ])
        if summary.get("balloon_breakeven") is not None:
            metrics_rows.append([
                cell("Vs. Annuity Break-even (Balloon)", bold=True),
                cell(fmt_pct(summary["balloon_breakeven"])),
            ])
        if summary.get("balloon_breakeven_abs") is not None:
            metrics_rows.append([
                cell("Absolute Break-even (Balloon)", bold=True),
                cell(fmt_pct(summary["balloon_breakeven_abs"])),
            ])
        # Inflation
        if summary.get("inflation_enabled") and summary.get("real_cost") is not None:
            metrics_rows.append([
                cell("Real Total Cost (PV)", bold=True),
                cell(fmt_money(summary["real_cost"], sym)),
            ])
            metrics_rows.append([
                cell("Inflation Discount", bold=True),
                cell(fmt_money(summary.get("inflation_savings", 0), sym)),
            ])
        # Risk metrics
        if summary.get("ltv") is not None:
            metrics_rows.append([
                cell("LTV (Loan-to-Value)", bold=True),
                cell(fmt_pct(summary["ltv"])),
            ])
        if summary.get("dscr") is not None:
            metrics_rows.append([
                cell("DSCR", bold=True),
                cell(f"{summary['dscr']:.2f}"),
            ])
        if summary.get("dti") is not None:
            metrics_rows.append([
                cell("DTI (Debt-to-Income)", bold=True),
                cell(fmt_pct(summary["dti"])),
            ])

    half_w = PAGE_W / 2 - 0.3*cm
    metrics_tbl = Table(metrics_rows, colWidths=[half_w, half_w])
    metrics_tbl.setStyle(TableStyle([
        ("FONTNAME",      (0,0),(-1,-1), F),
        ("FONTSIZE",      (0,0),(-1,-1), 10),
        ("ROWBACKGROUNDS",(0,0),(-1,-1), [LIGHTBLUE, LIGHTBLUE2]),
        ("GRID",          (0,0),(-1,-1), 0.4, BORDER_C),
        ("LEFTPADDING",   (0,0),(-1,-1), 8),
        ("RIGHTPADDING",  (0,0),(-1,-1), 8),
        ("TOPPADDING",    (0,0),(-1,-1), 5),
        ("BOTTOMPADDING", (0,0),(-1,-1), 5),
    ]))
    elems.append(metrics_tbl)
    elems.append(Spacer(1, 0.5*cm))

    # ── ТАБЛИЦА ПЛАТЕЖЕЙ ──────────────────────────────────────────────────────
    elems.append(Paragraph(t["section_schedule"], h2_s))
    elems.append(Spacer(1, 0.15*cm))

    cols     = list(df.columns)
    n_cols   = len(cols)
    total_key = (t.get("dep_total_row") if is_dep else t.get("total_row", "ИТОГО"))

    # Ширина колонок: первые 2 (Период, Дата) — фиксированные; остальные — поровну
    w_period = 1.4*cm
    w_date   = 2.4*cm
    w_num    = (PAGE_W - w_period - w_date) / max(n_cols - 2, 1)
    col_widths = [w_period, w_date] + [w_num] * (n_cols - 2)

    # Строка заголовков
    hdr_row = [hdr_cell(c) for c in cols]
    tbl_data = [hdr_row]

    for _, row in df.iterrows():
        vals     = list(row)
        is_total = str(vals[0]) == total_key
        row_cells = []
        for ci, v in enumerate(vals):
            if ci == 0:                          # Период
                row_cells.append(txt_cell(v, is_total))
            elif ci == 1:                        # Дата
                row_cells.append(txt_cell(v, is_total))
            else:                                # Числа
                is_red = isinstance(v, (int, float)) and v < 0
                row_cells.append(num_cell(v, is_total=is_total, is_red=is_red))
        tbl_data.append(row_cells)

    nr = len(tbl_data)
    pdf_t = Table(tbl_data, colWidths=col_widths, repeatRows=1)

    tbl_style = [
        # Заголовок
        ("BACKGROUND",    (0,0),  (-1,0),     MIDBLUE),
        ("FONTNAME",      (0,0),  (-1,0),     F_BOLD),
        ("FONTSIZE",      (0,0),  (-1,0),     tbl_fs),
        # Зебра (чётные/нечётные — без итоговой строки)
        ("ROWBACKGROUNDS",(0,1),  (-1, nr-2), [LIGHTBLUE2, LIGHTBLUE]),
        # Итоговая строка
        ("BACKGROUND",    (0,nr-1),(-1,nr-1), TOTAL_BG),
        ("FONTNAME",      (0,nr-1),(-1,nr-1), F_BOLD),
        # Сетка
        ("GRID",          (0,0),  (-1,-1),    0.3, BORDER_C),
        # Отступы
        ("LEFTPADDING",   (0,0),  (-1,-1),    4),
        ("RIGHTPADDING",  (0,0),  (-1,-1),    4),
        ("TOPPADDING",    (0,0),  (-1,-1),    3),
        ("BOTTOMPADDING", (0,0),  (-1,-1),    3),
        # Высота строк
        ("ROWHEIGHT",     (0,0),  (-1,-1),    tbl_fs + 6),
    ]
    pdf_t.setStyle(TableStyle(tbl_style))
    elems.append(pdf_t)
    elems.append(Spacer(1, 0.5*cm))

    # ── ДИСКЛЕЙМЕР ────────────────────────────────────────────────────────────
    elems.append(HRFlowable(width="100%", thickness=0.5,
                             color=BORDER_C, spaceBefore=4, spaceAfter=6))

    # Add Balloon Break-even to PDF if applicable (compact single line)
    be_rate_val = summary.get("balloon_breakeven")
    if be_rate_val is not None:
        be_s = make_style("pdf_be", size=9, bold=True,
                          color=rl_colors.HexColor("#D97706"), space_before=4)
        elems.append(Paragraph(
            f"📐 Inv. Break-even Rate (Balloon): {be_rate_val:.2f}%  "
            f"— Min. compound annual investment return to justify Balloon over Annuity.",
            be_s,
        ))
        elems.append(Spacer(1, 0.2*cm))

    disc_text = (
        "<b>Legal Disclaimer:</b> "
        "Results are for informational purposes only and do not constitute a public offer. "
        "Verify all terms with your financial institution. "
        "Yev Capital LoanLogic v3.0 / Bohdan Yevtushenko (MrCemper) — "
        "© 2026 Not liable for financial decisions based on these results."
    )
    elems.append(Paragraph(disc_text, disc_s))
    elems.append(Spacer(1, 0.2*cm))

    # ── ПОДПИСЬ ───────────────────────────────────────────────────────────────
    footer_text = (
        f"Yev Capital LoanLogic v3.0  ·  © 2026 Bohdan Yevtushenko (MrCemper)  ·  "
        f"{datetime.now().strftime('%d.%m.%Y %H:%M')}"
    )
    elems.append(Paragraph(footer_text, footer_s))

    doc.build(elems)
    return buf.getvalue()

# ─────────────────────────────────────────────────────────────────────────────
#  SESSION STATE
# ─────────────────────────────────────────────────────────────────────────────
def periods_from_dates(start: date, end: date, unit: str) -> int:
    """
    Считает количество ПОЛНЫХ периодов (unit) между start и end.
    Используется при альтернативном вводе срока через дату окончания.

    Raises:
      ValueError — если end <= start (нулевой или отрицательный диапазон),
        либо если интервал положителен, но короче одного полного периода
        (например, дата окончания всего на несколько дней позже старта при
        месячном unit). В обоих случаях явная ошибка лучше «фантомного»
        периода, который тихо округлил бы вверх до 1 и исказил срок.
    """
    from dateutil.relativedelta import relativedelta
    if end <= start:
        raise ValueError(
            f"End date ({end}) must be strictly later than start date ({start}). "
            f"Adjust the end date to a future point in time."
        )
    delta_map = {
        "weeks":     relativedelta(weeks=1),
        "months":    relativedelta(months=1),
        "quarters":  relativedelta(months=3),
        "halfyears": relativedelta(months=6),
        "years":     relativedelta(years=1),
    }
    if unit not in delta_map:
        raise ValueError(
            f"periods_from_dates: unknown unit {unit!r}. "
            f"Supported: {list(delta_map.keys())}.")
    delta = delta_map[unit]
    count, d = 0, start
    while d + delta <= end:
        d += delta
        count += 1
    if count < 1:
        raise ValueError(
            f"The interval from {start} to {end} is shorter than one '{unit}' "
            f"period, so no full period fits. Choose an end date at least one "
            f"{unit[:-1] if unit.endswith('s') else unit} later."
        )
    return count


def _state_defaults() -> dict:
    """Returns the canonical default values for all session-state fields.
    Single source of truth shared by init_state() (first-run seeding) and
    load_tpl() (resetting template fields a snapshot omits)."""
    from dateutil.relativedelta import relativedelta
    # Default end-date two years out, computed safely so a Feb-29 launch does
    # not raise (date(year+2, 2, 29) is invalid in a non-leap year). Using
    # relativedelta clamps Feb 29 → Feb 28 automatically.
    _default_end = date.today() + relativedelta(years=2)
    return {
        "lang": "en", "templates": {},
        "loan_amount": 100_000.0,
        "loan_term": 24, "term_unit": "months",
        "term_input_mode": "manual",
        "end_date": _default_end,
        "interest_rate": 6.5, "scheme": "annuity",
        "one_time_val": 0.0, "one_time_type": "pct",
        "monthly_val": 0.0, "monthly_type": "pct",
        "currency": "usd", "custom_symbol": "?",
        "calc_done": False, "schedule_df": None, "summary": {},
        "invest_sp500": False, "invest_custom": False, "invest_custom_rate": 8.0,
        "deposit_mode": "capitalize", "is_deposit": False,
        "start_date": date.today(),
        # ── Inflation Accounting ──────────────────────────────────────────────
        "inflation_enabled": False,
        "inflation_rate":    3.5,
        # ── Grace Period ──────────────────────────────────────────────────────
        "grace_enabled":   False,
        "grace_start":     1,
        "grace_duration":  3,
        "grace_type":      "interest_only",
        # ── Risk: LTV / DSCR / DTI ────────────────────────────────────────────
        "ltv_enabled":     False,
        "ltv_collateral":  150_000.0,
        "dscr_enabled":    False,
        "dscr_noi":        5_000.0,
        "dti_enabled":     False,
        "dti_income":      8_000.0,
        "dti_other_debts": 500.0,
        # ── Syndicated Loan (Multi-Tranche) ──────────────────────────────────
        "syndicated_enabled":  False,
        # Tranche A
        "synd_a_enabled": True,
        "synd_a_amount":  60_000.0, "synd_a_rate": 6.5,
        "synd_a_term":    24,       "synd_a_unit": "months",
        "synd_a_scheme":  "annuity",
        "synd_a_ot":      0.0,      "synd_a_mo":   0.0,
        "synd_a_offset":  0,
        # Tranche B
        "synd_b_enabled": True,
        "synd_b_amount":  30_000.0, "synd_b_rate": 7.5,
        "synd_b_term":    36,       "synd_b_unit": "months",
        "synd_b_scheme":  "annuity",
        "synd_b_ot":      0.0,      "synd_b_mo":   0.0,
        "synd_b_offset":  0,
        # Tranche C
        "synd_c_enabled": False,
        "synd_c_amount":  10_000.0, "synd_c_rate": 8.0,
        "synd_c_term":    12,       "synd_c_unit": "months",
        "synd_c_scheme":  "annuity",
        "synd_c_ot":      0.0,      "synd_c_mo":   0.0,
        "synd_c_offset":  0,
        # ── Audit Trail ───────────────────────────────────────────────────────
        "audit_log":           [],     # list of dicts with timestamp + change details
        "last_calc_snapshot":  None,   # snapshot of params on last Calculate press
        # ── Day-Count Convention ──────────────────────────────────────────────
        "day_count_enabled":   False,                # OFF by default per spec
        "day_count_method":    DAY_COUNT_DEFAULT,    # "30/360"
        # ── Theme (resolved lazily by get_active_theme() if not set) ─────────
        "theme":               None,                 # dict — filled by get_active_theme
        "theme_preset":        THEME_DEFAULT_KEY,
    }


def init_state():
    for k, v in _state_defaults().items():
        if k not in st.session_state:
            st.session_state[k] = v

# ─────────────────────────────────────────────────────────────────────────────
#  СИНДИЦИРОВАННЫЙ КРЕДИТ — Multi-Tranche Master Schedule
# ─────────────────────────────────────────────────────────────────────────────
def _term_to_months_local(term: int, unit: str) -> int:
    """
    Converts (term, unit) to integer months using ceil so we never lose a
    fractional period when normalizing. Does NOT clamp to a minimum: a zero
    or negative term is a genuine input error and must be caught by the
    caller (the syndicated loop raises a per-tranche error), not silently
    turned into a 1-month tranche.
    """
    return term_to_periods_in_base(term, unit, base_unit="months")


def calc_syndicated_master_schedule(
    tranches: list[dict],
    base_unit: str = "months",
    day_count: str | None = None,
    start_date: date | None = None,
) -> tuple[list[dict], dict, list[list[dict]]]:
    """
    Computes a consolidated (Master) schedule across multiple loan tranches
    that may be disbursed on DIFFERENT dates.

    Each tranche dict can optionally include:
      • "start_offset_months": int — offset from the global master start_date
                                      at which THIS tranche is disbursed.
                                      0 (default) = disbursed at master start.

    Tranches with start_offset > 0 contribute zero to master periods before
    their disbursement, and their amortization begins at master period
    `start_offset + 1`. The master's `n_periods` therefore equals
    max(tranche.start_offset + tranche.n) across all active tranches.
    """
    # Filter active tranches
    active = []
    for i, tr in enumerate(tranches):
        if tr.get("amount", 0) > 0:
            active.append((i, tr))

    if not active:
        return [], {"n_periods": 0, "total_principal": 0, "total_interest": 0,
                     "total_commission": 0, "total_payment": 0, "n_tranches_active": 0,
                     "tranche_errors": []}, []

    # Compute each tranche's schedule (in months)
    per_tranche = []
    tranche_errors = []
    master_end = 0  # max(start_offset + n_months)
    for tranche_id, tr in active:
        amt    = float(tr["amount"])
        rate   = float(tr["rate_pa"])
        n_orig = int(tr["term"])
        unit   = tr.get("unit", "months")
        scheme = tr.get("scheme", "annuity")
        n_months = _term_to_months_local(n_orig, unit)
        ot_val   = float(tr.get("ot_val", 0))
        mo_val   = float(tr.get("mo_val", 0))
        start_offset = max(0, int(tr.get("start_offset_months", 0)))

        # Explicitly reject a degenerate term as a surfaced per-tranche error
        # rather than letting it slip through as a silent 1-month tranche.
        if n_months < 1:
            tranche_errors.append((
                tranche_id,
                f"Tranche {chr(65 + tranche_id)}: term ({n_orig} {unit}) "
                f"is too short — must be at least one month.",
            ))
            continue

        # Compute the tranche's own start_date for day-count calculations
        tranche_start = start_date
        if start_date is not None and start_offset > 0:
            from dateutil.relativedelta import relativedelta
            tranche_start = start_date + relativedelta(months=start_offset)

        try:
            if scheme == "annuity":
                sched = calc_annuity(amt, n_months, rate, "months", mo_val,
                                      day_count=day_count, start_date=tranche_start)
            elif scheme == "classic":
                sched = calc_classic(amt, n_months, rate, "months", mo_val,
                                      day_count=day_count, start_date=tranche_start)
            elif scheme == "balloon":
                sched = calc_balloon(amt, n_months, rate, "months", mo_val,
                                      day_count=day_count, start_date=tranche_start)
            else:
                raise ValueError(
                    f"Unknown scheme {scheme!r} — expected one of "
                    f"'annuity', 'classic', 'balloon'.")
        except Exception as e:
            tranche_errors.append((
                tranche_id,
                f"Tranche {chr(65 + tranche_id)}: {type(e).__name__}: {e}",
            ))
            continue

        per_tranche.append({
            "id":            tranche_id,
            "letter":        chr(65 + tranche_id),
            "amount":        amt,
            "rate":          rate,
            "n":             n_months,
            "scheme":        scheme,
            "ot_comm":       ot_val,
            "schedule":      sched,
            "start_offset":  start_offset,
        })
        master_end = max(master_end, start_offset + n_months)

    if not per_tranche:
        return [], {"n_periods": 0, "total_principal": 0, "total_interest": 0,
                     "total_commission": 0, "total_payment": 0, "n_tranches_active": 0,
                     "tranche_errors": tranche_errors}, []

    # Build master schedule by summing across tranches with offset
    master = []
    for p in range(1, master_end + 1):
        bal_open = 0.0
        pay      = 0.0
        princ    = 0.0
        interest = 0.0
        comm     = 0.0
        bal_close = 0.0
        tranche_pmts = {}

        for tr in per_tranche:
            # Local period for this tranche: p - start_offset
            local_p = p - tr["start_offset"]
            if 1 <= local_p <= len(tr["schedule"]):
                row = tr["schedule"][local_p - 1]
                bal_open  += row["balance_open"]
                pay       += row["payment"]
                princ     += row["principal"]
                interest  += row["interest"]
                comm      += row["commission"]
                bal_close += row["balance_close"]
                tranche_pmts[tr["letter"]] = row["payment"]
            else:
                # Either tranche not yet disbursed (local_p < 1) or already matured
                tranche_pmts[tr["letter"]] = 0.0

        master.append({
            "period":          p,
            "balance_open":    round(bal_open, 2),
            "payment":         round(pay, 2),
            "principal":       round(princ, 2),
            "interest":        round(interest, 2),
            "commission":      round(comm, 2),
            "balance_close":   round(bal_close, 2),
            "tranche_payments": {k: round(v, 2) for k, v in tranche_pmts.items()},
        })

    max_n = master_end  # alias for backward compat

    # Compute totals
    total_principal = sum(tr["amount"] for tr in per_tranche)
    total_interest  = sum(r["interest"]   for r in master)
    total_comm_per  = sum(r["commission"] for r in master)
    total_ot_comm   = sum(tr["ot_comm"]   for tr in per_tranche)
    total_comm      = total_comm_per + total_ot_comm
    total_payment   = total_principal + total_interest + total_comm

    # First *non-zero* master payment — meaningful when tranches are staggered.
    # If all tranches have offset > 0, period 1 has no payment; downstream
    # risk metrics (DSCR, DTI on first payment) need the real first burden,
    # not a vacuous zero.
    first_pay_value = 0.0
    for row in master:
        if row["payment"] > 0:
            first_pay_value = row["payment"]
            break

    totals = {
        "n_periods":         max_n,
        "total_principal":   round(total_principal, 2),
        "total_interest":    round(total_interest, 2),
        "total_commission":  round(total_comm, 2),
        "total_payment":     round(total_payment, 2),
        "total_one_time_comm": round(total_ot_comm, 2),
        "n_tranches_active": len(per_tranche),
        "first_payment":     first_pay_value,
        "tranche_errors":    tranche_errors,
    }

    return master, totals, per_tranche


def calc_syndicated_blended_apr(per_tranche: list[dict],
                                  total_ot_comm: float = 0.0) -> float | None:
    """
    Blended APR across all tranches via IRR on the consolidated cash-flow.

    Each tranche may have a `start_offset` (in months) — the period at which
    its disbursement occurs and its schedule starts. The combined cash flow
    therefore places each tranche's negative disbursement (amount minus its
    own one-time commission) at its offset, and its scheduled payments at
    offset + i for i = 1..n_tranche.

    Cash flow convention (lender perspective, NPV = 0 form):
        CF_offset_k  = -(amount_k - ot_comm_k)       — net disbursement at offset
        CF_{offset_k + i} += payment_k_i             — periodic inflow

    NPV(r) = Σ_t CF_t / (1+r)^t = 0

    Robust solver:
      1. Newton's method (fast on well-shaped problems).
      2. Bracket on [-0.99, 5.0] per period, then bisection on sign change.
      3. Validate residual |NPV| < 1e-3 before returning.

    Returns:
      • float — annualised rate (% p.a.).
      • None  — if no valid root could be found within the bracket.
    """
    if not per_tranche:
        return None

    total_principal = sum(tr["amount"] for tr in per_tranche)
    if total_principal <= 0:
        return None

    # Build staggered cash-flow stream
    max_period = 0
    for tr in per_tranche:
        offset = int(tr.get("start_offset", 0))
        max_period = max(max_period, offset + int(tr["n"]))

    cfs = [0.0] * (max_period + 1)
    for tr in per_tranche:
        offset    = int(tr.get("start_offset", 0))
        amount    = float(tr["amount"])
        ot_k      = float(tr.get("ot_comm", 0.0))
        # Disbursement (negative outflow from lender) at this tranche's offset.
        # The one-time commission is netted out — it's a fee the borrower
        # effectively pays back at t=offset, so the actual disbursement amount
        # is (amount - ot_comm). This is consistent with the prior aggregate
        # convention but now placed correctly on the timeline.
        cfs[offset] += -(amount - ot_k)
        for i, row in enumerate(tr["schedule"], start=1):
            t = offset + i
            if t < len(cfs):
                cfs[t] += float(row["payment"])

    def npv(r: float) -> float:
        return sum(cfs[t] / (1.0 + r) ** t for t in range(len(cfs)))

    def dnpv(r: float) -> float:
        return sum(-t * cfs[t] / (1.0 + r) ** (t + 1)
                    for t in range(1, len(cfs)))

    # ── Phase 1: Newton with safe bounds ──────────────────────────────────────
    r_newton = 0.01
    converged = False
    for _ in range(200):
        f  = npv(r_newton)
        df = dnpv(r_newton)
        if abs(df) < 1e-12:
            break
        r_new = r_newton - f / df
        if r_new <= -0.99:
            r_new = -0.99
        if r_new > 5.0:
            r_new = 5.0
        if abs(r_new - r_newton) < 1e-10:
            r_newton = r_new
            converged = True
            break
        r_newton = r_new

    if converged and abs(npv(r_newton)) < 1e-3:
        if r_newton <= -1:
            return None
        return ((1.0 + r_newton) ** 12 - 1.0) * 100.0

    # ── Phase 2: bracket + bisection ──────────────────────────────────────────
    lo, hi = -0.99, 5.0
    f_lo, f_hi = npv(lo), npv(hi)
    if f_lo * f_hi > 0:
        return None  # no sign change → solver cannot locate root

    for _ in range(200):
        mid = (lo + hi) / 2
        f_mid = npv(mid)
        if abs(f_mid) < 1e-6:
            return ((1.0 + mid) ** 12 - 1.0) * 100.0
        if f_lo * f_mid < 0:
            hi, f_hi = mid, f_mid
        else:
            lo, f_lo = mid, f_mid
        if hi - lo < 1e-10:
            break

    r_final = (lo + hi) / 2
    if abs(npv(r_final)) > 1e-2:
        return None  # residual too large — declare failure
    if r_final <= -1:
        return None
    return ((1.0 + r_final) ** 12 - 1.0) * 100.0


# ─────────────────────────────────────────────────────────────────────────────
#  AUDIT TRAIL — логирование изменений ключевых параметров
# ─────────────────────────────────────────────────────────────────────────────
def _format_audit_value(field: str, value: float, sym: str = "$") -> str:
    """Форматирует значение поля для записи в аудит-лог."""
    if field == "amount":
        return f"{sym} {value:,.2f}".replace(",", "\u202f")
    elif field == "rate":
        return f"{value:.2f}%"
    elif field == "term":
        return f"{int(value)}"
    return str(value)


def record_audit_entry(t: dict, sym: str,
                        snapshot_old: dict | None,
                        snapshot_new: dict,
                        impact_old: dict | None = None,
                        impact_new: dict | None = None) -> None:
    """
    Сравнивает старый и новый snapshots ключевых параметров. Если есть
    изменения — добавляет запись в st.session_state.audit_log.

    snapshot keys: "amount", "rate", "term"
    impact  keys: "total_interest", "first_payment"

    Лог ограничен 50 последними записями (FIFO trim через list-slice),
    чтобы избежать неограниченного роста session_state в длинных сессиях.
    """
    AUDIT_LOG_MAX = 50

    if "audit_log" not in st.session_state or st.session_state.audit_log is None:
        st.session_state.audit_log = []

    # Первый расчёт — нет старого снимка, ничего не логируем
    if snapshot_old is None:
        return

    timestamp = datetime.now().strftime("%H:%M:%S")
    field_labels = {
        "amount": t.get("audit_field_amount", "Amount"),
        "rate":   t.get("audit_field_rate",   "Interest rate"),
        "term":   t.get("audit_field_term",   "Term"),
    }
    toggle_labels = {
        "grace":      t.get("audit_field_grace",      "Grace period"),
        "inflation":  t.get("audit_field_inflation",  "Inflation adj."),
        "day_count":  t.get("audit_field_day_count",  "Day-count"),
        "syndicated": t.get("audit_field_syndicated", "Syndicated mode"),
    }
    on_lbl  = t.get("audit_toggle_on",  "On")
    off_lbl = t.get("audit_toggle_off", "Off")

    changes = []
    for field in ("amount", "rate", "term"):
        old_v = snapshot_old.get(field)
        new_v = snapshot_new.get(field)
        if old_v is None or new_v is None:
            continue
        # Считаем "изменением" только реальное отличие
        if abs(float(old_v) - float(new_v)) > 1e-6:
            changes.append({
                "field":     field,
                "field_lbl": field_labels[field],
                "old_str":   _format_audit_value(field, old_v, sym),
                "new_str":   _format_audit_value(field, new_v, sym),
            })

    # Detect feature-toggle flips (On↔Off) so the audit trail also records
    # when the user enables/disables grace, inflation, day-count or syndicated.
    for field, lbl in toggle_labels.items():
        old_v = snapshot_old.get(field)
        new_v = snapshot_new.get(field)
        if old_v is None or new_v is None:
            continue
        if bool(old_v) != bool(new_v):
            changes.append({
                "field":     field,
                "field_lbl": lbl,
                "old_str":   on_lbl if old_v else off_lbl,
                "new_str":   on_lbl if new_v else off_lbl,
            })

    # Detect changes to other result-affecting string/value settings (scheme,
    # currency, term-input mode, day-count method, start date, commissions).
    # Their snapshot values are already display-ready strings, so we show them
    # verbatim. Missing keys (older snapshots) are skipped, never false-flagged.
    string_labels = {
        "scheme":           t.get("audit_field_scheme",     "Scheme"),
        "currency":         t.get("audit_field_currency",   "Currency"),
        "term_input_mode":  t.get("audit_field_term_mode",  "Term input mode"),
        "day_count_method": t.get("audit_field_dc_method",  "Day-count method"),
        "start_date":       t.get("audit_field_start_date", "Start date"),
        "one_time":         t.get("audit_field_one_time",   "One-time fee"),
        "monthly":          t.get("audit_field_monthly",    "Periodic fee"),
    }
    for field, lbl in string_labels.items():
        old_v = snapshot_old.get(field)
        new_v = snapshot_new.get(field)
        if old_v is None or new_v is None:
            continue
        if str(old_v) != str(new_v):
            changes.append({
                "field":     field,
                "field_lbl": lbl,
                "old_str":   str(old_v),
                "new_str":   str(new_v),
            })

    if not changes:
        return

    # Считаем impact на total_interest и first_payment (если данные доступны)
    impact_str = None
    if impact_old and impact_new:
        impact_parts = []
        old_int = impact_old.get("total_interest")
        new_int = impact_new.get("total_interest")
        if old_int is not None and new_int is not None and abs(new_int - old_int) > 1e-6:
            delta = new_int - old_int
            sign  = "+" if delta >= 0 else "-"
            delta_fmt = f"{sym} {abs(delta):,.2f}".replace(",", "\u202f")
            impact_parts.append(
                t.get("audit_impact_interest", "Total Interest impact: {delta}")
                 .format(delta=f"{sign}{delta_fmt}"))

        old_fp = impact_old.get("first_payment")
        new_fp = impact_new.get("first_payment")
        if old_fp is not None and new_fp is not None and abs(new_fp - old_fp) > 1e-6:
            d_fp = new_fp - old_fp
            sign_fp = "+" if d_fp >= 0 else "-"
            fp_fmt = f"{sym} {abs(d_fp):,.2f}".replace(",", "\u202f")
            impact_parts.append(
                t.get("audit_impact_first_payment", "First payment impact: {delta}")
                 .format(delta=f"{sign_fp}{fp_fmt}"))

        if impact_parts:
            impact_str = "  •  ".join(impact_parts)

    new_entry = {
        "timestamp": timestamp,
        "changes":   changes,
        "impact":    impact_str,
    }
    log = st.session_state.audit_log
    log.append(new_entry)
    # FIFO trim to keep memory and rerun-payload bounded
    if len(log) > AUDIT_LOG_MAX:
        st.session_state.audit_log = log[-AUDIT_LOG_MAX:]


# Canonical set of session-state fields captured by a template. Defined once
# so save_tpl and load_tpl agree, and so load_tpl can reset any field that a
# (possibly older / partial) template omits back to its default — preventing
# stale values from a previous configuration surviving a template load.
TEMPLATE_FIELDS = [
    "loan_amount", "loan_term", "term_unit", "interest_rate", "scheme",
    "one_time_val", "one_time_type", "monthly_val", "monthly_type",
    "currency", "custom_symbol", "deposit_mode", "is_deposit", "start_date",
    "term_input_mode", "end_date",
    # Grace period
    "grace_enabled", "grace_start", "grace_duration", "grace_type",
    # Inflation
    "inflation_enabled", "inflation_rate",
    # Risk metrics
    "ltv_enabled", "ltv_collateral",
    "dscr_enabled", "dscr_noi",
    "dti_enabled", "dti_income", "dti_other_debts",
    # Day-count convention
    "day_count_enabled", "day_count_method",
    # Investment-comparison inputs
    "invest_sp500", "invest_custom", "invest_custom_rate",
    # Syndicated mode (toggle + all three tranche configs)
    "syndicated_enabled",
    "synd_a_enabled", "synd_a_amount", "synd_a_rate", "synd_a_term",
    "synd_a_unit", "synd_a_scheme", "synd_a_ot", "synd_a_mo", "synd_a_offset",
    "synd_b_enabled", "synd_b_amount", "synd_b_rate", "synd_b_term",
    "synd_b_unit", "synd_b_scheme", "synd_b_ot", "synd_b_mo", "synd_b_offset",
    "synd_c_enabled", "synd_c_amount", "synd_c_rate", "synd_c_term",
    "synd_c_unit", "synd_c_scheme", "synd_c_ot", "synd_c_mo", "synd_c_offset",
]


def save_tpl(name):
    # Note: app-wide UI chrome (lang, theme, theme_preset) is intentionally
    # NOT part of a loan template — loading a saved loan should not silently
    # switch the interface language or color theme.
    # Only persist keys that actually exist in state (robust to future renames).
    snapshot = {k: st.session_state[k] for k in TEMPLATE_FIELDS if k in st.session_state}
    snapshot["saved_at"] = datetime.now().strftime("%d.%m.%Y %H:%M")
    st.session_state.templates[name] = snapshot

def load_tpl(name):
    snap = st.session_state.templates[name]
    # Reset every known template field that the snapshot does NOT contain back
    # to its init_state default first, so a partial/older template can't leave
    # stale values from the current configuration "behind" it after loading.
    defaults = _state_defaults()
    for k in TEMPLATE_FIELDS:
        if k not in snap and k in defaults:
            st.session_state[k] = defaults[k]
    # Then apply the template's own values.
    for k, v in snap.items():
        if k != "saved_at":
            st.session_state[k] = v

def del_tpl(name):
    st.session_state.templates.pop(name, None)

# ─────────────────────────────────────────────────────────────────────────────
#  ТЕМЫ — пресеты и кастомизация
# ─────────────────────────────────────────────────────────────────────────────
#  Каждая тема — словарь полей, которые подставляются в CSS-переменные.
#  Изменяешь пресет → весь UI меняется на лету (без st.rerun()).
#  Кастомные значения хранятся в browser localStorage (через JS-мост).
# ─────────────────────────────────────────────────────────────────────────────

THEME_PRESETS = {
    "dark_navy": {
        "name":            "Dark Navy",
        "bg":              "#0F172A",
        "bg_secondary":    "#1E293B",
        "bg_tertiary":     "#0F2444",
        "text":            "#E2E8F0",
        "text_muted":      "#94A3B8",
        "text_subtle":     "#64748B",
        "accent":          "#4FC3F7",
        "accent_strong":   "#1D4ED8",
        "accent_gradient": "linear-gradient(135deg,#1D4ED8,#0EA5E9)",
        "success":         "#10B981",
        "warning":         "#F59E0B",
        "danger":          "#DC2626",
        "border":          "#334155",
        "border_subtle":   "#1E2D3D",
        "input_text":      "#F1F5F9",
        "font_size":       1.0,     # multiplier
        "density":         1.0,     # multiplier for padding/spacing
        "radius":          8,       # px base border-radius
    },
    "high_contrast": {
        "name":            "High Contrast",
        "bg":              "#000000",
        "bg_secondary":    "#0A0A0A",
        "bg_tertiary":     "#141414",
        "text":            "#FFFFFF",
        "text_muted":      "#D4D4D4",
        "text_subtle":     "#A3A3A3",
        "accent":          "#FFD700",
        "accent_strong":   "#FFA500",
        "accent_gradient": "linear-gradient(135deg,#FFA500,#FFD700)",
        "success":         "#00FF7F",
        "warning":         "#FFD700",
        "danger":          "#FF4444",
        "border":          "#FFFFFF",
        "border_subtle":   "#404040",
        "input_text":      "#FFFFFF",
        "font_size":       1.1,
        "density":         1.1,
        "radius":          4,
    },
    "midnight": {
        "name":            "Midnight",
        "bg":              "#1A0B2E",
        "bg_secondary":    "#2D1B4E",
        "bg_tertiary":     "#1E0F3D",
        "text":            "#E9D5FF",
        "text_muted":      "#A78BFA",
        "text_subtle":     "#7C3AED",
        "accent":          "#C084FC",
        "accent_strong":   "#9333EA",
        "accent_gradient": "linear-gradient(135deg,#9333EA,#C084FC)",
        "success":         "#A78BFA",
        "warning":         "#F59E0B",
        "danger":          "#F472B6",
        "border":          "#5B21B6",
        "border_subtle":   "#3B0764",
        "input_text":      "#F5F3FF",
        "font_size":       1.0,
        "density":         1.0,
        "radius":          12,
    },
    "ocean": {
        "name":            "Ocean",
        "bg":              "#0C1929",
        "bg_secondary":    "#0E2235",
        "bg_tertiary":     "#0A3553",
        "text":            "#CFFAFE",
        "text_muted":      "#67E8F9",
        "text_subtle":     "#22D3EE",
        "accent":          "#06B6D4",
        "accent_strong":   "#0891B2",
        "accent_gradient": "linear-gradient(135deg,#0891B2,#06B6D4)",
        "success":         "#34D399",
        "warning":         "#FBBF24",
        "danger":          "#F87171",
        "border":          "#155E75",
        "border_subtle":   "#0E4C6B",
        "input_text":      "#ECFEFF",
        "font_size":       1.0,
        "density":         1.0,
        "radius":          10,
    },
}

THEME_DEFAULT_KEY = "dark_navy"


def build_css(theme: dict) -> str:
    """
    Builds the full CSS payload from a theme dict. All visible colors,
    paddings, radii and font sizes are driven by CSS-variables, which lets
    us swap themes (and live-edit via the sidebar) without touching the
    rest of the markup.

    Theme dict keys (see THEME_PRESETS for the full list of fields).
    Missing keys fall back to dark_navy values. Field values are also
    *defensively* validated/sanitized: a corrupt entry in localStorage
    must not be able to either crash build_css() or inject arbitrary CSS.
    """
    base = THEME_PRESETS[THEME_DEFAULT_KEY]
    t = {**base, **(theme or {})}

    # ── Safe numeric coercion ─────────────────────────────────────────────────
    # Coerce defensively: a non-numeric value (e.g. from a corrupt browser
    # storage entry) falls back to the default rather than crashing CSS build.
    def _safe_num(val, default: float) -> float:
        try:
            return float(val)
        except (TypeError, ValueError):
            return float(default)
    def _safe_int(val, default: int) -> int:
        try:
            return int(float(val))    # accept "8.0" strings
        except (TypeError, ValueError):
            return int(default)

    fs       = max(0.75, min(1.4, _safe_num(t.get("font_size"), 1.0)))
    density  = max(0.7,  min(1.4, _safe_num(t.get("density"),   1.0)))
    radius   = max(0,    min(24,  _safe_int(t.get("radius"),    8)))
    pad_md   = round(10 * density)
    pad_lg   = round(16 * density)
    pad_xs   = round(6  * density)

    # ── Safe color sanitization (CSS-injection defense) ───────────────────────
    # Color fields are interpolated raw into `:root { --app-bg: <value>; }`,
    # so an unsanitized string like `"#000; } body { display: none; } /*"`
    # could inject arbitrary CSS. Only well-formed hex colors (3/4/6/8 hex
    # digits) and the keyword "transparent" are accepted; anything else
    # falls back to the default-preset color for that field.
    import re as _re
    _HEX_RE = _re.compile(r'^#(?:[0-9a-fA-F]{3,4}|[0-9a-fA-F]{6}|[0-9a-fA-F]{8})$')

    def _safe_color(field: str) -> str:
        val = t.get(field, base[field])
        if isinstance(val, str):
            v = val.strip()
            if _HEX_RE.match(v) or v.lower() == "transparent":
                return v
        return base[field]

    # Gradient field can legitimately contain commas, spaces, and the word
    # "linear-gradient(...)" — sanitize separately by checking for known-bad
    # characters that would close the CSS declaration.
    def _safe_gradient() -> str:
        val = t.get("accent_gradient", base["accent_gradient"])
        if isinstance(val, str) and not any(c in val for c in ('{', '}', ';', '/*', '*/')):
            return val
        return base["accent_gradient"]

    bg              = _safe_color("bg")
    bg_secondary    = _safe_color("bg_secondary")
    bg_tertiary     = _safe_color("bg_tertiary")
    text            = _safe_color("text")
    text_muted      = _safe_color("text_muted")
    text_subtle     = _safe_color("text_subtle")
    accent          = _safe_color("accent")
    accent_strong   = _safe_color("accent_strong")
    accent_gradient = _safe_gradient()
    success         = _safe_color("success")
    warning         = _safe_color("warning")
    danger          = _safe_color("danger")
    border          = _safe_color("border")
    border_subtle   = _safe_color("border_subtle")
    input_text      = _safe_color("input_text")

    return f"""
<style>
:root{{
  --app-bg:           {bg};
  --app-bg-secondary: {bg_secondary};
  --app-bg-tertiary:  {bg_tertiary};
  --app-text:         {text};
  --app-text-muted:   {text_muted};
  --app-text-subtle:  {text_subtle};
  --app-accent:       {accent};
  --app-accent-strong:{accent_strong};
  --app-accent-grad:  {accent_gradient};
  --app-success:      {success};
  --app-warning:      {warning};
  --app-danger:       {danger};
  --app-border:       {border};
  --app-border-subtle:{border_subtle};
  --app-input-text:   {input_text};
  --app-radius:       {radius}px;
  --app-radius-lg:    {radius + 4}px;
  --app-pad-md:       {pad_md}px;
  --app-pad-lg:       {pad_lg}px;
  --app-pad-xs:       {pad_xs}px;
  --app-font-scale:   {fs};
}}
html{{font-size:calc(16px * var(--app-font-scale))}}
.stApp{{background:var(--app-bg);color:var(--app-text)}}
section[data-testid="stSidebar"]{{
  background:var(--app-bg-secondary)!important;
  border-right:1px solid var(--app-border)
}}
section[data-testid="stSidebar"] *{{color:var(--app-text)!important}}
div[data-testid="metric-container"]{{
  background:linear-gradient(135deg,var(--app-bg-secondary),var(--app-bg-tertiary));
  border:1px solid var(--app-border);
  border-radius:var(--app-radius-lg);
  padding:var(--app-pad-lg) calc(var(--app-pad-lg) + 2px);
  box-shadow:0 4px 20px rgba(0,0,0,.35)
}}
div[data-testid="metric-container"] label{{
  color:var(--app-text-muted)!important;
  font-size:.74rem!important;font-weight:600;
  text-transform:uppercase;letter-spacing:.06em
}}
div[data-testid="metric-container"] div[data-testid="metric-value"]{{
  color:var(--app-accent)!important;font-size:1.2rem!important;font-weight:700
}}
/* ── Шапка приложения ── */
.app-header{{
  background:linear-gradient(135deg,
    var(--app-bg-tertiary) 0%,
    var(--app-bg-secondary) 50%,
    var(--app-bg) 100%);
  border:1px solid var(--app-accent-strong);
  border-radius:var(--app-radius-lg);
  padding:24px 32px;margin-bottom:20px;text-align:center;
  box-shadow:0 8px 32px rgba(14,165,233,0.15)
}}
.app-header .brand{{
  font-size:1.9rem;font-weight:900;letter-spacing:.04em;margin:0;
  background:linear-gradient(90deg,var(--app-accent),var(--app-accent-strong));
  -webkit-background-clip:text;-webkit-text-fill-color:transparent;
  background-clip:text
}}
.app-header .slogan{{
  color:var(--app-text-subtle);font-size:.82rem;font-style:italic;
  letter-spacing:.12em;margin:6px 0 2px;font-weight:500
}}
.app-header p{{color:var(--app-text-muted);font-size:.86rem;margin:2px 0 0}}
/* ── Сайдбар брендинг ── */
.sidebar-brand{{
  text-align:center;padding:12px 8px 4px;
  border-bottom:1px solid var(--app-border);margin-bottom:8px
}}
.sidebar-brand .brand-name{{
  font-size:1rem;font-weight:800;
  background:linear-gradient(90deg,var(--app-accent),var(--app-accent-strong));
  -webkit-background-clip:text;-webkit-text-fill-color:transparent;
  background-clip:text;letter-spacing:.04em
}}
.sidebar-brand .brand-slogan{{
  font-size:.64rem;color:var(--app-text-subtle)!important;font-style:italic;
  letter-spacing:.08em;margin-top:2px
}}
/* ── Дисклеймер ── */
.disclaimer{{
  background:var(--app-bg);border:1px solid var(--app-border-subtle);
  border-radius:var(--app-radius);
  padding:var(--app-pad-md) calc(var(--app-pad-md) + 2px);
  margin-top:16px;font-size:.65rem;line-height:1.55;
  color:var(--app-text-subtle)!important
}}
.disclaimer b{{color:var(--app-text-muted)!important;font-size:.67rem}}
/* ── Кнопки ── */
.stButton>button{{
  background:var(--app-accent-grad);color:#fff;
  border:none;border-radius:var(--app-radius);
  font-weight:700;font-size:.88rem;
  padding:var(--app-pad-md) calc(var(--app-pad-lg) + 2px);
  transition:all .2s;box-shadow:0 4px 12px rgba(14,165,233,.3)
}}
.stButton>button:hover{{
  transform:translateY(-1px);box-shadow:0 6px 18px rgba(14,165,233,.5)
}}
.sec-title{{
  font-size:.95rem;font-weight:700;color:var(--app-accent);
  padding:5px 0 3px;border-bottom:1px solid var(--app-border);
  margin-bottom:12px;letter-spacing:.03em
}}
.scheme-info{{
  background:var(--app-bg-tertiary);border-left:3px solid var(--app-accent);
  border-radius:6px;padding:var(--app-pad-xs) calc(var(--app-pad-xs) + 5px);
  color:var(--app-text-muted);font-size:.8rem;margin-top:3px
}}
.invest-header{{
  background:linear-gradient(135deg,var(--app-bg-tertiary),var(--app-bg-secondary));
  border:1px solid var(--app-success);border-radius:var(--app-radius);
  padding:var(--app-pad-md) var(--app-pad-lg);margin:8px 0
}}
.invest-header p{{color:var(--app-success);font-weight:600;margin:0;font-size:.9rem}}
.savings-box{{
  background:linear-gradient(135deg,var(--app-bg-tertiary),var(--app-bg-secondary));
  border:1px solid var(--app-success);border-radius:var(--app-radius);
  padding:var(--app-pad-md) calc(var(--app-pad-lg) + 2px);
  color:var(--app-success);font-weight:600;
  text-align:center;margin-top:8px
}}
.deposit-box{{
  background:var(--app-bg-tertiary);border:1px solid var(--app-accent);
  border-radius:var(--app-radius);
  padding:var(--app-pad-md) var(--app-pad-lg);margin:6px 0
}}
.tpl-badge{{
  display:inline-block;background:var(--app-accent-strong);color:#fff;
  padding:2px 9px;border-radius:20px;font-size:.72rem;font-weight:600
}}
/* ── Подпись "Made by Yevtush" ── */
.made-by{{
  position:fixed;bottom:10px;right:14px;
  color:var(--app-border);font-size:.68rem;font-weight:700;
  letter-spacing:.05em;z-index:9999;
  pointer-events:none;user-select:none;
  background:var(--app-bg-secondary);
  padding:3px 10px;border-radius:20px;
  border:1px solid var(--app-border-subtle)
}}
/* Хедер: оставляем видимым (там кнопка сайдбара), но прозрачным */
footer{{visibility:hidden}}
#MainMenu{{visibility:hidden}}
header[data-testid="stHeader"]{{
  background:transparent!important;
  height:auto!important;
  min-height:2.5rem!important;
}}
/* Кнопка раскрытия сайдбара — всегда видна */
button[data-testid="stSidebarCollapsedControl"],
button[data-testid="collapsedControl"],
[data-testid="stSidebarCollapsedControl"],
[data-testid="collapsedControl"]{{
  visibility:visible!important;
  opacity:1!important;
  display:flex!important;
  background:var(--app-accent-strong)!important;
  border:1px solid var(--app-accent)!important;
  border-radius:var(--app-radius)!important;
  box-shadow:0 2px 8px rgba(0,0,0,.4)!important;
  z-index:99999!important;
}}
button[data-testid="stSidebarCollapsedControl"] svg,
button[data-testid="collapsedControl"] svg,
[data-testid="stSidebarCollapsedControl"] svg,
[data-testid="collapsedControl"] svg{{
  color:#FFFFFF!important;
  fill:#FFFFFF!important;
}}
/* ── Контрастный текст в полях ввода (mobile-fix + theme support) ── */
.stTextInput input,
.stNumberInput input,
.stTextArea textarea,
.stDateInput input,
.stTimeInput input,
div[data-baseweb="input"] input,
div[data-baseweb="select"] div,
div[data-baseweb="select"] span,
input[type="text"],
input[type="number"],
input[type="email"],
input[type="tel"],
input[type="search"],
textarea,
select{{
  color:var(--app-input-text)!important;
  -webkit-text-fill-color:var(--app-input-text)!important;
  opacity:1!important;
  font-weight:500!important;
  caret-color:var(--app-accent)!important;
}}
.stTextInput input::placeholder,
.stNumberInput input::placeholder,
.stTextArea textarea::placeholder,
input::placeholder,
textarea::placeholder{{
  color:var(--app-text-subtle)!important;
  -webkit-text-fill-color:var(--app-text-subtle)!important;
  opacity:.85!important;
}}
.stTextInput input:disabled,
.stNumberInput input:disabled,
input:disabled,
textarea:disabled,
select:disabled{{
  color:var(--app-text-muted)!important;
  -webkit-text-fill-color:var(--app-text-muted)!important;
  opacity:.7!important;
}}
div[data-baseweb="select"] > div{{color:var(--app-input-text)!important}}
div[data-baseweb="select"] [class*="placeholder"]{{color:var(--app-text-subtle)!important}}
</style>
"""


def _theme_persistence_js() -> str:
    """
    Returns an HTML snippet (rendered via streamlit.components.v1.html) that
    bridges browser localStorage to Streamlit's query_params.

    On first render: reads `yev_theme_v1` from localStorage and, if present,
    appends `?theme_b64=<base64-json>` to the URL, then triggers a one-time
    reload so Python can pick the value up via st.query_params. A sessionStorage
    guard prevents the reload from looping.

    Saving is handled separately by `_theme_save_js`, which writes the JSON
    back to localStorage whenever the theme changes.
    """
    return """
<script>
(function(){
  const KEY = 'yev_theme_v1';
  // On first load: if localStorage has saved theme AND query param is missing,
  // append ?theme_b64=... to the URL so Streamlit picks it up on the next run.
  try {
    const params = new URLSearchParams(window.parent.location.search);
    if (!params.has('theme_b64')) {
      const saved = window.parent.localStorage.getItem(KEY);
      if (saved) {
        const b64 = window.parent.btoa(unescape(encodeURIComponent(saved)));
        params.set('theme_b64', b64);
        const newUrl = window.parent.location.pathname + '?' + params.toString();
        window.parent.history.replaceState(null, '', newUrl);
        // Trigger a one-time reload so Streamlit re-reads query params
        if (!window.parent.sessionStorage.getItem('yev_theme_loaded')) {
          window.parent.sessionStorage.setItem('yev_theme_loaded', '1');
          window.parent.location.reload();
        }
      }
    }
  } catch(e) { /* localStorage may be blocked; degrade silently */ }
})();
</script>
"""


def _theme_save_js(theme_json: str) -> str:
    """JS snippet that saves a theme JSON into localStorage."""
    # Escape JSON for safe inlining in JS string literal
    import json as _json
    safe = _json.dumps(theme_json)  # double-encode for safe JS embedding
    return f"""
<script>
(function(){{
  try {{
    window.parent.localStorage.setItem('yev_theme_v1', {safe});
  }} catch(e) {{ /* blocked; theme will persist only this session */ }}
}})();
</script>
"""


def get_active_theme() -> dict:
    """
    Resolves the active theme:
      1. If `theme` is in st.session_state → use it.
      2. Else if `?theme_b64=...` query param exists (set by localStorage
         bridge) → decode and use, also write to session_state.
      3. Else → default preset.
    """
    if "theme" in st.session_state and isinstance(st.session_state.theme, dict):
        return st.session_state.theme

    # Try to recover from localStorage via query param
    try:
        qp = st.query_params
        b64 = qp.get("theme_b64")
        if b64:
            import base64, json as _json
            decoded = base64.b64decode(b64).decode("utf-8")
            theme = _json.loads(decoded)
            if isinstance(theme, dict):
                # Merge with default to fill missing fields
                merged = {**THEME_PRESETS[THEME_DEFAULT_KEY], **theme}
                st.session_state.theme = merged
                return merged
    except Exception:
        pass

    # Default
    default = dict(THEME_PRESETS[THEME_DEFAULT_KEY])
    st.session_state.theme = default
    return default


def render_glossary(t: dict):
    """
    Renders a language-aware financial glossary in the sidebar. Uses
    st.popover when available (Streamlit ≥ 1.32), otherwise falls back to
    st.expander so it still works on older Streamlit builds.

    The term list is built from TRANSLATIONS keys (gloss_*) so it follows the
    selected language; each entry is "**Term** — definition".
    """
    terms = [
        ("gloss_annuity_term",  "gloss_annuity_def"),
        ("gloss_bullet_term",   "gloss_bullet_def"),
        ("gloss_classic_term",  "gloss_classic_def"),
        ("gloss_dscr_term",     "gloss_dscr_def"),
        ("gloss_ltv_term",      "gloss_ltv_def"),
        ("gloss_dti_term",      "gloss_dti_def"),
        ("gloss_npv_term",      "gloss_npv_def"),
    ]

    def _body():
        for term_key, def_key in terms:
            term = t.get(term_key, term_key)
            definition = t.get(def_key, "")
            st.markdown(f"**{term}** — {definition}")

    title = t.get("glossary_title", "📚 Glossary")
    # st.popover was added in newer Streamlit; degrade gracefully if absent.
    popover = getattr(st, "popover", None)
    if callable(popover):
        with popover(title, use_container_width=True):
            _body()
    else:
        with st.expander(title, expanded=False):
            _body()


def render_theme_editor(t: dict):
    """Renders the theme editor in the sidebar."""
    import json as _json
    import streamlit.components.v1 as _components

    current_theme = get_active_theme()

    with st.expander(t.get("theme_section", "🎨 Theme"), expanded=False):
        # ── 1) Preset selector ────────────────────────────────────────────────
        preset_labels = {
            k: v["name"] for k, v in THEME_PRESETS.items()
        }
        preset_labels["__custom__"] = t.get("theme_custom", "Custom (edit below)")

        current_preset_key = st.session_state.get("theme_preset", THEME_DEFAULT_KEY)
        new_preset_key = st.selectbox(
            t.get("theme_preset_label", "Preset"),
            options=list(preset_labels.keys()),
            format_func=lambda k: preset_labels[k],
            index=list(preset_labels.keys()).index(current_preset_key)
                   if current_preset_key in preset_labels else 0,
            key="theme_preset_select",
            help=t.get("theme_preset_help", ""),
        )

        if new_preset_key != current_preset_key:
            st.session_state.theme_preset = new_preset_key
            if new_preset_key in THEME_PRESETS:
                st.session_state.theme = dict(THEME_PRESETS[new_preset_key])
            # __custom__: keep current colors, just allow editing
            _save_theme_to_localstorage(st.session_state.theme)
            st.rerun()

        # ── 2) Advanced editor (only if Custom) ───────────────────────────────
        if new_preset_key == "__custom__":
            st.caption(t.get("theme_custom_hint",
                              "Tweak any color. Changes apply instantly."))

            new_theme = dict(current_theme)

            color_fields = [
                ("bg",            "theme_bg",        "App Background"),
                ("bg_secondary",  "theme_bg_sec",    "Sidebar / Cards"),
                ("bg_tertiary",   "theme_bg_ter",    "Tertiary Background"),
                ("text",          "theme_text",      "Main Text"),
                ("text_muted",    "theme_text_m",    "Muted Text"),
                ("text_subtle",   "theme_text_s",    "Subtle Text"),
                ("accent",        "theme_accent",    "Accent / Links"),
                ("accent_strong", "theme_accent_s",  "Accent (Strong)"),
                ("success",       "theme_success",   "Success"),
                ("warning",       "theme_warning",   "Warning"),
                ("danger",        "theme_danger",    "Danger"),
                ("border",        "theme_border",    "Borders"),
                ("input_text",    "theme_input_txt", "Input Text"),
            ]

            for field, key, label in color_fields:
                new_theme[field] = st.color_picker(
                    t.get(f"theme_field_{field}", label),
                    value=current_theme.get(field, "#000000"),
                    key=key,
                )

            # Re-derive accent_gradient automatically when accent colors change
            new_theme["accent_gradient"] = (
                f"linear-gradient(135deg,{new_theme['accent_strong']},"
                f"{new_theme['accent']})"
            )

            st.markdown("---")

            new_theme["font_size"] = st.slider(
                t.get("theme_font_size", "Font Size"),
                min_value=0.75, max_value=1.4,
                value=float(current_theme.get("font_size", 1.0)),
                step=0.05, key="theme_font_size_slider",
                help=t.get("theme_font_size_help", "Scales all text proportionally"),
            )
            new_theme["density"] = st.slider(
                t.get("theme_density", "UI Density"),
                min_value=0.7, max_value=1.4,
                value=float(current_theme.get("density", 1.0)),
                step=0.05, key="theme_density_slider",
                help=t.get("theme_density_help", "Padding around UI elements"),
            )
            new_theme["radius"] = st.slider(
                t.get("theme_radius", "Border Radius"),
                min_value=0, max_value=24,
                value=int(current_theme.get("radius", 8)),
                step=1, key="theme_radius_slider",
                help=t.get("theme_radius_help", "Roundness of corners (px)"),
            )

            if new_theme != current_theme:
                st.session_state.theme = new_theme
                _save_theme_to_localstorage(new_theme)
                st.rerun()

        # ── 3) Reset & info ───────────────────────────────────────────────────
        if st.button(t.get("theme_reset", "↻ Reset to Default"),
                       key="btn_theme_reset",
                       use_container_width=True):
            st.session_state.theme = dict(THEME_PRESETS[THEME_DEFAULT_KEY])
            st.session_state.theme_preset = THEME_DEFAULT_KEY
            _save_theme_to_localstorage(st.session_state.theme)
            st.rerun()


def _save_theme_to_localstorage(theme: dict):
    """Pushes theme JSON to browser localStorage via an invisible JS snippet."""
    try:
        import json as _json
        import streamlit.components.v1 as _components
        js = _theme_save_js(_json.dumps(theme))
        _components.html(js, height=0)
    except Exception:
        # Component-rendering may fail in unusual harnesses; silently degrade.
        pass

# ─────────────────────────────────────────────────────────────────────────────
#  ГЛАВНАЯ ФУНКЦИЯ
# ─────────────────────────────────────────────────────────────────────────────
def main():
    st.set_page_config(page_title="Yev Capital LoanLogic", page_icon="💼",
                       layout="wide", initial_sidebar_state="expanded")
    init_state()

    # ── Theme: read from session_state / localStorage / default ───────────────
    active_theme = get_active_theme()
    st.markdown(build_css(active_theme), unsafe_allow_html=True)

    # ── localStorage bridge: on first load, recover saved theme via query param.
    # Rendered as zero-height iframe; runs JS that copies localStorage →
    # ?theme_b64=... and reloads once (sessionStorage guards re-reload loop).
    try:
        import streamlit.components.v1 as _components
        _components.html(_theme_persistence_js(), height=0)
    except Exception:
        pass

    # Постоянная подпись
    st.markdown('<div class="made-by">© 2026 Bohdan Yevtushenko · Yev Capital v3.0</div>', unsafe_allow_html=True)

    # ── Предупреждение о шрифте PDF (если кириллический шрифт не найден) ─────
    if PDF_FONT_WARN:
        st.warning(PDF_FONT_WARN)

    # ── Выбор языка ───────────────────────────────────────────────────────────
    lang_map = {"🇷🇺 Русский": "ru", "🇺🇦 Українська": "uk", "🇺🇸 English": "en"}
    with st.sidebar:
        # ── Брендинг в боковой панели ─────────────────────────────────────────
        st.markdown(
            '<div class="sidebar-brand">'
            '<div class="brand-name">Yev Capital</div>'
            '<div class="brand-slogan">Precision in every percent. Logic in every loan.</div>'
            '</div>',
            unsafe_allow_html=True,
        )

        lang_lbl = st.selectbox("🌐 Язык / Мова / Language",
                                 options=list(lang_map.keys()),
                                 index=list(lang_map.values()).index(st.session_state.lang))
        if lang_map[lang_lbl] != st.session_state.lang:
            st.session_state.lang = lang_map[lang_lbl]
            st.rerun()

        # ── Financial glossary (language-aware) ───────────────────────────────
        render_glossary(TRANSLATIONS[st.session_state.lang])

    t = TRANSLATIONS[st.session_state.lang]

    # ═══════════════════════════════════════════════════════════════════════════
    #  БОКОВАЯ ПАНЕЛЬ
    # ═══════════════════════════════════════════════════════════════════════════
    with st.sidebar:
        # ── Theme editor (presets + custom colors/sizes) ──────────────────────
        render_theme_editor(t)

        st.markdown(f"### {t['section_params']}")

        # Валюта
        cur_opts = {
            t["uah"]: "uah", t["usd"]: "usd", t["eur"]: "eur",
            t["rub"]: "rub", t["gbp"]: "gbp", t["jpy"]: "jpy",
            t["cad"]: "cad", t["aud"]: "aud", t["chf"]: "chf",
            t["cny"]: "cny", t["hkd"]: "hkd", t["custom"]: "custom",
        }
        # Safe index — fallback to usd if stored key not in current lang dict
        cur_keys = list(cur_opts.values())
        cur_idx  = cur_keys.index(st.session_state.currency) if st.session_state.currency in cur_keys else cur_keys.index("usd")
        cur_lbl  = st.selectbox(t["currency"], list(cur_opts.keys()), index=cur_idx)
        st.session_state.currency = cur_opts[cur_lbl]
        if st.session_state.currency == "custom":
            st.session_state.custom_symbol = st.text_input(
                t["custom_symbol"], value=st.session_state.custom_symbol, max_chars=6)
        sym = get_sym(st.session_state)

        st.divider()

        # ── Синдицированный кредит (Multi-Tranche) ────────────────────────────
        st.session_state.syndicated_enabled = st.checkbox(
            t.get("syndicated_toggle", "Syndicated Loan Mode"),
            value=st.session_state.syndicated_enabled,
            help=t.get("syndicated_help", ""),
            key="cb_syndicated",
        )

        if st.session_state.syndicated_enabled:
            st.markdown(f"**{t.get('syndicated_section', 'Syndicated Loan')}**")
            st.caption(t.get("syndicated_master_help", ""))

            # Available scheme options for tranches (no deposit)
            tranche_scheme_opts = {
                t["annuity"]:  "annuity",
                t["classic"]:  "classic",
                t["balloon"]:  "balloon",
            }
            unit_opts_local = {t["weeks"]:"weeks", t["months"]:"months",
                                t["quarters"]:"quarters", t["halfyears"]:"halfyears",
                                t["years"]:"years"}

            for letter, key_prefix in [("A", "synd_a"), ("B", "synd_b"), ("C", "synd_c")]:
                tranche_label = t.get("syndicated_tranche", "Tranche {n}").format(n=letter)
                with st.expander(tranche_label, expanded=(letter == "A")):
                    # Enable/Disable tranche
                    st.session_state[f"{key_prefix}_enabled"] = st.checkbox(
                        t.get("syndicated_tranche_enabled",
                               "Enable {tranche}").format(tranche=tranche_label),
                        value=st.session_state[f"{key_prefix}_enabled"],
                        key=f"cb_{key_prefix}_en",
                    )

                    if st.session_state[f"{key_prefix}_enabled"]:
                        # Amount + rate row
                        c1, c2 = st.columns([2, 1])
                        with c1:
                            st.session_state[f"{key_prefix}_amount"] = st.number_input(
                                t.get("syndicated_tranche_amount",
                                       "Amount ({tranche})").format(tranche=tranche_label),
                                min_value=0.0, max_value=1e14,
                                value=float(st.session_state[f"{key_prefix}_amount"]),
                                step=1_000.0, format="%.2f",
                                key=f"ni_{key_prefix}_amt",
                            )
                        with c2:
                            st.session_state[f"{key_prefix}_rate"] = st.number_input(
                                t.get("syndicated_tranche_rate",
                                       "Rate %").format(tranche=tranche_label),
                                min_value=0.0, max_value=999.9,
                                value=float(st.session_state[f"{key_prefix}_rate"]),
                                step=0.1, format="%.2f",
                                key=f"ni_{key_prefix}_rate",
                            )
                        # Term + unit row
                        c3, c4 = st.columns([1, 2])
                        with c3:
                            st.session_state[f"{key_prefix}_term"] = st.number_input(
                                t.get("syndicated_tranche_term",
                                       "Term").format(tranche=tranche_label),
                                min_value=1, max_value=1200,
                                value=int(st.session_state[f"{key_prefix}_term"]),
                                step=1, key=f"ni_{key_prefix}_term",
                            )
                        with c4:
                            cur_u = st.session_state[f"{key_prefix}_unit"]
                            cur_u_idx = (list(unit_opts_local.values()).index(cur_u)
                                          if cur_u in unit_opts_local.values() else 1)
                            u_lbl_t = st.selectbox(
                                t.get("syndicated_tranche_unit",
                                       "Unit").format(tranche=tranche_label),
                                options=list(unit_opts_local.keys()),
                                index=cur_u_idx,
                                key=f"sb_{key_prefix}_unit",
                            )
                            st.session_state[f"{key_prefix}_unit"] = unit_opts_local[u_lbl_t]
                        # Scheme
                        cur_sc = st.session_state[f"{key_prefix}_scheme"]
                        cur_sc_idx = (list(tranche_scheme_opts.values()).index(cur_sc)
                                       if cur_sc in tranche_scheme_opts.values() else 0)
                        sc_lbl_t = st.selectbox(
                            t.get("syndicated_tranche_scheme",
                                   "Scheme").format(tranche=tranche_label),
                            options=list(tranche_scheme_opts.keys()),
                            index=cur_sc_idx,
                            key=f"sb_{key_prefix}_scheme",
                        )
                        st.session_state[f"{key_prefix}_scheme"] = tranche_scheme_opts[sc_lbl_t]
                        # Commissions row
                        c5, c6 = st.columns(2)
                        with c5:
                            st.session_state[f"{key_prefix}_ot"] = st.number_input(
                                t.get("syndicated_tranche_ot_comm",
                                       "One-time").format(tranche=tranche_label),
                                min_value=0.0, max_value=1e10,
                                value=float(st.session_state[f"{key_prefix}_ot"]),
                                step=10.0, format="%.2f",
                                key=f"ni_{key_prefix}_ot",
                            )
                        with c6:
                            st.session_state[f"{key_prefix}_mo"] = st.number_input(
                                t.get("syndicated_tranche_mo_comm",
                                       "Periodic").format(tranche=tranche_label),
                                min_value=0.0, max_value=1e10,
                                value=float(st.session_state[f"{key_prefix}_mo"]),
                                step=1.0, format="%.2f",
                                key=f"ni_{key_prefix}_mo",
                            )
                        # Disbursement-offset row — для разных дат выдачи траншей
                        st.session_state[f"{key_prefix}_offset"] = st.number_input(
                            t.get("syndicated_tranche_offset",
                                   "Offset (months)").format(tranche=tranche_label),
                            min_value=0, max_value=240,
                            value=int(st.session_state.get(f"{key_prefix}_offset", 0)),
                            step=1, key=f"ni_{key_prefix}_offset",
                            help=t.get("syndicated_tranche_offset_help", ""),
                        )

            # Build tranches list (only enabled with positive amount)
            tranches_input = []
            zero_amount_tranches = []
            for letter, key_prefix in [("A", "synd_a"), ("B", "synd_b"), ("C", "synd_c")]:
                if not st.session_state[f"{key_prefix}_enabled"]:
                    continue
                if st.session_state[f"{key_prefix}_amount"] > 0:
                    tranches_input.append({
                        "amount":   st.session_state[f"{key_prefix}_amount"],
                        "rate_pa":  st.session_state[f"{key_prefix}_rate"],
                        "term":     st.session_state[f"{key_prefix}_term"],
                        "unit":     st.session_state[f"{key_prefix}_unit"],
                        "scheme":   st.session_state[f"{key_prefix}_scheme"],
                        "ot_val":   st.session_state[f"{key_prefix}_ot"],
                        "mo_val":   st.session_state[f"{key_prefix}_mo"],
                        "start_offset_months":
                                    int(st.session_state.get(f"{key_prefix}_offset", 0)),
                    })
                else:
                    # Enabled but amount is zero — flag it rather than silently
                    # dropping, so the user knows the tranche isn't counted.
                    zero_amount_tranches.append(letter)

            if zero_amount_tranches:
                st.warning(
                    t.get("syndicated_zero_amount_warn",
                           "Tranche(s) {letters} are enabled but have a zero "
                           "amount — they are not included. Set an amount or "
                           "disable them.").format(letters=", ".join(zero_amount_tranches)))

            if not tranches_input:
                st.error(t.get("syndicated_zero_error",
                                 "Enter at least one tranche."))
            else:
                # Compute totals (preview)
                _master, _totals, _per_tr = calc_syndicated_master_schedule(
                    tranches_input, base_unit="months"
                )

                # Сообщаем пользователю про сбойные транши явно
                tranche_errs = _totals.get("tranche_errors", [])
                if tranche_errs:
                    for _tid, _msg in tranche_errs:
                        st.error(_msg)

                if _totals["n_tranches_active"] > 0:
                    blended = calc_syndicated_blended_apr(
                        _per_tr, _totals.get("total_one_time_comm", 0))

                    # Sync loan amount (used by synd run_calculation). The
                    # blended APR is rendered locally below; it must NOT be
                    # written into st.session_state.interest_rate, because
                    # that field belongs to the single-loan input. Leaking it
                    # there would (a) shove a stale value into the next
                    # non-syndicated calc, and (b) potentially mislead the
                    # user about what rate they typed.
                    st.session_state.loan_amount = _totals["total_principal"]
                    # Store the blended APR in a separate, synd-scoped key so
                    # downstream renderers can pick it up without colliding.
                    st.session_state["_synd_blended_apr"] = blended

                    # Save tranches for later access in run_calculation/render
                    st.session_state["_syndicated_tranches"] = tranches_input

                    total_fmt = "{:,.2f}".format(
                        _totals["total_principal"]).replace(",", "\u202f")

                    blended_str = (f"{blended:.4f}%"
                                    if blended is not None else "N/A (IRR failed)")
                    units_label = t.get("months", "months")

                    st.markdown(
                        f"<div style='background:#0F2444;border:1px solid #1D4ED8;"
                        f"border-radius:8px;padding:8px 12px;margin-top:6px;'>"
                        f"<div style='color:#94A3B8;font-size:.78rem;text-transform:uppercase'>"
                        f"{t.get('syndicated_total', 'Total Syndicate')}: "
                        f"<b style='color:#4FC3F7'>{sym} {total_fmt}</b></div>"
                        f"<div style='color:#94A3B8;font-size:.78rem;text-transform:uppercase;margin-top:4px'>"
                        f"{t.get('syndicated_blended_rate', 'Blended APR')}: "
                        f"<b style='color:#FFD166'>{blended_str}</b></div>"
                        f"<div style='color:#64748B;font-size:.70rem;margin-top:3px'>"
                        f"{_totals['n_tranches_active']} active · "
                        f"{_totals['n_periods']} {units_label} (max horizon)</div>"
                        f"</div>",
                        unsafe_allow_html=True,
                    )
            st.divider()

        # Сумма (отображается только если syndicated mode выключен)
        if not st.session_state.syndicated_enabled:
            st.markdown(f"**{t['loan_amount']}**")
            amt_inp = st.number_input(
                t["loan_amount"], min_value=100.0, max_value=100_000_000_000_000.0,
                value=float(st.session_state.loan_amount), step=1_000.0,
                format="%.2f", label_visibility="collapsed")
            amt_sldr = st.slider(
                t["loan_amount_slider"], min_value=100, max_value=100_000_000,
                value=min(int(st.session_state.loan_amount), 100_000_000),
                step=10_000, label_visibility="collapsed")
            if abs(amt_inp - st.session_state.loan_amount) > 0.5:
                st.session_state.loan_amount = amt_inp
            elif amt_sldr != int(st.session_state.loan_amount):
                st.session_state.loan_amount = float(amt_sldr)
            display_amt = "{:,.2f}".format(st.session_state.loan_amount).replace(",", "\u202f")
            st.markdown(f"<div style='color:#4FC3F7;font-weight:700;font-size:1.05rem;'>"
                        f"{sym} {display_amt}</div>", unsafe_allow_html=True)

            st.divider()
        # else: amount and rate are set by syndicated_enabled block above

        # ── Срок кредита / вклада ─────────────────────────────────────────
        # Выбор единицы срока (нужна в обоих режимах)
        unit_opts = {t["weeks"]:"weeks", t["months"]:"months",
                      t["quarters"]:"quarters", t["halfyears"]:"halfyears",
                      t["years"]:"years"}
        u_lbl = st.selectbox(
            t["term_unit"],
            options=list(unit_opts.keys()),
            index=list(unit_opts.values()).index(st.session_state.term_unit),
        )
        st.session_state.term_unit = unit_opts[u_lbl]

        # Переключатель режима ввода срока
        mode_manual  = t.get("term_mode_manual",  "Manual (periods)")
        mode_enddate = t.get("term_mode_enddate",  "By End Date")
        term_mode = st.radio(
            "",
            options=[mode_manual, mode_enddate],
            index=0 if st.session_state.term_input_mode == "manual" else 1,
            horizontal=True,
            key="term_mode_radio",
            label_visibility="collapsed",
        )
        st.session_state.term_input_mode = (
            "manual" if term_mode == mode_manual else "enddate"
        )

        if st.session_state.term_input_mode == "manual":
            # Стандартный режим — числовое поле
            st.session_state.loan_term = st.number_input(
                t["loan_term"], min_value=1, max_value=1200,
                value=st.session_state.loan_term, step=1,
            )
        else:
            # Режим «дата окончания»
            # min_value на день позже start_date чтобы заведомо исключить end<=start
            min_end = st.session_state.start_date + timedelta(days=1)
            end_dt = st.date_input(
                t.get("end_date_label", "End Date"),
                value=max(st.session_state.end_date, min_end),
                min_value=min_end,
                max_value=date(2099, 12, 31),
                help=t.get("end_date_hint", ""),
                key="end_date_picker",
            )
            st.session_state.end_date = end_dt
            try:
                computed_n = periods_from_dates(
                    st.session_state.start_date,
                    end_dt,
                    st.session_state.term_unit,
                )
                st.session_state.loan_term = computed_n
                st.caption(
                    f"→ {computed_n} {st.session_state.term_unit}"
                )
            except ValueError as e:
                # The interval is too short to contain even one full period
                # (end<=start is already excluded by min_value above). Rather
                # than silently computing on a possibly-mismatched old term,
                # tell the user explicitly which term remains in effect so the
                # schedule below and the UI stay in sync conceptually.
                st.error(str(e))
                computed_n = st.session_state.loan_term
                st.caption(
                    t.get("end_date_fallback",
                           "↳ Keeping the previous term: {n} {unit}.").format(
                        n=computed_n, unit=st.session_state.term_unit))

        mo = term_to_months(st.session_state.loan_term, st.session_state.term_unit)
        cap = t.get("term_caption", "mo. / yrs")
        st.caption(
            f"≈ {mo:.1f} {cap.split('/')[0].strip()} "
            f"/ {mo/12:.2f} {cap.split('/')[-1].strip()}"
        )

        st.divider()

        # Ставка (скрыта в синдицированном режиме — там берётся WACC).
        # min_value allows a small negative band: negative-interest products
        # are exotic but real (some EU mortgages went sub-zero), and the
        # calculation core handles r < 0 correctly (constant annuity payment,
        # APR, schedules all verified). The previous 0.0 floor was an
        # interface-only restriction inconsistent with the math.
        if not st.session_state.syndicated_enabled:
            st.session_state.interest_rate = st.number_input(
                t["interest_rate"], min_value=-20.0, max_value=999.9,
                value=float(st.session_state.interest_rate), step=0.5, format="%.2f",
                help=t.get("help_negative_rate", ""))

        # Дата начала кредита / вклада
        st.session_state.start_date = st.date_input(
            t["start_date_label"],
            value=st.session_state.start_date,
            min_value=date(2000, 1, 1),
            max_value=date(2099, 12, 31),
            help=t["start_date_hint"],
        )

        # Схема
        scheme_opts = {t["annuity"]:"annuity", t["classic"]:"classic",
                       t["balloon"]:"balloon", t["deposit_scheme"]:"deposit"}
        cur_scheme_lbl = {v: k for k, v in scheme_opts.items()}.get(
            st.session_state.scheme, t["annuity"])
        sc_lbl = st.radio(t["calc_scheme"], list(scheme_opts.keys()),
                           index=list(scheme_opts.keys()).index(cur_scheme_lbl))
        st.session_state.scheme = scheme_opts[sc_lbl]

        # Пояснение схемы
        info_map = {"annuity":"scheme_annuity_info","classic":"scheme_classic_info",
                    "balloon":"scheme_balloon_info","deposit":"scheme_deposit_info"}
        st.markdown(f"<div class='scheme-info'>{t[info_map[st.session_state.scheme]]}</div>",
                    unsafe_allow_html=True)

        # Режим вклада (только если вклад)
        if st.session_state.scheme == "deposit":
            st.markdown("<div class='deposit-box'>", unsafe_allow_html=True)
            dep_opts = {t["deposit_capitalize"]:"capitalize", t["deposit_payout"]:"payout"}
            dep_lbl  = st.radio(t["deposit_mode_label"], list(dep_opts.keys()),
                                 index=list(dep_opts.values()).index(st.session_state.deposit_mode),
                                 key="dep_radio")
            st.session_state.deposit_mode = dep_opts[dep_lbl]
            st.session_state.is_deposit   = True
            st.markdown("</div>", unsafe_allow_html=True)
        else:
            st.session_state.is_deposit = False

        st.divider()

        # Комиссии (скрываем для вклада)
        if st.session_state.scheme != "deposit":
            st.markdown(f"**{t['section_commissions']}**")
            comm_type_opts = [t["pct_of_amount"], t["fixed_amount"]]

            ot_t_lbl = st.selectbox(t["one_time_type"], comm_type_opts,
                                     index=0 if st.session_state.one_time_type=="pct" else 1)
            st.session_state.one_time_type = "pct" if ot_t_lbl == t["pct_of_amount"] else "fixed"
            st.session_state.one_time_val  = st.number_input(
                t["one_time_comm"], min_value=0.0,
                max_value=100.0 if st.session_state.one_time_type=="pct" else 5e8,
                value=float(st.session_state.one_time_val),
                step=0.1 if st.session_state.one_time_type=="pct" else 500.0, format="%.2f")

            mo_t_lbl = st.selectbox(t["monthly_type"], comm_type_opts,
                                     index=0 if st.session_state.monthly_type=="pct" else 1)
            st.session_state.monthly_type = "pct" if mo_t_lbl == t["pct_of_amount"] else "fixed"
            st.session_state.monthly_val  = st.number_input(
                t["monthly_comm"], min_value=0.0,
                max_value=50.0 if st.session_state.monthly_type=="pct" else 5e7,
                value=float(st.session_state.monthly_val),
                step=0.05 if st.session_state.monthly_type=="pct" else 100.0, format="%.2f")
            st.divider()

        # ══════════════════════════════════════════════════════════════════════
        #  РАСШИРЕННЫЕ ОПЦИИ (st.expander блоки 3-7)
        # ══════════════════════════════════════════════════════════════════════
        is_loan_scheme = (st.session_state.scheme != "deposit")

        # ── Day-Count Convention (только для кредитов) ────────────────────────
        if is_loan_scheme:
            with st.expander(t.get("day_count_section",
                                     "📅 Day-Count Convention"),
                              expanded=False):
                st.session_state.day_count_enabled = st.checkbox(
                    t.get("day_count_toggle", "Use day-count convention"),
                    value=st.session_state.day_count_enabled,
                    help=t.get("day_count_toggle_help", ""),
                    key="cb_day_count",
                )
                if st.session_state.day_count_enabled:
                    cur = st.session_state.day_count_method
                    cur_idx = (DAY_COUNT_METHODS.index(cur)
                                if cur in DAY_COUNT_METHODS else 0)
                    st.session_state.day_count_method = st.selectbox(
                        t.get("day_count_method", "Day-count method"),
                        options=DAY_COUNT_METHODS,
                        index=cur_idx,
                        help=t.get("day_count_method_help", ""),
                        key="sb_day_count_method",
                    )
                    # Краткое описание выбранного метода
                    desc_key = {
                        "30/360":  "day_count_desc_30_360",
                        "30E/360": "day_count_desc_30E_360",
                        "ACT/360": "day_count_desc_ACT_360",
                        "ACT/365": "day_count_desc_ACT_365",
                        "ACT/ACT": "day_count_desc_ACT_ACT",
                    }.get(st.session_state.day_count_method)
                    if desc_key:
                        st.caption(t.get(desc_key, ""))
                    st.caption(
                        t.get("day_count_caption",
                               "Formula: I = P × r × (N / B). Current: **{method}**.")
                        .format(method=st.session_state.day_count_method)
                    )

        # ── Кредитные каникулы (только для кредитов) ──────────────────────────
        if is_loan_scheme:
            with st.expander(t.get("grace_expander", "Grace Period"), expanded=False):
                st.session_state.grace_enabled = st.checkbox(
                    t.get("grace_toggle", "Enable Grace Period"),
                    value=st.session_state.grace_enabled,
                    help=t.get("help_grace_toggle", ""),
                    key="cb_grace")
                if st.session_state.grace_enabled:
                    max_term = max(1, int(st.session_state.loan_term))
                    st.session_state.grace_start = st.number_input(
                        t.get("grace_start", "Start Period"),
                        min_value=1, max_value=max_term,
                        value=min(int(st.session_state.grace_start), max_term),
                        step=1, key="ni_grace_start",
                        help=t.get("help_grace_start", ""))
                    max_dur = max(1, max_term - st.session_state.grace_start + 1)
                    st.session_state.grace_duration = st.number_input(
                        t.get("grace_duration", "Duration (periods)"),
                        min_value=1, max_value=max_dur,
                        value=min(int(st.session_state.grace_duration), max_dur),
                        step=1, key="ni_grace_dur",
                        help=t.get("help_grace_duration", ""))
                    grace_type_opts = {
                        t.get("grace_interest_only", "Interest Only"): "interest_only",
                        t.get("grace_full_holiday",  "Full Holiday"):   "full_holiday",
                    }
                    cur_idx = list(grace_type_opts.values()).index(
                        st.session_state.grace_type) if st.session_state.grace_type in grace_type_opts.values() else 0
                    gt_lbl = st.radio(
                        t.get("grace_type", "Type"),
                        options=list(grace_type_opts.keys()),
                        index=cur_idx, horizontal=True, key="rd_grace_type",
                        help=t.get("help_grace_type", ""))
                    st.session_state.grace_type = grace_type_opts[gt_lbl]
                    if st.session_state.grace_type == "full_holiday":
                        st.info(t.get("info_full_holiday", ""), icon="⚠️")

        # ── Учёт инфляции ─────────────────────────────────────────────────────
        with st.expander(t.get("inflation_expander", "Inflation Accounting"),
                          expanded=False):
            st.session_state.inflation_enabled = st.checkbox(
                t.get("inflation_toggle", "Enable Inflation Adjustment"),
                value=st.session_state.inflation_enabled,
                help=t.get("help_inflation_toggle", ""),
                key="cb_inflation")
            if st.session_state.inflation_enabled:
                st.session_state.inflation_rate = st.number_input(
                    t.get("inflation_rate", "Expected Annual Inflation (%)"),
                    min_value=-50.0, max_value=200.0,
                    value=float(st.session_state.inflation_rate),
                    step=0.5, format="%.1f", key="ni_inflation",
                    help=t.get("help_inflation_rate", ""))
                st.caption(t.get("caption_inflation", ""))

        # ── DSCR (Business Analysis) ──────────────────────────────────────────
        if is_loan_scheme:
            with st.expander(t.get("dscr_expander", "Business Analysis (DSCR)"),
                              expanded=False):
                st.session_state.dscr_enabled = st.checkbox(
                    t.get("dscr_toggle", "Enable DSCR Analysis"),
                    value=st.session_state.dscr_enabled,
                    help=t.get("help_dscr_toggle", ""),
                    key="cb_dscr")
                if st.session_state.dscr_enabled:
                    st.session_state.dscr_noi = st.number_input(
                        t.get("dscr_noi", "Monthly NOI"),
                        min_value=0.0, max_value=1e10,
                        value=float(st.session_state.dscr_noi),
                        step=100.0, format="%.2f", key="ni_dscr",
                        help=t.get("help_dscr_noi", ""))
                    st.caption(t.get("caption_dscr", ""))

        # ── LTV (Loan-to-Value) ───────────────────────────────────────────────
        if is_loan_scheme:
            with st.expander(t.get("ltv_expander", "Collateral & LTV"),
                              expanded=False):
                st.session_state.ltv_enabled = st.checkbox(
                    t.get("ltv_toggle", "Enable LTV Analysis"),
                    value=st.session_state.ltv_enabled,
                    help=t.get("help_ltv_toggle", ""),
                    key="cb_ltv")
                if st.session_state.ltv_enabled:
                    st.session_state.ltv_collateral = st.number_input(
                        t.get("ltv_collateral", "Collateral / Asset Value"),
                        min_value=1.0, max_value=1e14,
                        value=float(st.session_state.ltv_collateral),
                        step=1_000.0, format="%.2f", key="ni_ltv",
                        help=t.get("help_ltv_collateral", ""))
                    st.caption(t.get("caption_ltv", ""))

        # ── DTI (Personal Solvency) ───────────────────────────────────────────
        if is_loan_scheme:
            with st.expander(t.get("dti_expander", "Personal Solvency (DTI)"),
                              expanded=False):
                st.session_state.dti_enabled = st.checkbox(
                    t.get("dti_toggle", "Enable DTI Analysis"),
                    value=st.session_state.dti_enabled,
                    help=t.get("help_dti_toggle", ""),
                    key="cb_dti")
                if st.session_state.dti_enabled:
                    st.session_state.dti_income = st.number_input(
                        t.get("dti_income", "Total Monthly Income"),
                        min_value=1.0, max_value=1e10,
                        value=float(st.session_state.dti_income),
                        step=100.0, format="%.2f", key="ni_dti_inc",
                        help=t.get("help_dti_income", ""))
                    st.session_state.dti_other_debts = st.number_input(
                        t.get("dti_other_debts", "Other Monthly Debt Payments"),
                        min_value=0.0, max_value=1e10,
                        value=float(st.session_state.dti_other_debts),
                        step=50.0, format="%.2f", key="ni_dti_deb",
                        help=t.get("help_dti_other_debts", ""))
                    st.caption(t.get("caption_dti", ""))

        # Кнопка расчёта
        if st.button(t["calc_btn"], use_container_width=True):
            # ── Снимок параметров ДО расчёта (для audit-trail) ───────────────
            new_snapshot = {
                "amount": float(st.session_state.loan_amount),
                "rate":   float(st.session_state.interest_rate),
                "term":   int(st.session_state.loan_term),
                # Feature toggles — so the audit trail also reflects when the
                # user switches major modes, not just the three numeric inputs.
                "grace":      bool(st.session_state.grace_enabled),
                "inflation":  bool(st.session_state.inflation_enabled),
                "day_count":  bool(st.session_state.day_count_enabled),
                "syndicated": bool(st.session_state.syndicated_enabled),
                # Other configuration that genuinely changes the result, so a
                # recalculation after editing any of these is recorded too.
                "scheme":          str(st.session_state.scheme),
                "currency":        str(st.session_state.currency),
                "term_input_mode": str(st.session_state.term_input_mode),
                "day_count_method": str(st.session_state.day_count_method),
                "start_date":      str(st.session_state.start_date),
                "one_time":        f"{float(st.session_state.one_time_val):.4f}"
                                    f"/{st.session_state.one_time_type}",
                "monthly":         f"{float(st.session_state.monthly_val):.4f}"
                                    f"/{st.session_state.monthly_type}",
            }
            old_snapshot = st.session_state.get("last_calc_snapshot")
            old_summary  = st.session_state.get("summary", {}) or None

            # Build syndicated tranches list if mode active
            synd_tranches = None
            if st.session_state.syndicated_enabled:
                synd_tranches = []
                for kp in ["synd_a", "synd_b", "synd_c"]:
                    if (st.session_state[f"{kp}_enabled"] and
                        st.session_state[f"{kp}_amount"] > 0):
                        synd_tranches.append({
                            "amount":   st.session_state[f"{kp}_amount"],
                            "rate_pa":  st.session_state[f"{kp}_rate"],
                            "term":     st.session_state[f"{kp}_term"],
                            "unit":     st.session_state[f"{kp}_unit"],
                            "scheme":   st.session_state[f"{kp}_scheme"],
                            "ot_val":   st.session_state[f"{kp}_ot"],
                            "mo_val":   st.session_state[f"{kp}_mo"],
                            "start_offset_months":
                                int(st.session_state.get(f"{kp}_offset", 0)),
                        })
                if not synd_tranches:
                    synd_tranches = None  # fallback

            try:
                df_d, smry = run_calculation(
                    st.session_state.loan_amount, st.session_state.loan_term,
                    st.session_state.interest_rate, st.session_state.term_unit,
                    st.session_state.scheme,
                    st.session_state.one_time_val, st.session_state.one_time_type,
                    st.session_state.monthly_val, st.session_state.monthly_type,
                    st.session_state.currency, st.session_state.custom_symbol, t,
                    deposit_mode=st.session_state.deposit_mode,
                    start_date=st.session_state.start_date,
                    grace_enabled=st.session_state.grace_enabled,
                    grace_start=st.session_state.grace_start,
                    grace_duration=st.session_state.grace_duration,
                    grace_type=st.session_state.grace_type,
                    inflation_enabled=st.session_state.inflation_enabled,
                    inflation_rate=st.session_state.inflation_rate,
                    ltv_enabled=st.session_state.ltv_enabled,
                    ltv_collateral=st.session_state.ltv_collateral,
                    dscr_enabled=st.session_state.dscr_enabled,
                    dscr_noi=st.session_state.dscr_noi,
                    dti_enabled=st.session_state.dti_enabled,
                    dti_income=st.session_state.dti_income,
                    dti_other_debts=st.session_state.dti_other_debts,
                    syndicated_tranches=synd_tranches,
                    day_count_enabled=st.session_state.day_count_enabled,
                    day_count_method=st.session_state.day_count_method,
                )
            except (ValueError, ZeroDivisionError, OverflowError) as calc_err:
                # Surface bad inputs (e.g. degenerate term, invalid rate) as a
                # friendly message instead of crashing the whole app with a
                # red Streamlit traceback.
                st.error(t.get("calc_error", "⚠️ Calculation could not be "
                                "completed with the given inputs.")
                          + f"\n\n`{calc_err}`")
                st.stop()

            st.session_state.schedule_df = df_d
            st.session_state.summary     = smry
            st.session_state.calc_done   = True

            # ── Запись в журнал изменений ────────────────────────────────────
            sym_for_audit = get_sym(st.session_state)
            record_audit_entry(
                t, sym_for_audit,
                snapshot_old=old_snapshot,
                snapshot_new=new_snapshot,
                impact_old=old_summary,
                impact_new=smry,
            )
            st.session_state.last_calc_snapshot = new_snapshot

        # ── Disclaimer (always English, AGPL-3.0) ───────────────────────────
        st.markdown(
            """
            <div class="disclaimer">
              <b>⚠️ Disclaimer</b><br><br>
              This application provides preliminary calculations for
              <b>informational purposes only</b> and does not constitute
              financial, legal, or investment advice. Always verify final
              loan or deposit terms directly with your financial institution.<br><br>
              The software is provided under the
              <b>GNU Affero General Public License v3.0 or later</b>
              and is distributed <b>WITHOUT ANY WARRANTY</b>; without even
              the implied warranty of MERCHANTABILITY or FITNESS FOR A
              PARTICULAR PURPOSE.<br><br>
              Copyright (c) 2026 Bohdan Yevtushenko (MrCemper)<br>
              <a href="https://github.com/Mr4Cemper/Yev-Capital-LoanLogic"
                 style="color:#93C5FD;text-decoration:none" target="_blank">
                https://github.com/Mr4Cemper/Yev-Capital-LoanLogic
              </a>
            </div>
            """,
            unsafe_allow_html=True,
        )

        # ── Source Code button ───────────────────────────────────────────────
        st.markdown(
            """
            <div style="margin-top:8px;text-align:center">
              <a href="https://github.com/Mr4Cemper/Yev-Capital-LoanLogic"
                 target="_blank"
                 style="display:inline-flex;align-items:center;gap:6px;
                        padding:5px 14px;
                        background:#1E293B;color:#93C5FD;
                        border:1px solid #334155;border-radius:6px;
                        font-size:.72rem;font-weight:600;
                        text-decoration:none;letter-spacing:.04em">
                &#128279; Source Code (GitHub)
              </a>
            </div>
            """,
            unsafe_allow_html=True,
        )

    # ═══════════════════════════════════════════════════════════════════════════
    #  ШАПКА
    # ═══════════════════════════════════════════════════════════════════════════
    st.markdown(
        '<div class="app-header">'
        '<div class="brand">Yev Capital LoanLogic</div>'
        '<div class="slogan">Precision in every percent. Logic in every loan.</div>'
        f'<p>{t["app_subtitle"]}</p>'
        '</div>',
        unsafe_allow_html=True,
    )

    # ─── Шаблоны ──────────────────────────────────────────────────────────────
    with st.expander(t["section_templates"], expanded=False):
        _render_templates(t)

    # ─── Нет результатов — приветствие ────────────────────────────────────────
    if not (st.session_state.calc_done and st.session_state.summary):
        wh2  = t.get("welcome_h2",   "Enter parameters and click")
        wcal = t.get("welcome_calc", "Calculate")
        wsub = t.get("welcome_sub",  "Annuity · Classic · Balloon · Deposit<br>Excel / PDF / Word / CSV · Investment Comparison")
        st.markdown(f"""
        <div style="text-align:center;padding:70px 40px;color:#475569">
          <div style="font-size:3.5rem"></div>
          <h2 style="color:#334155;margin-top:14px">
            {wh2}
            <span style="color:#4FC3F7">{wcal}</span>
          </h2>
          <p style="color:#475569;font-size:.88rem;margin-top:8px">{wsub}</p>
          <p style="color:#334155;font-size:.72rem;margin-top:24px">
            Yev Capital LoanLogic v3.0 ·
            © 2026 Bohdan Yevtushenko (MrCemper)
          </p>
        </div>""", unsafe_allow_html=True)
        return

    # ═══════════════════════════════════════════════════════════════════════════
    #  РЕЗУЛЬТАТЫ
    # ═══════════════════════════════════════════════════════════════════════════
    smry     = st.session_state.summary
    df_d     = st.session_state.schedule_df
    df_chart = smry["df_chart"]
    sym      = smry["sym"]
    is_dep   = smry.get("is_deposit", False)

    # ── Ветка ДЕПОЗИТ ─────────────────────────────────────────────────────────
    if is_dep:
        _render_deposit_results(t, smry, df_d, df_chart, sym)
        return

    # ── Ветка КРЕДИТ ──────────────────────────────────────────────────────────

    # ── HARD "PARTIAL / INVALID RESULT" BANNER ────────────────────────────────
    # Shown above all metrics so users cannot miss that the displayed numbers
    # don't reflect what they requested (grace failed, IRR diverged, etc.)
    if smry.get("partial_result"):
        problems = []
        if smry.get("grace_error"):
            problems.append(
                t.get("partial_grace", "Grace Period was NOT applied")
                + f" — {smry['grace_error']}"
            )
        if smry.get("effective_rate_error"):
            problems.append(
                t.get("partial_apr",
                       "Effective APR could not be computed")
            )
        st.error(
            f"⚠️ **{t.get('partial_result_banner', 'PARTIAL RESULT — TREAT AS UNRELIABLE')}**\n\n"
            + "\n\n".join(f"• {p}" for p in problems)
        )

    # 1. Метрики
    st.markdown(f"<div class='sec-title'>{t['section_results']}</div>", unsafe_allow_html=True)

    be         = smry.get("balloon_breakeven")
    be_lbl     = t.get("balloon_breakeven",     "Inv. Break-even Rate")
    be_tip     = t.get("balloon_breakeven_tip", "")
    is_balloon = smry.get("scheme_key") == "balloon"

    # Show 6-column layout when balloon break-even is available; otherwise 5 columns
    if is_balloon and be is not None:
        m1, m2, m3, m4, m5, m6 = st.columns(6)
    else:
        m1, m2, m3, m4, m5 = st.columns(5)
        m6 = None

    m1.metric(t["total_payment"],    fmt_money(smry["total_payment"],    sym),
               help=t.get("help_total_payment", ""))
    m2.metric(t["total_interest"],   fmt_money(smry["total_interest"],   sym),
               help=t.get("help_total_interest", ""))
    m3.metric(t["total_commission"], fmt_money(smry["total_commission"], sym))
    # APR / Effective rate — render "N/A" when IRR fails (no nominal-rate substitution)
    eff_val = smry.get("effective_rate")
    eff_err = smry.get("effective_rate_error")
    eff_display = fmt_pct(eff_val) if eff_val is not None else "N/A"
    m4.metric(t["effective_rate"], eff_display,
               help=t.get("help_eff_rate", ""))
    if eff_err:
        # Show a caption under the metric so the failure is visible.
        m4.caption(f"⚠️ {t.get('apr_failed_caption', 'APR computation failed.')}")
    m5.metric(t["monthly_payment"],  fmt_money(smry["first_payment"],    sym),
               help=t.get("help_first_payment", ""))
    if m6 is not None and be is not None:
        m6.metric(label=be_lbl, value=fmt_pct(be), help=be_tip or None)

    # Grace error banner — visible "result is partial" indicator
    if smry.get("grace_error"):
        st.error(
            f"⚠️ {t.get('grace_failed_banner', 'Grace Period could not be applied')}: "
            f"{smry['grace_error']}"
        )

    # Syndicated: surface any tranche that failed to compute, prominently —
    # totals below are computed ONLY from valid tranches, so users must know.
    synd_errs = smry.get("tranche_errors", [])
    if synd_errs:
        for tid, msg in synd_errs:
            st.error(f"⚠️ {msg}")
        st.caption(
            "ℹ️ Totals below reflect only the valid tranches. Resolve the "
            "errors above to include the missing tranche(s) in the consolidated debt."
        )

    st.markdown(
        f"<span style='color:#94A3B8;font-size:.81rem'>{t['overpayment_pct']}: "
        f"<b style='color:#FF6B6B'>{smry['overpay_pct']:.1f}%</b></span>",
        unsafe_allow_html=True)
    st.divider()

    # 2. Диаграммы
    st.markdown(f"<div class='sec-title'>{t['section_chart']}</div>", unsafe_allow_html=True)

    is_syndicated = smry.get("syndicated", False)
    # Day-Count comparison tab is meaningful only for single-loan (not deposit, not synd)
    show_dc_tab = (not is_syndicated
                    and not smry.get("is_deposit", False)
                    and smry.get("scheme_key") in ("annuity", "classic", "balloon"))

    if is_syndicated:
        tab1, tab2, tab3, tab4, tab_synd = st.tabs([
            "📊 " + t["chart_title"],
            "🥧 " + t["chart_pie_title"],
            t.get("tab_balance", "📉 Remaining Balance"),
            "⚖️ " + t["compare_schemes"],
            "🏦 " + t.get("syndicated_chart_title", "By Tranche"),
        ])
        tab_dc = None
    elif show_dc_tab:
        tab1, tab2, tab3, tab4, tab_dc = st.tabs([
            "📊 " + t["chart_title"],
            "🥧 " + t["chart_pie_title"],
            t.get("tab_balance", "📉 Remaining Balance"),
            "⚖️ " + t["compare_schemes"],
            "📅 " + t.get("daycount_compare_tab", "Day-Count Compare"),
        ])
        tab_synd = None
    else:
        tab1, tab2, tab3, tab4 = st.tabs([
            "📊 " + t["chart_title"],
            "🥧 " + t["chart_pie_title"],
            t.get("tab_balance", "📉 Remaining Balance"),
            "⚖️ " + t["compare_schemes"],
        ])
        tab_synd = None
        tab_dc   = None

    with tab1:
        st.plotly_chart(chart_bar(df_chart, t), use_container_width=True,
                        config={"displayModeBar": False})
    with tab2:
        st.plotly_chart(chart_pie(smry["principal"], smry["total_interest"],
                                  smry["total_commission"], t),
                        use_container_width=True, config={"displayModeBar": False})
    with tab3:
        st.plotly_chart(chart_balance(df_chart, t), use_container_width=True,
                        config={"displayModeBar": False})
    with tab4:
        # Compose the same kwargs the main loan used, so the comparison
        # reflects the actual configured loan (grace, day-count, inflation,
        # risk metrics) rather than a stripped-down baseline.
        ss = st.session_state
        common_kwargs = dict(
            grace_enabled = ss.get("grace_enabled", False),
            grace_start   = ss.get("grace_start", 1),
            grace_duration= ss.get("grace_duration", 0),
            grace_type    = ss.get("grace_type", "interest_only"),
            inflation_enabled = ss.get("inflation_enabled", False),
            inflation_rate    = ss.get("inflation_rate", 0.0),
            day_count_enabled = ss.get("day_count_enabled", False),
            day_count_method  = ss.get("day_count_method", DAY_COUNT_DEFAULT),
            start_date        = ss.get("start_date", date.today()),
        )

        # Some configurations are invalid for some schemes (e.g. grace at the
        # last period when n is small for balloon). Wrap each in try/except so
        # one failing scheme doesn't break the whole comparison.
        def _safe_compare_run(scheme):
            try:
                _, s = run_calculation(
                    ss.loan_amount, ss.loan_term,
                    ss.interest_rate, ss.term_unit,
                    scheme, ss.one_time_val, ss.one_time_type,
                    ss.monthly_val, ss.monthly_type,
                    ss.currency, ss.custom_symbol, t,
                    **common_kwargs)
                return s
            except Exception:
                return None

        sa = _safe_compare_run("annuity")
        sc = _safe_compare_run("classic")
        sb = _safe_compare_run("balloon")

        if sa is None or sc is None or sb is None:
            st.warning(t.get("compare_partial_fail",
                              "One or more scheme comparisons failed; see results for valid ones."))

        # Use valid results only. If a scheme failed, fall back to zeros so
        # the chart still renders cleanly with the others.
        ann_pay = sa["total_payment"]  if sa else 0
        ann_int = sa["total_interest"] if sa else 0
        cla_pay = sc["total_payment"]  if sc else 0
        cla_int = sc["total_interest"] if sc else 0
        bal_pay = sb["total_payment"]  if sb else None
        bal_int = sb["total_interest"] if sb else None

        st.plotly_chart(
            chart_compare(ann_pay, cla_pay, ann_int, cla_int, t,
                          bal_total=bal_pay, bal_int=bal_int),
            use_container_width=True, config={"displayModeBar": False})

        # ── Comparative summary message ───────────────────────────────────────
        # Show the savings line relative to the scheme the user actually picked.
        # If all three computed, also report how the picked scheme compares to
        # the other two.
        user_scheme = smry.get("scheme_key", "annuity")
        results_map = {"annuity": (sa, ann_pay), "classic": (sc, cla_pay),
                        "balloon": (sb, bal_pay)}
        own = results_map.get(user_scheme)

        if own and own[0] is not None:
            own_pay = own[1]
            # Build a list of (label, payment) for OTHER schemes that computed
            others = [(k, v[1]) for k, v in results_map.items()
                       if k != user_scheme and v[0] is not None and v[1] is not None]
            cheaper_than = [(k, p) for k, p in others if own_pay < p]
            more_than    = [(k, p) for k, p in others if own_pay > p]

            scheme_name_map = {
                "annuity": t.get("annuity_short", "Annuity"),
                "classic": t.get("classic_short", "Classic"),
                "balloon": t.get("balloon_short", "Balloon"),
            }

            # Show BOTH banners when the chosen scheme is cheaper than some
            # alternatives AND more expensive than others (mixed picture).
            # This avoids hiding half of the relevant comparison.
            if cheaper_than:
                parts = ", ".join(
                    f"{scheme_name_map[k]} ({fmt_money(p - own_pay, sym)})"
                    for k, p in cheaper_than
                )
                st.markdown(
                    f"<div class='savings-box'>💚 "
                    f"{t.get('compare_savings_vs', 'Savings vs')} {parts}"
                    f"</div>", unsafe_allow_html=True)
            if more_than:
                parts = ", ".join(
                    f"{scheme_name_map[k]} ({fmt_money(own_pay - p, sym)})"
                    for k, p in more_than
                )
                st.markdown(
                    f"<div style='background:var(--app-bg-tertiary);"
                    f"border:1px solid var(--app-warning);border-radius:10px;"
                    f"padding:11px 18px;color:var(--app-warning);font-weight:600;"
                    f"text-align:center;margin-top:8px'>"
                    f"⚠️ {t.get('compare_overpay_vs', 'Overpayment vs')} {parts}"
                    f"</div>", unsafe_allow_html=True)

        # ── Balloon Break-even (only when user picked Balloon) ────────────────
        # Showing this for annuity/classic just confuses users — it's only
        # relevant when one actually has the balloon's freed-up principal cash
        # to invest.
        if user_scheme == "balloon":
            be_rate = sb.get("balloon_breakeven") if sb else None
            if be_rate is not None:
                be_lbl = t.get("balloon_breakeven", "Inv. Break-even Rate")
                be_desc = t.get("balloon_breakeven_desc",
                                 "Min. investment yield to justify Balloon over Annuity")
                st.markdown(
                    f"<div style='background:var(--app-bg-tertiary);"
                    f"border:1px solid var(--app-accent-strong);"
                    f"border-radius:8px;padding:10px 16px;margin-top:10px;"
                    f"font-size:.85rem;color:var(--app-text-muted)'>"
                    f"📐 <b style='color:var(--app-accent)'>{be_lbl}:</b> "
                    f"<b style='color:var(--app-warning);font-size:1.05rem'>{be_rate:.2f}%</b>"
                    f"&nbsp;&nbsp;<span style='font-size:.75rem'>{be_desc}</span>"
                    f"</div>",
                    unsafe_allow_html=True,
                )

    # ── Day-Count Convention comparison ──────────────────────────────────────
    if tab_dc is not None:
        with tab_dc:
            st.markdown(
                f"<div style='color:var(--app-text-muted);font-size:.85rem;"
                f"margin-bottom:8px'>"
                f"{t.get('daycount_compare_caption', 'Same loan, five interest-day conventions side by side.')}"
                f"</div>",
                unsafe_allow_html=True,
            )
            # Need a start_date for the day-count math. If user didn't enable
            # day-count in sidebar, fall back to st.session_state.start_date
            # (always present in state defaults). The comparison is still
            # meaningful — it shows what the user would pay under each rule.
            sd = smry.get("start_date") or st.session_state.get("start_date") or date.today()
            scheme_key = smry.get("scheme_key", "annuity")
            mo_val = smry.get("commission_per_period", 0) or 0

            fig_dc, results = chart_daycount_compare(
                principal = smry["principal"],
                n         = smry["n_periods"],
                rate_pa   = smry["rate_pa"],
                unit      = smry.get("unit", "months"),
                scheme    = scheme_key,
                start_dt  = sd,
                mo_comm   = mo_val,
                t         = t,
            )

            if fig_dc is not None:
                st.plotly_chart(fig_dc, use_container_width=True,
                                 config={"displayModeBar": False})

                # Numeric difference summary (min vs max). Helps the user
                # quantify "how much does the choice of convention cost me".
                valid_vals = [v for v in results.values() if v is not None]
                if len(valid_vals) >= 2:
                    spread = max(valid_vals) - min(valid_vals)
                    best   = min(results, key=lambda k: results[k]
                                  if results[k] is not None else float('inf'))
                    worst  = max(results, key=lambda k: results[k]
                                  if results[k] is not None else float('-inf'))
                    st.markdown(
                        f"<div style='background:var(--app-bg-tertiary);"
                        f"border:1px solid var(--app-border);"
                        f"border-radius:8px;padding:10px 16px;margin-top:10px;"
                        f"font-size:.85rem;color:var(--app-text-muted)'>"
                        f"📌 <b style='color:var(--app-success)'>"
                        f"{t.get('daycount_best', 'Cheapest')}:</b> "
                        f"<b style='color:var(--app-text)'>{best}</b> "
                        f"({fmt_money(results[best], sym)}) "
                        f"&nbsp;·&nbsp; "
                        f"<b style='color:var(--app-danger)'>"
                        f"{t.get('daycount_worst', 'Most expensive')}:</b> "
                        f"<b style='color:var(--app-text)'>{worst}</b> "
                        f"({fmt_money(results[worst], sym)}) "
                        f"&nbsp;·&nbsp; "
                        f"<b style='color:var(--app-accent)'>"
                        f"{t.get('daycount_spread', 'Spread')}:</b> "
                        f"<b style='color:var(--app-text)'>{fmt_money(spread, sym)}</b>"
                        f"</div>",
                        unsafe_allow_html=True,
                    )
                # Per-method table
                rows = [{"Method": m,
                          t.get("total_interest", "Total Interest"):
                              (fmt_money(v, sym) if v is not None else "—")}
                         for m, v in results.items()]
                st.dataframe(pd.DataFrame(rows), use_container_width=True,
                              hide_index=True)
            else:
                st.warning(t.get("daycount_no_data",
                                  "Could not compute day-count comparison."))

    # ── Syndicated tranche breakdown ──────────────────────────────────────────
    if tab_synd is not None:
        with tab_synd:
            tranche_series = smry.get("tranche_payments_series", {})
            chart = chart_syndicated_tranches(tranche_series, t)
            if chart is not None:
                st.plotly_chart(chart, use_container_width=True,
                                config={"displayModeBar": False})
                st.caption(t.get("syndicated_chart_caption",
                                   "Each color = one tranche."))

                # Per-tranche summary table
                tranches_data = smry.get("tranches", [])
                if tranches_data:
                    summary_rows = []
                    for tr in tranches_data:
                        sched_sum = sum(r["payment"] for r in tr["schedule"])
                        sched_int = sum(r["interest"] for r in tr["schedule"])
                        summary_rows.append({
                            "Tranche":   tr["letter"],
                            "Amount":    fmt_money(tr["amount"], sym),
                            "Rate (%)":  f"{tr['rate']:.2f}",
                            "Term (mo)": tr["n"],
                            "Scheme":    tr["scheme"].capitalize(),
                            "Σ Payments":fmt_money(sched_sum, sym),
                            "Σ Interest":fmt_money(sched_int, sym),
                        })
                    st.dataframe(pd.DataFrame(summary_rows),
                                  use_container_width=True, hide_index=True)

    st.divider()

    # 3. Расширенная инвест-безубыточность
    _render_breakeven_panel(t, smry)

    # 4. Учёт инфляции (Real vs Nominal)
    if smry.get("inflation_enabled"):
        st.divider()
        _render_inflation_panel(t, smry, sym)

    # 5. Риск-аналитика (LTV / DSCR / DTI)
    if any(smry.get(k) is not None for k in ("ltv", "dscr", "dti")):
        st.divider()
        _render_risk_panel(t, smry, sym)

    st.divider()

    # 6. Сравнение с инвестициями (S&P 500 / custom yield)
    _render_invest_loan(t, smry, df_chart, sym)
    st.divider()

    # 7. Анализ рефинансирования (модуль)
    _render_refinance_panel(t, smry, sym)
    st.divider()

    # 8. Email отправка отчёта
    _render_email_panel(t, df_d, smry, sym, is_deposit=False)
    st.divider()

    # 9. График платежей
    _render_schedule(t, df_d, smry, sym)

    # 10. Audit Trail (журнал изменений)
    _render_audit_trail(t)


# ─────────────────────────────────────────────────────────────────────────────
#  РЕНДЕР РЕЗУЛЬТАТОВ ДЕПОЗИТА
# ─────────────────────────────────────────────────────────────────────────────
def _render_deposit_results(t, smry, df_d, df_chart, sym):
    """Полный рендер раздела результатов для режима вклада/депозита."""
    principal     = smry["principal"]
    final_balance = smry["final_balance"]
    total_earned  = smry["total_earned"]
    total_payout  = smry["total_payout"]
    eff_rate      = smry["effective_rate"]
    mode          = smry["deposit_mode"]

    # ── 1. Метрики депозита ───────────────────────────────────────────────────
    st.markdown(f"<div class='sec-title'>{t['section_results']}</div>", unsafe_allow_html=True)

    d1, d2, d3, d4 = st.columns(4)
    d1.metric(t["dep_initial"],      fmt_money(principal,     sym))
    d2.metric(t["dep_final_balance"], fmt_money(final_balance, sym))
    d3.metric(t["dep_total_earned"],  fmt_money(total_earned,  sym))
    d4.metric(t["dep_rate_label"],    fmt_pct(eff_rate))

    # Дополнительная строка
    gain_pct = total_earned / principal * 100 if principal > 0 else 0
    mode_lbl = t["deposit_capitalize"] if mode == "capitalize" else t["deposit_payout"]
    st.markdown(
        f"<span style='color:#94A3B8;font-size:.81rem'>"
        f"Прибыль: <b style='color:#06D6A0'>+{gain_pct:.2f}%</b> &nbsp;·&nbsp; "
        f"Режим: <b style='color:#4FC3F7'>{mode_lbl}</b>"
        f"</span>",
        unsafe_allow_html=True,
    )
    if mode == "payout" and total_payout > 0:
        st.markdown(
            f"<span style='color:#94A3B8;font-size:.81rem'>"
            f"{t['dep_period_payout']}: <b style='color:#FFD166'>{fmt_money(total_payout/len(df_d.dropna(subset=[t['dep_balance_close']])), sym)}</b>"
            f"</span>",
            unsafe_allow_html=True,
        )
    st.divider()

    # ── 2. Диаграммы депозита ─────────────────────────────────────────────────
    st.markdown(f"<div class='sec-title'>{t['section_chart']}</div>", unsafe_allow_html=True)
    tab1, tab2, tab3, tab4 = st.tabs([
        "📈 " + t["dep_growth_title"],
        "📊 " + t["chart_title"],
        "🥧 " + t["chart_pie_title"],
        "⚖️ " + t["dep_vs_invest"],
    ])
    with tab1:
        st.plotly_chart(chart_dep_growth(df_chart, principal, t),
                        use_container_width=True, config={"displayModeBar": False})
    with tab2:
        st.plotly_chart(chart_dep_bar(df_chart, principal, t),
                        use_container_width=True, config={"displayModeBar": False})
    with tab3:
        st.plotly_chart(chart_dep_pie(principal, total_earned, t),
                        use_container_width=True, config={"displayModeBar": False})
    with tab4:
        st.plotly_chart(
            chart_dep_compare_modes(
                principal, st.session_state.loan_term,
                st.session_state.interest_rate, st.session_state.term_unit, t, sym),
            use_container_width=True, config={"displayModeBar": False})

    st.divider()

    # ── 3. Сравнение депозита с альтернативой ─────────────────────────────────
    _render_invest_deposit(t, smry, df_chart, sym)
    st.divider()

    # ── 4. Email отправка отчёта ──────────────────────────────────────────────
    _render_email_panel(t, df_d, smry, sym, is_deposit=True)
    st.divider()

    # ── 5. Таблица ────────────────────────────────────────────────────────────
    _render_schedule(t, df_d, smry, sym, is_deposit=True)

    # ── 6. Audit Trail (журнал изменений) ─────────────────────────────────────
    _render_audit_trail(t)


# ─────────────────────────────────────────────────────────────────────────────
#  СРАВНЕНИЕ С ИНВЕСТИЦИЯМИ — КРЕДИТ
# ─────────────────────────────────────────────────────────────────────────────
# ─────────────────────────────────────────────────────────────────────────────
#  EMAIL / SMTP — отправка отчёта
# ─────────────────────────────────────────────────────────────────────────────
def _load_smtp_config() -> dict | None:
    """
    Загружает SMTP-конфигурацию из st.secrets.

    Поддерживает оба синтаксиса доступа Streamlit:
      • st.secrets["smtp"]["server"]   — словарный
      • st.secrets.smtp.server         — атрибутный

    Возвращает dict {server, port, login, password, sender} или None,
    если конфигурация отсутствует или невалидна.
    Никогда не выбрасывает исключение наружу.
    """
    try:
        # Пытаемся достать секцию `smtp` любым из двух способов
        try:
            smtp_cfg = st.secrets["smtp"]
        except (KeyError, AttributeError, TypeError):
            smtp_cfg = getattr(st.secrets, "smtp", None)

        if smtp_cfg is None:
            return None

        # Универсальный getter — работает и со словарём, и с объектом
        def _get(cfg, key, default=None):
            if hasattr(cfg, "__getitem__"):
                try:
                    return cfg[key]
                except (KeyError, TypeError):
                    pass
            return getattr(cfg, key, default)

        server   = _get(smtp_cfg, "server")
        port     = _get(smtp_cfg, "port")
        login    = _get(smtp_cfg, "login")
        password = _get(smtp_cfg, "password")
        sender   = _get(smtp_cfg, "sender", login)

        # Проверяем что обязательные поля — реальные непустые строки
        if not server or not login or not password or port is None:
            return None
        if not isinstance(server, str) or not isinstance(login, str):
            return None
        if "Mock" in str(server) or "Mock" in str(login):
            return None

        return {
            "server":   str(server).strip(),
            "port":     int(port),
            "login":    str(login).strip(),
            "password": str(password),
            "sender":   str(sender).strip() if sender else str(login).strip(),
        }
    except Exception:
        return None


def is_smtp_configured() -> bool:
    """True если SMTP-секреты валидно настроены."""
    return _load_smtp_config() is not None


def is_valid_email(addr: str) -> bool:
    """Простая валидация email через regex."""
    import re
    if not addr or not addr.strip():
        return False
    return bool(re.match(
        r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$',
        addr.strip(),
    ))


def send_report_email(recipient: str, subject: str, body: str,
                       attachment_bytes: bytes, attachment_name: str,
                       attachment_mime: str) -> tuple[bool, str]:
    """
    Отправляет email с вложением через SMTP.
    Конфигурация берётся из st.secrets["smtp"] (или st.secrets.smtp).

    Returns: (success: bool, message: str)
    """
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.application import MIMEApplication

    # Валидация email
    if not recipient or not recipient.strip():
        return False, "no_recipient"
    if not is_valid_email(recipient):
        return False, "invalid_email"

    # Загрузка SMTP-секретов
    cfg = _load_smtp_config()
    if cfg is None:
        return False, "no_secrets"

    server, port = cfg["server"], cfg["port"]
    login, password, sender = cfg["login"], cfg["password"], cfg["sender"]

    # Header-injection defense: strip CR/LF (and surrounding whitespace) from
    # any value placed into an email header. Without this, a newline in the
    # subject could inject additional headers (e.g. "Subject: x\nBcc: ...").
    def _sanitize_header(value: str) -> str:
        return re.sub(r"[\r\n]+", " ", str(value)).strip()

    safe_recipient = _sanitize_header(recipient)
    safe_subject   = _sanitize_header(subject) if subject else ""
    # Re-validate the recipient AFTER sanitizing, in case stripping changed it.
    if not is_valid_email(safe_recipient):
        return False, "invalid_email"

    # Сборка письма
    msg            = MIMEMultipart()
    msg["From"]    = sender
    msg["To"]      = safe_recipient
    msg["Subject"] = safe_subject

    if body:
        # Body goes into the message payload (not a header), so newlines are
        # fine here — only headers are injection-sensitive.
        msg.attach(MIMEText(body, "plain", "utf-8"))

    # If an attachment was intended (a filename/MIME was supplied) but the
    # bytes are empty or missing, the export upstream likely failed. Sending an
    # attachment-less "report" email would mislead the recipient, so refuse.
    if attachment_name and not attachment_bytes:
        return False, "empty_attachment"

    if attachment_bytes:
        part = MIMEApplication(attachment_bytes, _subtype=attachment_mime.split("/")[-1])
        part.add_header("Content-Disposition", "attachment",
                         filename=attachment_name)
        msg.attach(part)

    # Отправка
    try:
        if port == 465:
            with smtplib.SMTP_SSL(server, port, timeout=30) as smtp:
                smtp.login(login, password)
                smtp.send_message(msg)
        else:
            with smtplib.SMTP(server, port, timeout=30) as smtp:
                smtp.ehlo()
                smtp.starttls()
                smtp.ehlo()
                smtp.login(login, password)
                smtp.send_message(msg)
        return True, "sent"
    except smtplib.SMTPAuthenticationError:
        return False, "Authentication failed — check login/password"
    except smtplib.SMTPRecipientsRefused:
        return False, "Recipient address refused by server"
    except smtplib.SMTPServerDisconnected:
        return False, "Server disconnected"
    except smtplib.SMTPConnectError:
        return False, "Cannot connect to SMTP server"
    except (TimeoutError, OSError) as e:
        return False, f"Network error: {e}"
    except Exception as e:
        return False, f"Unexpected error: {e}"


def _render_email_panel(t, df_d, smry, sym, is_deposit: bool = False):
    """
    UI блок отправки отчёта по email.

    Логика гейтов:
      • SMTP не настроен → кнопка disabled, info-баннер + "How to setup" expander
      • Email не валидный → кнопка disabled, мягкое info-уведомление
      • Всё ок → кнопка активна
    """
    # Уникальные ключи виджетов (loan vs deposit)
    suffix = "_dep" if is_deposit else "_loan"

    with st.expander(t.get("email_section", "📧 Send Report via Email"),
                      expanded=False):
        st.caption(t.get("email_caption", ""))

        smtp_ok = is_smtp_configured()

        # ── Если SMTP не настроен — баннер + инструкция ──────────────────────
        if not smtp_ok:
            st.info(
                t.get("email_disabled_warning",
                       "Email service is not configured. "
                       "Please add SMTP credentials to secrets."),
                icon="ℹ️",
            )

        # Поля ввода — рендерим всегда (но кнопку заблокируем при необходимости)
        col_a, col_b = st.columns([2, 1])
        with col_a:
            recipient = st.text_input(
                t.get("email_recipient", "Recipient Email"),
                value="",
                placeholder="user@example.com",
                key=f"email_to{suffix}",
                help=t.get("email_recipient_help", ""),
                disabled=not smtp_ok,
            )
        with col_b:
            fmt_opts = {
                t.get("email_format_pdf",  "PDF"):           "pdf",
                t.get("email_format_xlsx", "Excel (XLSX)"):  "xlsx",
                t.get("email_format_docx", "Word (DOCX)"):   "docx",
            }
            fmt_lbl = st.selectbox(
                t.get("email_format", "Attachment Format"),
                options=list(fmt_opts.keys()),
                key=f"email_fmt{suffix}",
                disabled=not smtp_ok,
            )
            fmt_choice = fmt_opts[fmt_lbl]

        subject = st.text_input(
            t.get("email_subject", "Subject"),
            value=t.get("email_default_subject", "Yev Capital LoanLogic Report"),
            key=f"email_subj{suffix}",
            disabled=not smtp_ok,
        )
        body = st.text_area(
            t.get("email_message", "Message"),
            value=t.get("email_default_body", ""),
            height=110,
            key=f"email_body{suffix}",
            disabled=not smtp_ok,
        )

        # ── Гейтинг кнопки: smtp + валидный email ────────────────────────────
        email_valid = is_valid_email(recipient)
        can_send    = smtp_ok and email_valid

        # Подсказка пользователю если email невалидный (но SMTP настроен)
        if smtp_ok and recipient and not email_valid:
            st.caption("⚠️ " + t.get("email_invalid_warning",
                                       "Enter a valid email address."))

        if st.button(
            t.get("email_send_btn", "📤 Send Report"),
            key=f"btn_send_email{suffix}",
            disabled=not can_send,
            type="primary" if can_send else "secondary",
        ):
            with st.spinner(t.get("email_sending", "Sending...")):
                # Готовим вложение
                ts = datetime.now().strftime("%Y%m%d_%H%M")
                prefix = "dep" if is_deposit else "loan"
                try:
                    if fmt_choice == "pdf":
                        attach = export_pdf(df_d, smry, t, sym)
                        name   = f"{prefix}_{ts}.pdf"
                        mime   = "application/pdf"
                    elif fmt_choice == "xlsx":
                        attach = export_excel(df_d, smry, t, sym)
                        name   = f"{prefix}_{ts}.xlsx"
                        mime   = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    else:  # docx
                        attach = export_docx(df_d, smry, t, sym)
                        name   = f"{prefix}_{ts}.docx"
                        mime   = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                except Exception as e:
                    st.error(t.get("email_error", "Send failed: {error}").format(error=e))
                    return

                ok, msg = send_report_email(
                    recipient=recipient,
                    subject=subject,
                    body=body,
                    attachment_bytes=attach,
                    attachment_name=name,
                    attachment_mime=mime,
                )

                if ok:
                    st.success(t.get("email_success",
                                      "Report sent to {email}").format(email=recipient))
                else:
                    if msg == "no_recipient":
                        st.error(t.get("email_no_recipient",
                                        "Please enter recipient email"))
                    elif msg == "invalid_email":
                        st.error(t.get("email_invalid", "Invalid email address"))
                    elif msg == "no_secrets":
                        st.warning(t.get("email_disabled_warning",
                                          "SMTP not configured"))
                    elif msg == "empty_attachment":
                        st.error(t.get("email_empty_attachment",
                                        "The report file is empty (export may "
                                        "have failed). Nothing was sent."))
                    else:
                        st.error(t.get("email_error",
                                        "Send failed: {error}").format(error=msg))

        # ── "How to setup Email" — отдельный вложенный expander ──────────────
        with st.expander(t.get("email_setup_title", "📖 How to setup Email"),
                          expanded=not smtp_ok):
            st.markdown(t.get("email_setup_steps", ""))


# ─────────────────────────────────────────────────────────────────────────────
#  МОДУЛЬ РЕФИНАНСИРОВАНИЯ
# ─────────────────────────────────────────────────────────────────────────────
def _render_refinance_panel(t, smry, sym):
    """Самостоятельный модуль анализа рефинансирования.
    Все поля вводятся пользователем независимо от текущего расчёта."""
    with st.expander(t.get("refi_section", "🔄 Refinancing Analysis"), expanded=False):
        st.caption(t.get("refi_caption", ""))
        st.info(t.get("refi_annuity_note",
                       "ℹ️ This analysis models both the current and the new "
                       "loan as **annuity** (equal-payment) loans. If your "
                       "current loan uses a different scheme (classic or "
                       "balloon), treat the comparison as approximate."))

        # ── Текущий кредит ───────────────────────────────────────────────────
        st.markdown(f"**{t.get('refi_current_block', 'Current Loan')}**")
        col_a, col_b = st.columns(2)
        with col_a:
            cur_balance = st.number_input(
                t.get("refi_current_balance", "Outstanding Balance"),
                min_value=0.01, max_value=1e14,
                value=float(st.session_state.get("refi_cur_balance",
                                                  smry.get("principal", 100_000.0))),
                step=1000.0, format="%.2f", key="refi_cur_balance",
                help=t.get("refi_help_balance", ""))
            st.caption(t.get("refi_balance_hint",
                              "↳ Enter the amount you still owe today, not the "
                              "original loan amount. (Pre-filled with the "
                              "original principal as a starting point.)"))
            cur_rate = st.number_input(
                t.get("refi_current_rate", "Current Rate (% annual)"),
                min_value=-20.0, max_value=100.0,
                value=float(st.session_state.get("refi_cur_rate",
                                                  smry.get("rate_pa", 12.0))),
                step=0.1, format="%.2f", key="refi_cur_rate")
        with col_b:
            remaining = st.number_input(
                t.get("refi_remaining_term", "Remaining Term (months)"),
                min_value=1, max_value=600,
                value=int(st.session_state.get("refi_remaining", 24)),
                step=1, key="refi_remaining")
            penalty_type_opts = {
                t.get("refi_penalty_pct",   "Penalty (% of balance)"): "pct",
                t.get("refi_penalty_fixed", "Penalty (fixed)"):         "fixed",
            }
            cur_pen_type = st.session_state.get("refi_penalty_type", "pct")
            pt_idx = list(penalty_type_opts.values()).index(cur_pen_type) \
                     if cur_pen_type in penalty_type_opts.values() else 0
            pt_lbl = st.radio(
                t.get("refi_penalty_type", "Penalty Type"),
                options=list(penalty_type_opts.keys()),
                index=pt_idx, horizontal=True, key="refi_penalty_type_radio",
                help=t.get("refi_help_penalty", ""))
            st.session_state["refi_penalty_type"] = penalty_type_opts[pt_lbl]
            penalty_input = st.number_input(
                t.get("refi_penalty", "Early Closure Penalty"),
                min_value=0.0,
                max_value=100.0 if st.session_state["refi_penalty_type"] == "pct" else 1e10,
                value=float(st.session_state.get("refi_penalty_val",
                            2.0 if st.session_state["refi_penalty_type"] == "pct" else 1000.0)),
                step=0.1 if st.session_state["refi_penalty_type"] == "pct" else 100.0,
                format="%.2f", key="refi_penalty_val")

        st.divider()

        # ── Новый кредит ─────────────────────────────────────────────────────
        st.markdown(f"**{t.get('refi_new_block', 'New Loan')}**")
        col_c, col_d, col_e = st.columns(3)
        with col_c:
            new_rate = st.number_input(
                t.get("refi_new_rate", "New Rate (% annual)"),
                min_value=-20.0, max_value=100.0,
                value=float(st.session_state.get("refi_new_rate",
                                                  max(cur_rate - 2.0, 0.5))),
                step=0.1, format="%.2f", key="refi_new_rate")
        with col_d:
            new_term = st.number_input(
                t.get("refi_new_term", "New Loan Term (months)"),
                min_value=1, max_value=600,
                value=int(st.session_state.get("refi_new_term", remaining)),
                step=1, key="refi_new_term")
        with col_e:
            new_fees = st.number_input(
                t.get("refi_new_fees", "Origination Fees"),
                min_value=0.0, max_value=1e10,
                value=float(st.session_state.get("refi_new_fees", 1000.0)),
                step=100.0, format="%.2f", key="refi_new_fees",
                help=t.get("refi_help_fees", ""))

        # ── NPV discount rate ────────────────────────────────────────────────
        # Annual rate at which future savings/costs are discounted. With 0%
        # the NPV comparison reduces to a nominal sum; in practice the
        # opportunity-cost-of-capital is non-zero, so the input is exposed
        # explicitly rather than buried inside calc_refinance_analysis.
        refi_discount_rate = st.number_input(
            t.get("refi_discount_rate", "NPV Discount Rate (% annual)"),
            min_value=0.0, max_value=50.0,
            value=float(st.session_state.get("refi_discount_rate", 5.0)),
            step=0.1, format="%.2f", key="refi_discount_rate",
            help=t.get("refi_discount_rate_help",
                        "Annual rate used to discount future cash flows when "
                        "computing NPV. Use your alternative-investment yield."))

        # ── Расчёт ────────────────────────────────────────────────────────────
        if st.button(t.get("refi_calculate", "Calculate Refinancing"),
                      key="btn_refi_calc"):
            # Преобразуем штраф в абсолютную сумму
            if st.session_state["refi_penalty_type"] == "pct":
                penalty_amt = cur_balance * penalty_input / 100.0
            else:
                penalty_amt = penalty_input

            try:
                result = calc_refinance_analysis(
                    current_balance=cur_balance,
                    current_rate_pa=cur_rate,
                    remaining_months=int(remaining),
                    penalty=penalty_amt,
                    new_rate_pa=new_rate,
                    new_term_months=int(new_term),
                    new_fees=new_fees,
                    discount_rate_pa=refi_discount_rate,
                )
            except Exception as e:
                st.error(f"Calculation error: {e}")
                return

            st.divider()

            # Метрики
            r1, r2, r3 = st.columns(3)
            r1.metric(
                t.get("refi_current_payment", "Current Payment"),
                fmt_money(result["current_payment"], sym),
            )
            r2.metric(
                t.get("refi_new_payment", "New Payment"),
                fmt_money(result["new_payment"], sym),
            )
            delta_str = fmt_money(result["monthly_savings"], sym)
            delta_arr = "+" if result["monthly_savings"] > 0 else ""
            r3.metric(
                t.get("refi_monthly_savings", "Monthly Savings"),
                delta_str,
                delta=f"{delta_arr}{result['monthly_savings']/result['current_payment']*100:.1f}%"
                       if result["current_payment"] > 0 else "—",
            )

            r4, r5, r6 = st.columns(3)
            r4.metric(
                t.get("refi_total_costs", "Total Switching Costs"),
                fmt_money(result["total_costs"], sym),
                help=f"Penalty: {fmt_money(penalty_amt, sym)} + Fees: {fmt_money(new_fees, sym)}",
            )
            be_display = (f"{result['breakeven_months']:.1f} months"
                          if result["breakeven_months"] is not None else "N/A")
            r5.metric(
                t.get("refi_breakeven_months", "Break-even Point"),
                be_display,
                help=t.get("refi_help_breakeven", ""),
            )
            r6.metric(
                t.get("refi_total_savings", "Net Savings (NPV-based)"),
                fmt_money(result["npv_savings"], sym),
                help=t.get("refi_help_npv", ""),
            )

            # Вердикт
            if result["verdict"] == "worth":
                st.success(
                    t.get("refi_worth_it",
                          "✓ Refinancing is worthwhile. Pays back in {months:.1f} months.")
                    .format(months=result["breakeven_months"])
                )
            elif result["verdict"] == "long_payback":
                st.warning(t.get("refi_long_payback",
                                  "⚠️ Payback period exceeds new loan term."))
            elif result["verdict"] == "longer_term_trap":
                st.warning(t.get("refi_longer_term_trap",
                                  "⚠️ Longer-term trap: monthly drops but total cost rises."))
            elif result["verdict"] == "shorter_term_win":
                st.success(t.get("refi_shorter_term_win",
                                  "✓ Higher monthly payment, but lower total cost in "
                                  "present-value terms — economically worthwhile."))
            else:
                st.error(t.get("refi_not_worth",
                                "✗ Refinancing not worthwhile — new payment is higher."))


# ─────────────────────────────────────────────────────────────────────────────
#  РАСШИРЕННАЯ ИНВЕСТ-БЕЗУБЫТОЧНОСТЬ
# ─────────────────────────────────────────────────────────────────────────────
def _render_breakeven_panel(t, smry):
    """Renders the Investment Break-even Analysis panel.

    Shown metrics depend on the chosen scheme:
      • All schemes: Universal break-even (vs total interest).
      • Annuity:     also "Annuity vs Classic" pairwise (when does investing
                     the early-period savings make annuity competitive).
      • Classic:     just the universal — classic is the cheapest baseline.
      • Balloon:     adds "Vs Annuity" (existing legacy), "Vs Classic" (new),
                     and "Absolute".
    """
    universal_be = smry.get("universal_breakeven")
    balloon_be   = smry.get("balloon_breakeven")
    balloon_abs  = smry.get("balloon_breakeven_abs")
    be_vs_cla    = smry.get("breakeven_vs_classic")
    be_vs_ann    = smry.get("breakeven_vs_annuity")
    scheme       = smry.get("scheme_key")
    is_balloon   = scheme == "balloon"
    is_annuity   = scheme == "annuity"

    # Nothing meaningful to show
    if (universal_be is None and balloon_be is None and balloon_abs is None
            and be_vs_cla is None and be_vs_ann is None):
        return

    st.markdown(
        f"<div class='sec-title'>{t.get('invest_breakeven_section', 'Investment Break-even')}</div>",
        unsafe_allow_html=True,
    )

    def _fmt_be(v):
        """Format break-even rate or N/A if mathematically impossible."""
        if v is None:
            return "N/A"
        return fmt_pct(v)

    # Layout: 2 columns for annuity (Universal + vs Classic),
    #         3 columns for balloon (Universal + Vs Annuity + Vs Classic + Abs → 4-col)
    #         1 column for classic.
    if is_balloon:
        ncols = 4 if be_vs_cla is not None else 3
        cols = st.columns(ncols)
        ci = 0
        if universal_be is not None:
            cols[ci].metric(
                label=t.get("invest_breakeven_universal", "Universal Break-even Rate"),
                value=_fmt_be(universal_be),
                help=t.get("invest_breakeven_universal_tip", ""),
            )
            ci += 1
        cols[ci].metric(
            label=t.get("invest_breakeven_vs_ann", "Vs. Annuity Break-even"),
            value=_fmt_be(balloon_be),
            help=t.get("invest_breakeven_vs_ann_tip", ""),
        )
        ci += 1
        if be_vs_cla is not None and ci < ncols:
            cols[ci].metric(
                label=t.get("invest_breakeven_vs_cla_balloon",
                             "Vs. Classic Break-even"),
                value=_fmt_be(be_vs_cla),
                help=t.get("invest_breakeven_vs_cla_balloon_tip",
                            "Min. investment yield at which Balloon (with all its deferred principal) "
                            "matches Classic by maturity."),
            )
            ci += 1
        if ci < ncols:
            cols[ci].metric(
                label=t.get("invest_breakeven_abs", "Absolute Break-even"),
                value=_fmt_be(balloon_abs),
                help=t.get("invest_breakeven_abs_tip", ""),
            )
        st.caption(t.get("balloon_be_caption", ""))

    elif is_annuity:
        # Show Universal + Annuity-vs-Classic (if computed)
        if be_vs_cla is not None:
            c1, c2 = st.columns(2)
            if universal_be is not None:
                c1.metric(
                    label=t.get("invest_breakeven_universal", "Universal Break-even Rate"),
                    value=_fmt_be(universal_be),
                    help=t.get("invest_breakeven_universal_tip", ""),
                )
            c2.metric(
                label=t.get("invest_breakeven_vs_cla_annuity",
                             "Annuity Vs. Classic Break-even"),
                value=_fmt_be(be_vs_cla),
                help=t.get("invest_breakeven_vs_cla_annuity_tip",
                            "Min. annual yield at which investing the early-period "
                            "savings of Annuity-over-Classic compensates for the "
                            "later-period overpayment by maturity."),
            )
            st.caption(t.get("annuity_be_caption",
                              "If you can earn ≥ this rate on the early-period cash "
                              "you save with Annuity, you finish at least as well off "
                              "as you would have with Classic."))
        elif universal_be is not None:
            cols = st.columns(1)
            cols[0].metric(
                label=t.get("invest_breakeven_universal", "Universal Break-even Rate"),
                value=_fmt_be(universal_be),
                help=t.get("invest_breakeven_universal_tip", ""),
            )
    else:
        # Classic — just universal
        if universal_be is not None:
            cols = st.columns(1)
            cols[0].metric(
                label=t.get("invest_breakeven_universal", "Universal Break-even Rate"),
                value=_fmt_be(universal_be),
                help=t.get("invest_breakeven_universal_tip", ""),
            )


# ─────────────────────────────────────────────────────────────────────────────
#  ИНФЛЯЦИЯ — Real vs Nominal Cost
# ─────────────────────────────────────────────────────────────────────────────
def _render_inflation_panel(t, smry, sym):
    """Рендерит блок Real vs Nominal Cost при включённом учёте инфляции."""
    if not smry.get("inflation_enabled") or smry.get("real_cost") is None:
        return

    st.markdown(
        f"<div class='sec-title'>{t.get('inflation_section', 'Real vs Nominal Cost')}</div>",
        unsafe_allow_html=True,
    )

    nominal  = smry["total_payment"]
    real     = smry["real_cost"]
    savings  = smry.get("inflation_savings") or (nominal - real)
    inf_rate = smry.get("inflation_rate") or 0.0

    c1, c2, c3 = st.columns(3)
    c1.metric(
        label=t.get("nominal_cost", "Nominal Total Cost"),
        value=fmt_money(nominal, sym),
        help=t.get("help_nominal", ""),
    )
    c2.metric(
        label=t.get("real_cost", "Real Total Cost (PV)"),
        value=fmt_money(real, sym),
        help=t.get("help_real_cost_long", "") + "\n\n" + t.get("help_real", ""),
    )
    delta_pct = -savings / nominal * 100 if nominal > 0 else 0
    c3.metric(
        label=t.get("inflation_savings", "Inflation Discount"),
        value=fmt_money(savings, sym),
        delta=f"{delta_pct:.1f}%" if nominal > 0 else "—",
        help=t.get("help_disc", ""),
    )

    inflation_note = t.get('inflation_note',
        "Real cost discounts future payments to today's purchasing power.")
    rate_lbl = t.get('inflation_rate', 'Inflation')
    full_caption = t.get("caption_inflation_full",
        "💡 {note} Inflation 'eats' the real value of money over time.").format(note=inflation_note)
    st.caption(f"{full_caption}  ({rate_lbl}: {inf_rate:.1f}%)")


# ─────────────────────────────────────────────────────────────────────────────
#  РИСК-АНАЛИТИКА  —  LTV / DSCR / DTI
# ─────────────────────────────────────────────────────────────────────────────
def calc_credit_health_score(ltv: float | None,
                              dscr: float | None,
                              dti: float | None) -> float | None:
    """
    Combines the available risk metrics into a single 0–100 credit-health
    score (higher = healthier). Only the metrics that are actually provided
    (non-None) contribute; the result is their equal-weighted average. Returns
    None when none of the three is available, so the caller can hide the gauge.

    Each metric maps through a PIECEWISE-LINEAR curve whose break-points are
    aligned with the qualitative status tiers (ltv_status / dscr_status /
    dti_status) and with the A/B/C/D rank cut-offs (A≥80, B≥60, C≥40, D<40).
    This guarantees the gauge and the risk tiles tell the SAME story — e.g. an
    LTV that the tile calls "Standard" lands in the B band, not in C.

    Tier-boundary → score anchors (then linearly interpolated between them):
      • LTV  (lower better): 60→90 (Safe/Standard line, top of A),
                              80→70 (Standard/High line, mid B),
                              95→45 (High/Critical line, mid C),
                              ≥110→0. ≤60 saturates toward 100.
      • DSCR (higher better): 1.25→90 (Safe line, A),
                              1.00→65 (Warning line, B),
                              0.75→45 (deep warning, C),
                              ≤0.50→0. ≥1.50 saturates toward 100.
      • DTI  (lower better): 28→90 (Excellent line, A),
                              36→70 (Good line, B),
                              43→50 (Acceptable line, C),
                              ≥60→0. ≤20 saturates toward 100.
    """
    def _interp(x, points):
        """Piecewise-linear interpolation. `points` is a list of (x, score)
        sorted by x ascending; clamps outside the range."""
        if x <= points[0][0]:
            return points[0][1]
        if x >= points[-1][0]:
            return points[-1][1]
        for (x0, y0), (x1, y1) in zip(points, points[1:]):
            if x0 <= x <= x1:
                if x1 == x0:
                    return y1
                frac = (x - x0) / (x1 - x0)
                return y0 + frac * (y1 - y0)
        return points[-1][1]

    sub_scores = []

    if ltv is not None:
        # Lower LTV is better → descending score. Anchored to ltv_status tiers.
        sub_scores.append(_interp(ltv, [
            (40.0, 100.0), (60.0, 90.0), (80.0, 70.0),
            (95.0, 45.0), (110.0, 0.0),
        ]))

    if dscr is not None:
        # Higher DSCR is better → ascending score. Anchored to dscr_status tiers.
        sub_scores.append(_interp(dscr, [
            (0.50, 0.0), (0.75, 45.0), (1.00, 65.0),
            (1.25, 90.0), (1.50, 100.0),
        ]))

    if dti is not None:
        # Lower DTI is better → descending score. Anchored to dti_status tiers.
        sub_scores.append(_interp(dti, [
            (20.0, 100.0), (28.0, 90.0), (36.0, 70.0),
            (43.0, 50.0), (60.0, 0.0),
        ]))

    if not sub_scores:
        return None
    return sum(sub_scores) / len(sub_scores)


def credit_health_rank(score: float | None, t: dict) -> tuple[str, str]:
    """
    Maps a 0–100 score to (rank_label, color_hex).
      80–100 → A (green), 60–79 → B (lime), 40–59 → C (orange), 0–39 → D (red).
    None → neutral 'N/A'.
    """
    if score is None:
        return t.get("credit_rank_na", "N/A"), "#64748B"
    if score >= 80:
        return t.get("credit_rank_a", "Rank A — Excellent"), "#10B981"
    if score >= 60:
        return t.get("credit_rank_b", "Rank B — Good"),      "#A3E635"
    if score >= 40:
        return t.get("credit_rank_c", "Rank C — Fair"),      "#F59E0B"
    return t.get("credit_rank_d", "Rank D — High Risk"),     "#DC2626"


def chart_credit_health_gauge(score: float, t: dict):
    """
    Builds a Plotly gauge (speedometer) for the 0–100 credit-health score,
    themed to match the app's other charts. Coloured bands mark the A/B/C/D
    zones; the needle/value shows the current score.
    """
    # Rank on the SAME value that is displayed (rounded to 1 dp). Otherwise a
    # raw score of e.g. 79.96 would print "80.0" (looks like Rank A) while the
    # bar/title were coloured for Rank B — a confusing boundary mismatch.
    shown = round(score, 1)
    rank_label, rank_color = credit_health_rank(shown, t)
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=shown,
        number=dict(suffix=" / 100", font=dict(color=C["text"], size=26)),
        title=dict(
            text=f"{t.get('credit_health_title', 'Credit Health')}<br>"
                 f"<span style='font-size:0.8em;color:{rank_color}'>{rank_label}</span>",
            font=dict(color=C["text"], size=15),
        ),
        gauge=dict(
            axis=dict(range=[0, 100], tickwidth=1, tickcolor=C["muted"],
                       tickfont=dict(color=C["muted"], size=10)),
            bar=dict(color=rank_color, thickness=0.28),
            bgcolor=C["card"],
            borderwidth=1, bordercolor=C["grid"],
            steps=[
                dict(range=[0, 40],   color="rgba(220,38,38,0.25)"),    # D
                dict(range=[40, 60],  color="rgba(245,158,11,0.25)"),   # C
                dict(range=[60, 80],  color="rgba(163,230,53,0.22)"),   # B
                dict(range=[80, 100], color="rgba(16,185,129,0.25)"),   # A
            ],
            threshold=dict(
                line=dict(color=rank_color, width=4),
                thickness=0.78, value=shown,
            ),
        ),
    ))
    fig.update_layout(
        paper_bgcolor=C["bg"], font=dict(color=C["text"]),
        margin=dict(l=20, r=20, t=70, b=10), height=300,
    )
    return fig


def _render_risk_panel(t, smry, sym):
    """Рендерит блок риск-аналитики, если хотя бы один индикатор включён."""
    ltv  = smry.get("ltv")
    dscr = smry.get("dscr")
    dti  = smry.get("dti")

    has_any = any(v is not None for v in (ltv, dscr, dti))
    if not has_any:
        return

    st.markdown(
        f"<div class='sec-title'>{t.get('risk_section', 'Risk Analytics')}</div>",
        unsafe_allow_html=True,
    )

    metrics = []
    if ltv is not None:  metrics.append("ltv")
    if dscr is not None: metrics.append("dscr")
    if dti is not None:  metrics.append("dti")

    cols = st.columns(len(metrics))

    for i, m in enumerate(metrics):
        if m == "ltv":
            label, color = ltv_status(ltv, t)
            display_val = fmt_pct(ltv) if ltv is not None else "N/A"
            cols[i].metric(
                label=t.get("ltv_label", "Loan-to-Value (LTV)"),
                value=display_val,
                help=t.get("help_ltv_metric", ""),
            )
            cols[i].markdown(
                f"<div style='text-align:center;margin-top:-8px;"
                f"font-size:.78rem;font-weight:700;color:{color}'>"
                f"● {label}</div>",
                unsafe_allow_html=True,
            )
            if ltv is not None and ltv > 80.0:
                cols[i].caption(f"⚠️ {t.get('ltv_warning', 'High LTV')}")

        elif m == "dscr":
            label, color = dscr_status(dscr, t)
            display_val = f"{dscr:.2f}" if dscr is not None and dscr > 0 else "N/A"
            cols[i].metric(
                label=t.get("dscr_label", "DSCR"),
                value=display_val,
                help=t.get("help_dscr_metric", ""),
            )
            cols[i].markdown(
                f"<div style='text-align:center;margin-top:-8px;"
                f"font-size:.78rem;font-weight:700;color:{color}'>"
                f"● {label}</div>",
                unsafe_allow_html=True,
            )

        elif m == "dti":
            label, color = dti_status(dti, t)
            display_val = fmt_pct(dti) if dti is not None and dti > 0 else "N/A"
            cols[i].metric(
                label=t.get("dti_label", "DTI"),
                value=display_val,
                help=t.get("help_dti_metric", ""),
            )
            cols[i].markdown(
                f"<div style='text-align:center;margin-top:-8px;"
                f"font-size:.78rem;font-weight:700;color:{color}'>"
                f"● {label}</div>",
                unsafe_allow_html=True,
            )

    # ── Credit Health speedometer ────────────────────────────────────────────
    # A single 0–100 gauge summarising the available risk metrics. Shown only
    # when at least one of LTV/DSCR/DTI is present (same condition as the tiles).
    score = calc_credit_health_score(ltv, dscr, dti)
    if score is not None:
        st.plotly_chart(chart_credit_health_gauge(score, t),
                         use_container_width=True)
        contributing = ", ".join(
            lbl for lbl, val in (("LTV", ltv), ("DSCR", dscr), ("DTI", dti))
            if val is not None)
        st.caption(t.get("credit_health_caption",
                          "Score combines the metrics you provided ({metrics}). "
                          "Higher is healthier.").format(metrics=contributing))


def _render_invest_loan(t, smry, df_chart, sym):
    """
    Для кредита: «Что если вместо кредитных платежей инвестировать эти деньги?»
    Сравниваем: накопленный портфель из кредитных платежей vs остаток долга.
    """
    st.markdown(f"<div class='sec-title'>{t['invest_section']}</div>", unsafe_allow_html=True)
    st.caption(t["invest_explanation"])

    col_a, col_b = st.columns(2)
    with col_a:
        sp_on = st.checkbox(t["invest_sp500"], value=st.session_state.invest_sp500, key="cb_sp500")
    with col_b:
        cu_on = st.checkbox(t["invest_custom"], value=st.session_state.invest_custom, key="cb_custom")

    if sp_on and cu_on:
        st.warning(t["invest_only_one"])
        cu_on = False
    st.session_state.invest_sp500  = sp_on
    st.session_state.invest_custom = cu_on

    if not (sp_on or cu_on):
        return

    if sp_on:
        yield_pct   = SP500_BENCHMARK_RATE
        yield_label = f"S&P 500 — {SP500_BENCHMARK_RATE:g}% {t['invest_rate_label']}"
        st.caption(t.get("sp500_disclaimer", ""))
    else:
        yield_pct = st.slider(
            t["invest_custom_rate"], min_value=0.5, max_value=100.0,
            value=float(st.session_state.invest_custom_rate), step=0.5, key="invest_sldr")
        st.session_state.invest_custom_rate = yield_pct
        yield_label = f"{yield_pct:.1f}% {t['invest_rate_label']}"

    payments       = smry["payments"]
    invest_vals    = calc_investment(payments, yield_pct, smry["unit"])
    total_invested = sum(payments)
    final_val      = invest_vals[-1]
    net_gain       = final_val - total_invested

    ic1, ic2, ic3 = st.columns(3)
    ic1.metric(t["invest_total_paid"],    fmt_money(total_invested, sym))
    ic2.metric(t["invest_portfolio_val"], fmt_money(final_val,      sym))
    delta_str = f"+{net_gain/total_invested*100:.1f}%" if total_invested > 0 else "—"
    ic3.metric(t["invest_net_gain"], fmt_money(net_gain, sym), delta=delta_str)

    st.plotly_chart(chart_invest(df_chart, invest_vals, t, yield_label),
                    use_container_width=True, config={"displayModeBar": False})


# ─────────────────────────────────────────────────────────────────────────────
#  СРАВНЕНИЕ С ИНВЕСТИЦИЯМИ — ДЕПОЗИТ
# ─────────────────────────────────────────────────────────────────────────────
def _render_invest_deposit(t, smry, df_chart, sym):
    """
    Для депозита: «Что если вместо этого депозита вложить под другой %?»
    Сравниваем рост тела вклада под депозитной ставкой vs под альтернативной.
    """
    st.markdown(f"<div class='sec-title'>{t['dep_invest_section']}</div>", unsafe_allow_html=True)
    st.caption(t["dep_invest_caption"])

    col_a, col_b = st.columns(2)
    with col_a:
        sp_on = st.checkbox(t["invest_sp500"], value=st.session_state.invest_sp500, key="cb_sp500")
    with col_b:
        cu_on = st.checkbox(t["invest_custom"], value=st.session_state.invest_custom, key="cb_custom")

    if sp_on and cu_on:
        st.warning(t["invest_only_one"])
        cu_on = False
    st.session_state.invest_sp500  = sp_on
    st.session_state.invest_custom = cu_on

    if not (sp_on or cu_on):
        return

    if sp_on:
        alt_yield   = SP500_BENCHMARK_RATE
        yield_label = f"S&P 500 — {SP500_BENCHMARK_RATE:g}% {t['invest_rate_label']}"
        st.caption(t.get("sp500_disclaimer", ""))
    else:
        alt_yield = st.slider(
            t["invest_custom_rate"], min_value=0.5, max_value=100.0,
            value=float(st.session_state.invest_custom_rate), step=0.5, key="invest_sldr")
        st.session_state.invest_custom_rate = alt_yield
        yield_label = f"{alt_yield:.1f}% {t['invest_rate_label']}"

    principal = smry["principal"]
    n         = len(smry["schedule"])
    unit      = smry["unit"]
    ppy       = periods_per_year(unit)

    # Рост депозита по периодам (balance_close из schedule)
    dep_vals = [r["balance_close"] for r in smry["schedule"]]

    # Рост альтернативных инвестиций: principal * (1 + r_per_period)^t
    r_alt    = alt_yield / 100 / ppy
    alt_vals = [principal * (1 + r_alt) ** i for i in range(1, n + 1)]

    dep_final = dep_vals[-1]
    alt_final = alt_vals[-1]
    diff      = alt_final - dep_final

    # Метрики
    dm1, dm2, dm3 = st.columns(3)
    dm1.metric(t["dep_invest_yours"], fmt_money(dep_final, sym))
    dm2.metric(t["dep_invest_alt"],   fmt_money(alt_final, sym))
    diff_color = "#06D6A0" if diff >= 0 else "#FF6B6B"
    dm3.metric(
        t["dep_invest_diff"],
        fmt_money(abs(diff), sym),
        delta=f"{'Альтернатива' if diff > 0 else 'Депозит'} выгоднее"
        if diff != 0 else "Одинаково",
    )

    # График
    fig = chart_dep_vs_alternative(df_chart, alt_vals, dep_vals, t, yield_label, sym)
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})


# ─────────────────────────────────────────────────────────────────────────────
#  ТАБЛИЦА + СКАЧИВАНИЕ (общая для кредита и депозита)
# ─────────────────────────────────────────────────────────────────────────────
# ─────────────────────────────────────────────────────────────────────────────
#  AUDIT TRAIL — UI блок в нижней части страницы
# ─────────────────────────────────────────────────────────────────────────────
def _render_audit_trail(t: dict) -> None:
    """Рендерит журнал изменений в expander внизу страницы."""
    log = st.session_state.get("audit_log", [])

    # expander закрыт по умолчанию чтобы не "толкал" интерфейс вниз
    with st.expander(t.get("audit_section", "📋 Audit Trail"),
                      expanded=False):
        st.caption(t.get("audit_caption", ""))

        if not log:
            st.info(t.get("audit_empty", "No entries yet."), icon="ℹ️")
            return

        # Кнопка очистки
        col_clear, _ = st.columns([1, 4])
        with col_clear:
            if st.button(t.get("audit_clear", "Clear Trail"),
                          key="btn_clear_audit"):
                st.session_state.audit_log = []
                st.rerun()

        # Записи — самые свежие сверху, моноширинный шрифт
        # Используем st.markdown с inline-стилями для чистого моноширинного блока
        rendered_lines = []
        for entry in reversed(log):
            ts = entry.get("timestamp", "")
            for ch in entry.get("changes", []):
                line = (
                    f"<div style='font-family:Consolas,\"Courier New\",monospace;"
                    f"font-size:.82rem;color:#E2E8F0;"
                    f"padding:6px 12px;margin:3px 0;"
                    f"background:#0F172A;border-left:3px solid #4FC3F7;"
                    f"border-radius:4px;'>"
                    f"<span style='color:#7DD3FC;font-weight:700'>{ts}</span> &nbsp; "
                    f"{t.get('audit_changed_to', '{field} changed from {old} to {new}').format(field=ch['field_lbl'], old=ch['old_str'], new=ch['new_str'])}"
                )
                if entry.get("impact"):
                    line += (
                        f"<br><span style='color:#94A3B8;font-size:.74rem;"
                        f"margin-left:0;display:block;margin-top:2px'>"
                        f"&nbsp;&nbsp;→ {entry['impact']}</span>"
                    )
                line += "</div>"
                rendered_lines.append(line)

        st.markdown("\n".join(rendered_lines), unsafe_allow_html=True)


def _render_schedule(t, df_d, smry, sym, is_deposit=False):
    """Рендер таблицы графика платежей с кнопками скачивания."""
    st.markdown(f"<div class='sec-title'>{t['section_schedule']}</div>", unsafe_allow_html=True)

    # If the calculation produced no schedule (e.g. a syndicated run where all
    # tranches failed), there is nothing to tabulate or export. Surface the
    # error and skip the table/export buttons rather than crashing the export
    # functions on an empty DataFrame.
    if df_d is None or getattr(df_d, "empty", True) or len(df_d.columns) == 0:
        err = smry.get("synd_empty_error") or t.get(
            "calc_error",
            "⚠️ No schedule to display — the calculation produced no result.")
        st.warning(err)
        return

    # Пояснения колонок
    with st.expander(t["col_explain"], expanded=False):
        if is_deposit:
            st.markdown(
                f"{t['dep_tooltip_interest']}\n\n"
                f"{t['dep_tooltip_balance']}\n\n"
                f"{t['dep_tooltip_payout']}"
            )
        else:
            st.markdown(
                f"{t['tooltip_payment']}\n\n"
                f"{t['tooltip_principal']}\n\n"
                f"{t['tooltip_interest']}"
            )

    # Кнопки скачивания
    prefix   = "dep" if is_deposit else "loan"
    dc1, dc2, dc3, dc4, dc5 = st.columns(5)
    ts = datetime.now().strftime('%Y%m%d_%H%M')

    with dc1:
        st.download_button(
            t["download_excel"],
            data=export_excel(df_d, smry, t, sym),
            file_name=f"{prefix}_{ts}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    with dc2:
        st.download_button(
            t["download_pdf"],
            data=export_pdf(df_d, smry, t, sym),
            file_name=f"{prefix}_{ts}.pdf",
            mime="application/pdf")
    with dc3:
        st.download_button(
            t["download_docx"],
            data=export_docx(df_d, smry, t, sym),
            file_name=f"{prefix}_{ts}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with dc4:
        csv_bytes = df_d.to_csv(index=False, encoding="utf-8-sig").encode("utf-8-sig")
        st.download_button(
            t["download_csv"],
            data=csv_bytes,
            file_name=f"{prefix}_{ts}.csv",
            mime="text/csv")
    with dc5:
        flat_csv_bytes = export_flat_csv(df_d, smry, t, sym, is_deposit=is_deposit)
        st.download_button(
            t.get("download_csv_flat", "⬇️ Flat CSV (ERP)"),
            data=flat_csv_bytes,
            file_name=f"{prefix}_{ts}_flat.csv",
            mime="text/csv",
            help=t.get("download_csv_flat_help",
                        "Clean data for ERP import: numbers only, ISO dates, no totals."),
        )

    # Таблица
    st.dataframe(df_d, use_container_width=True,
                 height=min(640, 42 + 36*len(df_d)))


# ─────────────────────────────────────────────────────────────────────────────
#  ШАБЛОНЫ
# ─────────────────────────────────────────────────────────────────────────────
def _render_templates(t):
    st.markdown(f"**{t['template_examples']}**")
    cols4 = st.columns(4)
    presets = {
        t["mortgage"]:  dict(loan_amount=3_000_000, loan_term=20, term_unit="years",
                              interest_rate=9.5,  scheme="annuity",  is_deposit=False),
        t["car_loan"]:  dict(loan_amount=500_000,  loan_term=5,  term_unit="years",
                              interest_rate=14.0, scheme="annuity",  is_deposit=False),
        t["consumer"]:  dict(loan_amount=100_000,  loan_term=24, term_unit="months",
                              interest_rate=22.0, scheme="classic",  is_deposit=False),
        t["deposit"]:   dict(loan_amount=200_000,  loan_term=12, term_unit="months",
                              interest_rate=12.0, scheme="deposit",  is_deposit=True,
                              deposit_mode="capitalize"),
    }
    for col, (name, params) in zip(cols4, presets.items()):
        if col.button(name, use_container_width=True, key=f"pre_{name}"):
            for k, v in params.items():
                st.session_state[k] = v
            st.rerun()

    st.divider()

    # Сохранение
    cn, cs = st.columns([3, 1])
    with cn:
        tpl_name = st.text_input(t["template_name"], placeholder=t["enter_name"],
                                  label_visibility="collapsed", key="tpl_inp")
    with cs:
        if st.button(t["save_template"], use_container_width=True, key="btn_save"):
            clean = tpl_name.strip()
            if clean:
                # Warn (don't silently clobber) if a template with this name
                # already exists. The user confirms the overwrite by pressing
                # Save a second time; the pending name is tracked in state.
                exists = clean in st.session_state.templates
                pending = st.session_state.get("_tpl_overwrite_pending")
                if exists and pending != clean:
                    st.session_state["_tpl_overwrite_pending"] = clean
                    st.warning(t["template_overwrite_warn"].format(name=clean))
                else:
                    save_tpl(clean)
                    st.session_state.pop("_tpl_overwrite_pending", None)
                    st.success(t["template_saved"])
            else:
                st.warning(t["template_name_empty"])

    # Список
    if st.session_state.templates:
        st.markdown(f"**{t['load_template']}**")
        for name, tpl in list(st.session_state.templates.items()):
            c1, c2, c3 = st.columns([4, 2, 1])
            # Escape the user-supplied name before embedding into HTML to
            # prevent markup injection / interface breakage.
            safe_name = html.escape(str(name))
            safe_saved = html.escape(str(tpl.get("saved_at", "")))
            c1.markdown(
                f"<span class='tpl-badge'>💾 {safe_name}</span> "
                f"<span style='color:#64748B;font-size:.7rem'>{safe_saved}</span>",
                unsafe_allow_html=True)
            # Stable, collision-free widget key from a hash of the name —
            # avoids DuplicateWidgetID and oversized keys from long/unicode names.
            name_key = hashlib.md5(str(name).encode("utf-8")).hexdigest()[:12]
            if c2.button(t["load_template"], key=f"load_{name_key}",
                          use_container_width=True):
                load_tpl(name)
                # Surface confirmation on the *next* run (after st.rerun), since
                # rerun would otherwise wipe the success message immediately.
                st.session_state["_tpl_flash"] = ("loaded", name)
                st.rerun()
            if c3.button(t["delete_template"], key=f"del_{name_key}",
                          use_container_width=True):
                del_tpl(name)
                st.session_state["_tpl_flash"] = ("deleted", name)
                st.rerun()
    else:
        st.info(t["no_templates"])

    # Flash message from a prior load/delete that triggered st.rerun.
    flash = st.session_state.pop("_tpl_flash", None)
    if flash:
        kind, fname = flash
        if kind == "loaded":
            st.success(t["template_loaded"])
        elif kind == "deleted":
            st.info(t.get("template_deleted", t["no_templates"]))


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    main()